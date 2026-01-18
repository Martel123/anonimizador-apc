"""
Legal Document Anonymizer Module (Enhanced)
Hybrid PII detection using regex + spaCy NER for Peruvian legal documents.
Supports three substitution modes: Tokens, Asterisks, Synthetic Data.
"""

import re
import os
import json
import uuid
import random
import logging
from datetime import datetime
from typing import Dict, List, Tuple, Any, Optional, Set
from collections import defaultdict

ALLOWED_EXTENSIONS_ANON = {'.docx', '.pdf'}
MAX_FILE_SIZE_MB = 10
MAX_PDF_PAGES = 50

NLP_MODEL = None

def get_nlp():
    """Lazy load spaCy model."""
    global NLP_MODEL
    if NLP_MODEL is None:
        try:
            import spacy
            NLP_MODEL = spacy.load("es_core_news_md")
            logging.info("spaCy Spanish model loaded successfully")
        except Exception as e:
            logging.warning(f"Could not load spaCy model: {e}. NER detection disabled.")
            NLP_MODEL = False
    return NLP_MODEL if NLP_MODEL else None


PLACEHOLDER_PATTERN = re.compile(r'\{\{[^}]+\}\}')

DNI_PATTERN = re.compile(r'\b[0-9]{8}\b')
RUC_PATTERN = re.compile(r'\b[12][0-9]{10}\b')
EMAIL_PATTERN = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b')
PHONE_PATTERNS = [
    re.compile(r'\+51\s?9[0-9]{8}\b'),
    re.compile(r'\b9[0-9]{8}\b'),
    re.compile(r'\b0[1-9][0-9]\s?[0-9]{6,7}\b'),
    re.compile(r'\([0-9]{2,3}\)\s?[0-9]{6,7}\b'),
]
EXPEDIENTE_PATTERN = re.compile(r'\b[0-9]{5}-[0-9]{4}-[0-9]-[0-9]{4}-[A-Z]{2}-[A-Z]{2}-[0-9]{2}\b', re.IGNORECASE)
CASILLA_PATTERN = re.compile(r'\bcasilla\s+(?:electr[oó]nica\s+)?(?:n[°oº]?\s*)?[0-9]+\b', re.IGNORECASE)
JUZGADO_PATTERN = re.compile(r'\b(?:[0-9]+[°ºo]?\s*)?juzgado\s+(?:de\s+)?(?:paz\s+letrado|familia|civil|penal|laboral|mixto|comercial)[^.]*', re.IGNORECASE)

ACTA_PATTERN = re.compile(r'\b(?:acta\s+(?:de\s+)?(?:conciliaci[oó]n|audiencia|constataci[oó]n|inspecci[oó]n)?)\s*(?:n[°oº]?\s*)?[0-9]+[-/]?[0-9]*\b', re.IGNORECASE)

ENTIDAD_PATTERNS = [
    re.compile(r'\b(?:banco|caja|financiera|cooperativa)\s+[A-Za-záéíóúñÁÉÍÓÚÑ\s]+(?:\s+S\.?A\.?(?:\.?C\.?)?)?', re.IGNORECASE),
    re.compile(r'\b[A-Z][A-Za-záéíóúñÁÉÍÓÚÑ\s&]+\s+(?:S\.A\.C?\.?|S\.R\.L\.?|E\.I\.R\.L\.?|S\.A\.A\.?)\b', re.IGNORECASE),
    re.compile(r'\b(?:notaría|notaria)\s+(?:de\s+)?[A-Za-záéíóúñÁÉÍÓÚÑ\s]+\b', re.IGNORECASE),
]

CUENTA_PATTERNS = [
    re.compile(r'\b(?:cuenta|cta\.?)\s*(?:de\s+ahorros?|corriente)?\s*(?:n[°oº]?\s*)?[0-9]{10,20}\b', re.IGNORECASE),
    re.compile(r'\b(?:CCI|cci)\s*[:.]?\s*[0-9]{20}\b'),
    re.compile(r'\b[0-9]{3}[-\s]?[0-9]{3}[-\s]?[0-9]{10,14}\b'),
]

PLACA_PATTERN = re.compile(r'\b(?:placa\s+(?:de\s+rodaje\s+)?)?[A-Z]{3}[-\s]?[0-9]{3}\b', re.IGNORECASE)

FIRMA_PATTERNS = [
    re.compile(r'(?:firma(?:do)?|firmante|suscrit[oa])[:.\s]+[^\n]{0,50}', re.IGNORECASE),
    re.compile(r'_{5,}', re.IGNORECASE),
    re.compile(r'/s/\s*[A-Za-záéíóúñÁÉÍÓÚÑ\s]+', re.IGNORECASE),
    re.compile(r'(?:FIRMADO\s+(?:POR|DIGITALMENTE))[^\n]*', re.IGNORECASE),
]

SELLO_PATTERNS = [
    re.compile(r'(?:sello|sellado)[:.\s]+[^\n]{0,50}', re.IGNORECASE),
    re.compile(r'\[SELLO[^\]]*\]', re.IGNORECASE),
]

HUELLA_PATTERNS = [
    re.compile(r'(?:huella\s+(?:digital|dactilar))[:.\s]*[^\n]{0,30}', re.IGNORECASE),
    re.compile(r'\[HUELLA[^\]]*\]', re.IGNORECASE),
]

ADDRESS_KEYWORDS = [
    r'\bAv(?:enida)?\.?\s+',
    r'\bJr(?:\.|irón)?\s+',
    r'\bCalle\s+',
    r'\bPsje(?:\.|Pasaje)?\s+',
    r'\bMz(?:\.|anzana)?\s+',
    r'\bLt(?:\.|ote)?\s+',
    r'\bDpto(?:\.|Departamento)?\s+',
    r'\bUrb(?:\.|anizaci[oó]n)?\s+',
    r'\bAA\.?HH\.?\s+',
    r'\bP\.?J\.?\s+',
    r'\bDistrito\s+(?:de\s+)?',
    r'\bProvincia\s+(?:de\s+)?',
]

ADDRESS_PATTERN = re.compile(
    r'(' + '|'.join(ADDRESS_KEYWORDS) + r')[A-Za-záéíóúñÁÉÍÓÚÑ0-9\s,.\-°º#]+(?=[\.\n,;]|$)',
    re.IGNORECASE
)

NAME_CONTEXT_KEYWORDS = [
    r'(?:señor|señora|sr\.|sra\.)\s+',
    r'(?:don|doña)\s+',
    r'(?:el|la)\s+(?:demandante|demandado|demandada)\s+',
    r'(?:el|la)\s+(?:solicitante|invitado|invitada)\s+',
    r'identificad[oa]\s+con\s+(?:DNI|documento)',
    r'(?:abogad[oa]|letrad[oa])\s+',
    r'(?:el|la)\s+(?:menor|menores?)\s+',
    r'(?:madre|padre|hijo|hija)\s+',
    r'(?:cónyuge|esposo|esposa)\s+',
    r'(?:testigo|perito)\s+',
]

NAME_PATTERN = re.compile(
    r'(?:' + '|'.join(NAME_CONTEXT_KEYWORDS) + r')([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+){1,4})',
    re.IGNORECASE
)

UPPERCASE_NAME_PATTERN = re.compile(r'\b([A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,}){1,4})\b')

MONEY_PATTERN = re.compile(r'(?:S/\.?\s*|US\$\s*|\$\s*|PEN\s+|USD\s+)[0-9]{1,3}(?:[,\'][0-9]{3})*(?:\.[0-9]{2})?', re.IGNORECASE)

EXCLUDED_WORDS = {
    'SEÑOR', 'SEÑORA', 'JUEZ', 'JUEZA', 'DEMANDA', 'DEMANDANTE', 'DEMANDADO',
    'FISCAL', 'FISCAL', 'CÓDIGO', 'CIVIL', 'PENAL', 'PROCESAL', 'CONSTITUCIONAL',
    'ARTÍCULO', 'ARTICULO', 'INCISO', 'NUMERAL', 'RESOLUCIÓN', 'RESOLUCION',
    'DECRETO', 'LEY', 'REGLAMENTO', 'EXPEDIENTE', 'JUZGADO', 'SALA', 'CORTE',
    'SUPREMA', 'SUPERIOR', 'PODER', 'JUDICIAL', 'REPÚBLICA', 'REPUBLICA',
    'PERÚ', 'PERU', 'ESTADO', 'CONSTITUCIÓN', 'CONSTITUCION', 'LIMA', 'CALLAO',
    'AUTO', 'SENTENCIA', 'APELACIÓN', 'APELACION', 'CASACIÓN', 'CASACION',
    'RECURSO', 'ESCRITO', 'RAZÓN', 'RAZON', 'SOCIAL', 'DENUNCIA', 'DENUNCIA',
    'MINISTERIO', 'PÚBLICO', 'PUBLICO', 'DEFENSORÍA', 'DEFENSORIA', 'PUEBLO',
    'QUE', 'DEL', 'LOS', 'LAS', 'POR', 'CON', 'SIN', 'PARA', 'DESDE', 'HASTA',
    'SOBRE', 'ANTE', 'CONTRA', 'ENTRE', 'MEDIANTE', 'SEGÚN', 'SEGUN',
}

PERUVIAN_FIRST_NAMES = [
    'Juan', 'Carlos', 'José', 'Luis', 'Miguel', 'Pedro', 'Jorge', 'Fernando',
    'Roberto', 'Ricardo', 'Manuel', 'Francisco', 'Antonio', 'Eduardo', 'Daniel',
    'María', 'Ana', 'Rosa', 'Carmen', 'Patricia', 'Elizabeth', 'Claudia',
    'Sandra', 'Martha', 'Gloria', 'Teresa', 'Luz', 'Isabel', 'Silvia', 'Beatriz',
    'Alejandro', 'Adrián', 'Ángel', 'Arturo', 'César', 'Diego', 'Ernesto',
    'Felipe', 'Gabriel', 'Héctor', 'Iván', 'Javier', 'Martín', 'Óscar', 'Pablo',
    'Raúl', 'Sergio', 'Víctor', 'Walter', 'Alberto', 'Alfredo', 'Andrés',
    'Andrea', 'Angélica', 'Carla', 'Carolina', 'Cecilia', 'Diana', 'Elena',
    'Fabiola', 'Gabriela', 'Irma', 'Julia', 'Karen', 'Laura', 'Liliana',
    'Lorena', 'Lucía', 'Margarita', 'Mónica', 'Natalia', 'Norma', 'Paola',
    'Pilar', 'Rocío', 'Sonia', 'Susana', 'Verónica', 'Victoria', 'Yolanda'
]

PERUVIAN_LAST_NAMES = [
    'García', 'Rodríguez', 'Martínez', 'López', 'González', 'Hernández',
    'Pérez', 'Sánchez', 'Ramírez', 'Torres', 'Flores', 'Rivera', 'Gómez',
    'Díaz', 'Reyes', 'Cruz', 'Morales', 'Ortiz', 'Gutiérrez', 'Chávez',
    'Rojas', 'Mendoza', 'Vargas', 'Castro', 'Jiménez', 'Ruiz', 'Vásquez',
    'Medina', 'Paredes', 'Espinoza', 'Quispe', 'Huamán', 'Ccama', 'Mamani',
    'Condori', 'Apaza', 'Cusi', 'Chambi', 'Lima', 'Ramos', 'Castillo',
    'Silva', 'Fernández', 'Campos', 'Delgado', 'Vega', 'Aguilar', 'Salazar',
    'Herrera', 'Navarro', 'Ponce', 'Villanueva', 'Acosta', 'Miranda', 'León'
]

PERUVIAN_STREETS = [
    'Las Flores', 'Los Pinos', 'San Martín', 'Grau', 'Bolognesi', 'Tacna',
    'Arequipa', 'Lima', 'Cusco', 'La Marina', 'Colonial', 'Venezuela',
    'Brasil', 'Argentina', 'Uruguay', 'Petit Thouars', 'Salaverry', 'Javier Prado',
    'Angamos', 'Benavides', 'Larco', 'Pardo', 'Sucre', 'La Paz', 'Los Álamos'
]

PERUVIAN_DISTRICTS = [
    'Miraflores', 'San Isidro', 'Surco', 'La Molina', 'San Borja', 'Barranco',
    'Jesús María', 'Lince', 'Magdalena', 'Pueblo Libre', 'San Miguel',
    'Chorrillos', 'Villa El Salvador', 'San Juan de Lurigancho', 'Comas',
    'Los Olivos', 'Independencia', 'Ate', 'Santa Anita', 'El Agustino'
]


class SubstitutionMode:
    TOKENS = 'tokens'


class EntityMapping:
    """Maintains consistent mapping of entities to token placeholders."""
    
    def __init__(self, mode: str = SubstitutionMode.TOKENS):
        self.mode = SubstitutionMode.TOKENS
        self.mappings: Dict[str, Dict[str, str]] = defaultdict(dict)
        self.counters: Dict[str, int] = defaultdict(int)
        self.reverse_mappings: Dict[str, Dict[str, str]] = defaultdict(dict)
    
    def get_substitute(self, entity_type: str, value: str) -> str:
        """Get or create a token placeholder for a value."""
        normalized = value.strip()
        normalized_key = normalized.upper()
        
        if normalized_key in self.mappings[entity_type]:
            return self.mappings[entity_type][normalized_key]
        
        self.counters[entity_type] += 1
        substitute = f"{{{{{entity_type}_{self.counters[entity_type]}}}}}"
        
        self.mappings[entity_type][normalized_key] = substitute
        self.reverse_mappings[entity_type][substitute] = {
            'original_masked': self._mask_value(value, entity_type),
            'original': None
        }
        
        return substitute
    
    def _mask_value(self, value: str, entity_type: str) -> str:
        """Create a masked version of the value for the report."""
        if len(value) <= 4:
            return '*' * len(value)
        
        if entity_type in ['DNI', 'RUC']:
            return value[:2] + '*' * (len(value) - 4) + value[-2:]
        elif entity_type == 'EMAIL':
            parts = value.split('@')
            if len(parts) == 2:
                user = parts[0][:2] + '***'
                return f"{user}@{parts[1]}"
        elif entity_type == 'TELEFONO':
            return value[:3] + '***' + value[-2:]
        elif entity_type in ['NOMBRE', 'PERSONA']:
            words = value.split()
            if len(words) >= 2:
                return words[0][:2] + '*** ' + words[-1][:2] + '***'
            return value[:2] + '***'
        
        if len(value) > 8:
            return value[:3] + '...' + value[-3:]
        return value[:2] + '***'
    
    def get_summary(self) -> Dict[str, int]:
        """Get count of entities by type."""
        return {k: len(v) for k, v in self.mappings.items() if v}
    
    def get_replacements_for_report(self) -> Dict[str, List[Dict[str, str]]]:
        """Get list of replacements for the report."""
        result = {}
        for entity_type, substitutes in self.reverse_mappings.items():
            if substitutes:
                result[entity_type] = [
                    {"placeholder": sub, "masked_original": data['original_masked']}
                    for sub, data in substitutes.items()
                ]
        return result
    
    def get_mapping_dict(self) -> Dict[str, Dict[str, str]]:
        """Get full mapping dictionary for export (original -> substitute)."""
        result = {}
        for entity_type, mapping in self.mappings.items():
            if mapping:
                result[entity_type] = dict(mapping)
        return result
    
    def to_dict(self) -> Dict[str, Any]:
        """Serialize mapping state to dictionary for storage."""
        return {
            'mode': self.mode,
            'mappings': dict(self.mappings),
            'counters': dict(self.counters),
            'reverse_mappings': {k: dict(v) for k, v in self.reverse_mappings.items()}
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'EntityMapping':
        """Restore mapping state from dictionary."""
        mapping = cls()
        mapping.mappings = defaultdict(dict, {k: dict(v) for k, v in data.get('mappings', {}).items()})
        mapping.counters = defaultdict(int, data.get('counters', {}))
        mapping.reverse_mappings = defaultdict(dict, {k: dict(v) for k, v in data.get('reverse_mappings', {}).items()})
        return mapping


def find_existing_placeholders(text: str) -> List[Tuple[int, int]]:
    """Find all existing {{...}} placeholders to exclude from detection."""
    positions = []
    for match in PLACEHOLDER_PATTERN.finditer(text):
        positions.append((match.start(), match.end()))
    return positions


def is_in_placeholder(start: int, end: int, placeholder_positions: List[Tuple[int, int]]) -> bool:
    """Check if a position overlaps with an existing placeholder."""
    for p_start, p_end in placeholder_positions:
        if not (end <= p_start or start >= p_end):
            return True
    return False


def detect_entities_regex(text: str, placeholder_positions: List[Tuple[int, int]]) -> List[Tuple[str, str, int, int, float]]:
    """
    Detect PII entities using regex patterns.
    Returns list of (entity_type, value, start, end, confidence).
    """
    entities = []
    
    for match in EXPEDIENTE_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            entities.append(('EXPEDIENTE', match.group(), match.start(), match.end(), 0.95))
    
    for match in CASILLA_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            entities.append(('CASILLA', match.group(), match.start(), match.end(), 0.95))
    
    for match in JUZGADO_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            entities.append(('JUZGADO', match.group(), match.start(), match.end(), 0.90))
    
    for match in RUC_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            value = match.group()
            if not _is_money_context(text, match.start(), match.end()):
                entities.append(('RUC', value, match.start(), match.end(), 0.90))
    
    ruc_positions = {(e[2], e[3]) for e in entities if e[0] == 'RUC'}
    for match in DNI_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            value = match.group()
            if (match.start(), match.end()) not in ruc_positions:
                if not _is_money_context(text, match.start(), match.end()):
                    if not _is_date_context(text, match.start(), match.end()):
                        entities.append(('DNI', value, match.start(), match.end(), 0.85))
    
    for match in EMAIL_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            entities.append(('EMAIL', match.group(), match.start(), match.end(), 0.95))
    
    for pattern in PHONE_PATTERNS:
        for match in pattern.finditer(text):
            if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
                value = match.group()
                if not any(e[2] == match.start() for e in entities):
                    entities.append(('TELEFONO', value, match.start(), match.end(), 0.85))
    
    for match in ADDRESS_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            value = match.group().strip()
            if len(value) > 10:
                entities.append(('DIRECCION', value, match.start(), match.end(), 0.75))
    
    for match in NAME_PATTERN.finditer(text):
        name = match.group(1).strip()
        name_start = match.start() + match.group().index(name)
        name_end = name_start + len(name)
        if not is_in_placeholder(name_start, name_end, placeholder_positions):
            if len(name) > 5 and ' ' in name:
                if not _is_excluded_name(name):
                    entities.append(('PERSONA', name, name_start, name_end, 0.80))
    
    for match in UPPERCASE_NAME_PATTERN.finditer(text):
        value = match.group(1)
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            words = value.split()
            if 2 <= len(words) <= 5:
                if not _is_excluded_name(value):
                    if not any(e[2] <= match.start() < e[3] or e[2] < match.end() <= e[3] for e in entities):
                        entities.append(('PERSONA', value, match.start(), match.end(), 0.60))
    
    for match in ACTA_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            entities.append(('ACTA', match.group(), match.start(), match.end(), 0.90))
    
    for pattern in ENTIDAD_PATTERNS:
        for match in pattern.finditer(text):
            if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
                value = match.group().strip()
                if len(value) > 5:
                    if not any(e[2] == match.start() for e in entities):
                        entities.append(('ENTIDAD', value, match.start(), match.end(), 0.85))
    
    for pattern in CUENTA_PATTERNS:
        for match in pattern.finditer(text):
            if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
                if not any(e[2] == match.start() for e in entities):
                    entities.append(('CUENTA', match.group(), match.start(), match.end(), 0.90))
    
    for match in PLACA_PATTERN.finditer(text):
        if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
            value = match.group()
            if re.search(r'[A-Z]{3}[-\s]?[0-9]{3}', value, re.IGNORECASE):
                entities.append(('PLACA', value, match.start(), match.end(), 0.90))
    
    for pattern in FIRMA_PATTERNS:
        for match in pattern.finditer(text):
            if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
                if not any(e[2] == match.start() for e in entities):
                    entities.append(('FIRMA', match.group(), match.start(), match.end(), 0.70))
    
    for pattern in SELLO_PATTERNS:
        for match in pattern.finditer(text):
            if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
                if not any(e[2] == match.start() for e in entities):
                    entities.append(('SELLO', match.group(), match.start(), match.end(), 0.70))
    
    for pattern in HUELLA_PATTERNS:
        for match in pattern.finditer(text):
            if not is_in_placeholder(match.start(), match.end(), placeholder_positions):
                if not any(e[2] == match.start() for e in entities):
                    entities.append(('HUELLA', match.group(), match.start(), match.end(), 0.70))
    
    return entities


def detect_entities_ner(text: str, placeholder_positions: List[Tuple[int, int]]) -> List[Tuple[str, str, int, int, float]]:
    """
    Detect entities using spaCy NER.
    Returns list of (entity_type, value, start, end, confidence).
    """
    nlp = get_nlp()
    if nlp is None:
        return []
    
    entities = []
    
    try:
        doc = nlp(text)
        
        for ent in doc.ents:
            if is_in_placeholder(ent.start_char, ent.end_char, placeholder_positions):
                continue
            
            if ent.label_ == 'PER':
                if len(ent.text) > 3 and ' ' in ent.text:
                    if not _is_excluded_name(ent.text):
                        confidence = 0.75
                        context_start = max(0, ent.start_char - 50)
                        context = text[context_start:ent.start_char].lower()
                        if any(kw in context for kw in ['identificado', 'demandante', 'demandado', 'señor', 'señora', 'don', 'doña', 'abogado', 'menor', 'madre', 'padre']):
                            confidence = 0.90
                        entities.append(('PERSONA', ent.text, ent.start_char, ent.end_char, confidence))
            
            elif ent.label_ == 'LOC':
                if len(ent.text) > 5:
                    context_start = max(0, ent.start_char - 30)
                    context = text[context_start:ent.start_char].lower()
                    if any(kw in context for kw in ['domicilio', 'reside', 'ubicado', 'dirección', 'av.', 'jr.', 'calle']):
                        entities.append(('DIRECCION', ent.text, ent.start_char, ent.end_char, 0.70))
    
    except Exception as e:
        logging.error(f"NER detection error: {e}")
    
    return entities


def _is_excluded_name(value: str) -> bool:
    """Check if value should be excluded (legal terms, common words)."""
    words = value.upper().split()
    if all(w in EXCLUDED_WORDS for w in words):
        return True
    if len(words) == 1 and words[0] in EXCLUDED_WORDS:
        return True
    return False


def _is_money_context(text: str, start: int, end: int) -> bool:
    """Check if the number is in a money context."""
    context_start = max(0, start - 20)
    context = text[context_start:start].lower()
    money_indicators = ['s/', 's/.', 'us$', '$', 'soles', 'dólares', 'dolares', 'monto', 'suma', 'pago', 'pen', 'usd']
    return any(ind in context for ind in money_indicators)


def _is_date_context(text: str, start: int, end: int) -> bool:
    """Check if the number looks like a date."""
    value = text[start:end]
    context_end = min(len(text), end + 15)
    after = text[end:context_end].lower().strip()
    
    if after.startswith(('/', '-')) and len(after) > 1 and after[1:3].isdigit():
        return True
    
    try:
        if int(value) > 31000000:
            return False
    except:
        pass
    
    return False


def detect_entities_hybrid(text: str) -> Tuple[List[Tuple[str, str, int, int, float]], List[Tuple[str, str, int, int, float]]]:
    """
    Hybrid entity detection using both regex and NER.
    Returns (confirmed_entities, needs_review_entities).
    """
    placeholder_positions = find_existing_placeholders(text)
    
    regex_entities = detect_entities_regex(text, placeholder_positions)
    ner_entities = detect_entities_ner(text, placeholder_positions)
    
    all_entities = []
    seen_positions = set()
    
    for entity in regex_entities:
        pos_key = (entity[2], entity[3])
        if pos_key not in seen_positions:
            all_entities.append(entity)
            seen_positions.add(pos_key)
    
    for entity in ner_entities:
        pos_key = (entity[2], entity[3])
        overlaps = False
        for seen_start, seen_end in seen_positions:
            if not (entity[3] <= seen_start or entity[2] >= seen_end):
                overlaps = True
                break
        if not overlaps:
            all_entities.append(entity)
            seen_positions.add(pos_key)
    
    all_entities.sort(key=lambda x: (x[2], -(x[3] - x[2])))
    
    confirmed = []
    needs_review = []
    
    for entity in all_entities:
        entity_type, value, start, end, confidence = entity
        
        if entity_type == 'PERSONA' and confidence < 0.80:
            needs_review.append(entity)
        else:
            confirmed.append(entity)
    
    return confirmed, needs_review


def replace_entities(text: str, entities: List[Tuple[str, str, int, int, float]], mapping: EntityMapping) -> str:
    """Replace all detected entities with their substitutes."""
    non_overlapping = []
    for entity in entities:
        overlaps = False
        for existing in non_overlapping:
            if not (entity[3] <= existing[2] or entity[2] >= existing[3]):
                overlaps = True
                break
        if not overlaps:
            non_overlapping.append(entity)
    
    non_overlapping.sort(key=lambda x: -x[2])
    
    result = text
    for entity_type, value, start, end, _ in non_overlapping:
        substitute = mapping.get_substitute(entity_type, value)
        result = result[:start] + substitute + result[end:]
    
    return result


def post_verification(text: str, original_entities: List[Tuple[str, str, int, int, float]]) -> List[Tuple[str, str, int, int, float]]:
    """
    Post-verification scan for any remaining PERSON entities.
    Returns entities that were missed in initial pass.
    """
    placeholder_positions = find_existing_placeholders(text)
    ner_entities = detect_entities_ner(text, placeholder_positions)
    
    missed = []
    for entity in ner_entities:
        if entity[0] == 'PERSONA':
            value_upper = entity[1].upper()
            if not any(value_upper == e[1].upper() for e in original_entities):
                if not entity[1].startswith('{{'):
                    missed.append(entity)
    
    return missed


def final_pii_scan(text: str) -> List[Dict[str, Any]]:
    """
    Final scan to detect any remaining PII in anonymized text.
    Returns list of remaining PII for blocking download if found.
    """
    placeholder_positions = find_existing_placeholders(text)
    remaining_pii = []
    
    regex_entities = detect_entities_regex(text, placeholder_positions)
    for entity in regex_entities:
        if not entity[1].startswith('{{'):
            remaining_pii.append({
                'type': entity[0],
                'value': entity[1],
                'start': entity[2],
                'end': entity[3]
            })
    
    ner_entities = detect_entities_ner(text, placeholder_positions)
    for entity in ner_entities:
        if not entity[1].startswith('{{') and entity[4] >= 0.80:
            if not any(r['value'] == entity[1] for r in remaining_pii):
                remaining_pii.append({
                    'type': entity[0],
                    'value': entity[1],
                    'start': entity[2],
                    'end': entity[3]
                })
    
    return remaining_pii


def replace_in_paragraph_run_aware(para, value: str, substitute: str) -> bool:
    """
    Replace value with substitute in paragraph, handling text split across runs.
    Returns True if replacement was made.
    """
    full_text = para.text
    if value not in full_text:
        return False
    
    for run in para.runs:
        if value in run.text:
            run.text = run.text.replace(value, substitute)
    
    if value in para.text:
        run_texts = []
        run_indices = []
        for i, run in enumerate(para.runs):
            run_texts.append(run.text)
            run_indices.append(i)
        
        combined = ''.join(run_texts)
        if value in combined:
            start_idx = combined.find(value)
            end_idx = start_idx + len(value)
            
            current_pos = 0
            for i, run_text in enumerate(run_texts):
                run_start = current_pos
                run_end = current_pos + len(run_text)
                
                if run_start < end_idx and run_end > start_idx:
                    overlap_start = max(start_idx - run_start, 0)
                    overlap_end = min(end_idx - run_start, len(run_text))
                    
                    new_text = run_text[:overlap_start]
                    if run_start <= start_idx < run_end:
                        new_text += substitute
                    new_text += run_text[overlap_end:]
                    para.runs[i].text = new_text
                
                current_pos = run_end
    
    return True


def replace_value_in_docx(doc, value: str, substitute: str, replace_all: bool = True) -> int:
    """
    Replace value with substitute throughout the DOCX document.
    Handles paragraphs, tables, headers, footers.
    Returns count of replacements made.
    """
    count = 0
    
    for para in doc.paragraphs:
        if replace_in_paragraph_run_aware(para, value, substitute):
            count += 1
            if not replace_all:
                return count
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_paragraph_run_aware(para, value, substitute):
                        count += 1
                        if not replace_all:
                            return count
    
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                if replace_in_paragraph_run_aware(para, value, substitute):
                    count += 1
                    if not replace_all:
                        return count
        if section.footer:
            for para in section.footer.paragraphs:
                if replace_in_paragraph_run_aware(para, value, substitute):
                    count += 1
                    if not replace_all:
                        return count
    
    return count


def anonymize_text(text: str, mode: str = SubstitutionMode.TOKENS, strict_mode: bool = True) -> Tuple[str, Dict[str, Any], EntityMapping, List[Dict]]:
    """
    Anonymize text content.
    Returns (anonymized_text, summary, mapping, needs_review_list).
    """
    mapping = EntityMapping(mode=mode)
    
    confirmed, needs_review = detect_entities_hybrid(text)
    
    anonymized = replace_entities(text, confirmed, mapping)
    
    if strict_mode:
        missed = post_verification(anonymized, confirmed)
        needs_review.extend(missed)
    
    needs_review_list = [
        {
            'id': str(uuid.uuid4())[:8],
            'type': e[0],
            'value': e[1],
            'start': e[2],
            'end': e[3],
            'confidence': e[4],
            'context': text[max(0, e[2]-30):min(len(text), e[3]+30)]
        }
        for e in needs_review
    ]
    
    summary = {
        'entities_found': mapping.get_summary(),
        'total_entities': sum(mapping.get_summary().values()),
        'replacements': mapping.get_replacements_for_report(),
        'mode': mode,
        'strict_mode': strict_mode,
        'needs_review_count': len(needs_review_list)
    }
    
    return anonymized, summary, mapping, needs_review_list


def anonymize_docx(file_path: str, mode: str = SubstitutionMode.TOKENS, strict_mode: bool = True) -> Tuple[Any, Dict[str, Any], EntityMapping, List[Dict]]:
    """
    Anonymize a DOCX file.
    Returns the anonymized Document object, summary dict, entity mapping, and needs_review list.
    """
    from docx import Document
    
    doc = Document(file_path)
    
    all_text_parts = []
    for para in doc.paragraphs:
        all_text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text_parts.append(para.text)
    
    full_text = '\n'.join(all_text_parts)
    
    mapping = EntityMapping(mode=mode)
    confirmed, needs_review = detect_entities_hybrid(full_text)
    
    def replace_in_paragraph(para, entities_list):
        for entity_type, value, _, _, _ in entities_list:
            if value in para.text:
                substitute = mapping.get_substitute(entity_type, value)
                for run in para.runs:
                    if value in run.text:
                        run.text = run.text.replace(value, substitute)
                if value in para.text:
                    para.text = para.text.replace(value, substitute)
    
    for para in doc.paragraphs:
        replace_in_paragraph(para, confirmed)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, confirmed)
    
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                replace_in_paragraph(para, confirmed)
        if section.footer:
            for para in section.footer.paragraphs:
                replace_in_paragraph(para, confirmed)
    
    needs_review_list = [
        {
            'id': str(uuid.uuid4())[:8],
            'type': e[0],
            'value': e[1],
            'start': e[2],
            'end': e[3],
            'confidence': e[4],
            'context': full_text[max(0, e[2]-30):min(len(full_text), e[3]+30)]
        }
        for e in needs_review
    ]
    
    summary = {
        'entities_found': mapping.get_summary(),
        'total_entities': sum(mapping.get_summary().values()),
        'replacements': mapping.get_replacements_for_report(),
        'mode': mode,
        'strict_mode': strict_mode,
        'needs_review_count': len(needs_review_list)
    }
    
    return doc, summary, mapping, needs_review_list


def anonymize_pdf(file_path: str, mode: str = SubstitutionMode.TOKENS, strict_mode: bool = True) -> Tuple[Optional[str], Dict[str, Any], EntityMapping, bool, List[Dict]]:
    """
    Anonymize a PDF file.
    Returns (text, summary, mapping, is_scanned, needs_review_list).
    For scanned PDFs, returns is_scanned=True.
    """
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF processing")
    
    reader = PdfReader(file_path)
    
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError(f"El PDF excede el límite de {MAX_PDF_PAGES} páginas")
    
    all_text = []
    for page in reader.pages:
        text = page.extract_text() or ''
        all_text.append(text)
    
    full_text = '\n\n'.join(all_text)
    
    if len(full_text.strip()) < 100:
        return None, {}, EntityMapping(), True, []
    
    anonymized_text, summary, mapping, needs_review = anonymize_text(full_text, mode, strict_mode)
    summary['page_count'] = len(reader.pages)
    
    return anonymized_text, summary, mapping, False, needs_review


def apply_review_decisions(text: str, decisions: Dict[str, bool], pending_entities: List[Dict], mapping: EntityMapping) -> str:
    """
    Apply user review decisions.
    decisions: {entity_id: True (anonymize) or False (keep original)}
    """
    entities_to_apply = []
    
    for entity in pending_entities:
        entity_id = entity['id']
        if decisions.get(entity_id, False):
            entities_to_apply.append((
                entity['type'],
                entity['value'],
                entity['start'],
                entity['end'],
                entity['confidence']
            ))
    
    entities_to_apply.sort(key=lambda x: -x[2])
    
    result = text
    for entity_type, value, start, end, _ in entities_to_apply:
        if value in result:
            substitute = mapping.get_substitute(entity_type, value)
            result = result.replace(value, substitute, 1)
    
    return result


def save_anonymized_docx(doc, output_path: str):
    """Save the anonymized DOCX document."""
    doc.save(output_path)


def extract_docx_text(file_path: str) -> str:
    """Extract all text from a DOCX file for preview."""
    from docx import Document
    doc = Document(file_path)
    
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = ' '.join(p.text for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(' | '.join(row_text))
    
    return '\n'.join(text_parts)


def create_anonymized_pdf(text: str, output_path: str):
    """Create a new PDF with the anonymized text."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        spaceAfter=12
    )
    
    story = []
    paragraphs = text.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            clean_text = para.replace('\n', ' ').strip()
            clean_text = clean_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(clean_text, normal_style))
            story.append(Spacer(1, 6))
    
    doc.build(story)


def generate_report(summary: Dict[str, Any], original_filename: str, file_type: str) -> Dict[str, Any]:
    """Generate a detailed anonymization report."""
    report = {
        'fecha_procesamiento': datetime.now().isoformat(),
        'archivo_original': original_filename,
        'tipo_archivo': file_type,
        'modo_sustitucion': summary.get('mode', 'tokens'),
        'modo_estricto': summary.get('strict_mode', True),
        'resumen': {
            'total_entidades_detectadas': summary.get('total_entities', 0),
            'entidades_por_tipo': summary.get('entities_found', {}),
            'entidades_pendientes_revision': summary.get('needs_review_count', 0),
        },
        'reemplazos': summary.get('replacements', {}),
        'advertencias': [],
        'version': '2.0.0'
    }
    
    if 'DIRECCION' in summary.get('entities_found', {}):
        report['advertencias'].append(
            'Las direcciones fueron detectadas mediante heurísticas. '
            'Revise el documento para confirmar la correcta anonimización.'
        )
    
    if 'PERSONA' in summary.get('entities_found', {}):
        report['advertencias'].append(
            'Los nombres fueron detectados mediante NER y patrones. '
            'Pueden existir nombres adicionales no detectados.'
        )
    
    if summary.get('needs_review_count', 0) > 0:
        report['advertencias'].append(
            f'Se detectaron {summary["needs_review_count"]} entidades que requieren revisión manual.'
        )
    
    return report


def generate_report_txt(report: Dict[str, Any]) -> str:
    """Generate a text version of the report."""
    lines = [
        "=" * 60,
        "REPORTE DE ANONIMIZACIÓN",
        "=" * 60,
        "",
        f"Fecha de procesamiento: {report['fecha_procesamiento']}",
        f"Archivo original: {report['archivo_original']}",
        f"Tipo de archivo: {report['tipo_archivo']}",
        f"Modo de sustitución: {report.get('modo_sustitucion', 'tokens')}",
        f"Modo estricto: {'Sí' if report.get('modo_estricto', True) else 'No'}",
        "",
        "-" * 40,
        "RESUMEN DE ENTIDADES DETECTADAS",
        "-" * 40,
        f"Total de entidades: {report['resumen']['total_entidades_detectadas']}",
        f"Pendientes de revisión: {report['resumen'].get('entidades_pendientes_revision', 0)}",
        "",
    ]
    
    for entity_type, count in report['resumen']['entidades_por_tipo'].items():
        lines.append(f"  {entity_type}: {count}")
    
    lines.extend(["", "-" * 40, "REEMPLAZOS REALIZADOS", "-" * 40])
    
    for entity_type, replacements in report['reemplazos'].items():
        lines.append(f"\n{entity_type}:")
        for r in replacements:
            lines.append(f"  {r['placeholder']} <- {r['masked_original']}")
    
    if report['advertencias']:
        lines.extend(["", "-" * 40, "ADVERTENCIAS", "-" * 40])
        for adv in report['advertencias']:
            lines.append(f"  * {adv}")
    
    lines.extend(["", "=" * 60])
    
    return '\n'.join(lines)


def generate_mapping_csv(mapping: EntityMapping) -> str:
    """Generate CSV content for the mapping dictionary."""
    import csv
    import io
    
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Tipo', 'Original (enmascarado)', 'Sustitución'])
    
    for entity_type, substitutes in mapping.reverse_mappings.items():
        for substitute, data in substitutes.items():
            writer.writerow([entity_type, data['original_masked'], substitute])
    
    return output.getvalue()


def cleanup_old_files(directory: str, max_age_minutes: int = 30):
    """Remove files older than max_age_minutes from directory."""
    import time
    
    if not os.path.exists(directory):
        return
    
    current_time = time.time()
    max_age_seconds = max_age_minutes * 60
    
    for filename in os.listdir(directory):
        filepath = os.path.join(directory, filename)
        if os.path.isfile(filepath):
            file_age = current_time - os.path.getmtime(filepath)
            if file_age > max_age_seconds:
                try:
                    os.remove(filepath)
                    logging.info(f"Cleaned up old file: {filepath}")
                except Exception as e:
                    logging.error(f"Error removing file {filepath}: {e}")
