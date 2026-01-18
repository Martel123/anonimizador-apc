"""
Robust Anonymizer Pipeline
Hardened module with fallbacks at every stage to ensure the pipeline never crashes.
All errors are controlled and returned with proper error codes.
"""

import os
import re
import uuid
import logging
import traceback
import zipfile
import tempfile
import shutil
from typing import Dict, List, Tuple, Any, Optional
from datetime import datetime


ERROR_CODES = {
    'UPLOAD_ERROR': 'Error al subir el archivo',
    'SIZE_ERROR': 'Archivo demasiado grande',
    'FORMAT_ERROR': 'Formato de archivo inválido',
    'PARSE_ERROR': 'No se pudo leer el documento',
    'DOCX_PARSE_ERROR': 'No se pudo leer el archivo DOCX',
    'PDF_PARSE_ERROR': 'No se pudo leer el archivo PDF',
    'PDF_SCANNED': 'PDF escaneado sin texto (OCR no incluido)',
    'NLP_ERROR': 'Error en el procesamiento de lenguaje',
    'DETECTION_ERROR': 'Error en la detección de datos sensibles',
    'ANONYMIZE_ERROR': 'Error al anonimizar el documento',
    'OUTPUT_ERROR': 'Error al generar el documento de salida',
    'UNKNOWN_ERROR': 'Error desconocido'
}

MAX_FILE_SIZE_MB = 10
MAX_PDF_PAGES = 50
ALLOWED_EXTENSIONS = {'.docx', '.pdf'}


def generate_error_id() -> str:
    """Generate a short unique error ID for tracking."""
    return uuid.uuid4().hex[:8].upper()


def create_controlled_error(code: str, error_id: str, details: str = None, stage: int = 0) -> Dict:
    """Create a controlled error response."""
    message = ERROR_CODES.get(code, ERROR_CODES['UNKNOWN_ERROR'])
    return {
        'ok': False,
        'error_id': error_id,
        'code': code,
        'stage': stage,
        'message_public': f"{message}. Error ID: {error_id}",
        'details': details
    }


def log_error(error_id: str, code: str, file_type: str, file_size: int, stage: int, error: Exception):
    """Log error without exposing document content."""
    logging.error(f"[{error_id}] ANONYMIZER_ERROR | code={code} | stage=ETAPA_{stage} | type={file_type} | size={file_size}")
    logging.error(f"[{error_id}] Exception: {type(error).__name__}: {str(error)[:500]}")
    logging.error(f"[{error_id}] Traceback:\n{traceback.format_exc()}")


def validate_file_content(file_data: bytes, ext: str) -> Tuple[bool, Optional[str]]:
    """
    STAGE 0: Validate file content matches expected format.
    Returns (is_valid, error_code).
    """
    try:
        if ext == '.docx':
            if len(file_data) < 4:
                return False, 'FORMAT_ERROR'
            if file_data[:4] != b'PK\x03\x04':
                return False, 'FORMAT_ERROR'
            try:
                import io
                with zipfile.ZipFile(io.BytesIO(file_data)) as zf:
                    if 'word/document.xml' not in zf.namelist():
                        return False, 'FORMAT_ERROR'
            except zipfile.BadZipFile:
                return False, 'FORMAT_ERROR'
            return True, None
        
        elif ext == '.pdf':
            if len(file_data) < 5:
                return False, 'FORMAT_ERROR'
            if not file_data[:5].startswith(b'%PDF-'):
                return False, 'FORMAT_ERROR'
            return True, None
        
        return False, 'FORMAT_ERROR'
    except Exception:
        return False, 'FORMAT_ERROR'


def extract_text_docx_primary(file_path: str) -> Tuple[Optional[str], Optional[Any]]:
    """Primary DOCX text extractor using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text_parts.append(para.text)
        return '\n'.join(text_parts), doc
    except Exception as e:
        logging.warning(f"Primary DOCX extractor failed: {e}")
        return None, None


def extract_text_docx_fallback(file_path: str) -> Tuple[Optional[str], None]:
    """Fallback DOCX text extractor using mammoth."""
    try:
        import mammoth
        with open(file_path, 'rb') as f:
            result = mammoth.extract_raw_text(f)
            return result.value, None
    except Exception as e:
        logging.warning(f"Fallback DOCX extractor (mammoth) failed: {e}")
        return None, None


def extract_text_docx_simple(file_path: str) -> Tuple[Optional[str], None]:
    """Ultra-simple DOCX extractor by parsing XML directly."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        
        with zipfile.ZipFile(file_path) as zf:
            xml_content = zf.read('word/document.xml')
        
        root = ET.fromstring(xml_content)
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        texts = []
        for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                texts.append(t.text)
        
        return ' '.join(texts), None
    except Exception as e:
        logging.warning(f"Simple DOCX extractor failed: {e}")
        return None, None


def extract_text_pdf_primary(file_path: str) -> Tuple[Optional[str], int]:
    """Primary PDF text extractor using PyMuPDF (fitz)."""
    try:
        import fitz
        doc = fitz.open(file_path)
        page_count = len(doc)
        
        if page_count > MAX_PDF_PAGES:
            doc.close()
            raise ValueError(f"PDF excede el límite de {MAX_PDF_PAGES} páginas")
        
        text_parts = []
        for page in doc:
            text = page.get_text()
            text_parts.append(text)
        doc.close()
        
        full_text = '\n\n'.join(text_parts)
        return full_text, page_count
    except ValueError:
        raise
    except Exception as e:
        logging.warning(f"Primary PDF extractor (PyMuPDF) failed: {e}")
        return None, 0


def extract_text_pdf_fallback(file_path: str) -> Tuple[Optional[str], int]:
    """Fallback PDF text extractor using pdfplumber."""
    try:
        import pdfplumber
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            
            if page_count > MAX_PDF_PAGES:
                raise ValueError(f"PDF excede el límite de {MAX_PDF_PAGES} páginas")
            
            text_parts = []
            for page in pdf.pages:
                text = page.extract_text() or ''
                text_parts.append(text)
        
        full_text = '\n\n'.join(text_parts)
        return full_text, page_count
    except ValueError:
        raise
    except Exception as e:
        logging.warning(f"Fallback PDF extractor (pdfplumber) failed: {e}")
        return None, 0


def extract_text_pdf_legacy(file_path: str) -> Tuple[Optional[str], int]:
    """Legacy PDF text extractor using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        page_count = len(reader.pages)
        
        if page_count > MAX_PDF_PAGES:
            raise ValueError(f"PDF excede el límite de {MAX_PDF_PAGES} páginas")
        
        text_parts = []
        for page in reader.pages:
            text = page.extract_text() or ''
            text_parts.append(text)
        
        full_text = '\n\n'.join(text_parts)
        return full_text, page_count
    except ValueError:
        raise
    except Exception as e:
        logging.warning(f"Legacy PDF extractor (PyPDF2) failed: {e}")
        return None, 0


def extract_text_with_fallback(file_path: str, ext: str) -> Dict:
    """
    STAGE 1: Extract text with multiple fallback extractors.
    Returns a dict with text, doc object (for DOCX), page_count, and success status.
    """
    result = {
        'success': False,
        'text': None,
        'doc': None,
        'page_count': 0,
        'extractor_used': None,
        'is_scanned': False,
        'error_code': None
    }
    
    try:
        if ext == '.docx':
            text, doc = extract_text_docx_primary(file_path)
            if text is not None:
                result['text'] = text
                result['doc'] = doc
                result['extractor_used'] = 'python-docx'
                result['success'] = True
                return result
            
            text, _ = extract_text_docx_fallback(file_path)
            if text is not None:
                result['text'] = text
                result['extractor_used'] = 'mammoth'
                result['success'] = True
                return result
            
            text, _ = extract_text_docx_simple(file_path)
            if text is not None:
                result['text'] = text
                result['extractor_used'] = 'xml-direct'
                result['success'] = True
                return result
            
            result['error_code'] = 'DOCX_PARSE_ERROR'
            return result
        
        elif ext == '.pdf':
            try:
                text, page_count = extract_text_pdf_primary(file_path)
                if text is not None:
                    result['text'] = text
                    result['page_count'] = page_count
                    result['extractor_used'] = 'pymupdf'
                    result['success'] = True
                    
                    if len(text.strip()) < 100:
                        result['is_scanned'] = True
                        result['error_code'] = 'PDF_SCANNED'
                        result['success'] = False
                    return result
            except ValueError as ve:
                result['error_code'] = 'SIZE_ERROR'
                result['text'] = str(ve)
                return result
            
            try:
                text, page_count = extract_text_pdf_fallback(file_path)
                if text is not None:
                    result['text'] = text
                    result['page_count'] = page_count
                    result['extractor_used'] = 'pdfplumber'
                    result['success'] = True
                    
                    if len(text.strip()) < 100:
                        result['is_scanned'] = True
                        result['error_code'] = 'PDF_SCANNED'
                        result['success'] = False
                    return result
            except ValueError as ve:
                result['error_code'] = 'SIZE_ERROR'
                result['text'] = str(ve)
                return result
            
            try:
                text, page_count = extract_text_pdf_legacy(file_path)
                if text is not None:
                    result['text'] = text
                    result['page_count'] = page_count
                    result['extractor_used'] = 'pypdf2'
                    result['success'] = True
                    
                    if len(text.strip()) < 100:
                        result['is_scanned'] = True
                        result['error_code'] = 'PDF_SCANNED'
                        result['success'] = False
                    return result
            except ValueError as ve:
                result['error_code'] = 'SIZE_ERROR'
                result['text'] = str(ve)
                return result
            
            result['error_code'] = 'PDF_PARSE_ERROR'
            return result
        
        result['error_code'] = 'FORMAT_ERROR'
        return result
    
    except Exception as e:
        logging.error(f"Text extraction failed completely: {e}")
        result['error_code'] = 'PARSE_ERROR'
        return result


def detect_pii_regex_only(text: str) -> Tuple[List, List]:
    """
    STAGE 2 FALLBACK: Pure regex-based PII detection without any NLP.
    Returns (confirmed_entities, needs_review_entities).
    """
    import anonymizer as anon
    
    confirmed = []
    needs_review = []
    
    try:
        placeholder_positions = anon.find_existing_placeholders(text)
        
        try:
            regex_entities = anon.detect_entities_regex(text, placeholder_positions)
            for e in regex_entities:
                if e[4] >= 0.80:
                    confirmed.append(e)
                else:
                    needs_review.append(e)
        except Exception as ex:
            logging.warning(f"Regex detection failed: {ex}")
        
        try:
            dir_entities = anon.detect_direccion_enhanced(text, placeholder_positions)
            for e in dir_entities:
                if e[4] >= 0.80:
                    confirmed.append(e)
                else:
                    needs_review.append(e)
        except Exception as ex:
            logging.warning(f"Address detection failed: {ex}")
        
        try:
            tel_entities = anon.detect_telefono_enhanced(text, placeholder_positions)
            for e in tel_entities:
                if e[4] >= 0.80:
                    confirmed.append(e)
                else:
                    needs_review.append(e)
        except Exception as ex:
            logging.warning(f"Phone detection failed: {ex}")
        
        try:
            persona_entities = anon.detect_persona_aggressive(text, placeholder_positions)
            for e in persona_entities:
                if e[4] >= 0.80:
                    confirmed.append(e)
                else:
                    needs_review.append(e)
        except Exception as ex:
            logging.warning(f"Person detection failed: {ex}")
        
    except Exception as e:
        logging.error(f"Regex-only detection had errors: {e}")
    
    try:
        confirmed = anon.deduplicate_entities(confirmed)
        needs_review = anon.deduplicate_entities(needs_review)
    except:
        pass
    
    return confirmed, needs_review


def detect_pii_with_fallback(text: str) -> Dict:
    """
    STAGE 2: Detect PII using 4-layer detection system.
    Returns a dict with confirmed entities, needs_review entities, and status.
    """
    result = {
        'success': True,
        'confirmed': [],
        'needs_review': [],
        'detector_used': None,
        'warning': None,
        'metadata': {}
    }
    
    try:
        # Use new 4-layer detection system
        from detector_capas import detect_all_pii
        
        entities, metadata = detect_all_pii(text)
        result['metadata'] = metadata
        
        # Convert Entity objects to tuple format for compatibility
        confirmed = []
        for e in entities:
            # Format: (type, value, start, end, confidence)
            confirmed.append((e.type, e.value, e.start, e.end, e.confidence))
        
        result['confirmed'] = confirmed
        result['needs_review'] = []  # New system handles all entities as confirmed
        
        # Determine detector description
        if metadata.get('spacy_used'):
            result['detector_used'] = 'capas (spacy+regex+heuristic)'
        else:
            result['detector_used'] = 'capas (regex+heuristic)'
        
        return result
        
    except Exception as e:
        logging.warning(f"New detector failed, falling back to legacy: {e}")
        
        # Fallback to legacy detection
        try:
            import anonymizer as anon
            confirmed, needs_review = anon.detect_entities_hybrid(text)
            result['confirmed'] = confirmed
            result['needs_review'] = needs_review
            result['detector_used'] = 'legacy (hybrid)'
            result['warning'] = 'Usando detector legacy'
            return result
        except Exception as e2:
            logging.warning(f"Legacy detection also failed: {e2}")
        
        # Final fallback to regex-only
        try:
            confirmed, needs_review = detect_pii_regex_only(text)
            result['confirmed'] = confirmed
            result['needs_review'] = needs_review
            result['detector_used'] = 'regex-only'
            result['warning'] = 'Usando solo patrones regex'
            return result
        except Exception as e3:
            logging.error(f"All detection methods failed: {e3}")
            result['success'] = False
            result['warning'] = 'Error en detección de datos sensibles'
            return result


def anonymize_text_robust(text: str, confirmed_entities: List, mapping) -> Tuple[str, bool]:
    """
    STAGE 3: Apply anonymization replacements.
    Returns (anonymized_text, success).
    """
    try:
        import anonymizer as anon
        anonymized = anon.replace_entities(text, confirmed_entities, mapping)
        return anonymized, True
    except Exception as e:
        logging.error(f"Text anonymization failed: {e}")
        return text, False


def anonymize_docx_robust(doc, confirmed_entities: List, mapping) -> bool:
    """
    STAGE 3: Apply anonymization to DOCX document.
    Returns success status.
    """
    try:
        import anonymizer as anon
        for entity_type, value, _, _, _ in confirmed_entities:
            substitute = mapping.get_substitute(entity_type, value)
            try:
                anon.replace_value_in_docx(doc, value, substitute, replace_all=True)
            except Exception as e:
                logging.warning(f"Failed to replace '{value[:20]}...' in DOCX: {e}")
        return True
    except Exception as e:
        logging.error(f"DOCX anonymization failed: {e}")
        return False


def post_scan_final(anonymized_text: str, original_confirmed: List) -> Dict:
    """
    STAGE 4: Final scan to check if any PII might have been missed.
    Uses the new 4-layer detection system on the final output.
    Returns scan result with needs_review flag.
    """
    result = {
        'needs_review': False,
        'potential_issues': [],
        'warning': None
    }
    
    try:
        # Use new 4-layer post-scan
        from detector_capas import post_scan_final as new_post_scan
        
        needs_review, detected = new_post_scan(anonymized_text)
        if needs_review:
            result['needs_review'] = True
            # Convert to tuple format for compatibility
            for item in detected:
                # Create pseudo-entity tuples: (type, 'DETECTED', 0, 0, 1.0)
                result['potential_issues'].append(
                    (item['type'], f"POSIBLE_{item['type']}", 0, 0, 1.0)
                )
        
    except Exception as e:
        logging.warning(f"New post-scan failed, trying legacy: {e}")
        
        try:
            import anonymizer as anon
            missed = anon.post_verification(anonymized_text, original_confirmed)
            if missed:
                result['needs_review'] = True
                result['potential_issues'] = missed
        except Exception as e2:
            logging.warning(f"Legacy post-scan also failed: {e2}")
            result['warning'] = 'Post-verificación no disponible'
    
    return result


def process_document_robust(file_path: str, ext: str, file_size: int, strict_mode: bool = True, generate_mapping: bool = False) -> Dict:
    """
    Main robust processing pipeline.
    Processes a document through all stages with fallbacks.
    Returns a comprehensive result dict.
    """
    error_id = generate_error_id()
    file_type = ext.replace('.', '').upper()
    
    result = {
        'ok': True,
        'error_id': error_id,
        'stage_reached': 0,
        'warnings': [],
        'text': None,
        'doc': None,
        'summary': {},
        'mapping': None,
        'needs_review': [],
        'page_count': 0,
        'extractor_used': None,
        'detector_used': None
    }
    
    try:
        with open(file_path, 'rb') as f:
            file_data = f.read()
        
        is_valid, error_code = validate_file_content(file_data, ext)
        if not is_valid:
            log_error(error_id, error_code, file_type, file_size, 0, Exception("Invalid file format"))
            return create_controlled_error(error_code, error_id, "Archivo inválido o corrupto", 0)
        
        result['stage_reached'] = 1
        
        extraction = extract_text_with_fallback(file_path, ext)
        
        if not extraction['success']:
            if extraction['is_scanned']:
                log_error(error_id, 'PDF_SCANNED', file_type, file_size, 1, Exception("Scanned PDF"))
                return create_controlled_error('PDF_SCANNED', error_id, "El PDF no contiene texto extraíble", 1)
            
            error_code = extraction.get('error_code', 'PARSE_ERROR')
            log_error(error_id, error_code, file_type, file_size, 1, Exception(f"Extraction failed: {error_code}"))
            return create_controlled_error(error_code, error_id, "No se pudo extraer texto", 1)
        
        result['text'] = extraction['text']
        result['doc'] = extraction['doc']
        result['page_count'] = extraction.get('page_count', 0)
        result['extractor_used'] = extraction['extractor_used']
        result['stage_reached'] = 2
        
        detection = detect_pii_with_fallback(extraction['text'])
        
        if detection.get('warning'):
            result['warnings'].append(detection['warning'])
        
        result['detector_used'] = detection['detector_used']
        confirmed = detection['confirmed']
        needs_review_entities = detection['needs_review']
        result['stage_reached'] = 3
        
        import anonymizer as anon
        mapping = anon.EntityMapping(mode='tokens')
        result['mapping'] = mapping
        
        if ext == '.docx' and result['doc'] is not None:
            success = anonymize_docx_robust(result['doc'], confirmed, mapping)
            if not success:
                result['warnings'].append('Algunos reemplazos en DOCX pueden no haberse aplicado correctamente')
            
            text_parts = []
            for para in result['doc'].paragraphs:
                text_parts.append(para.text)
            result['anonymized_text'] = '\n'.join(text_parts)
        else:
            anonymized_text, success = anonymize_text_robust(extraction['text'], confirmed, mapping)
            result['anonymized_text'] = anonymized_text
            if not success:
                result['warnings'].append('La anonimización puede estar incompleta')
        
        result['stage_reached'] = 4
        
        if strict_mode:
            post_scan = post_scan_final(result.get('anonymized_text', result['text']), confirmed)
            if post_scan['needs_review']:
                for issue in post_scan['potential_issues']:
                    needs_review_entities.append(issue)
            if post_scan.get('warning'):
                result['warnings'].append(post_scan['warning'])
        
        result['needs_review'] = [
            {
                'id': str(uuid.uuid4())[:8],
                'type': e[0],
                'value': e[1],
                'start': e[2],
                'end': e[3],
                'confidence': e[4],
                'context': extraction['text'][max(0, e[2]-30):min(len(extraction['text']), e[3]+30)]
            }
            for e in needs_review_entities
        ]
        
        result['summary'] = {
            'entities_found': mapping.get_summary(),
            'total_entities': sum(mapping.get_summary().values()),
            'replacements': mapping.get_replacements_for_report(),
            'mode': 'tokens',
            'strict_mode': strict_mode,
            'needs_review_count': len(result['needs_review'])
        }
        
        return result
        
    except Exception as e:
        log_error(error_id, 'UNKNOWN_ERROR', file_type, file_size, result.get('stage_reached', 0), e)
        return create_controlled_error('UNKNOWN_ERROR', error_id, str(e)[:200], result.get('stage_reached', 0))


def save_output_robust(result: Dict, job_id: str, output_dir: str, original_filename: str, ext: str, generate_mapping: bool = False) -> Dict:
    """
    Save output files robustly with error handling.
    """
    output = {
        'ok': True,
        'output_path': None,
        'report_path': None,
        'text_path': None,
        'mapping_path': None,
        'output_type': ext.replace('.', ''),
        'error': None
    }
    
    try:
        import anonymizer as anon
        
        if ext == '.docx' and result.get('doc') is not None:
            output_filename = f"{job_id}_anonimizado.docx"
            output_path = os.path.join(output_dir, output_filename)
            try:
                result['doc'].save(output_path)
                output['output_path'] = output_path
            except Exception as e:
                logging.error(f"Failed to save DOCX: {e}")
                txt_path = os.path.join(output_dir, f"{job_id}_anonimizado.txt")
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(result.get('anonymized_text', result.get('text', '')))
                output['output_path'] = txt_path
                output['output_type'] = 'txt'
                output['warning'] = 'Documento guardado como texto plano'
        
        elif ext == '.pdf':
            try:
                output_filename = f"{job_id}_anonimizado.pdf"
                output_path = os.path.join(output_dir, output_filename)
                anon.create_anonymized_pdf(result.get('anonymized_text', result.get('text', '')), output_path)
                output['output_path'] = output_path
            except Exception as e:
                logging.error(f"Failed to create PDF: {e}")
                txt_path = os.path.join(output_dir, f"{job_id}_anonimizado.txt")
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(result.get('anonymized_text', result.get('text', '')))
                output['output_path'] = txt_path
                output['output_type'] = 'txt'
                output['warning'] = 'Documento guardado como texto plano'
            
            text_path = os.path.join(output_dir, f"{job_id}_text.txt")
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(result.get('anonymized_text', result.get('text', '')))
            output['text_path'] = text_path
        
        report = anon.generate_report(result.get('summary', {}), original_filename, ext.upper())
        report_json_path = os.path.join(output_dir, f"{job_id}_reporte.json")
        
        import json
        with open(report_json_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        output['report_path'] = report_json_path
        
        report_txt = anon.generate_report_txt(report)
        report_txt_path = os.path.join(output_dir, f"{job_id}_reporte.txt")
        with open(report_txt_path, 'w', encoding='utf-8') as f:
            f.write(report_txt)
        
        if generate_mapping and result.get('mapping'):
            try:
                mapping_csv = anon.generate_mapping_csv(result['mapping'])
                mapping_path = os.path.join(output_dir, f"{job_id}_mapping.csv")
                with open(mapping_path, 'w', encoding='utf-8') as f:
                    f.write(mapping_csv)
                output['mapping_path'] = mapping_path
            except Exception as e:
                logging.warning(f"Failed to generate mapping CSV: {e}")
        
        output['report'] = report
        
    except Exception as e:
        logging.error(f"Output save failed: {e}")
        output['ok'] = False
        output['error'] = str(e)
    
    return output
