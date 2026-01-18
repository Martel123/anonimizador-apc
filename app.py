import os
import csv
import json
import logging
import requests
import resend
from functools import wraps
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, send_file, flash, jsonify, session, g, abort
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from openai import OpenAI
import anonymizer as anon_module

from models import db, User, DocumentRecord, Plantilla, Modelo, Estilo, CampoPlantilla, Tenant, Case, CaseAssignment, CaseDocument, Task, FinishedDocument, ImagenModelo, CaseAttachment, ModeloTabla, ReviewSession, ReviewIssue, TwoFALog, EstiloDocumento, PricingConfig, PricingAddon, CheckoutSession, Subscription, ActivationToken, TaskDocument, TaskReminder, CalendarEvent, EventAttendee, UserArgumentationStyle, ArgumentationSession, ArgumentationMessage, ArgumentationJob, AgentSession, AgentMessage, LegalStrategy, CostEstimate, CaseEvent, CaseType, CaseCustomField, CaseCustomFieldValue, AuditLog, TipoActa, FormResponse
import qrcode
import threading
import queue
import re
import time
from io import BytesIO
import base64
import uuid

log_level = logging.DEBUG if os.environ.get("FLASK_DEBUG", "false").lower() == "true" else logging.INFO
logging.basicConfig(level=log_level)

app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET")
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Por favor inicia sesión para acceder a esta página.'
login_manager.login_message_category = 'info'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


MODELOS = {
    "aumento_alimentos": {
        "nombre": "Aumento de alimentos",
        "plantilla": "aumento_alimentos.txt",
        "carpeta_estilos": "aumento_alimentos"
    }
}

PLAN_CONFIG_DEFAULT = {
    'basico': {
        'nombre': 'Plan Básico',
        'max_usuarios': 2,
        'max_documentos_mes': 50,
        'max_plantillas': 5,
        'precio_mensual': 29.99,
        'features': ['feature_generate', 'feature_models', 'feature_terminados', 'feature_auditoria_basic']
    },
    'medio': {
        'nombre': 'Plan Medio',
        'max_usuarios': 5,
        'max_documentos_mes': 150,
        'max_plantillas': 15,
        'precio_mensual': 59.99,
        'features': ['feature_generate', 'feature_models', 'feature_terminados', 'feature_casos', 
                     'feature_tareas', 'feature_calendario', 'feature_historial', 'feature_auditoria_basic',
                     'feature_onboarding', 'feature_argumentacion_ia']
    },
    'avanzado': {
        'nombre': 'Plan Avanzado',
        'max_usuarios': 8,
        'max_documentos_mes': 9999,
        'max_plantillas': 9999,
        'precio_mensual': 99.99,
        'features': ['feature_generate', 'feature_models', 'feature_terminados', 'feature_casos',
                     'feature_tareas', 'feature_calendario', 'feature_historial', 'feature_estadisticas',
                     'feature_auditoria_basic', 'feature_auditoria_avanzada', 'feature_onboarding',
                     'feature_argumentacion_ia', 'feature_api_access', 'feature_soporte_prioritario', 'feature_agente_ia']
    }
}


def get_dynamic_plan_config():
    """Obtiene la configuración de planes desde la base de datos o usa los defaults."""
    try:
        from models import PlanConfiguration
        db_plans = PlanConfiguration.get_all_plans_dict()
        if db_plans:
            return db_plans
    except Exception as e:
        logging.debug(f"Using default plan config: {e}")
    return PLAN_CONFIG_DEFAULT


def get_plan_config(tenant):
    """Obtiene la configuración del plan del tenant."""
    plans = get_dynamic_plan_config()
    if not tenant:
        return plans.get('basico', PLAN_CONFIG_DEFAULT['basico'])
    plan = getattr(tenant, 'plan', 'basico') or 'basico'
    return plans.get(plan, plans.get('basico', PLAN_CONFIG_DEFAULT['basico']))


PLAN_CONFIG = PLAN_CONFIG_DEFAULT


def tenant_can_add_user(tenant):
    """Verifica si el tenant puede agregar más usuarios según su plan."""
    if not tenant:
        return False
    plan_config = get_plan_config(tenant)
    current_users = User.query.filter_by(tenant_id=tenant.id, activo=True).count()
    return current_users < plan_config['max_usuarios']


def get_tenant_user_limit(tenant):
    """Obtiene el límite de usuarios del tenant según su plan."""
    plan_config = get_plan_config(tenant)
    return plan_config['max_usuarios']


def has_feature(tenant, feature_name):
    """Verifica si el Centro tiene acceso a una feature según su plan."""
    if not tenant:
        return False
    plan_config = get_plan_config(tenant)
    return feature_name in plan_config.get('features', [])


def require_feature(feature_name):
    """Decorador para verificar que el Centro tiene acceso a una feature."""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for('login'))
            tenant = get_current_tenant()
            if not tenant:
                flash("No tienes acceso a esta funcionalidad.", "error")
                return redirect(url_for('dashboard'))
            if not has_feature(tenant, feature_name):
                flash("Esta funcionalidad no está disponible en tu plan actual. Contacta al administrador para actualizar tu plan.", "warning")
                return redirect(url_for('dashboard'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator


def log_audit(tenant_id, evento, descripcion=None, extra_data=None, user_id=None):
    """Registra un evento de auditoría."""
    try:
        if user_id is None and current_user.is_authenticated:
            user_id = current_user.id
        
        ip_address = request.remote_addr if request else None
        
        audit = AuditLog(
            tenant_id=tenant_id,
            actor_user_id=user_id,
            evento=evento,
            descripcion=descripcion,
            extra_data=extra_data,
            ip_address=ip_address
        )
        db.session.add(audit)
        db.session.commit()
    except Exception as e:
        logging.error(f"Error al registrar auditoría: {e}")
        db.session.rollback()


def create_default_tipos_acta(tenant_id):
    """Crea los tipos de acta predeterminados para un Centro nuevo."""
    from models import TipoActa
    for orden, tipo_default in enumerate(TipoActa.TIPOS_DEFAULT):
        tipo = TipoActa(
            tenant_id=tenant_id,
            nombre=tipo_default['nombre'],
            categoria=tipo_default['categoria'],
            icono=tipo_default['icono'],
            orden=orden,
            activo=True
        )
        db.session.add(tipo)
    db.session.commit()


BASE_PERSISTENT = os.environ.get("PERSISTENT_DIR", os.path.dirname(os.path.abspath(__file__)))

CARPETA_MODELOS = os.path.join(BASE_PERSISTENT, "modelos_legales")
CARPETA_ESTILOS = os.path.join(BASE_PERSISTENT, "estilos_estudio")
CARPETA_RESULTADOS = os.path.join(BASE_PERSISTENT, "Resultados")
CARPETA_PLANTILLAS_SUBIDAS = os.path.join(BASE_PERSISTENT, "plantillas_subidas")
CARPETA_ESTILOS_SUBIDOS = os.path.join(BASE_PERSISTENT, "estilos_subidos")
CARPETA_IMAGENES_MODELOS = os.path.join(BASE_PERSISTENT, "static", "imagenes_modelos")
ALLOWED_IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}

CAMPOS_LARGOS = ['HECHOS', 'FUNDAMENTOS', 'DETALLE', 'NARRACION', 'DESCRIPCION', 
                 'PETITORIO', 'PRETENSION', 'CONCLUSION', 'OBSERVACIONES', 'ANTECEDENTES']


def extract_placeholders_from_docx(docx_path):
    """Extrae todos los placeholders {{CAMPO}} de un archivo docx."""
    import re
    from docx import Document
    
    placeholders = set()
    pattern = r'\{\{([^}]+)\}\}'
    
    try:
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            matches = re.findall(pattern, para.text)
            for match in matches:
                placeholders.add(match.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = re.findall(pattern, para.text)
                        for match in matches:
                            placeholders.add(match.strip())
        
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    matches = re.findall(pattern, para.text)
                    for match in matches:
                        placeholders.add(match.strip())
            if section.footer:
                for para in section.footer.paragraphs:
                    matches = re.findall(pattern, para.text)
                    for match in matches:
                        placeholders.add(match.strip())
    except Exception as e:
        logging.error(f"Error extrayendo placeholders de {docx_path}: {e}")
        return []
    
    return sorted(list(placeholders))


def extract_placeholders_from_text(text):
    """Extrae placeholders {{CAMPO}} de texto plano."""
    import re
    pattern = r'\{\{([^}]+)\}\}'
    matches = re.findall(pattern, text)
    return sorted(list(set(m.strip() for m in matches)))


def is_campo_largo(campo_nombre):
    """Determina si un campo debe usar textarea."""
    campo_upper = campo_nombre.upper()
    for largo in CAMPOS_LARGOS:
        if largo in campo_upper:
            return True
    return False


def generate_qr_code_base64(data):
    """Genera un QR code y lo retorna como base64."""
    import qrcode
    import io
    import base64
    
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(data)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    
    return base64.b64encode(buffer.getvalue()).decode('utf-8')

argumentation_job_queue = queue.Queue()
argumentation_worker_started = False


def detect_document_sections(texto):
    """Detecta secciones del documento legal usando patrones regex mejorados."""
    secciones = {
        'hechos': {'inicio': None, 'fin': None, 'texto': None},
        'fundamentos': {'inicio': None, 'fin': None, 'texto': None},
        'petitorio': {'inicio': None, 'fin': None, 'texto': None}
    }
    
    patrones = {
        'hechos': [
            r'(?i)(I+\.?\s*)?(DE\s+LOS\s+)?HECHOS\s*:?',
            r'(?i)PRIMERO\s*[:.\-]',
            r'(?i)ANTECEDENTES\s*:?',
            r'(?i)EXPOSICI[OÓ]N\s+DE\s+(LOS\s+)?HECHOS',
            r'(?i)RELACI[OÓ]N\s+DE\s+(LOS\s+)?HECHOS',
            r'(?i)NARRACI[OÓ]N\s+DE\s+(LOS\s+)?HECHOS'
        ],
        'fundamentos': [
            r'(?i)(I+\.?\s*)?FUNDAMENTOS?\s+(DE\s+)?(DERECHO|JUR[IÍ]DICO|HECHO\s+Y\s+DERECHO)',
            r'(?i)FUNDAMENTACI[OÓ]N\s+JUR[IÍ]DICA',
            r'(?i)BASE\s+LEGAL',
            r'(?i)MARCO\s+(LEGAL|JUR[IÍ]DICO|NORMATIVO)',
            r'(?i)CONSIDERACIONES\s+(DE\s+)?DERECHO',
            r'(?i)AMPARO\s+LEGAL'
        ],
        'petitorio': [
            r'(?i)PETITORIO\s*:?',
            r'(?i)POR\s+(LO\s+)?TANTO\s*:?',
            r'(?i)POR\s+ESTAS?\s+CONSIDERACIONES?\s*:?',
            r'(?i)SOLICITO\s*:?',
            r'(?i)SE\s+SIRVA\s*:?',
            r'(?i)PEDIMENTO\s*:?',
            r'(?i)CONCLUSI[OÓ]N\s*:?'
        ]
    }
    
    lineas = texto.split('\n')
    posiciones = []
    
    for seccion, patrones_sec in patrones.items():
        for i, linea in enumerate(lineas):
            linea_norm = linea.strip()
            if len(linea_norm) > 200:
                continue
            for patron in patrones_sec:
                if re.search(patron, linea_norm):
                    posiciones.append((seccion, i))
                    break
            else:
                continue
            break
    
    posiciones.sort(key=lambda x: x[1])
    
    for idx, (seccion, inicio) in enumerate(posiciones):
        fin = posiciones[idx + 1][1] if idx + 1 < len(posiciones) else len(lineas)
        secciones[seccion]['inicio'] = inicio
        secciones[seccion]['fin'] = fin
        secciones[seccion]['texto'] = '\n'.join(lineas[inicio:fin])
    
    return secciones


def detect_intent(instrucciones):
    """Detecta si el usuario hace una pregunta o pide una modificación.
    Prioriza preguntas cuando hay signos de interrogación o frases interrogativas."""
    instrucciones_lower = instrucciones.lower().strip()
    
    if '?' in instrucciones or '¿' in instrucciones:
        return 'explanation'
    
    frases_pregunta_explicita = [
        'cuales son', 'cuáles son', 'que hiciste', 'qué hiciste', 
        'que cambiaste', 'qué cambiaste', 'que modificaste', 'qué modificaste',
        'que parte', 'qué parte', 'que partes', 'qué partes',
        'quiero ver', 'quiero saber', 'muestrame', 'muéstrame',
        'dime que', 'dime qué', 'explica que', 'explica qué',
        'cuales fueron', 'cuáles fueron', 'que son los cambios', 'qué son los cambios',
        'los cambios que hiciste', 'los cambios que realizaste',
        'me puedes decir', 'me puedes explicar', 'puedes decirme', 'puedes explicarme'
    ]
    
    for frase in frases_pregunta_explicita:
        if frase in instrucciones_lower:
            return 'explanation'
    
    frases_pregunta_inicio = [
        'qué ', 'que ', 'cuál ', 'cual ', 'cómo ', 'como ', 'por qué ', 'por que ',
        'dónde ', 'donde ', 'cuándo ', 'cuando ', 'quién ', 'quien ',
        'explica', 'explicame', 'explícame', 'dime ', 'cuéntame', 'cuentame'
    ]
    
    for frase in frases_pregunta_inicio:
        if instrucciones_lower.startswith(frase):
            return 'explanation'
    
    preguntas_keywords = [
        'explica', 'explicame', 'explícame', 'por qué', 'por que', 'qué pasa si', 'que pasa si',
        'cómo puedo', 'como puedo', 'ayúdame a entender', 'ayudame a entender',
        'qué argumento', 'que argumento', 'está bien', 'esta bien', 'es correcto',
        'qué opinas', 'que opinas', 'crees que', 'debería', 'deberia',
        'puedo agregar', 'puedo añadir', 'sugieres', 'recomiendas', 'significa',
        'cuál es', 'cual es', 'qué es', 'que es', 'dime', 'cuéntame', 'cuentame',
        'los cambios', 'que hiciste', 'que modificaste', 'que cambiaste'
    ]
    
    modificacion_keywords = [
        'añade', 'anade', 'agrega', 'elimina', 'quita', 'borra',
        'reescribe', 'modifica', 'cambia a', 'cambia el', 'cambia la', 'cambia los',
        'mejora el', 'mejora la', 'mejora los', 'refuerza', 'amplía', 'amplia', 
        'reduce', 'resume', 'desarrolla', 'incluye', 'incorpora', 'expande', 
        'reestructura', 'pon ', 'coloca', 'escribe'
    ]
    
    score_pregunta = sum(1 for kw in preguntas_keywords if kw in instrucciones_lower)
    score_modificacion = sum(1 for kw in modificacion_keywords if kw in instrucciones_lower)
    
    if score_pregunta >= score_modificacion and score_pregunta > 0:
        return 'explanation'
    return 'rewrite'


def extract_section_text(texto_completo, section):
    """Extrae solo la sección solicitada del documento."""
    if section == 'full':
        return texto_completo
    
    secciones = detect_document_sections(texto_completo)
    
    if secciones.get(section, {}).get('texto'):
        return secciones[section]['texto']
    
    return texto_completo


def merge_section_result(texto_original, section, texto_mejorado):
    """Reinserta la sección mejorada en el documento original.
    Si no se detecta la sección, devuelve el documento completo mejorado (texto_mejorado)
    pero SOLO si es claramente un documento completo, sino conserva el original."""
    if section == 'full':
        return texto_mejorado
    
    secciones = detect_document_sections(texto_original)
    
    if not secciones.get(section, {}).get('inicio'):
        len_original = len(texto_original.strip())
        len_mejorado = len(texto_mejorado.strip())
        
        if len_mejorado > len_original * 0.7:
            return texto_mejorado
        else:
            logging.warning(f"Section '{section}' not found in document, returning improved text")
            return texto_mejorado
    
    lineas = texto_original.split('\n')
    inicio = secciones[section]['inicio']
    fin = secciones[section]['fin']
    
    resultado = lineas[:inicio] + texto_mejorado.split('\n') + lineas[fin:]
    return '\n'.join(resultado)


def process_argumentation_job(job_id):
    """Procesa un job de argumentación en segundo plano."""
    with app.app_context():
        job = ArgumentationJob.query.get(job_id)
        if not job or job.status != 'queued':
            return
        
        start_time = time.time()
        
        try:
            job.status = 'processing'
            job.started_at = datetime.utcnow()
            db.session.commit()
            
            sesion = ArgumentationSession.query.get(job.session_id)
            if not sesion:
                job.status = 'failed'
                job.error_message = 'Sesión no encontrada'
                job.completed_at = datetime.utcnow()
                db.session.commit()
                return
            
            documento_actual = sesion.ultima_version_mejorada or sesion.documento_original
            
            extraction_start = time.time()
            texto_a_procesar = extract_section_text(documento_actual, job.section)
            job.extraction_ms = int((time.time() - extraction_start) * 1000)
            
            estilo_instrucciones = ""
            for e in UserArgumentationStyle.ESTILOS_PREDEFINIDOS:
                if e['nombre'] == job.estilo:
                    estilo_instrucciones = e['instrucciones']
                    break
            
            if not estilo_instrucciones:
                estilo_custom = UserArgumentationStyle.query.filter_by(
                    user_id=job.user_id,
                    nombre=job.estilo,
                    activo=True
                ).first()
                if estilo_custom:
                    estilo_instrucciones = estilo_custom.instrucciones
            
            ia_start = time.time()
            client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=120.0)
            
            if job.job_type == 'explanation':
                system_prompt = f"""Eres un asistente juridico experto. El usuario te hace una consulta sobre un documento legal.
Responde de forma clara, concisa y profesional. No reescribas el documento completo, solo responde la pregunta.
Cita articulos o jurisprudencia relevante si aplica.
Estilo de comunicacion: {job.estilo}
{estilo_instrucciones}"""
                user_content = f"Documento de referencia:\n{texto_a_procesar[:15000]}\n\nPregunta del usuario:\n{job.instructions}"
            else:
                section_name = ArgumentationJob.SECTIONS.get(job.section, job.section)
                system_prompt = f"""Actua como un asistente juridico especializado en redaccion y argumentacion.

Tu tarea es modificar directamente el texto del documento juridico, aplicando EXACTAMENTE las instrucciones del usuario.

SECCION A TRABAJAR: {section_name}

REGLAS ESTRICTAS:
1. Manten intactos todos los datos facticos (nombres, DNIs, fechas, montos, direcciones)
2. Si encuentras incoherencias factuales, solo senala el error; no inventes datos nuevos
3. Puedes anadir parrafos completos si el usuario lo pide
4. Puedes eliminar fragmentos si el usuario lo pide
5. Puedes reorganizar la logica argumentativa si ayuda a la claridad
6. Respeta la estructura de la seccion
7. No inventes hechos ni articulos falsos
8. Aplica el estilo solicitado: {job.estilo}
9. {estilo_instrucciones}

INSTRUCCIONES DEL USUARIO:
{job.instructions}

Devuelve SOLO la seccion modificada, sin comentarios meta, lista para usar."""
                user_content = f"Texto a mejorar:\n\n{texto_a_procesar[:18000]}"
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ],
                temperature=0.4,
                max_tokens=6000
            )
            
            resultado = response.choices[0].message.content
            job.ia_ms = int((time.time() - ia_start) * 1000)
            
            if job.job_type == 'rewrite' and job.section != 'full':
                resultado_final = merge_section_result(documento_actual, job.section, resultado)
            else:
                resultado_final = resultado
            
            job.result_text = resultado_final
            
            mensaje_usuario = ArgumentationMessage(
                session_id=sesion.id,
                role="user",
                content=job.instructions,
                estilo_aplicado=job.estilo
            )
            db.session.add(mensaje_usuario)
            
            mensaje_ia = ArgumentationMessage(
                session_id=sesion.id,
                role="assistant",
                content=resultado_final,
                estilo_aplicado=job.estilo,
                message_type=job.job_type
            )
            db.session.add(mensaje_ia)
            
            if job.job_type == 'rewrite':
                sesion.ultima_version_mejorada = resultado_final
                sesion.estilo_usado = job.estilo
                sesion.updated_at = datetime.utcnow()
            
            job.status = 'done'
            job.completed_at = datetime.utcnow()
            job.total_ms = int((time.time() - start_time) * 1000)
            db.session.commit()
            
        except Exception as e:
            logging.error(f"Error processing argumentation job {job_id}: {e}")
            job.status = 'failed'
            job.error_message = str(e)[:500]
            job.completed_at = datetime.utcnow()
            job.total_ms = int((time.time() - start_time) * 1000)
            db.session.commit()


def argumentation_worker():
    """Worker que procesa jobs de argumentación en segundo plano."""
    while True:
        try:
            job_id = argumentation_job_queue.get(timeout=5)
            if job_id:
                process_argumentation_job(job_id)
        except queue.Empty:
            continue
        except Exception as e:
            logging.error(f"Error in argumentation worker: {e}")


def start_argumentation_worker():
    """Inicia el worker de argumentación si no está corriendo."""
    global argumentation_worker_started
    if not argumentation_worker_started:
        worker_thread = threading.Thread(target=argumentation_worker, daemon=True)
        worker_thread.start()
        argumentation_worker_started = True
        logging.info("Argumentation worker started")


def get_resend_credentials():
    """Get Resend API credentials from Replit connector."""
    try:
        hostname = os.environ.get('REPLIT_CONNECTORS_HOSTNAME')
        token = os.environ.get('REPL_IDENTITY')
        if not token:
            token = os.environ.get('WEB_REPL_RENEWAL')
            if token:
                token = 'depl ' + token
        else:
            token = 'repl ' + token
        
        if not hostname or not token:
            logging.warning("Missing REPLIT_CONNECTORS_HOSTNAME or token for Resend")
            return None, None
            
        response = requests.get(
            f'https://{hostname}/api/v2/connection?include_secrets=true&connector_names=resend',
            headers={'Accept': 'application/json', 'X_REPLIT_TOKEN': token}
        )
        data = response.json()
        if data.get('items'):
            settings = data['items'][0].get('settings', {})
            api_key = settings.get('api_key')
            # Use verified subdomain for sending emails
            from_email = "notificaciones@notificaciones.apcjuridica.com"
            logging.info(f"Resend credentials loaded, from_email: {from_email}")
            return api_key, from_email
    except Exception as e:
        logging.error(f"Error getting Resend credentials: {e}")
    return None, None


def send_notification_email(to_email, subject, html_content):
    """Send email notification via Resend."""
    try:
        logging.info(f"Attempting to send email to {to_email} with subject: {subject}")
        api_key, from_email = get_resend_credentials()
        if not api_key or not from_email:
            logging.warning(f"Resend not configured - api_key: {bool(api_key)}, from_email: {bool(from_email)}")
            return False
        
        logging.info(f"Resend credentials obtained. From: {from_email}")
        resend.api_key = api_key
        result = resend.Emails.send({
            "from": from_email,
            "to": [to_email],
            "subject": subject,
            "html": html_content
        })
        logging.info(f"Email sent successfully to {to_email}. Result: {result}")
        return True
    except Exception as e:
        logging.error(f"Error sending email to {to_email}: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return False


def check_and_send_notifications(tenant_id):
    """Check documents and send notifications for urgent/importante deadlines."""
    try:
        docs = FinishedDocument.query.filter(
            FinishedDocument.tenant_id == tenant_id,
            FinishedDocument.plazo_entrega.isnot(None),
            FinishedDocument.case_id.isnot(None)
        ).all()
        
        for doc in docs:
            priority = doc.get_priority_status()
            
            if priority == 'urgente' and not doc.sent_urgente_notification:
                sent_any = False
                if doc.case:
                    for assignment in doc.case.assignments:
                        if assignment.user and assignment.user.email:
                            if send_notification_email(
                                assignment.user.email,
                                f"URGENTE: Documento '{doc.nombre}' vence pronto",
                                f"<h2>Documento Urgente</h2><p>El documento <strong>{doc.nombre}</strong> tiene plazo de entrega en menos de 24 horas.</p><p>Expediente: {doc.numero_expediente or 'N/A'}</p>"
                            ):
                                sent_any = True
                if sent_any:
                    doc.sent_urgente_notification = True
                    db.session.commit()
                
            elif priority == 'importante' and not doc.sent_importante_notification:
                sent_any = False
                if doc.case:
                    for assignment in doc.case.assignments:
                        if assignment.user and assignment.user.email:
                            if send_notification_email(
                                assignment.user.email,
                                f"Importante: Documento '{doc.nombre}' vence en 2 dias",
                                f"<h2>Recordatorio Importante</h2><p>El documento <strong>{doc.nombre}</strong> tiene plazo de entrega en menos de 48 horas.</p><p>Expediente: {doc.numero_expediente or 'N/A'}</p>"
                            ):
                                sent_any = True
                if sent_any:
                    doc.sent_importante_notification = True
                    db.session.commit()
    except Exception as e:
        logging.error(f"Error checking notifications: {e}")


import re

def extract_text_from_docx(file_path):
    """Extract all text from a Word document."""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)


def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or '')
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"Error extracting PDF: {e}")
        return ""


def extract_text_from_txt(file_path):
    """Extract text from TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logging.error(f"Error reading TXT: {e}")
        return ""


ALLOWED_EXTENSIONS = {'.docx', '.pdf', '.txt'}


def extract_text_from_file(file_path):
    """Extract text from file based on extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    return ""


def detect_placeholders_from_text(text):
    """
    Detect placeholders in text that represent fields to fill.
    Detects each occurrence of dots or underscores as a separate field.
    """
    campos = []
    campo_counter = {}
    
    pattern_curly_double = re.findall(r'\{\{([^}]+)\}\}', text)
    pattern_curly_single = re.findall(r'\{([^{}]+)\}', text)
    pattern_brackets_double = re.findall(r'\[\[([^\]]+)\]\]', text)
    pattern_brackets_single = re.findall(r'\[([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s_]+)\]', text)
    
    all_patterns = pattern_curly_double + pattern_curly_single + pattern_brackets_double + pattern_brackets_single
    for p in all_patterns:
        cleaned = p.strip().replace('_', ' ')
        if cleaned and len(cleaned) > 1:
            campos.append(cleaned)
    
    dot_pattern = r'([A-Za-zÁÉÍÓÚáéíóúÑñ\s\.\,\-°]+?)(?:N[°º]?\s*)?[\.…]{3,}|_{3,}'
    
    matches = list(re.finditer(dot_pattern, text))
    
    for match in matches:
        full_match = match.group(0)
        context = match.group(1) if match.group(1) else ""
        
        context = context.strip()
        context = re.sub(r'^[,\.\s]+', '', context)
        context = re.sub(r'[,\.\s]+$', '', context)
        
        if len(context) < 3:
            start = max(0, match.start() - 50)
            before_text = text[start:match.start()]
            words = re.findall(r'[A-Za-zÁÉÍÓÚáéíóúÑñ]+(?:\s+[A-Za-zÁÉÍÓÚáéíóúÑñ]+)*', before_text)
            if words:
                context = words[-1] if len(words[-1]) > 2 else ' '.join(words[-2:]) if len(words) > 1 else words[-1]
        
        if 'N°' in full_match or 'N.' in context or 'Nº' in full_match:
            if 'D.N.I' in context or 'DNI' in context:
                context = 'Número de DNI'
            elif 'celular' in context.lower() or 'teléfono' in context.lower():
                context = 'Número de celular'
            elif 'expediente' in context.lower():
                context = 'Número de expediente'
            else:
                context = 'Número'
        
        if context and len(context) >= 2:
            base_name = context[:100]
            if base_name.lower() in campo_counter:
                campo_counter[base_name.lower()] += 1
                campos.append(f"{base_name} {campo_counter[base_name.lower()]}")
            else:
                campo_counter[base_name.lower()] = 1
                campos.append(base_name)
    
    lines = text.split('\n')
    for line in lines:
        if ':' in line and re.search(r':\s*$', line.strip()):
            parts = line.split(':')
            if parts[0].strip() and len(parts[0].strip()) > 2:
                campo_name = parts[0].strip()
                if campo_name.lower() not in campo_counter:
                    campos.append(campo_name)
                    campo_counter[campo_name.lower()] = 1
    
    final_campos = []
    for c in campos:
        if c and len(c.strip()) > 1:
            final_campos.append(c.strip())
    
    return final_campos


def campo_to_key(campo_name):
    """Convert a campo name to a valid key."""
    key = campo_name.lower().strip()
    key = re.sub(r'[áàäâ]', 'a', key)
    key = re.sub(r'[éèëê]', 'e', key)
    key = re.sub(r'[íìïî]', 'i', key)
    key = re.sub(r'[óòöô]', 'o', key)
    key = re.sub(r'[úùüû]', 'u', key)
    key = re.sub(r'[ñ]', 'n', key)
    key = re.sub(r'[^a-z0-9]+', '_', key)
    key = re.sub(r'_+', '_', key)
    key = key.strip('_')
    return key[:50] if key else 'campo'


def detect_placeholders_with_context(text):
    """
    Detect placeholders in text and return them with position context.
    Returns list of dicts with: nombre, etiqueta, tipo, start, end, contexto, match_text
    """
    campos = []
    campo_counter = {}
    seen_positions = set()
    
    patterns = [
        (r'\{\{([^}]+)\}\}', 'curly_double'),
        (r'\{([^{}]+)\}', 'curly_single'),
        (r'\[\[([^\]]+)\]\]', 'bracket_double'),
        (r'\[([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s_]+)\]', 'bracket_single'),
    ]
    
    for pattern, pattern_type in patterns:
        for match in re.finditer(pattern, text):
            start, end = match.start(), match.end()
            if any(s <= start < e or s < end <= e for s, e in seen_positions):
                continue
            seen_positions.add((start, end))
            
            captured = match.group(1).strip().replace('_', ' ')
            if captured and len(captured) > 1:
                context_start = max(0, start - 30)
                context_end = min(len(text), end + 30)
                contexto = text[context_start:context_end]
                
                campos.append({
                    'nombre': campo_to_key(captured),
                    'etiqueta': captured,
                    'tipo': 'text',
                    'start': start,
                    'end': end,
                    'contexto': contexto,
                    'match_text': match.group(0),
                    'pattern_type': pattern_type
                })
    
    dot_pattern = r'([A-Za-zÁÉÍÓÚáéíóúÑñ\s\.\,\-°]+?)(?:N[°º]?\s*)?([\.…]{3,}|_{3,})'
    
    for match in re.finditer(dot_pattern, text):
        start, end = match.start(), match.end()
        if any(s <= start < e or s < end <= e for s, e in seen_positions):
            continue
        seen_positions.add((start, end))
        
        full_match = match.group(0)
        context = match.group(1) if match.group(1) else ""
        placeholder_chars = match.group(2) if match.group(2) else ""
        
        context = context.strip()
        context = re.sub(r'^[,\.\s]+', '', context)
        context = re.sub(r'[,\.\s]+$', '', context)
        
        if len(context) < 3:
            before_start = max(0, match.start() - 50)
            before_text = text[before_start:match.start()]
            words = re.findall(r'[A-Za-zÁÉÍÓÚáéíóúÑñ]+(?:\s+[A-Za-zÁÉÍÓÚáéíóúÑñ]+)*', before_text)
            if words:
                context = words[-1] if len(words[-1]) > 2 else ' '.join(words[-2:]) if len(words) > 1 else words[-1]
        
        if 'N°' in full_match or 'N.' in context or 'Nº' in full_match:
            if 'D.N.I' in context or 'DNI' in context:
                context = 'Número de DNI'
            elif 'celular' in context.lower() or 'teléfono' in context.lower():
                context = 'Número de celular'
            elif 'expediente' in context.lower():
                context = 'Número de expediente'
            else:
                context = 'Número'
        
        if context and len(context) >= 2:
            base_name = context[:100]
            if base_name.lower() in campo_counter:
                campo_counter[base_name.lower()] += 1
                display_name = f"{base_name} {campo_counter[base_name.lower()]}"
            else:
                campo_counter[base_name.lower()] = 1
                display_name = base_name
            
            context_start = max(0, start - 30)
            context_end = min(len(text), end + 30)
            contexto = text[context_start:context_end]
            
            campos.append({
                'nombre': campo_to_key(display_name),
                'etiqueta': display_name,
                'tipo': 'text',
                'start': start,
                'end': end,
                'contexto': contexto,
                'match_text': full_match,
                'pattern_type': 'dots_underscores'
            })
    
    lines = text.split('\n')
    current_pos = 0
    for line in lines:
        if ':' in line and re.search(r':\s*$', line.strip()):
            parts = line.split(':')
            if parts[0].strip() and len(parts[0].strip()) > 2:
                campo_name = parts[0].strip()
                if campo_name.lower() not in campo_counter:
                    start = current_pos + line.find(campo_name)
                    end = current_pos + len(line)
                    
                    if not any(s <= start < e or s < end <= e for s, e in seen_positions):
                        seen_positions.add((start, end))
                        context_start = max(0, current_pos - 10)
                        context_end = min(len(text), current_pos + len(line) + 10)
                        
                        campos.append({
                            'nombre': campo_to_key(campo_name),
                            'etiqueta': campo_name,
                            'tipo': 'text',
                            'start': start,
                            'end': end,
                            'contexto': text[context_start:context_end],
                            'match_text': line.strip(),
                            'pattern_type': 'colon_field'
                        })
                        campo_counter[campo_name.lower()] = 1
        current_pos += len(line) + 1
    
    campos.sort(key=lambda x: x['start'])
    return campos


def generate_highlighted_html(text, campos):
    """
    Generate HTML with highlighted placeholders for preview.
    Properly escapes content to prevent XSS attacks.
    Iterates in forward order to preserve campo index synchronization.
    """
    from markupsafe import escape
    
    if not campos:
        escaped = str(escape(text))
        return escaped.replace('\n', '<br>')
    
    sorted_campos = sorted(enumerate(campos), key=lambda x: x[1]['start'])
    
    segments = []
    last_pos = 0
    
    for original_idx, campo in sorted_campos:
        start = campo['start']
        end = campo['end']
        
        if start > last_pos:
            before_text = text[last_pos:start]
            segments.append(str(escape(before_text)))
        
        match_text = text[start:end]
        escaped_match = str(escape(match_text))
        escaped_etiqueta = str(escape(campo['etiqueta']))
        escaped_nombre = str(escape(campo['nombre']))
        
        colors = ['bg-yellow-200', 'bg-green-200', 'bg-blue-200', 'bg-pink-200', 'bg-purple-200', 'bg-orange-200']
        color = colors[original_idx % len(colors)]
        
        highlighted = f'<span class="placeholder-highlight {color} px-1 rounded cursor-pointer hover:ring-2 hover:ring-blue-500" data-campo-index="{original_idx}" data-campo-nombre="{escaped_nombre}" title="{escaped_etiqueta}">{escaped_match}</span>'
        segments.append(highlighted)
        
        last_pos = end
    
    if last_pos < len(text):
        remaining_text = text[last_pos:]
        segments.append(str(escape(remaining_text)))
    
    result = ''.join(segments)
    result = result.replace('\n', '<br>')
    return result


def super_admin_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not current_user.is_super_admin():
            flash("Acceso restringido a super administradores.", "error")
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated_function


def admin_estudio_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not current_user.is_super_admin() and not current_user.is_admin_estudio():
            flash("Acceso restringido a administradores.", "error")
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated_function


def coordinador_or_admin_required(f):
    """Permite acceso a coordinadores, admin_estudio y super_admin."""
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not current_user.is_super_admin() and not current_user.is_admin_estudio() and not current_user.is_coordinador():
            flash("Acceso restringido a coordinadores y administradores.", "error")
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated_function


def get_current_tenant():
    if not current_user.is_authenticated:
        return None
    if current_user.is_super_admin() and 'impersonate_tenant_id' in session:
        return Tenant.query.get(session['impersonate_tenant_id'])
    return current_user.tenant


def get_tenant_id():
    tenant = get_current_tenant()
    return tenant.id if tenant else None


def get_resultados_folder(tenant=None):
    if tenant:
        folder = os.path.join(CARPETA_RESULTADOS, f"tenant_{tenant.id}")
    else:
        folder = CARPETA_RESULTADOS
    os.makedirs(folder, exist_ok=True)
    return folder


def cargar_plantilla(nombre_archivo, tenant_id=None):
    key = nombre_archivo.replace('.txt', '')
    
    if tenant_id:
        plantilla_db = Plantilla.query.filter_by(key=key, tenant_id=tenant_id, activa=True).first()
        if plantilla_db:
            return plantilla_db.contenido
    
    ruta = os.path.join(CARPETA_MODELOS, nombre_archivo)
    if os.path.exists(ruta):
        with open(ruta, "r", encoding="utf-8") as f:
            return f.read()
    return ""


def cargar_estilos(carpeta_estilos, tenant_id=None):
    if tenant_id:
        estilos_db = Estilo.query.filter_by(plantilla_key=carpeta_estilos, tenant_id=tenant_id, activo=True).all()
        if estilos_db:
            return "\n\n---\n\n".join([e.contenido for e in estilos_db])
    
    ruta_carpeta = os.path.join(CARPETA_ESTILOS, carpeta_estilos)
    estilos = []
    if os.path.exists(ruta_carpeta):
        for archivo in os.listdir(ruta_carpeta):
            if archivo.endswith(".txt"):
                ruta_archivo = os.path.join(ruta_carpeta, archivo)
                with open(ruta_archivo, "r", encoding="utf-8") as f:
                    estilos.append(f.read())
    return "\n\n---\n\n".join(estilos)


def construir_prompt(plantilla, estilos, datos_caso, campos_dinamicos=None, datos_tablas=None):
    datos_str = ""
    if campos_dinamicos and len(campos_dinamicos) > 0:
        for campo in campos_dinamicos:
            valor = datos_caso.get(campo.nombre_campo, '{{FALTA_DATO}}')
            datos_str += f"- {campo.etiqueta}: {valor}\n"
    else:
        datos_str = f"""- Invitado: {datos_caso.get('invitado', '{{FALTA_DATO}}')}
- Demandante: {datos_caso.get('demandante1', '{{FALTA_DATO}}')}
- DNI Demandante: {datos_caso.get('dni_demandante1', '{{FALTA_DATO}}')}
- Argumento 1: {datos_caso.get('argumento1', '{{FALTA_DATO}}')}
- Argumento 2: {datos_caso.get('argumento2', '{{FALTA_DATO}}')}
- Argumento 3: {datos_caso.get('argumento3', '{{FALTA_DATO}}')}
- Conclusión: {datos_caso.get('conclusion', '{{FALTA_DATO}}')}"""
    
    tablas_str = ""
    if datos_tablas:
        for tabla_nombre, tabla_info in datos_tablas.items():
            tablas_str += f"\n\n[TABLA: {tabla_nombre}]\n"
            columnas = tabla_info.get('columnas', [])
            filas = tabla_info.get('filas', [])
            if columnas:
                tablas_str += "| " + " | ".join(columnas) + " |\n"
                tablas_str += "|" + "|".join(["---"] * len(columnas)) + "|\n"
            for fila in filas:
                tablas_str += "| " + " | ".join([str(fila.get(col, '')) for col in columnas]) + " |\n"
            if tabla_info.get('total'):
                tablas_str += f"TOTAL: {tabla_info['total']}\n"
    
    prompt = f"""Eres un abogado experto del estudio jurídico especializado en derecho de familia.

══════════════════════════════════════════════════════════════
VOCABULARIO Y FRASES FORMALES OBLIGATORIAS:
══════════════════════════════════════════════════════════════
{estilos if estilos else "(No hay ejemplos de estilo disponibles)"}

══════════════════════════════════════════════════════════════
PLANTILLA BASE:
══════════════════════════════════════════════════════════════
{plantilla if plantilla else "(No hay plantilla disponible)"}

══════════════════════════════════════════════════════════════
DATOS DEL CASO:
══════════════════════════════════════════════════════════════
{datos_str}{tablas_str}

══════════════════════════════════════════════════════════════
INSTRUCCIONES:
══════════════════════════════════════════════════════════════
1. Usa las frases formales del vocabulario cuando corresponda a cada sección.
2. Si el vocabulario incluye citas legales, incorpóralas en la fundamentación jurídica.
3. Estructura el documento con secciones numeradas si corresponde (PRIMERO, SEGUNDO...).
4. Mantén tono formal y respetuoso.
5. Si falta un dato, conserva {{{{FALTA_DATO}}}}.
6. Montos en números y letras: S/1,000.00 (MIL CON 00/100 SOLES).
7. Usa mayúsculas para énfasis en términos legales.
8. Redacta el documento completo sin explicaciones adicionales.
9. Si hay tablas de datos (gastos, honorarios, etc.), incluye la tabla formateada en el documento usando el formato:
   [[TABLA:{{'tabla_nombre'}}]]
   La tabla será insertada automáticamente en esa ubicación."""
    return prompt


def generar_con_ia(prompt):
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Eres un abogado experto en redacción de documentos legales."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        return response.choices[0].message.content
    except Exception as e:
        logging.error(f"Error al generar con IA: {e}")
        return None


def get_tenant_logo_path(tenant):
    if tenant and tenant.logo_path:
        logo_path = os.path.join("static", "tenants", tenant.slug, tenant.logo_path)
        if os.path.exists(logo_path):
            return logo_path
    default_logo = os.path.join("static", "logo_estudio.png")
    if os.path.exists(default_logo):
        return default_logo
    return None


def agregar_tabla_word(doc, tabla_nombre, tabla_data):
    """Add a formatted table to the Word document."""
    from docx.shared import Inches
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    columnas = tabla_data.get('columnas', [])
    filas = tabla_data.get('filas', [])
    total = tabla_data.get('total')
    mostrar_total = tabla_data.get('mostrar_total', False)
    
    if not columnas or not filas:
        return
    
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run(tabla_nombre.upper())
    run_titulo.bold = True
    run_titulo.font.name = 'Times New Roman'
    run_titulo.font.size = Pt(11)
    p_titulo.paragraph_format.space_before = Pt(12)
    p_titulo.paragraph_format.space_after = Pt(6)
    
    num_filas = len(filas) + 1
    if mostrar_total and total:
        num_filas += 1
    
    table = doc.add_table(rows=num_filas, cols=len(columnas))
    table.style = 'Table Grid'
    
    header_row = table.rows[0]
    for i, col in enumerate(columnas):
        cell = header_row.cells[i]
        cell.text = col
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6E6E6"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    for row_idx, fila in enumerate(filas):
        row = table.rows[row_idx + 1]
        for col_idx, col in enumerate(columnas):
            cell = row.cells[col_idx]
            cell.text = str(fila.get(col, ''))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    
    if mostrar_total and total:
        total_row = table.rows[-1]
        total_row.cells[0].text = 'TOTAL'
        for paragraph in total_row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        
        total_row.cells[-1].text = str(total)
        for paragraph in total_row.cells[-1].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        
        for cell in total_row.cells:
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()


def guardar_docx(texto, nombre_archivo, tenant=None, datos_tablas=None):
    doc = Document()
    
    estilo_doc = None
    font_name = 'Times New Roman'
    font_size = 12
    line_spacing = 1.5
    
    if tenant:
        estilo_doc = EstiloDocumento.query.filter_by(tenant_id=tenant.id).first()
        if estilo_doc:
            font_name = estilo_doc.fuente
            font_size = estilo_doc.tamano_base
            line_spacing = estilo_doc.interlineado
    
    sections = doc.sections
    for section in sections:
        if estilo_doc:
            section.top_margin = Cm(estilo_doc.margen_superior)
            section.bottom_margin = Cm(estilo_doc.margen_inferior)
            section.left_margin = Cm(estilo_doc.margen_izquierdo)
            section.right_margin = Cm(estilo_doc.margen_derecho)
        else:
            section.top_margin = Cm(3.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)
        
        logo_path = get_tenant_logo_path(tenant)
        if logo_path and os.path.exists(logo_path):
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_para.add_run()
            run.add_picture(logo_path, width=Cm(4))
            
            if tenant:
                info_lines = tenant.get_header_info()
            else:
                info_lines = [
                    "Autorizado su funcionamiento por Resolución Directoral N.º 3562-2022-JUS/DGDPAJ-DCMA",
                    "Dirección: Av. Javier Prado Este 255, oficina 701. Distrito de San Isidro, Lima-Perú",
                    "Teléfono (01) – 6757575 / 994647890",
                    "Página web: www.abogadasperu.com"
                ]
            
            for linea in info_lines:
                info_para = header.add_paragraph()
                info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                info_run = info_para.add_run(linea)
                info_run.font.name = 'Times New Roman'
                info_run.font.size = Pt(9)
                info_para.paragraph_format.space_after = Pt(0)
                info_para.paragraph_format.space_before = Pt(0)
    
    titulos_principales = ['SUMILLA:', 'PETITORIO:', 'HECHOS:', 'FUNDAMENTOS', 'ANEXOS:', 
                          'POR TANTO:', 'VÍA PROCEDIMENTAL:', 'CONTRACAUTELA:',
                          'FUNDAMENTACION JURÍDICA:', 'FUNDAMENTACIÓN JURÍDICA:']
    titulos_secundarios = ['PRIMERO:', 'SEGUNDO:', 'TERCERO:', 'CUARTO:', 'QUINTO:',
                          'SEXTO:', 'SÉPTIMO:', 'OCTAVO:', 'NOVENO:', 'DÉCIMO:',
                          'DATOS DEL SOLICITANTE:', 'DATOS DE LOS SOLICITANTES:',
                          'NOMBRE Y DIRECCIÓN DEL DEMANDADO:', 'NOMBRE DEL INVITADO',
                          'OTRAS PERSONAS CON DERECHO ALIMENTARIO']
    encabezados = ['SEÑOR JUEZ', 'SEÑORA JUEZ', 'SEÑOR:', 'SEÑORA:', 'PRESENTE']
    
    import re
    tabla_pattern = re.compile(r'\[\[TABLA:([^\]]+)\]\]')
    imagen_pattern = re.compile(r'\{\{IMAGEN:([^}]+)\}\}')
    tablas_insertadas = set()
    
    for parrafo in texto.split("\n"):
        linea = parrafo.strip()
        if not linea:
            continue
        
        tabla_match = tabla_pattern.search(linea)
        if tabla_match and datos_tablas:
            tabla_ref = tabla_match.group(1).strip()
            for tabla_nombre, tabla_data in datos_tablas.items():
                if tabla_nombre.lower() in tabla_ref.lower() or tabla_ref.lower() in tabla_nombre.lower():
                    if tabla_nombre not in tablas_insertadas:
                        agregar_tabla_word(doc, tabla_nombre, tabla_data)
                        tablas_insertadas.add(tabla_nombre)
                    break
            continue
        
        imagen_match = imagen_pattern.search(linea)
        if imagen_match:
            imagen_key = imagen_match.group(1).strip().lower()
            imagen_key_norm = imagen_key.replace(' ', '_').replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u').replace('ñ', 'n')
            
            imagen_path = None
            if tenant:
                campo_imagen = CampoPlantilla.query.filter_by(
                    tenant_id=tenant.id,
                    tipo='file'
                ).filter(CampoPlantilla.nombre_campo.ilike(f'%{imagen_key_norm}%')).first()
                
                if campo_imagen and campo_imagen.archivo_path:
                    full_path = os.path.join(CARPETA_IMAGENES_MODELOS, campo_imagen.archivo_path)
                    if os.path.exists(full_path):
                        imagen_path = full_path
            
            if imagen_path:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    run.add_picture(imagen_path, width=Cm(8))
                except Exception as e:
                    logging.error(f"Error insertando imagen {imagen_key}: {e}")
                    p.add_run(f"[Imagen: {imagen_key}]")
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Imagen: {imagen_key} no encontrada]")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.italic = True
            continue
        
        p = doc.add_paragraph()
        run = p.add_run(linea)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        
        es_titulo_principal = any(linea.upper().startswith(t.upper()) for t in titulos_principales)
        es_titulo_secundario = any(linea.upper().startswith(t.upper()) for t in titulos_secundarios)
        es_encabezado = any(linea.upper().startswith(t.upper()) for t in encabezados)
        
        if es_titulo_principal:
            run.bold = True
            run.font.size = Pt(font_size)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
        elif es_titulo_secundario:
            run.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif es_encabezado:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
        elif linea.startswith('_____'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
        elif linea.upper().startswith('D.N.I'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
        
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = line_spacing
    
    if datos_tablas:
        for tabla_nombre, tabla_data in datos_tablas.items():
            if tabla_nombre not in tablas_insertadas:
                agregar_tabla_word(doc, tabla_nombre, tabla_data)
    
    folder = get_resultados_folder(tenant)
    ruta = os.path.join(folder, nombre_archivo)
    doc.save(ruta)
    return ruta


def validar_dato(valor):
    if not valor or valor.strip() == "":
        return "{{FALTA_DATO}}"
    return valor.strip()


def extraer_datos_tablas(form_data, tipo_documento, tenant_id):
    """Extract table data from form submission based on model tables."""
    datos_tablas = {}
    
    if not tenant_id:
        return datos_tablas
    
    modelo_db = Modelo.query.filter_by(key=tipo_documento, tenant_id=tenant_id).first()
    if not modelo_db:
        return datos_tablas
    
    tablas = ModeloTabla.query.filter_by(modelo_id=modelo_db.id, tenant_id=tenant_id).all()
    
    def normalize_field_name(text):
        text = text.lower().replace(' ', '_')
        text = text.replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
        text = text.replace('ñ', 'n')
        text = ''.join(c for c in text if c.isalnum() or c == '_')
        text = text.replace('.', '')
        return text
    
    for tabla in tablas:
        tabla_nombre = normalize_field_name(tabla.nombre)
        
        columnas = tabla.columnas if tabla.columnas else []
        filas = []
        
        for fila_idx in range(tabla.num_filas):
            fila_data = {}
            fila_has_data = False
            
            for col in columnas:
                col_nombre = normalize_field_name(col)
                
                campo_key = f"tabla_{tabla_nombre}_{fila_idx}_{col_nombre}"
                valor = form_data.get(campo_key, '').strip()
                fila_data[col] = valor
                if valor:
                    fila_has_data = True
            
            if fila_has_data:
                filas.append(fila_data)
        
        total_valor = None
        if tabla.mostrar_total and tabla.columna_total:
            col_total_nombre = normalize_field_name(tabla.columna_total)
            campo_total_key = f"tabla_{tabla_nombre}_total_{col_total_nombre}"
            total_valor = form_data.get(campo_total_key, '').strip()
        
        if filas:
            datos_tablas[tabla.nombre] = {
                'columnas': columnas,
                'filas': filas,
                'total': total_valor,
                'mostrar_total': tabla.mostrar_total
            }
    
    return datos_tablas


@app.context_processor
def inject_tenant():
    return {
        'current_tenant': get_current_tenant() if current_user.is_authenticated else None,
        'is_impersonating': 'impersonate_tenant_id' in session and current_user.is_authenticated and current_user.is_super_admin()
    }


@app.context_processor
def inject_current_year():
    return {'current_year': datetime.now().year}


ANONYMIZER_TEMP_DIR = os.path.join(os.getcwd(), 'temp_anonymizer')
ANONYMIZER_OUTPUT_DIR = os.path.join(os.getcwd(), 'anonymizer_output')
os.makedirs(ANONYMIZER_TEMP_DIR, exist_ok=True)
os.makedirs(ANONYMIZER_OUTPUT_DIR, exist_ok=True)


@app.route("/")
def index():
    return redirect(url_for('anonymizer_home'))


@app.route("/anonymizer")
def anonymizer_home():
    anon_module.cleanup_old_files(ANONYMIZER_TEMP_DIR, max_age_minutes=30)
    anon_module.cleanup_old_files(ANONYMIZER_OUTPUT_DIR, max_age_minutes=30)
    return render_template("anonymizer_home.html")


@app.route("/anonymizer/process", methods=["POST"])
def anonymizer_process():
    if 'documento' not in request.files:
        flash("Por favor selecciona un documento.", "error")
        return redirect(url_for('anonymizer_home'))
    
    archivo = request.files['documento']
    if archivo.filename == '':
        flash("No se seleccionó ningún archivo.", "error")
        return redirect(url_for('anonymizer_home'))
    
    ext = os.path.splitext(archivo.filename)[1].lower()
    if ext not in anon_module.ALLOWED_EXTENSIONS_ANON:
        flash("Formato no permitido. Solo se aceptan archivos DOCX y PDF.", "error")
        return redirect(url_for('anonymizer_home'))
    
    if archivo.content_length and archivo.content_length > anon_module.MAX_FILE_SIZE_MB * 1024 * 1024:
        flash(f"El archivo excede el tamaño máximo de {anon_module.MAX_FILE_SIZE_MB} MB.", "error")
        return redirect(url_for('anonymizer_home'))
    
    modo = request.form.get('modo', 'tokens')
    if modo not in ['tokens', 'asterisks', 'synthetic']:
        modo = 'tokens'
    strict_mode = 'strict_mode' in request.form
    generate_mapping = 'generate_mapping' in request.form
    
    job_id = uuid.uuid4().hex
    original_filename = secure_filename(archivo.filename)
    temp_path = os.path.join(ANONYMIZER_TEMP_DIR, f"{job_id}_{original_filename}")
    archivo.save(temp_path)
    
    try:
        needs_review = []
        
        if ext == '.docx':
            doc, summary, mapping, needs_review = anon_module.anonymize_docx(temp_path, mode=modo, strict_mode=strict_mode)
            output_filename = f"{job_id}_anonimizado.docx"
            output_path = os.path.join(ANONYMIZER_OUTPUT_DIR, output_filename)
            anon_module.save_anonymized_docx(doc, output_path)
            output_type = 'docx'
        else:
            text, summary, mapping, is_scanned, needs_review = anon_module.anonymize_pdf(temp_path, mode=modo, strict_mode=strict_mode)
            if is_scanned:
                flash("El PDF parece ser escaneado y no contiene texto extraíble. Por ahora solo soportamos PDFs con texto.", "warning")
                os.remove(temp_path)
                return redirect(url_for('anonymizer_home'))
            output_filename = f"{job_id}_anonimizado.pdf"
            output_path = os.path.join(ANONYMIZER_OUTPUT_DIR, output_filename)
            anon_module.create_anonymized_pdf(text, output_path)
            text_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_text.txt")
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(text)
            output_type = 'pdf'
        
        report = anon_module.generate_report(summary, original_filename, ext.upper())
        report_json_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_reporte.json")
        with open(report_json_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        report_txt = anon_module.generate_report_txt(report)
        report_txt_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_reporte.txt")
        with open(report_txt_path, 'w', encoding='utf-8') as f:
            f.write(report_txt)
        
        if generate_mapping:
            mapping_csv = anon_module.generate_mapping_csv(mapping)
            mapping_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_mapping.csv")
            with open(mapping_path, 'w', encoding='utf-8') as f:
                f.write(mapping_csv)
        
        job_state = {
            'original_filename': original_filename,
            'output_type': output_type,
            'modo': modo,
            'strict_mode': strict_mode,
            'generate_mapping': generate_mapping,
            'needs_review': needs_review,
            'summary': summary,
            'mapping_state': mapping.to_dict()
        }
        state_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_state.json")
        with open(state_path, 'w', encoding='utf-8') as f:
            json.dump(job_state, f, ensure_ascii=False)
        
        os.remove(temp_path)
        
        if strict_mode and needs_review:
            return render_template("anonymizer_review.html",
                                 job_id=job_id,
                                 pending_count=len(needs_review),
                                 pending_entities=needs_review,
                                 confirmed_count=summary.get('total_entities', 0),
                                 entities_summary=summary.get('entities_found', {}))
        
        return render_template("anonymizer_result.html",
                             job_id=job_id,
                             total_entities=summary.get('total_entities', 0),
                             entities_summary=summary.get('entities_found', {}),
                             replacements=summary.get('replacements', {}),
                             warnings=report.get('advertencias', []),
                             output_type=output_type,
                             mode=modo,
                             has_mapping=generate_mapping)
    
    except ValueError as e:
        flash(str(e), "error")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        return redirect(url_for('anonymizer_home'))
    except Exception as e:
        logging.error(f"Error en anonimización: {e}")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        flash("Error al procesar el documento. Por favor intenta de nuevo.", "error")
        return redirect(url_for('anonymizer_home'))


@app.route("/anonymizer/review/<job_id>/apply", methods=["POST"])
def anonymizer_apply_review(job_id):
    if not re.match(r'^[a-f0-9]{32}$', job_id):
        flash("ID de trabajo inválido.", "error")
        return redirect(url_for('anonymizer_home'))
    
    state_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_state.json")
    if not os.path.exists(state_path):
        flash("La sesión ha expirado. Por favor procesa el documento de nuevo.", "error")
        return redirect(url_for('anonymizer_home'))
    
    with open(state_path, 'r', encoding='utf-8') as f:
        job_state = json.load(f)
    
    decisions = {}
    for key, value in request.form.items():
        if key.startswith('decisions[') and key.endswith(']'):
            entity_id = key[10:-1]
            decisions[entity_id] = True
    
    pending_entities = job_state.get('needs_review', [])
    approved_entities = [e for e in pending_entities if decisions.get(e['id'], False)]
    approved_count = len(approved_entities)
    
    output_type = job_state.get('output_type', 'docx')
    modo = job_state.get('modo', 'tokens')
    
    mapping = anon_module.EntityMapping.from_dict(job_state.get('mapping_state', {'mode': modo}))
    
    if approved_entities:
        for e in approved_entities:
            _ = mapping.get_substitute(e['type'], e['value'])
        
        if output_type == 'docx':
            docx_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_anonimizado.docx")
            if os.path.exists(docx_path):
                from docx import Document
                doc = Document(docx_path)
                for e in approved_entities:
                    substitute = mapping.get_substitute(e['type'], e['value'])
                    for para in doc.paragraphs:
                        if e['value'] in para.text:
                            for run in para.runs:
                                if e['value'] in run.text:
                                    run.text = run.text.replace(e['value'], substitute)
                            if e['value'] in para.text:
                                para.text = para.text.replace(e['value'], substitute)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if e['value'] in para.text:
                                        para.text = para.text.replace(e['value'], substitute)
                doc.save(docx_path)
        else:
            pdf_text_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_text.txt")
            pdf_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_anonimizado.pdf")
            if os.path.exists(pdf_text_path):
                with open(pdf_text_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                for e in approved_entities:
                    substitute = mapping.get_substitute(e['type'], e['value'])
                    text = text.replace(e['value'], substitute)
                anon_module.create_anonymized_pdf(text, pdf_path)
                with open(pdf_text_path, 'w', encoding='utf-8') as f:
                    f.write(text)
    
    updated_summary = job_state['summary'].copy()
    updated_summary['total_entities'] = updated_summary.get('total_entities', 0) + approved_count
    updated_summary['entities_found'] = mapping.get_summary()
    updated_summary['replacements'] = mapping.get_replacements_for_report()
    
    job_state['summary'] = updated_summary
    job_state['mapping_state'] = mapping.to_dict()
    job_state['needs_review'] = [e for e in pending_entities if not decisions.get(e['id'], False)]
    with open(state_path, 'w', encoding='utf-8') as f:
        json.dump(job_state, f, ensure_ascii=False)
    
    report = anon_module.generate_report(updated_summary, job_state.get('original_filename', ''), output_type.upper())
    if approved_count > 0:
        report['advertencias'].append(f"Se anonimizaron {approved_count} entidades adicionales tras la revisión manual.")
    
    report_json_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_reporte.json")
    with open(report_json_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    report_txt = anon_module.generate_report_txt(report)
    report_txt_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_reporte.txt")
    with open(report_txt_path, 'w', encoding='utf-8') as f:
        f.write(report_txt)
    
    if job_state.get('generate_mapping', False):
        mapping_csv = anon_module.generate_mapping_csv(mapping)
        mapping_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_mapping.csv")
        with open(mapping_path, 'w', encoding='utf-8') as f:
            f.write(mapping_csv)
    
    return render_template("anonymizer_result.html",
                         job_id=job_id,
                         total_entities=updated_summary.get('total_entities', 0),
                         entities_summary=updated_summary.get('entities_found', {}),
                         replacements=updated_summary.get('replacements', {}),
                         warnings=report.get('advertencias', []),
                         output_type=output_type,
                         mode=modo,
                         has_mapping=job_state.get('generate_mapping', False))


@app.route("/anonymizer/review/<job_id>/skip")
def anonymizer_skip_review(job_id):
    if not re.match(r'^[a-f0-9]{32}$', job_id):
        flash("ID de trabajo inválido.", "error")
        return redirect(url_for('anonymizer_home'))
    
    state_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_state.json")
    if not os.path.exists(state_path):
        flash("La sesión ha expirado. Por favor procesa el documento de nuevo.", "error")
        return redirect(url_for('anonymizer_home'))
    
    with open(state_path, 'r', encoding='utf-8') as f:
        job_state = json.load(f)
    
    report_path = os.path.join(ANONYMIZER_OUTPUT_DIR, f"{job_id}_reporte.json")
    warnings = []
    if os.path.exists(report_path):
        with open(report_path, 'r', encoding='utf-8') as f:
            report = json.load(f)
            warnings = report.get('advertencias', [])
    
    return render_template("anonymizer_result.html",
                         job_id=job_id,
                         total_entities=job_state['summary'].get('total_entities', 0),
                         entities_summary=job_state['summary'].get('entities_found', {}),
                         replacements=job_state['summary'].get('replacements', {}),
                         warnings=warnings,
                         output_type=job_state.get('output_type', 'docx'),
                         mode=job_state.get('modo', 'tokens'),
                         has_mapping=job_state.get('generate_mapping', False))


@app.route("/anonymizer/download/<job_id>/<file_type>")
def anonymizer_download(job_id, file_type):
    if not re.match(r'^[a-f0-9]{32}$', job_id):
        flash("ID de trabajo inválido.", "error")
        return redirect(url_for('anonymizer_home'))
    
    if file_type == 'document':
        for ext in ['docx', 'pdf']:
            filename = f"{job_id}_anonimizado.{ext}"
            filepath = os.path.join(ANONYMIZER_OUTPUT_DIR, filename)
            if os.path.exists(filepath):
                return send_file(filepath, as_attachment=True, download_name=f"documento_anonimizado.{ext}")
    elif file_type == 'report':
        filename = f"{job_id}_reporte.txt"
        filepath = os.path.join(ANONYMIZER_OUTPUT_DIR, filename)
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True, download_name="reporte_anonimizacion.txt")
    elif file_type == 'report_json':
        filename = f"{job_id}_reporte.json"
        filepath = os.path.join(ANONYMIZER_OUTPUT_DIR, filename)
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True, download_name="reporte_anonimizacion.json")
    elif file_type == 'mapping':
        filename = f"{job_id}_mapping.csv"
        filepath = os.path.join(ANONYMIZER_OUTPUT_DIR, filename)
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True, download_name="diccionario_mapeo.csv")
    
    flash("Archivo no encontrado. Es posible que haya expirado.", "error")
    return redirect(url_for('anonymizer_home'))


@app.route("/dashboard-app")
@login_required
def dashboard_redirect():
    return redirect(url_for('dashboard'))


@app.route("/demo")
def demo():
    """Demo page with cost simulator."""
    # Initialize pricing defaults if needed
    PricingConfig.init_defaults()
    
    pricing = {
        'price_per_seat': float(PricingConfig.get_value('price_per_seat', '69.00')),
        'currency': PricingConfig.get_value('currency', 'USD'),
        'currency_symbol': PricingConfig.get_value('currency_symbol', '$'),
        'trial_days': int(PricingConfig.get_value('trial_days', '14')),
        'platform_name': PricingConfig.get_value('platform_name', 'LegalDoc Pro'),
        'min_seats': int(PricingConfig.get_value('min_seats', '1')),
        'max_seats': int(PricingConfig.get_value('max_seats', '100')),
    }
    
    addons = PricingAddon.query.filter_by(activo=True).order_by(PricingAddon.orden).all()
    
    return render_template("demo.html", pricing=pricing, addons=addons)


@app.route("/checkout/start")
def checkout_start():
    """Start checkout process."""
    PricingConfig.init_defaults()
    
    pricing = {
        'price_per_seat': float(PricingConfig.get_value('price_per_seat', '69.00')),
        'currency': PricingConfig.get_value('currency', 'USD'),
        'currency_symbol': PricingConfig.get_value('currency_symbol', '$'),
        'trial_days': int(PricingConfig.get_value('trial_days', '14')),
        'platform_name': PricingConfig.get_value('platform_name', 'LegalDoc Pro'),
        'min_seats': int(PricingConfig.get_value('min_seats', '1')),
        'max_seats': int(PricingConfig.get_value('max_seats', '100')),
    }
    
    addons = PricingAddon.query.filter_by(activo=True).order_by(PricingAddon.orden).all()
    
    return render_template("checkout_start.html", pricing=pricing, addons=addons)


@app.route("/checkout/create-session", methods=["POST"])
def checkout_create_session():
    """Create a checkout session and process trial signup."""
    from datetime import timedelta
    import secrets
    import re
    from email_validator import validate_email, EmailNotValidError
    
    nombre_estudio = request.form.get("nombre_estudio", "").strip()
    admin_nombre = request.form.get("admin_nombre", "").strip()
    admin_email = request.form.get("admin_email", "").strip().lower()
    seats = request.form.get("seats", type=int, default=1)
    addon_ids = request.form.getlist("addons")
    
    # Basic validation
    if not nombre_estudio or not admin_nombre or not admin_email:
        flash("Por favor completa todos los campos requeridos.", "error")
        return redirect(url_for('checkout_start'))
    
    # Validate email format
    try:
        validate_email(admin_email)
    except EmailNotValidError:
        flash("Por favor ingresa un email válido.", "error")
        return redirect(url_for('checkout_start'))
    
    # Get pricing config
    PricingConfig.init_defaults()
    min_seats = int(PricingConfig.get_value('min_seats', '1'))
    max_seats = int(PricingConfig.get_value('max_seats', '100'))
    price_per_seat = float(PricingConfig.get_value('price_per_seat', '69.00'))
    trial_days = int(PricingConfig.get_value('trial_days', '14'))
    
    # Validate seat bounds
    if seats < min_seats or seats > max_seats:
        seats = max(min_seats, min(seats, max_seats))
    
    # Check if email already exists
    existing_user = User.query.filter_by(email=admin_email).first()
    if existing_user:
        flash("Este email ya está registrado. Por favor inicia sesión.", "error")
        return redirect(url_for('login'))
    
    # Calculate total with validated addons
    subtotal = seats * price_per_seat
    addons_total = 0
    selected_addons = []
    
    if addon_ids:
        try:
            addon_int_ids = [int(aid) for aid in addon_ids]
            addons = PricingAddon.query.filter(
                PricingAddon.id.in_(addon_int_ids), 
                PricingAddon.activo == True
            ).all()
            for addon in addons:
                addons_total += float(addon.precio)
                selected_addons.append({'id': addon.id, 'nombre': addon.nombre, 'precio': float(addon.precio)})
        except (ValueError, TypeError):
            pass
    
    total_amount = subtotal + addons_total
    
    try:
        # Create checkout session
        session_id = secrets.token_urlsafe(32)
        checkout = CheckoutSession(
            session_id=session_id,
            nombre_estudio=nombre_estudio,
            admin_nombre=admin_nombre,
            admin_email=admin_email,
            seats=seats,
            addons=selected_addons,
            subtotal=subtotal,
            total_amount=total_amount,
            currency=PricingConfig.get_value('currency', 'USD'),
            status='pending',
            expires_at=datetime.utcnow() + timedelta(hours=24)
        )
        db.session.add(checkout)
        
        # Generate slug from studio name
        slug_base = re.sub(r'[^a-z0-9]+', '-', nombre_estudio.lower()).strip('-')
        if not slug_base:
            slug_base = 'estudio'
        slug = slug_base
        counter = 1
        while Tenant.query.filter_by(slug=slug).first():
            slug = f"{slug_base}-{counter}"
            counter += 1
        
        # Create tenant
        tenant = Tenant(
            nombre=nombre_estudio,
            slug=slug,
            subscription_status='trial',
            activo=True
        )
        db.session.add(tenant)
        db.session.flush()
        
        # Create admin user with temporary password
        temp_password = secrets.token_urlsafe(16)
        user = User(
            username=admin_nombre,
            email=admin_email,
            role='admin_estudio',
            tenant_id=tenant.id,
            password_set=False,
            first_login_completed=False,
            onboarding_completed=False,
            activo=True
        )
        user.set_password(temp_password)
        db.session.add(user)
        db.session.flush()
        
        # Create subscription
        subscription = Subscription(
            tenant_id=tenant.id,
            seats_purchased=seats,
            seats_used=1,
            status='trial',
            plan_type='monthly',
            price_per_seat=price_per_seat,
            currency=PricingConfig.get_value('currency', 'USD'),
            current_period_start=datetime.utcnow(),
            trial_ends_at=datetime.utcnow() + timedelta(days=trial_days)
        )
        db.session.add(subscription)
        
        # Update checkout session
        checkout.status = 'trial_started'
        checkout.tenant_id = tenant.id
        
        # Create activation token for password setup
        activation_token = secrets.token_urlsafe(32)
        activation = ActivationToken(
            user_id=user.id,
            token=activation_token,
            tipo='set_password',
            expires_at=datetime.utcnow() + timedelta(hours=48)
        )
        db.session.add(activation)
        
        db.session.commit()
        
        # Send activation email (after commit)
        activation_url = url_for('activate_account', token=activation_token, _external=True)
        platform_name = PricingConfig.get_value('platform_name', 'LegalDoc Pro')
        email_html = f'''
        <h2>¡Bienvenido a {platform_name}!</h2>
        <p>Hola {admin_nombre},</p>
        <p>Tu estudio <strong>{nombre_estudio}</strong> ha sido creado exitosamente.</p>
        <p>Tu prueba gratis de {trial_days} días comienza ahora.</p>
        <p>Para acceder a tu cuenta, primero debes establecer tu contraseña:</p>
        <p><a href="{activation_url}" style="display:inline-block;padding:12px 24px;background-color:#3B82F6;color:white;text-decoration:none;border-radius:6px;font-weight:bold;">Establecer mi contraseña</a></p>
        <p>Este enlace expira en 48 horas.</p>
        <p>Saludos,<br>El equipo de {platform_name}</p>
        '''
        send_notification_email(admin_email, f"Activa tu cuenta en {platform_name}", email_html)
        
        return redirect(url_for('checkout_success', session_id=session_id))
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error creating checkout session: {e}")
        flash("Ocurrió un error al procesar tu registro. Por favor intenta de nuevo.", "error")
        return redirect(url_for('checkout_start'))


@app.route("/checkout/success")
def checkout_success():
    """Show checkout success page."""
    session_id = request.args.get('session_id')
    if not session_id:
        return redirect(url_for('index'))
    
    checkout = CheckoutSession.query.filter_by(session_id=session_id).first()
    if not checkout:
        return redirect(url_for('index'))
    
    PricingConfig.init_defaults()
    platform_name = PricingConfig.get_value('platform_name', 'LegalDoc Pro')
    trial_days = int(PricingConfig.get_value('trial_days', '14'))
    
    return render_template("checkout_success.html", 
                          checkout=checkout, 
                          platform_name=platform_name,
                          trial_days=trial_days)


@app.route("/api/culqi/charge", methods=["POST"])
def culqi_create_charge():
    """Create a charge using Culqi token. Amount is calculated server-side for security."""
    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'error': 'No data provided'}), 400
    
    token_id = data.get('token_id')
    session_id = data.get('session_id')
    
    if not token_id or not session_id:
        return jsonify({'success': False, 'error': 'Missing required fields'}), 400
    
    # Load checkout session from database - amount is calculated server-side
    checkout = CheckoutSession.query.filter_by(session_id=session_id).first()
    if not checkout:
        logging.warning(f"Invalid checkout session attempted: {session_id}")
        return jsonify({'success': False, 'error': 'Sesión de pago inválida'}), 400
    
    if checkout.status == 'paid':
        return jsonify({'success': False, 'error': 'Esta sesión ya fue pagada'}), 400
    
    if checkout.status not in ['pending', 'trial']:
        return jsonify({'success': False, 'error': 'Estado de sesión inválido'}), 400
    
    # Get server-side values (never trust client for amounts)
    PricingConfig.init_defaults()
    currency = PricingConfig.get_value('currency', 'PEN')
    platform_name = PricingConfig.get_value('platform_name', 'LegalDoc Pro')
    
    amount_cents = int(checkout.total_amount * 100)  # Convert to cents
    email = checkout.admin_email
    description = f'Suscripción {platform_name} - {checkout.nombre_estudio}'
    
    culqi_private_key = os.environ.get('CULQI_PRIVATE_KEY')
    if not culqi_private_key:
        logging.error("CULQI_PRIVATE_KEY not configured")
        return jsonify({'success': False, 'error': 'Payment not configured'}), 500
    
    try:
        headers = {
            'Authorization': f'Bearer {culqi_private_key}',
            'Content-Type': 'application/json'
        }
        
        charge_data = {
            'amount': amount_cents,
            'currency_code': currency,
            'email': email,
            'source_id': token_id,
            'description': description
        }
        
        logging.info(f"Creating Culqi charge for session {session_id}: {amount_cents} {currency}")
        
        response = requests.post(
            'https://api.culqi.com/v2/charges',
            headers=headers,
            json=charge_data
        )
        
        result = response.json()
        
        if response.status_code in [200, 201]:
            charge_id = result.get('id')
            
            checkout.culqi_charge_id = charge_id
            checkout.culqi_token_id = token_id
            checkout.status = 'paid'
            checkout.paid_at = datetime.utcnow()
            
            if checkout.tenant:
                checkout.tenant.subscription_status = 'active'
                subscription = Subscription.query.filter_by(tenant_id=checkout.tenant_id).first()
                if subscription:
                    subscription.status = 'active'
            
            db.session.commit()
            logging.info(f"Culqi charge successful: {charge_id} for session {session_id}")
            
            return jsonify({
                'success': True,
                'charge_id': charge_id,
                'message': 'Pago procesado exitosamente'
            })
        else:
            error_message = result.get('user_message', result.get('merchant_message', 'Error procesando pago'))
            logging.error(f"Culqi charge failed for session {session_id}: {result}")
            return jsonify({'success': False, 'error': error_message}), 400
            
    except Exception as e:
        logging.error(f"Error creating Culqi charge for session {session_id}: {e}")
        return jsonify({'success': False, 'error': 'Error procesando el pago'}), 500


@app.route("/checkout/payment/<session_id>")
def checkout_payment(session_id):
    """Show payment form for a checkout session."""
    checkout = CheckoutSession.query.filter_by(session_id=session_id).first()
    if not checkout:
        flash("Sesión de checkout no encontrada.", "error")
        return redirect(url_for('checkout_start'))
    
    if checkout.status == 'paid':
        return redirect(url_for('checkout_success', session_id=session_id))
    
    culqi_public_key = os.environ.get('CULQI_PUBLIC_KEY', '')
    
    PricingConfig.init_defaults()
    pricing = {
        'currency': PricingConfig.get_value('currency', 'PEN'),
        'currency_symbol': PricingConfig.get_value('currency_symbol', 'S/'),
        'platform_name': PricingConfig.get_value('platform_name', 'LegalDoc Pro'),
    }
    
    return render_template("checkout_payment.html",
                          checkout=checkout,
                          pricing=pricing,
                          culqi_public_key=culqi_public_key)


@app.route("/activate/<token>")
def activate_account(token):
    """Show password setup page for activation token."""
    activation = ActivationToken.query.filter_by(token=token).first()
    
    if not activation:
        flash("Enlace de activación inválido.", "error")
        return redirect(url_for('login'))
    
    if not activation.is_valid():
        flash("Este enlace de activación ha expirado o ya fue usado.", "error")
        return redirect(url_for('login'))
    
    return render_template("activate.html", token=token, user=activation.user)


@app.route("/activate/<token>/set-password", methods=["POST"])
def activate_set_password(token):
    """Set password for activated account."""
    activation = ActivationToken.query.filter_by(token=token).first()
    
    if not activation or not activation.is_valid():
        flash("Enlace de activación inválido o expirado.", "error")
        return redirect(url_for('login'))
    
    password = request.form.get("password", "")
    password_confirm = request.form.get("password_confirm", "")
    
    if len(password) < 8:
        flash("La contraseña debe tener al menos 8 caracteres.", "error")
        return redirect(url_for('activate_account', token=token))
    
    if password != password_confirm:
        flash("Las contraseñas no coinciden.", "error")
        return redirect(url_for('activate_account', token=token))
    
    # Set password
    user = activation.user
    user.set_password(password)
    user.password_set = True
    activation.mark_used()
    
    db.session.commit()
    
    # Log in the user
    login_user(user)
    user.last_login = datetime.utcnow()
    db.session.commit()
    
    flash("¡Contraseña establecida exitosamente!", "success")
    
    # If 2FA is required, redirect to setup
    if user.twofa_required and not user.twofa_enabled:
        return redirect(url_for('setup_2fa'))
    
    # If onboarding not completed, redirect there
    if not user.onboarding_completed:
        return redirect(url_for('onboarding'))
    
    return redirect(url_for('dashboard'))


@app.route("/onboarding")
@login_required
def onboarding():
    """Onboarding wizard for new users."""
    tenant = get_current_tenant()
    if not tenant:
        return redirect(url_for('dashboard'))
    
    subscription = Subscription.query.filter_by(tenant_id=tenant.id).first()
    
    return render_template("onboarding.html", 
                          tenant=tenant, 
                          subscription=subscription)


@app.route("/onboarding/complete", methods=["POST"])
@login_required
def onboarding_complete():
    """Mark onboarding as complete."""
    current_user.onboarding_completed = True
    current_user.first_login_completed = True
    db.session.commit()
    
    flash("¡Bienvenido a la plataforma!", "success")
    return redirect(url_for('dashboard'))


@app.route("/dashboard")
@login_required
def dashboard():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    
    from datetime import timedelta
    today = datetime.now()
    week_ago = today - timedelta(days=7)
    month_ago = today - timedelta(days=30)
    
    total_documentos = 0
    docs_este_mes = 0
    docs_semana = 0
    documentos_recientes = []
    
    if tenant_id:
        total_documentos = DocumentRecord.query.filter_by(tenant_id=tenant_id).count()
        docs_este_mes = DocumentRecord.query.filter(
            DocumentRecord.tenant_id == tenant_id,
            DocumentRecord.fecha >= month_ago
        ).count()
        docs_semana = DocumentRecord.query.filter(
            DocumentRecord.tenant_id == tenant_id,
            DocumentRecord.fecha >= week_ago
        ).count()
        documentos_recientes = DocumentRecord.query.filter_by(tenant_id=tenant_id).order_by(
            DocumentRecord.fecha.desc()
        ).limit(5).all()
    
    total_plantillas = len(MODELOS)
    estilos_disponibles = 0
    if tenant_id:
        total_plantillas += Plantilla.query.filter_by(tenant_id=tenant_id, activa=True).count()
        estilos_disponibles = Estilo.query.filter_by(tenant_id=tenant_id, activo=True).count()
    
    total_usuarios = 0
    usuarios_activos = 0
    if tenant_id:
        total_usuarios = User.query.filter_by(tenant_id=tenant_id).count()
        usuarios_activos = User.query.filter_by(tenant_id=tenant_id, activo=True).count()
    
    promedio_diario = round(docs_semana / 7, 1) if docs_semana > 0 else 0
    
    tipo_mas_usado = "-"
    if tenant_id and total_documentos > 0:
        result = db.session.query(
            DocumentRecord.tipo_documento, 
            db.func.count(DocumentRecord.id).label('count')
        ).filter_by(tenant_id=tenant_id).group_by(
            DocumentRecord.tipo_documento
        ).order_by(db.desc('count')).first()
        if result:
            tipo_mas_usado = result[0][:20] + "..." if len(result[0]) > 20 else result[0]
    
    casos_activos = 0
    casos_pendientes = 0
    tareas_pendientes = 0
    tareas_vencidas = 0
    casos_recientes = []
    tareas_urgentes = []
    
    if tenant_id:
        casos_activos = Case.query.filter(
            Case.tenant_id == tenant_id,
            Case.estado.in_(['en_proceso', 'en_espera'])
        ).count()
        casos_pendientes = Case.query.filter_by(tenant_id=tenant_id, estado='por_comenzar').count()
        tareas_pendientes = Task.query.filter_by(tenant_id=tenant_id, estado='pendiente').count()
        tareas_vencidas = Task.query.filter(
            Task.tenant_id == tenant_id,
            Task.estado.notin_(['completado', 'cancelado']),
            Task.fecha_vencimiento.isnot(None),
            Task.fecha_vencimiento < today
        ).count()
        casos_recientes = Case.query.filter_by(tenant_id=tenant_id).order_by(
            Case.updated_at.desc()
        ).limit(5).all()
        tareas_urgentes = Task.query.filter(
            Task.tenant_id == tenant_id,
            Task.estado.notin_(['completado', 'cancelado'])
        ).order_by(Task.fecha_vencimiento.asc().nullslast()).limit(5).all()
    
    stats = {
        'total_documentos': total_documentos,
        'docs_este_mes': docs_este_mes,
        'docs_semana': docs_semana,
        'casos_activos': casos_activos,
        'casos_pendientes': casos_pendientes,
        'tareas_pendientes': tareas_pendientes,
        'tareas_vencidas': tareas_vencidas,
        'total_plantillas': total_plantillas,
        'estilos_disponibles': estilos_disponibles,
        'total_usuarios': total_usuarios,
        'usuarios_activos': usuarios_activos,
        'promedio_diario': promedio_diario,
        'tipo_mas_usado': tipo_mas_usado
    }
    
    return render_template("dashboard.html", stats=stats, documentos_recientes=documentos_recientes, casos_recientes=casos_recientes, tareas_urgentes=tareas_urgentes)


@app.route("/generador")
@login_required
def generador():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    
    modelos_completos = dict(MODELOS)
    if tenant_id:
        plantillas_db = Plantilla.query.filter_by(tenant_id=tenant_id, activa=True).all()
        for p in plantillas_db:
            if p.key not in modelos_completos:
                modelos_completos[p.key] = {
                    "nombre": p.nombre,
                    "plantilla": f"{p.key}.txt",
                    "carpeta_estilos": p.carpeta_estilos or p.key
                }
        
        if current_user.can_manage_cases() or current_user.role in ['admin_estudio', 'super_admin', 'coordinador']:
            modelos_usuario = Modelo.query.filter_by(tenant_id=tenant_id, activa=True).all()
        else:
            modelos_usuario = Modelo.query.filter_by(tenant_id=tenant_id, created_by_id=current_user.id, activa=True).all()
        
        for m in modelos_usuario:
            if m.key not in modelos_completos:
                modelos_completos[m.key] = {
                    "nombre": m.nombre,
                    "plantilla": f"{m.key}.txt",
                    "carpeta_estilos": m.carpeta_estilos or m.key,
                    "es_modelo_usuario": True
                }
    
    return render_template("generador.html", modelos=modelos_completos, tenant=tenant)


@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == "POST":
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        
        user = User.query.filter_by(email=email, activo=True).first()
        if user and user.check_password(password):
            if user.twofa_enabled:
                session['pending_2fa_user_id'] = user.id
                next_page = request.args.get('next')
                if next_page:
                    session['next_after_2fa'] = next_page
                return redirect(url_for('verificar_2fa'))
            
            if user.requires_2fa() and not user.twofa_enabled:
                user.last_login = datetime.utcnow()
                db.session.commit()
                login_user(user)
                flash("Tu rol requiere autenticación de dos factores. Por favor configúrala ahora.", "warning")
                return redirect(url_for('activar_2fa'))
            
            user.last_login = datetime.utcnow()
            db.session.commit()
            login_user(user)
            flash("Sesión iniciada correctamente.", "success")
            next_page = request.args.get('next')
            return redirect(next_page or url_for('index'))
        else:
            flash("Email o contraseña incorrectos.", "error")
    
    return render_template("login.html")


@app.route("/forgot_password", methods=["GET", "POST"])
def forgot_password():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        
        if not email:
            flash("Por favor ingresa tu correo electrónico.", "error")
            return render_template("forgot_password.html")
        
        user = User.query.filter_by(email=email, activo=True).first()
        
        if user:
            existing_tokens = ActivationToken.query.filter_by(
                user_id=user.id, 
                tipo='password_reset', 
                used=False
            ).all()
            for t in existing_tokens:
                t.used = True
            db.session.commit()
            
            token = ActivationToken.create_token(user.id, tipo='password_reset', hours=24)
            
            try:
                app_base_url = os.environ.get('APP_BASE_URL', request.host_url.rstrip('/'))
                reset_url = f"{app_base_url}/reset_password/{token.token}"
                
                import resend
                resend.api_key = os.environ.get('RESEND_API_KEY')
                mail_from = os.environ.get('MAIL_FROM', 'noreply@resend.dev')
                
                html_content = f'''
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                    <h2 style="color: #333;">Recuperación de Contraseña</h2>
                    <p>Hola {user.username},</p>
                    <p>Recibimos una solicitud para restablecer tu contraseña. Haz clic en el siguiente enlace para crear una nueva contraseña:</p>
                    <p style="margin: 20px 0;">
                        <a href="{reset_url}" style="background-color: #2563eb; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block;">
                            Restablecer Contraseña
                        </a>
                    </p>
                    <p>Este enlace expirará en 24 horas.</p>
                    <p>Si no solicitaste restablecer tu contraseña, puedes ignorar este correo.</p>
                    <hr style="border: none; border-top: 1px solid #eee; margin: 20px 0;">
                    <p style="color: #666; font-size: 12px;">Centro de Conciliación - Sistema de Gestión</p>
                </div>
                '''
                
                resend.Emails.send({
                    "from": f"Centro de Conciliación <{mail_from}>",
                    "to": [user.email],
                    "subject": "Recuperación de Contraseña",
                    "html": html_content
                })
                
                log_audit(
                    tenant_id=user.tenant_id,
                    user_id=user.id,
                    accion='password_reset_requested',
                    detalle=f'Solicitud de recuperación de contraseña para {user.email}'
                )
            except Exception as e:
                logging.error(f"Error sending password reset email: {e}")
        
        flash("Si el correo existe en nuestro sistema, recibirás instrucciones para restablecer tu contraseña.", "success")
        return redirect(url_for('login'))
    
    return render_template("forgot_password.html")


@app.route("/reset_password/<token>", methods=["GET", "POST"])
def reset_password(token):
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    activation = ActivationToken.query.filter_by(token=token, tipo='password_reset').first()
    
    if not activation or not activation.is_valid():
        flash("El enlace de recuperación es inválido o ha expirado.", "error")
        return redirect(url_for('forgot_password'))
    
    user = activation.user
    
    if request.method == "POST":
        password = request.form.get("password", "")
        password_confirm = request.form.get("password_confirm", "")
        
        if not password or len(password) < 8:
            flash("La contraseña debe tener al menos 8 caracteres.", "error")
            return render_template("reset_password.html", token=token)
        
        if password != password_confirm:
            flash("Las contraseñas no coinciden.", "error")
            return render_template("reset_password.html", token=token)
        
        user.set_password(password)
        activation.mark_used()
        
        log_audit(
            tenant_id=user.tenant_id,
            user_id=user.id,
            accion='password_reset_completed',
            detalle=f'Contraseña restablecida exitosamente para {user.email}'
        )
        
        flash("Tu contraseña ha sido actualizada. Ya puedes iniciar sesión.", "success")
        return redirect(url_for('login'))
    
    return render_template("reset_password.html", token=token)


@app.route("/mi-cuenta/cambiar-password", methods=["GET", "POST"])
@login_required
def cambiar_password():
    """Permite al usuario logueado cambiar su contraseña."""
    if request.method == "POST":
        current_password = request.form.get("current_password", "")
        new_password = request.form.get("new_password", "")
        confirm_password = request.form.get("confirm_password", "")
        
        if not current_password or not new_password or not confirm_password:
            flash("Todos los campos son obligatorios.", "error")
            return render_template("cambiar_password.html")
        
        if not current_user.check_password(current_password):
            flash("La contraseña actual es incorrecta.", "error")
            return render_template("cambiar_password.html")
        
        if len(new_password) < 8:
            flash("La nueva contraseña debe tener al menos 8 caracteres.", "error")
            return render_template("cambiar_password.html")
        
        if new_password != confirm_password:
            flash("Las contraseñas nuevas no coinciden.", "error")
            return render_template("cambiar_password.html")
        
        if current_password == new_password:
            flash("La nueva contraseña debe ser diferente a la actual.", "error")
            return render_template("cambiar_password.html")
        
        current_user.set_password(new_password)
        db.session.commit()
        
        log_audit(
            tenant_id=current_user.tenant_id,
            user_id=current_user.id,
            accion='password_changed',
            detalle='El usuario cambió su contraseña'
        )
        
        flash("Tu contraseña ha sido actualizada correctamente.", "success")
        return redirect(url_for('dashboard'))
    
    return render_template("cambiar_password.html")


@app.route("/registro", methods=["GET", "POST"])
def registro():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    return render_template("registro.html")


@app.route("/registro_estudio", methods=["GET", "POST"])
def registro_estudio():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == "POST":
        nombre_estudio = request.form.get("nombre_estudio", "").strip()
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        password_confirm = request.form.get("password_confirm", "")
        
        if not nombre_estudio or not username or not email or not password:
            flash("Todos los campos son obligatorios.", "error")
            return render_template("registro_estudio.html")
        
        if password != password_confirm:
            flash("Las contraseñas no coinciden.", "error")
            return render_template("registro_estudio.html")
        
        if len(password) < 6:
            flash("La contraseña debe tener al menos 6 caracteres.", "error")
            return render_template("registro_estudio.html")
        
        if User.query.filter_by(email=email).first():
            flash("Ya existe una cuenta con este email.", "error")
            return render_template("registro_estudio.html")
        
        if User.query.filter_by(username=username).first():
            flash("Ya existe un usuario con este nombre. Por favor elige otro.", "error")
            return render_template("registro_estudio.html")
        
        slug = nombre_estudio.lower().replace(" ", "-").replace(".", "")[:50]
        base_slug = slug
        counter = 1
        while Tenant.query.filter_by(slug=slug).first():
            slug = f"{base_slug}-{counter}"
            counter += 1
        
        tenant = Tenant(
            nombre=nombre_estudio,
            slug=slug,
            activo=True
        )
        db.session.add(tenant)
        db.session.flush()
        
        is_first_user = User.query.count() == 0
        user = User(
            username=username,
            email=email,
            tenant_id=tenant.id,
            role='super_admin' if is_first_user else 'admin_estudio',
            activo=True
        )
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        
        login_user(user)
        flash(f"Estudio '{nombre_estudio}' creado exitosamente. Eres el administrador.", "success")
        return redirect(url_for('index'))
    
    return render_template("registro_estudio.html")


@app.route("/logout")
@login_required
def logout():
    if 'impersonate_tenant_id' in session:
        del session['impersonate_tenant_id']
    if 'pending_2fa_user_id' in session:
        del session['pending_2fa_user_id']
    logout_user()
    flash("Sesión cerrada correctamente.", "success")
    return redirect(url_for('index'))


@app.route("/seguridad")
@login_required
def seguridad():
    """Security settings page for 2FA management."""
    requires_2fa = current_user.requires_2fa()
    backup_codes_count = len(current_user.twofa_backup_codes_hashed) if current_user.twofa_backup_codes_hashed else 0
    return render_template("seguridad.html", 
                          requires_2fa=requires_2fa,
                          backup_codes_count=backup_codes_count)


@app.route("/seguridad/activar_2fa", methods=["GET", "POST"])
@login_required
def activar_2fa():
    """Setup 2FA for the user."""
    if current_user.twofa_enabled:
        flash("Ya tienes la autenticación de dos factores activada.", "info")
        return redirect(url_for('seguridad'))
    
    if request.method == "POST":
        codigo = request.form.get("codigo", "").strip()
        
        if not codigo:
            flash("Ingresa el código de 6 dígitos.", "error")
            return redirect(url_for('activar_2fa'))
        
        if current_user.verify_totp(codigo):
            current_user.twofa_enabled = True
            current_user.twofa_last_verified_at = datetime.utcnow()
            backup_codes = current_user.generate_backup_codes()
            
            TwoFALog.log_event(
                user_id=current_user.id,
                event_type='setup',
                success=True,
                ip=request.remote_addr,
                user_agent=request.headers.get('User-Agent', '')[:500]
            )
            db.session.commit()
            
            session['show_backup_codes'] = backup_codes
            return redirect(url_for('mostrar_backup_codes'))
        else:
            TwoFALog.log_event(
                user_id=current_user.id,
                event_type='setup',
                success=False,
                ip=request.remote_addr,
                user_agent=request.headers.get('User-Agent', '')[:500],
                details='Invalid TOTP code during setup'
            )
            db.session.commit()
            flash("Código incorrecto. Verifica que el código coincida con tu app.", "error")
            return redirect(url_for('activar_2fa'))
    
    current_user.generate_totp_secret()
    db.session.commit()
    
    totp_uri = current_user.get_totp_uri()
    qr = qrcode.make(totp_uri)
    buffer = BytesIO()
    qr.save(buffer, format='PNG')
    qr_base64 = base64.b64encode(buffer.getvalue()).decode()
    
    secret = current_user.get_totp_secret()
    
    return render_template("activar_2fa.html", 
                          qr_code=qr_base64, 
                          secret=secret)


@app.route("/seguridad/backup_codes")
@login_required
def mostrar_backup_codes():
    """Show backup codes after 2FA setup."""
    backup_codes = session.pop('show_backup_codes', None)
    if not backup_codes:
        flash("No hay códigos de respaldo para mostrar.", "info")
        return redirect(url_for('seguridad'))
    
    return render_template("backup_codes.html", backup_codes=backup_codes)


@app.route("/seguridad/regenerar_backup", methods=["POST"])
@login_required
def regenerar_backup_codes():
    """Regenerate backup codes."""
    if not current_user.twofa_enabled:
        flash("Primero debes activar la autenticación de dos factores.", "error")
        return redirect(url_for('seguridad'))
    
    password = request.form.get("password", "")
    if not current_user.check_password(password):
        flash("Contraseña incorrecta.", "error")
        return redirect(url_for('seguridad'))
    
    backup_codes = current_user.generate_backup_codes()
    
    TwoFALog.log_event(
        user_id=current_user.id,
        event_type='regenerate_backup',
        success=True,
        ip=request.remote_addr,
        user_agent=request.headers.get('User-Agent', '')[:500]
    )
    db.session.commit()
    
    session['show_backup_codes'] = backup_codes
    return redirect(url_for('mostrar_backup_codes'))


@app.route("/seguridad/desactivar_2fa", methods=["POST"])
@login_required
def desactivar_2fa():
    """Disable 2FA for the user."""
    if not current_user.twofa_enabled:
        flash("La autenticación de dos factores no está activada.", "info")
        return redirect(url_for('seguridad'))
    
    if current_user.requires_2fa():
        flash("Tu rol requiere autenticación de dos factores. No puedes desactivarla.", "error")
        return redirect(url_for('seguridad'))
    
    password = request.form.get("password", "")
    if not current_user.check_password(password):
        flash("Contraseña incorrecta.", "error")
        return redirect(url_for('seguridad'))
    
    current_user.disable_2fa()
    
    TwoFALog.log_event(
        user_id=current_user.id,
        event_type='disable',
        success=True,
        ip=request.remote_addr,
        user_agent=request.headers.get('User-Agent', '')[:500]
    )
    db.session.commit()
    
    flash("Autenticación de dos factores desactivada.", "success")
    return redirect(url_for('seguridad'))


@app.route("/verificar_2fa", methods=["GET", "POST"])
def verificar_2fa():
    """Second step of login: verify 2FA code."""
    if 'pending_2fa_user_id' not in session:
        return redirect(url_for('login'))
    
    user_id = session['pending_2fa_user_id']
    user = User.query.get(user_id)
    
    if not user:
        session.pop('pending_2fa_user_id', None)
        return redirect(url_for('login'))
    
    failed_attempts = TwoFALog.count_recent_failures(user_id)
    if failed_attempts >= 5:
        flash("Demasiados intentos fallidos. Intenta de nuevo en 15 minutos.", "error")
        return render_template("verificar_2fa.html", locked=True)
    
    if request.method == "POST":
        codigo = request.form.get("codigo", "").strip().replace(" ", "").replace("-", "")
        use_backup = request.form.get("use_backup") == "1"
        
        if not codigo:
            flash("Ingresa un código.", "error")
            return render_template("verificar_2fa.html", locked=False)
        
        success = False
        event_type = 'verify_attempt'
        
        if use_backup:
            codigo = codigo.upper()
            if "-" not in codigo and len(codigo) == 8:
                codigo = f"{codigo[:4]}-{codigo[4:]}"
            success = user.verify_backup_code(codigo)
            event_type = 'backup_used'
        else:
            success = user.verify_totp(codigo)
        
        TwoFALog.log_event(
            user_id=user_id,
            event_type=event_type,
            success=success,
            ip=request.remote_addr,
            user_agent=request.headers.get('User-Agent', '')[:500]
        )
        
        if success:
            user.twofa_last_verified_at = datetime.utcnow()
            user.last_login = datetime.utcnow()
            db.session.commit()
            
            session.pop('pending_2fa_user_id', None)
            login_user(user)
            
            flash("Sesión iniciada correctamente.", "success")
            next_page = session.pop('next_after_2fa', None)
            return redirect(next_page or url_for('index'))
        else:
            db.session.commit()
            remaining = 5 - failed_attempts - 1
            flash(f"Código incorrecto. Te quedan {remaining} intentos.", "error")
            return render_template("verificar_2fa.html", locked=False)
    
    return render_template("verificar_2fa.html", locked=False)


@app.route("/admin/reset_2fa/<int:user_id>", methods=["POST"])
@coordinador_or_admin_required
def reset_2fa_usuario(user_id):
    """Reset 2FA for a user in the same tenant."""
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    target_user = User.query.get_or_404(user_id)
    
    if not current_user.is_super_admin():
        if target_user.tenant_id != tenant.id:
            flash("No puedes modificar usuarios de otro estudio.", "error")
            return redirect(url_for("admin_usuarios"))
        if target_user.is_super_admin():
            flash("No puedes resetear 2FA de un super administrador.", "error")
            return redirect(url_for("admin_usuarios"))
    
    target_user.disable_2fa()
    
    TwoFALog.log_event(
        user_id=target_user.id,
        event_type='reset',
        success=True,
        ip=request.remote_addr,
        user_agent=request.headers.get('User-Agent', '')[:500],
        reset_by_id=current_user.id,
        details=f'Reset by {current_user.email}'
    )
    db.session.commit()
    
    flash(f"2FA reseteado para {target_user.username}. El usuario deberá configurarlo nuevamente.", "success")
    return redirect(url_for("admin_usuarios"))


@app.route("/procesar_ia", methods=["POST"])
@login_required
def procesar_ia():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    tipo_documento = request.form.get("tipo_documento")
    
    plantilla_db = None
    if tenant_id:
        plantilla_db = Plantilla.query.filter_by(key=tipo_documento, tenant_id=tenant_id, activa=True).first()
    
    if tipo_documento in MODELOS:
        modelo = MODELOS[tipo_documento]
    elif plantilla_db:
        modelo = {
            "nombre": plantilla_db.nombre,
            "plantilla": f"{tipo_documento}.txt",
            "carpeta_estilos": plantilla_db.carpeta_estilos or tipo_documento
        }
    else:
        flash("Tipo de documento no válido.", "error")
        return redirect(url_for("index"))
    
    if tenant_id:
        campos_dinamicos = CampoPlantilla.query.filter_by(plantilla_key=tipo_documento, tenant_id=tenant_id).order_by(CampoPlantilla.orden).all()
    else:
        campos_dinamicos = []
    
    if campos_dinamicos:
        datos_caso = {}
        archivos_subidos = {}
        for campo in campos_dinamicos:
            if campo.tipo == 'file':
                archivo = request.files.get(campo.nombre_campo)
                if archivo and archivo.filename:
                    from werkzeug.utils import secure_filename
                    import uuid
                    filename = secure_filename(archivo.filename)
                    unique_filename = f"{uuid.uuid4().hex}_{filename}"
                    upload_folder = os.path.join('archivos_campos', f'tenant_{tenant_id}')
                    os.makedirs(upload_folder, exist_ok=True)
                    filepath = os.path.join(upload_folder, unique_filename)
                    archivo.save(filepath)
                    datos_caso[campo.nombre_campo] = f"[Archivo: {filename}]"
                    archivos_subidos[campo.nombre_campo] = filepath
                else:
                    datos_caso[campo.nombre_campo] = "[Sin archivo]"
            else:
                datos_caso[campo.nombre_campo] = validar_dato(request.form.get(campo.nombre_campo, ""))
    else:
        datos_caso = {
            "invitado": validar_dato(request.form.get("invitado", "")),
            "demandante1": validar_dato(request.form.get("demandante1", "")),
            "dni_demandante1": validar_dato(request.form.get("dni_demandante1", "")),
            "argumento1": validar_dato(request.form.get("argumento1", "")),
            "argumento2": validar_dato(request.form.get("argumento2", "")),
            "argumento3": validar_dato(request.form.get("argumento3", "")),
            "conclusion": validar_dato(request.form.get("conclusion", ""))
        }
    
    datos_tablas = extraer_datos_tablas(request.form, tipo_documento, tenant_id)
    
    plantilla = cargar_plantilla(modelo["plantilla"], tenant_id)
    estilos = cargar_estilos(modelo["carpeta_estilos"], tenant_id)
    prompt = construir_prompt(plantilla, estilos, datos_caso, campos_dinamicos if campos_dinamicos else None, datos_tablas)
    
    texto_generado = generar_con_ia(prompt)
    
    if not texto_generado:
        flash("Error al generar el documento. Verifica tu API key de OpenAI.", "error")
        return redirect(url_for("index"))
    
    fecha_actual = datetime.now()
    nombre_archivo = f"{tipo_documento}_{fecha_actual.strftime('%Y%m%d_%H%M%S')}.docx"
    
    guardar_docx(texto_generado, nombre_archivo, tenant, datos_tablas)
    
    demandante_campo = datos_caso.get("demandante1") or datos_caso.get("nombre_demandante") or datos_caso.get("demandante") or "Sin nombre"
    if demandante_campo == "{{FALTA_DATO}}":
        demandante_campo = "Sin nombre"
    
    record = DocumentRecord(
        user_id=current_user.id,
        tenant_id=tenant_id,
        fecha=fecha_actual,
        tipo_documento=modelo["nombre"],
        tipo_documento_key=tipo_documento,
        demandante=demandante_campo,
        archivo=nombre_archivo,
        texto_generado=texto_generado,
        datos_caso=datos_caso
    )
    db.session.add(record)
    db.session.commit()
    
    flash(f"Documento generado exitosamente: {nombre_archivo}", "success")
    return redirect(url_for("descargar", nombre_archivo=nombre_archivo))


@app.route("/preview", methods=["POST"])
@login_required
def preview():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    tipo_documento = request.form.get("tipo_documento")
    
    plantilla_db = None
    if tenant_id:
        plantilla_db = Plantilla.query.filter_by(key=tipo_documento, tenant_id=tenant_id, activa=True).first()
    
    if tipo_documento in MODELOS:
        modelo = MODELOS[tipo_documento]
    elif plantilla_db:
        modelo = {
            "nombre": plantilla_db.nombre,
            "plantilla": f"{tipo_documento}.txt",
            "carpeta_estilos": plantilla_db.carpeta_estilos or tipo_documento
        }
    else:
        flash("Tipo de documento no válido.", "error")
        return redirect(url_for("index"))
    
    if tenant_id:
        campos_dinamicos = CampoPlantilla.query.filter_by(plantilla_key=tipo_documento, tenant_id=tenant_id).order_by(CampoPlantilla.orden).all()
    else:
        campos_dinamicos = []
    
    if campos_dinamicos:
        datos_caso = {}
        archivos_subidos = {}
        for campo in campos_dinamicos:
            if campo.tipo == 'file':
                archivo = request.files.get(campo.nombre_campo)
                if archivo and archivo.filename:
                    from werkzeug.utils import secure_filename
                    import uuid
                    filename = secure_filename(archivo.filename)
                    unique_filename = f"{uuid.uuid4().hex}_{filename}"
                    upload_folder = os.path.join('archivos_campos', f'tenant_{tenant_id}')
                    os.makedirs(upload_folder, exist_ok=True)
                    filepath = os.path.join(upload_folder, unique_filename)
                    archivo.save(filepath)
                    datos_caso[campo.nombre_campo] = f"[Archivo: {filename}]"
                    archivos_subidos[campo.nombre_campo] = filepath
                else:
                    datos_caso[campo.nombre_campo] = "[Sin archivo]"
            else:
                datos_caso[campo.nombre_campo] = validar_dato(request.form.get(campo.nombre_campo, ""))
    else:
        datos_caso = {
            "invitado": validar_dato(request.form.get("invitado", "")),
            "demandante1": validar_dato(request.form.get("demandante1", "")),
            "dni_demandante1": validar_dato(request.form.get("dni_demandante1", "")),
            "argumento1": validar_dato(request.form.get("argumento1", "")),
            "argumento2": validar_dato(request.form.get("argumento2", "")),
            "argumento3": validar_dato(request.form.get("argumento3", "")),
            "conclusion": validar_dato(request.form.get("conclusion", ""))
        }
    
    datos_tablas = extraer_datos_tablas(request.form, tipo_documento, tenant_id)
    
    plantilla = cargar_plantilla(modelo["plantilla"], tenant_id)
    estilos = cargar_estilos(modelo["carpeta_estilos"], tenant_id)
    prompt = construir_prompt(plantilla, estilos, datos_caso, campos_dinamicos if campos_dinamicos else None, datos_tablas)
    
    texto_generado = generar_con_ia(prompt)
    
    if not texto_generado:
        flash("Error al generar el preview. Verifica tu API key de OpenAI.", "error")
        return redirect(url_for("index"))
    
    return render_template("preview.html", 
                          texto=texto_generado, 
                          datos_caso=datos_caso,
                          datos_tablas=datos_tablas,
                          tipo_documento=tipo_documento,
                          modelo=modelo)


@app.route("/guardar_desde_preview", methods=["POST"])
@login_required
def guardar_desde_preview():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    tipo_documento = request.form.get("tipo_documento")
    texto_editado = request.form.get("texto_editado")
    
    plantilla_db = None
    if tenant_id:
        plantilla_db = Plantilla.query.filter_by(key=tipo_documento, tenant_id=tenant_id, activa=True).first()
    
    if tipo_documento in MODELOS:
        modelo = MODELOS[tipo_documento]
    elif plantilla_db:
        modelo = {
            "nombre": plantilla_db.nombre,
            "plantilla": f"{tipo_documento}.txt",
            "carpeta_estilos": plantilla_db.carpeta_estilos or tipo_documento
        }
    else:
        flash("Tipo de documento no válido.", "error")
        return redirect(url_for("index"))
    
    datos_caso_str = request.form.get("datos_caso", "{}")
    try:
        datos_caso = json.loads(datos_caso_str)
    except:
        datos_caso = {}
    
    datos_tablas_str = request.form.get("datos_tablas", "{}")
    try:
        datos_tablas = json.loads(datos_tablas_str)
    except:
        datos_tablas = {}
    
    fecha_actual = datetime.now()
    nombre_archivo = f"{tipo_documento}_{fecha_actual.strftime('%Y%m%d_%H%M%S')}.docx"
    
    guardar_docx(texto_editado, nombre_archivo, tenant, datos_tablas if datos_tablas else None)
    
    demandante_campo = datos_caso.get("demandante1") or datos_caso.get("nombre_demandante") or datos_caso.get("demandante") or "Sin nombre"
    if demandante_campo == "{{FALTA_DATO}}":
        demandante_campo = "Sin nombre"
    
    record = DocumentRecord(
        user_id=current_user.id,
        tenant_id=tenant_id,
        fecha=fecha_actual,
        tipo_documento=modelo["nombre"],
        tipo_documento_key=tipo_documento,
        demandante=demandante_campo,
        archivo=nombre_archivo,
        texto_generado=texto_editado,
        datos_caso=datos_caso
    )
    db.session.add(record)
    db.session.commit()
    
    flash(f"Documento guardado exitosamente: {nombre_archivo}", "success")
    return redirect(url_for("descargar", nombre_archivo=nombre_archivo))


@app.route("/editar/<int:doc_id>", methods=["GET", "POST"])
@login_required
def editar_documento(doc_id):
    record = DocumentRecord.query.get_or_404(doc_id)
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    
    if not tenant_id:
        flash("Necesitas un contexto de estudio para editar documentos.", "error")
        return redirect(url_for("historial"))
    
    if record.tenant_id != tenant_id:
        flash("No tienes permiso para editar este documento.", "error")
        return redirect(url_for("historial"))
    
    if request.method == "POST":
        texto_editado = request.form.get("texto_editado")
        
        folder = get_resultados_folder(tenant)
        ruta = os.path.join(folder, record.archivo)
        
        doc = Document()
        for parrafo in texto_editado.split("\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo)
        doc.save(ruta)
        
        record.texto_generado = texto_editado
        record.fecha = datetime.now()
        db.session.commit()
        
        flash("Documento actualizado exitosamente.", "success")
        return redirect(url_for("historial"))
    
    return render_template("editar.html", record=record)


@app.route("/descargar/<nombre_archivo>")
@login_required
def descargar(nombre_archivo):
    safe_filename = secure_filename(nombre_archivo)
    if not safe_filename or safe_filename != nombre_archivo:
        flash("Nombre de archivo no válido.", "error")
        return redirect(url_for("index"))
    
    if not safe_filename.endswith(".docx"):
        flash("Tipo de archivo no permitido.", "error")
        return redirect(url_for("index"))
    
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    
    if current_user.is_super_admin() and tenant_id:
        record = DocumentRecord.query.filter_by(archivo=safe_filename, tenant_id=tenant_id).first()
    elif current_user.is_super_admin() and not tenant_id:
        record = None
    elif current_user.is_admin and tenant_id:
        record = DocumentRecord.query.filter_by(archivo=safe_filename, tenant_id=tenant_id).first()
    else:
        record = DocumentRecord.query.filter_by(archivo=safe_filename, user_id=current_user.id, tenant_id=tenant_id).first()
    
    if not record:
        flash("Documento no encontrado o no tienes permiso para accederlo.", "error")
        return redirect(url_for("historial"))
    
    doc_tenant = Tenant.query.get(record.tenant_id) if record.tenant_id else None
    folder = get_resultados_folder(doc_tenant)
    ruta_completa = os.path.join(os.path.abspath(folder), safe_filename)
    
    if not os.path.exists(ruta_completa):
        old_path = os.path.join(os.path.abspath(CARPETA_RESULTADOS), safe_filename)
        if os.path.exists(old_path):
            ruta_completa = old_path
            folder = CARPETA_RESULTADOS
        else:
            flash("Archivo no encontrado.", "error")
            return redirect(url_for("index"))
    
    return send_from_directory(
        os.path.abspath(folder), 
        safe_filename, 
        as_attachment=True
    )


@app.route("/historial")
@login_required
def historial():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    
    search = request.args.get('search', '').strip()
    tipo_filter = request.args.get('tipo', '').strip()
    fecha_desde = request.args.get('fecha_desde', '').strip()
    fecha_hasta = request.args.get('fecha_hasta', '').strip()
    
    if current_user.is_super_admin() and tenant_id:
        query = DocumentRecord.query.filter_by(tenant_id=tenant_id)
    elif current_user.is_super_admin() and not tenant_id:
        query = DocumentRecord.query.filter(DocumentRecord.id < 0)
    elif current_user.is_admin and tenant_id:
        query = DocumentRecord.query.filter_by(tenant_id=tenant_id)
    else:
        query = DocumentRecord.query.filter_by(user_id=current_user.id, tenant_id=tenant_id)
    
    if search:
        query = query.filter(
            db.or_(
                DocumentRecord.demandante.ilike(f'%{search}%'),
                DocumentRecord.tipo_documento.ilike(f'%{search}%')
            )
        )
    
    if tipo_filter:
        query = query.filter(DocumentRecord.tipo_documento_key == tipo_filter)
    
    if fecha_desde:
        try:
            fecha_desde_dt = datetime.strptime(fecha_desde, '%Y-%m-%d')
            query = query.filter(DocumentRecord.fecha >= fecha_desde_dt)
        except ValueError:
            pass
    
    if fecha_hasta:
        try:
            fecha_hasta_dt = datetime.strptime(fecha_hasta, '%Y-%m-%d')
            fecha_hasta_dt = fecha_hasta_dt.replace(hour=23, minute=59, second=59)
            query = query.filter(DocumentRecord.fecha <= fecha_hasta_dt)
        except ValueError:
            pass
    
    documentos = query.order_by(DocumentRecord.fecha.desc()).all()
    
    return render_template("historial.html", 
                          documentos=documentos, 
                          modelos=MODELOS,
                          search=search,
                          tipo_filter=tipo_filter,
                          fecha_desde=fecha_desde,
                          fecha_hasta=fecha_hasta)


@app.route("/super_admin")
@super_admin_required
def super_admin():
    tenants = Tenant.query.order_by(Tenant.created_at.desc()).all()
    
    stats = []
    for t in tenants:
        doc_count = DocumentRecord.query.filter_by(tenant_id=t.id).count()
        user_count = User.query.filter_by(tenant_id=t.id).count()
        last_doc = DocumentRecord.query.filter_by(tenant_id=t.id).order_by(DocumentRecord.fecha.desc()).first()
        stats.append({
            'tenant': t,
            'docs': doc_count,
            'users': user_count,
            'last_activity': last_doc.fecha if last_doc else None
        })
    
    total_docs = DocumentRecord.query.count()
    total_users = User.query.count()
    total_tenants = Tenant.query.count()
    
    return render_template("super_admin.html",
                          stats=stats,
                          total_docs=total_docs,
                          total_users=total_users,
                          total_tenants=total_tenants)


@app.route("/super_admin/impersonate/<int:tenant_id>")
@super_admin_required
def impersonate_tenant(tenant_id):
    tenant = Tenant.query.get_or_404(tenant_id)
    session['impersonate_tenant_id'] = tenant_id
    flash(f"Ahora estás viendo como: {tenant.nombre}", "info")
    return redirect(url_for('index'))


@app.route("/super_admin/stop_impersonate")
@super_admin_required
def stop_impersonate():
    if 'impersonate_tenant_id' in session:
        del session['impersonate_tenant_id']
    flash("Volviste a tu vista de super administrador.", "info")
    return redirect(url_for('super_admin'))


@app.route("/super_admin/eliminar_estudio/<int:tenant_id>", methods=["POST"])
@super_admin_required
def eliminar_estudio(tenant_id):
    tenant = Tenant.query.get_or_404(tenant_id)
    nombre = tenant.nombre
    
    User.query.filter_by(tenant_id=tenant_id).delete()
    DocumentRecord.query.filter_by(tenant_id=tenant_id).delete()
    FinishedDocument.query.filter_by(tenant_id=tenant_id).delete()
    Plantilla.query.filter_by(tenant_id=tenant_id).delete()
    Estilo.query.filter_by(tenant_id=tenant_id).delete()
    CampoPlantilla.query.filter_by(tenant_id=tenant_id).delete()
    Modelo.query.filter_by(tenant_id=tenant_id).delete()
    Case.query.filter_by(tenant_id=tenant_id).delete()
    Task.query.filter_by(tenant_id=tenant_id).delete()
    ReviewSession.query.filter_by(tenant_id=tenant_id).delete()
    
    db.session.delete(tenant)
    db.session.commit()
    
    flash(f"Centro '{nombre}' eliminado correctamente junto con todos sus datos.", "success")
    log_audit(tenant_id, 'CENTRO_DELETED', f'Centro eliminado: {nombre}')
    return redirect(url_for('super_admin'))


@app.route("/super_admin/crear_centro", methods=["GET", "POST"])
@super_admin_required
def super_admin_crear_centro():
    """Crear un nuevo Centro de Conciliación."""
    if request.method == "POST":
        nombre = request.form.get('nombre', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        plan = request.form.get('plan', 'basico')
        
        dynamic_plans = get_dynamic_plan_config()
        if not nombre or not email or not password:
            flash("Nombre, email y contraseña son requeridos.", "error")
            return render_template("super_admin_crear_centro.html", plan_config=dynamic_plans)
        
        if User.query.filter_by(email=email).first():
            flash("Ya existe un usuario con ese email.", "error")
            return render_template("super_admin_crear_centro.html", plan_config=dynamic_plans)
        
        slug = nombre.lower().replace(' ', '_').replace('.', '')[:50]
        slug = re.sub(r'[^a-z0-9_]', '', slug)
        base_slug = slug
        counter = 1
        while Tenant.query.filter_by(slug=slug).first():
            slug = f"{base_slug}_{counter}"
            counter += 1
        
        tenant = Tenant(
            nombre=nombre,
            slug=slug,
            plan=plan,
            activo=True,
            subscription_status='active',
            onboarding_completed=False
        )
        db.session.add(tenant)
        db.session.flush()
        
        user = User(
            username=nombre.split()[0] if ' ' in nombre else nombre,
            email=email,
            tenant_id=tenant.id,
            role='admin_estudio',
            activo=True
        )
        user.set_password(password)
        db.session.add(user)
        
        create_default_tipos_acta(tenant.id)
        
        db.session.commit()
        
        log_audit(tenant.id, 'USER_CREATED', f'Centro creado con admin: {email}', {'plan': plan})
        
        dynamic_plans = get_dynamic_plan_config()
        plan_name = dynamic_plans.get(plan, {}).get('nombre', plan)
        flash(f"Centro '{nombre}' creado exitosamente con plan {plan_name}.", "success")
        return redirect(url_for('super_admin'))
    
    return render_template("super_admin_crear_centro.html", plan_config=get_dynamic_plan_config())


@app.route("/super_admin/editar_centro/<int:tenant_id>", methods=["GET", "POST"])
@super_admin_required
def super_admin_editar_centro(tenant_id):
    """Editar un Centro de Conciliación existente."""
    tenant = Tenant.query.get_or_404(tenant_id)
    
    if request.method == "POST":
        action = request.form.get('action', 'update')
        
        if action == 'update':
            tenant.nombre = request.form.get('nombre', tenant.nombre).strip()
            tenant.direccion = request.form.get('direccion', '').strip()
            tenant.telefono = request.form.get('telefono', '').strip()
            tenant.pagina_web = request.form.get('pagina_web', '').strip()
            tenant.resolucion_directoral = request.form.get('resolucion_directoral', '').strip()
            db.session.commit()
            log_audit(tenant_id, 'CENTRO_UPDATED', f'Centro actualizado: {tenant.nombre}')
            flash("Centro actualizado correctamente.", "success")
        
        elif action == 'change_plan':
            old_plan = tenant.plan
            new_plan = request.form.get('plan', 'basico')
            dynamic_plans = get_dynamic_plan_config()
            if new_plan in dynamic_plans:
                tenant.plan = new_plan
                db.session.commit()
                log_audit(tenant_id, 'PLAN_CHANGED', f'Plan cambiado de {old_plan} a {new_plan}', 
                         {'old_plan': old_plan, 'new_plan': new_plan})
                flash(f"Plan cambiado a {dynamic_plans[new_plan]['nombre']}.", "success")
        
        elif action == 'toggle_active':
            tenant.activo = not tenant.activo
            db.session.commit()
            estado = "activado" if tenant.activo else "desactivado"
            log_audit(tenant_id, 'CENTRO_UPDATED', f'Centro {estado}: {tenant.nombre}')
            flash(f"Centro {estado}.", "success")
        
        elif action == 'add_user':
            username = request.form.get('username', '').strip()
            email = request.form.get('user_email', '').strip()
            password = request.form.get('user_password', '')
            role = request.form.get('user_role', 'usuario_estudio')
            
            if not username or not email or not password:
                flash("Todos los campos son requeridos para crear usuario.", "error")
            elif User.query.filter_by(email=email).first():
                flash("Ya existe un usuario con ese email.", "error")
            else:
                if not tenant_can_add_user(tenant):
                    plan_cfg = get_plan_config(tenant)
                    flash(f"El Centro ha alcanzado el límite de usuarios de su plan ({plan_cfg['max_usuarios']}).", "error")
                else:
                    user = User(
                        username=username,
                        email=email,
                        tenant_id=tenant_id,
                        role=role,
                        activo=True
                    )
                    user.set_password(password)
                    db.session.add(user)
                    db.session.commit()
                    log_audit(tenant_id, 'USER_CREATED', f'Usuario creado: {email}', {'role': role})
                    flash(f"Usuario '{username}' creado exitosamente.", "success")
        
        return redirect(url_for('super_admin_editar_centro', tenant_id=tenant_id))
    
    usuarios = User.query.filter_by(tenant_id=tenant_id).all()
    plan_config = get_plan_config(tenant)
    usuarios_actuales = User.query.filter_by(tenant_id=tenant_id, activo=True).count()
    puede_agregar = usuarios_actuales < plan_config['max_usuarios']
    
    all_plans = get_dynamic_plan_config()
    return render_template("super_admin_editar_centro.html", 
                          tenant=tenant, 
                          usuarios=usuarios,
                          plan_config=plan_config,
                          all_plans=all_plans,
                          usuarios_actuales=usuarios_actuales,
                          puede_agregar=puede_agregar)


@app.route("/super_admin/planes", methods=["GET", "POST"])
@super_admin_required
def super_admin_planes():
    """Gestionar planes y sus limitaciones."""
    from models import PlanConfiguration
    import json
    
    PlanConfiguration.initialize_defaults()
    
    if request.method == "POST":
        action = request.form.get('action', '')
        
        if action == 'update_plan':
            plan_id = request.form.get('plan_id')
            plan = PlanConfiguration.query.get(plan_id)
            if plan:
                plan.nombre = request.form.get('nombre', plan.nombre).strip()
                plan.max_usuarios = int(request.form.get('max_usuarios', 2))
                plan.max_documentos_mes = int(request.form.get('max_documentos_mes', 50))
                plan.max_plantillas = int(request.form.get('max_plantillas', 5))
                plan.precio_mensual = float(request.form.get('precio_mensual', 0))
                plan.descripcion = request.form.get('descripcion', '').strip()
                
                features = request.form.getlist('features')
                plan.set_features_list(features)
                
                db.session.commit()
                flash(f"Plan '{plan.nombre}' actualizado correctamente.", "success")
        
        elif action == 'create_plan':
            plan_key = request.form.get('plan_key', '').strip().lower()
            plan_key = re.sub(r'[^a-z0-9_]', '', plan_key)
            
            if not plan_key:
                flash("La clave del plan es requerida.", "error")
            elif PlanConfiguration.query.filter_by(plan_key=plan_key).first():
                flash("Ya existe un plan con esa clave.", "error")
            else:
                max_orden = db.session.query(db.func.max(PlanConfiguration.orden)).scalar() or 0
                plan = PlanConfiguration(
                    plan_key=plan_key,
                    nombre=request.form.get('nombre', 'Nuevo Plan').strip(),
                    max_usuarios=int(request.form.get('max_usuarios', 2)),
                    max_documentos_mes=int(request.form.get('max_documentos_mes', 50)),
                    max_plantillas=int(request.form.get('max_plantillas', 5)),
                    precio_mensual=float(request.form.get('precio_mensual', 0)),
                    descripcion=request.form.get('descripcion', '').strip(),
                    features=json.dumps(request.form.getlist('features')),
                    orden=max_orden + 1,
                    activo=True
                )
                db.session.add(plan)
                db.session.commit()
                flash(f"Plan '{plan.nombre}' creado exitosamente.", "success")
        
        elif action == 'toggle_plan':
            plan_id = request.form.get('plan_id')
            plan = PlanConfiguration.query.get(plan_id)
            if plan and plan.plan_key not in ['basico', 'medio', 'avanzado']:
                plan.activo = not plan.activo
                db.session.commit()
                estado = "activado" if plan.activo else "desactivado"
                flash(f"Plan '{plan.nombre}' {estado}.", "success")
            else:
                flash("No se pueden desactivar los planes predeterminados.", "error")
        
        elif action == 'delete_plan':
            plan_id = request.form.get('plan_id')
            plan = PlanConfiguration.query.get(plan_id)
            if plan and plan.plan_key not in ['basico', 'medio', 'avanzado']:
                tenants_using = Tenant.query.filter_by(plan=plan.plan_key).count()
                if tenants_using > 0:
                    flash(f"No se puede eliminar el plan. {tenants_using} centro(s) lo están usando.", "error")
                else:
                    nombre = plan.nombre
                    db.session.delete(plan)
                    db.session.commit()
                    flash(f"Plan '{nombre}' eliminado.", "success")
            else:
                flash("No se pueden eliminar los planes predeterminados.", "error")
        
        return redirect(url_for('super_admin_planes'))
    
    planes = PlanConfiguration.query.order_by(PlanConfiguration.orden).all()
    features_disponibles = PlanConfiguration.FEATURES_DISPONIBLES
    
    tenants_por_plan = {}
    for plan in planes:
        tenants_por_plan[plan.plan_key] = Tenant.query.filter_by(plan=plan.plan_key).count()
    
    return render_template("super_admin_planes.html",
                          planes=planes,
                          features_disponibles=features_disponibles,
                          tenants_por_plan=tenants_por_plan)


@app.route("/system/pricing", methods=["GET", "POST"])
@super_admin_required
def system_pricing():
    """Panel de super admin para gestionar precios y addons."""
    if request.method == "POST":
        action = request.form.get("action")
        
        if action == "update_config":
            configs = [
                ('platform_name', request.form.get('platform_name', 'LegalDoc Pro')),
                ('currency', request.form.get('currency', 'USD')),
                ('currency_symbol', request.form.get('currency_symbol', '$')),
                ('price_per_seat', request.form.get('price_per_seat', '69.00')),
                ('min_seats', request.form.get('min_seats', '1')),
                ('max_seats', request.form.get('max_seats', '100')),
                ('trial_days', request.form.get('trial_days', '14')),
            ]
            for key, value in configs:
                PricingConfig.set_value(key, value, user_id=current_user.id)
            flash("Configuración de precios actualizada.", "success")
        
        elif action == "create_addon":
            nombre = request.form.get('addon_nombre', '').strip()
            descripcion = request.form.get('addon_descripcion', '').strip()
            precio = request.form.get('addon_precio', '0')
            tipo = request.form.get('addon_tipo', 'monthly')
            
            if nombre:
                addon = PricingAddon(
                    nombre=nombre,
                    descripcion=descripcion,
                    precio=float(precio),
                    tipo=tipo,
                    currency=PricingConfig.get_value('currency', 'USD')
                )
                db.session.add(addon)
                db.session.commit()
                flash(f"Complemento '{nombre}' creado.", "success")
            else:
                flash("El nombre del complemento es requerido.", "error")
        
        elif action == "toggle_addon":
            addon_id = request.form.get('addon_id')
            addon = PricingAddon.query.get(addon_id)
            if addon:
                addon.activo = not addon.activo
                db.session.commit()
                estado = "activado" if addon.activo else "desactivado"
                flash(f"Complemento '{addon.nombre}' {estado}.", "success")
        
        elif action == "delete_addon":
            addon_id = request.form.get('addon_id')
            addon = PricingAddon.query.get(addon_id)
            if addon:
                nombre = addon.nombre
                db.session.delete(addon)
                db.session.commit()
                flash(f"Complemento '{nombre}' eliminado.", "success")
        
        elif action == "update_addon":
            addon_id = request.form.get('addon_id')
            addon = PricingAddon.query.get(addon_id)
            if addon:
                addon.nombre = request.form.get('edit_nombre', addon.nombre).strip()
                addon.descripcion = request.form.get('edit_descripcion', '').strip()
                try:
                    addon.precio = float(request.form.get('edit_precio', addon.precio))
                except ValueError:
                    flash("El precio debe ser un número válido.", "error")
                    return redirect(url_for('system_pricing'))
                addon.tipo = request.form.get('edit_tipo', addon.tipo)
                addon.currency = request.form.get('edit_currency', addon.currency)
                db.session.commit()
                flash(f"Complemento '{addon.nombre}' actualizado.", "success")
        
        return redirect(url_for('system_pricing'))
    
    config = PricingConfig.get_pricing()
    config_list = PricingConfig.query.order_by(PricingConfig.key).all()
    addons = PricingAddon.query.order_by(PricingAddon.orden, PricingAddon.id).all()
    
    return render_template("system/pricing.html",
                          config=config,
                          config_list=config_list,
                          addons=addons)


@app.route("/admin")
@coordinador_or_admin_required
def admin():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    plantillas = Plantilla.query.filter_by(tenant_id=tenant.id).all()
    estilos = Estilo.query.filter_by(tenant_id=tenant.id).all()
    usuarios = User.query.filter_by(tenant_id=tenant.id).all()
    total_docs = DocumentRecord.query.filter_by(tenant_id=tenant.id).count()
    
    return render_template("admin.html", 
                          plantillas=plantillas, 
                          estilos=estilos,
                          usuarios=usuarios,
                          total_docs=total_docs,
                          modelos=MODELOS,
                          tenant=tenant)


@app.route("/admin/modelos")
@coordinador_or_admin_required
def admin_modelos():
    """Admin view to see all models from all users in the tenant."""
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    todos_modelos = Modelo.query.filter_by(tenant_id=tenant.id).all()
    usuarios = {u.id: u for u in User.query.filter_by(tenant_id=tenant.id).all()}
    
    return render_template("admin_modelos.html", 
                          modelos=todos_modelos, 
                          usuarios=usuarios,
                          modelos_sistema=MODELOS,
                          tenant=tenant)


@app.route("/admin/estilos")
@coordinador_or_admin_required
def admin_estilos():
    """Admin view to see all styles from all users in the tenant."""
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    todos_estilos = Estilo.query.filter_by(tenant_id=tenant.id).all()
    usuarios = {u.id: u for u in User.query.filter_by(tenant_id=tenant.id).all()}
    
    return render_template("admin_estilos.html", 
                          estilos=todos_estilos, 
                          usuarios=usuarios,
                          tenant=tenant)


@app.route("/configurar_estudio", methods=["GET", "POST"])
@coordinador_or_admin_required
def configurar_estudio():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    if request.method == "POST":
        tenant.nombre = request.form.get("nombre", "").strip() or tenant.nombre
        tenant.resolucion_directoral = request.form.get("resolucion_directoral", "").strip()
        tenant.direccion = request.form.get("direccion", "").strip()
        tenant.telefono = request.form.get("telefono", "").strip()
        tenant.pagina_web = request.form.get("pagina_web", "").strip()
        tenant.pais = request.form.get("pais", "").strip()
        tenant.ciudad = request.form.get("ciudad", "").strip()
        tenant.areas_practica = request.form.get("areas_practica", "").strip()
        
        if 'logo' in request.files:
            file = request.files['logo']
            if file and file.filename:
                filename = secure_filename(file.filename)
                tenant_folder = os.path.join("static", "tenants", tenant.slug)
                os.makedirs(tenant_folder, exist_ok=True)
                filepath = os.path.join(tenant_folder, filename)
                file.save(filepath)
                tenant.logo_path = filename
        
        db.session.commit()
        flash("Configuración del estudio actualizada.", "success")
        return redirect(url_for("admin"))
    
    return render_template("configurar_estudio.html", tenant=tenant)


@app.route("/configurar_apariencia", methods=["GET", "POST"])
@admin_estudio_required
def configurar_apariencia():
    """Configura colores y branding del estudio."""
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    if request.method == "POST":
        color_primario = request.form.get("color_primario", "").strip()
        color_secundario = request.form.get("color_secundario", "").strip()
        
        if color_primario and len(color_primario) == 7 and color_primario.startswith('#'):
            tenant.color_primario = color_primario.upper()
        if color_secundario and len(color_secundario) == 7 and color_secundario.startswith('#'):
            tenant.color_secundario = color_secundario.upper()
        
        if 'logo' in request.files:
            file = request.files['logo']
            if file and file.filename:
                filename = secure_filename(file.filename)
                tenant_folder = os.path.join("static", "tenants", tenant.slug)
                os.makedirs(tenant_folder, exist_ok=True)
                filepath = os.path.join(tenant_folder, filename)
                file.save(filepath)
                tenant.logo_path = filename
        
        db.session.commit()
        flash("Apariencia del estudio actualizada.", "success")
        return redirect(url_for("configurar_apariencia"))
    
    return render_template("configurar_apariencia.html", tenant=tenant)


@app.route("/estilos_documentos", methods=["GET", "POST"])
@admin_estudio_required
def estilos_documentos():
    """Configura el estilo de los documentos generados."""
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    estilo = EstiloDocumento.get_or_create(tenant.id)
    
    if request.method == "POST":
        estilo.fuente = request.form.get("fuente", "Times New Roman")
        estilo.tamano_base = int(request.form.get("tamano_base", 12))
        estilo.interlineado = float(request.form.get("interlineado", 1.5))
        estilo.margen_superior = float(request.form.get("margen_superior", 2.5))
        estilo.margen_inferior = float(request.form.get("margen_inferior", 2.5))
        estilo.margen_izquierdo = float(request.form.get("margen_izquierdo", 3.0))
        estilo.margen_derecho = float(request.form.get("margen_derecho", 2.5))
        
        db.session.commit()
        flash("Estilo de documentos actualizado.", "success")
        return redirect(url_for("estilos_documentos"))
    
    return render_template("estilos_documentos.html", 
                         estilo=estilo, 
                         fuentes_permitidas=EstiloDocumento.FUENTES_PERMITIDAS)


@app.route("/preferencias_usuario", methods=["GET", "POST"])
@login_required
def preferencias_usuario():
    """Ruta deshabilitada."""
    abort(404)


@app.route("/admin/tipos-caso", methods=["GET", "POST"])
@coordinador_or_admin_required
def admin_tipos_caso():
    """Ruta deshabilitada."""
    abort(404)


def _admin_tipos_caso_disabled():
    """Gestionar tipos de caso y campos personalizados - DESHABILITADO."""
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    if request.method == "POST":
        action = request.form.get("action")
        
        if action == "create_type":
            nombre = request.form.get("nombre", "").strip()
            descripcion = request.form.get("descripcion", "").strip()
            icono = request.form.get("icono", "fa-folder")
            color = request.form.get("color", "blue")
            
            if not nombre:
                flash("El nombre del tipo de caso es obligatorio.", "error")
            elif CaseType.query.filter_by(tenant_id=tenant.id, nombre=nombre).first():
                flash("Ya existe un tipo de caso con ese nombre.", "error")
            else:
                case_type = CaseType(
                    tenant_id=tenant.id,
                    nombre=nombre,
                    descripcion=descripcion,
                    icono=icono,
                    color=color
                )
                db.session.add(case_type)
                db.session.commit()
                flash(f"Tipo de caso '{nombre}' creado exitosamente.", "success")
        
        elif action == "update_type":
            type_id = request.form.get("type_id")
            case_type = CaseType.query.get(type_id)
            if case_type and case_type.tenant_id == tenant.id:
                case_type.nombre = request.form.get("nombre", case_type.nombre).strip()
                case_type.descripcion = request.form.get("descripcion", "").strip()
                case_type.icono = request.form.get("icono", case_type.icono)
                case_type.color = request.form.get("color", case_type.color)
                db.session.commit()
                flash(f"Tipo de caso actualizado.", "success")
        
        elif action == "delete_type":
            type_id = request.form.get("type_id")
            case_type = CaseType.query.get(type_id)
            if case_type and case_type.tenant_id == tenant.id:
                if case_type.cases.count() > 0:
                    flash("No se puede eliminar un tipo de caso que tiene casos asociados.", "error")
                else:
                    nombre = case_type.nombre
                    db.session.delete(case_type)
                    db.session.commit()
                    flash(f"Tipo de caso '{nombre}' eliminado.", "success")
        
        elif action == "add_field":
            type_id = request.form.get("type_id")
            case_type = CaseType.query.get(type_id) if type_id else None
            
            nombre = request.form.get("field_nombre", "").strip().lower().replace(" ", "_")
            label = request.form.get("field_label", "").strip()
            tipo = request.form.get("field_tipo", "text")
            placeholder = request.form.get("field_placeholder", "").strip()
            opciones = request.form.get("field_opciones", "").strip()
            requerido = request.form.get("field_requerido") == "on"
            
            if not nombre or not label:
                flash("El nombre y la etiqueta del campo son obligatorios.", "error")
            else:
                field = CaseCustomField(
                    tenant_id=tenant.id,
                    case_type_id=type_id if type_id else None,
                    nombre=nombre,
                    label=label,
                    tipo=tipo,
                    placeholder=placeholder,
                    opciones=opciones,
                    requerido=requerido
                )
                db.session.add(field)
                db.session.commit()
                flash(f"Campo '{label}' agregado.", "success")
        
        elif action == "delete_field":
            field_id = request.form.get("field_id")
            field = CaseCustomField.query.get(field_id)
            if field and field.tenant_id == tenant.id:
                label = field.label
                db.session.delete(field)
                db.session.commit()
                flash(f"Campo '{label}' eliminado.", "success")
        
        elif action == "toggle_type":
            type_id = request.form.get("type_id")
            case_type = CaseType.query.get(type_id)
            if case_type and case_type.tenant_id == tenant.id:
                case_type.activo = not case_type.activo
                db.session.commit()
                estado = "activado" if case_type.activo else "desactivado"
                flash(f"Tipo de caso {estado}.", "success")
        
        return redirect(url_for("admin_tipos_caso"))
    
    tipos = CaseType.query.filter_by(tenant_id=tenant.id).order_by(CaseType.orden, CaseType.nombre).all()
    campos_generales = CaseCustomField.query.filter_by(tenant_id=tenant.id, case_type_id=None).order_by(CaseCustomField.orden).all()
    
    return render_template("admin_tipos_caso.html",
                          tipos=tipos,
                          campos_generales=campos_generales,
                          tipos_campo=CaseCustomField.TIPOS)


@app.route("/api/admin/tipos-caso/<int:type_id>/campos")
@coordinador_or_admin_required
def get_tipo_campos(type_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/admin/usuarios", methods=["GET", "POST"])
@coordinador_or_admin_required
def admin_usuarios():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    is_admin = current_user.role in ['super_admin', 'admin_estudio']
    
    plan_config = get_plan_config(tenant)
    usuarios_actuales = User.query.filter_by(tenant_id=tenant.id, activo=True).count()
    puede_agregar = usuarios_actuales < plan_config['max_usuarios']
    
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        role = request.form.get("role", "usuario_estudio")
        
        if not username or not email or not password:
            flash("Todos los campos son obligatorios.", "error")
        elif not puede_agregar:
            flash(f"Has alcanzado el límite de {plan_config['max_usuarios']} usuarios para tu plan {plan_config['nombre']}. Contacta al administrador para actualizar tu plan.", "error")
        elif User.query.filter_by(email=email).first():
            flash("Ya existe un usuario con ese email.", "error")
        else:
            allowed_roles = ['admin_estudio', 'coordinador', 'usuario_estudio'] if is_admin else ['coordinador', 'usuario_estudio']
            if role not in allowed_roles:
                role = 'usuario_estudio'
            
            user = User(
                username=username,
                email=email,
                tenant_id=tenant.id,
                role=role,
                activo=True
            )
            user.set_password(password)
            db.session.add(user)
            db.session.commit()
            flash(f"Usuario {username} creado exitosamente.", "success")
            usuarios_actuales += 1
            puede_agregar = usuarios_actuales < plan_config['max_usuarios']
    
    usuarios = User.query.filter_by(tenant_id=tenant.id).all()
    return render_template("admin_usuarios.html", 
                          usuarios=usuarios, 
                          tenant=tenant, 
                          is_admin=is_admin,
                          plan_config=plan_config,
                          usuarios_actuales=usuarios_actuales,
                          puede_agregar=puede_agregar)


@app.route("/admin/usuario/toggle/<int:user_id>", methods=["POST"])
@coordinador_or_admin_required
def toggle_usuario(user_id):
    tenant = get_current_tenant()
    user = User.query.get_or_404(user_id)
    
    if user.tenant_id != tenant.id:
        flash("No tienes permiso para modificar este usuario.", "error")
        return redirect(url_for("admin_usuarios"))
    
    if user.id == current_user.id:
        flash("No puedes desactivar tu propia cuenta.", "error")
        return redirect(url_for("admin_usuarios"))
    
    is_admin = current_user.role in ['super_admin', 'admin_estudio']
    if not is_admin and user.role == 'admin_estudio':
        flash("No puedes modificar usuarios administradores.", "error")
        return redirect(url_for("admin_usuarios"))
    
    user.activo = not user.activo
    db.session.commit()
    status = "activado" if user.activo else "desactivado"
    flash(f"Usuario {user.username} {status}.", "success")
    return redirect(url_for("admin_usuarios"))


@app.route("/admin/convertir", methods=["GET", "POST"])
@coordinador_or_admin_required
def convertir_documento():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    documento_convertido = None
    campos_detectados = 0
    
    if request.method == "POST":
        archivo = request.files.get('archivo')
        
        if not archivo or not archivo.filename or not archivo.filename.endswith('.docx'):
            flash("Debes subir un archivo Word (.docx).", "error")
            return render_template("convertir_documento.html")
        
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document(archivo)
            campo_num = 0
            
            dot_pattern = re.compile(r'[\.…]{4,}|_{4,}')
            
            for para in doc.paragraphs:
                text = para.text
                if dot_pattern.search(text):
                    new_text = text
                    for match in dot_pattern.finditer(text):
                        campo_num += 1
                        new_text = new_text.replace(match.group(), f'{{{{campo_{campo_num}}}}}', 1)
                    
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_text
                    else:
                        para.text = new_text
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text
                            if dot_pattern.search(text):
                                new_text = text
                                for match in dot_pattern.finditer(text):
                                    campo_num += 1
                                    new_text = new_text.replace(match.group(), f'{{{{campo_{campo_num}}}}}', 1)
                                
                                for run in para.runs:
                                    run.text = ""
                                if para.runs:
                                    para.runs[0].text = new_text
                                else:
                                    para.text = new_text
            
            if campo_num == 0:
                flash("No se encontraron espacios con puntos o guiones para convertir.", "error")
                return render_template("convertir_documento.html")
            
            convertidos_folder = os.path.join("documentos_convertidos", f"tenant_{tenant.id}")
            os.makedirs(convertidos_folder, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            original_name = secure_filename(archivo.filename)
            output_name = f"convertido_{timestamp}_{original_name}"
            output_path = os.path.join(convertidos_folder, output_name)
            
            doc.save(output_path)
            
            campos_detectados = campo_num
            documento_convertido = output_name
            
            flash(f"Documento convertido exitosamente. Se reemplazaron {campo_num} campos.", "success")
            
        except Exception as e:
            logging.error(f"Error al convertir documento: {e}")
            flash("Error al procesar el documento. Verifica que sea un archivo Word válido.", "error")
    
    return render_template("convertir_documento.html", 
                         documento_convertido=documento_convertido,
                         campos_detectados=campos_detectados)


@app.route("/admin/convertir/descargar/<nombre_archivo>")
@coordinador_or_admin_required
def descargar_convertido(nombre_archivo):
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    convertidos_folder = os.path.join("documentos_convertidos", f"tenant_{tenant.id}")
    file_path = os.path.join(convertidos_folder, secure_filename(nombre_archivo))
    
    if not os.path.exists(file_path):
        flash("El archivo no existe.", "error")
        return redirect(url_for("convertir_documento"))
    
    return send_file(file_path, as_attachment=True, download_name=nombre_archivo)


@app.route("/admin/plantilla", methods=["GET", "POST"])
@coordinador_or_admin_required
def admin_plantilla():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    plantilla_id = request.args.get('id', type=int)
    plantilla = Plantilla.query.filter_by(id=plantilla_id, tenant_id=tenant.id).first() if plantilla_id else None
    campos_detectados = []
    
    if plantilla_id and not plantilla:
        flash("No tienes permiso para editar esta plantilla.", "error")
        return redirect(url_for("admin"))
    
    if request.method == "POST":
        key = request.form.get("key", "").strip()
        nombre = request.form.get("nombre", "").strip()
        
        archivo = request.files.get('archivo_word')
        contenido = ""
        archivo_path = None
        
        if archivo and archivo.filename and archivo.filename.endswith('.docx'):
            tenant_folder = os.path.join(CARPETA_PLANTILLAS_SUBIDAS, f"tenant_{tenant.id}")
            os.makedirs(tenant_folder, exist_ok=True)
            
            safe_name = secure_filename(archivo.filename)
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            archivo_name = f"{timestamp}_{safe_name}"
            archivo_path = os.path.join(tenant_folder, archivo_name)
            archivo.save(archivo_path)
            
            contenido = extract_text_from_docx(archivo_path)
            campos_detectados = detect_placeholders_from_text(contenido)
        elif plantilla:
            contenido = plantilla.contenido
        
        if not key or not nombre:
            flash("La clave y el nombre son obligatorios.", "error")
            return render_template("admin_plantilla.html", plantilla=plantilla, campos_detectados=campos_detectados)
        
        if not contenido and not plantilla:
            flash("Debes subir un archivo Word con la plantilla.", "error")
            return render_template("admin_plantilla.html", plantilla=plantilla, campos_detectados=campos_detectados)
        
        if archivo_path and contenido and len(contenido.strip()) < 50:
            flash("El documento Word parece estar vacío o tiene muy poco contenido.", "error")
            if archivo_path and os.path.exists(archivo_path):
                os.remove(archivo_path)
            return render_template("admin_plantilla.html", plantilla=plantilla, campos_detectados=campos_detectados)
        
        if plantilla:
            plantilla.key = key
            plantilla.nombre = nombre
            if contenido:
                plantilla.contenido = contenido
            if archivo_path:
                plantilla.archivo_original = archivo_path
                nuevos_campos = 0
                if campos_detectados:
                    for i, campo_name in enumerate(campos_detectados):
                        campo_key = campo_to_key(campo_name)
                        existing_campo = CampoPlantilla.query.filter_by(
                            plantilla_key=key, 
                            nombre_campo=campo_key, 
                            tenant_id=tenant.id
                        ).first()
                        if not existing_campo:
                            max_orden = db.session.query(db.func.max(CampoPlantilla.orden)).filter_by(
                                plantilla_key=key, tenant_id=tenant.id
                            ).scalar() or 0
                            campo = CampoPlantilla(
                                plantilla_key=key,
                                nombre_campo=campo_key[:100],
                                etiqueta=campo_name[:200] if len(campo_name) <= 200 else campo_name[:197] + "...",
                                tipo='text',
                                requerido=True,
                                orden=max_orden + i + 1,
                                tenant_id=tenant.id
                            )
                            db.session.add(campo)
                            nuevos_campos += 1
                if nuevos_campos > 0:
                    flash(f"Plantilla actualizada. Se detectaron {nuevos_campos} campos nuevos.", "success")
                else:
                    flash("Plantilla actualizada exitosamente.", "success")
            else:
                flash("Plantilla actualizada exitosamente.", "success")
        else:
            existing = Plantilla.query.filter_by(key=key, tenant_id=tenant.id).first()
            if existing:
                flash("Ya existe una plantilla con esta clave.", "error")
                return render_template("admin_plantilla.html", plantilla=plantilla, campos_detectados=campos_detectados)
            
            plantilla = Plantilla(
                key=key, 
                nombre=nombre, 
                contenido=contenido, 
                archivo_original=archivo_path,
                carpeta_estilos=key,
                tenant_id=tenant.id
            )
            db.session.add(plantilla)
            db.session.flush()
            
            if campos_detectados:
                for i, campo_name in enumerate(campos_detectados):
                    campo_key = campo_to_key(campo_name)
                    existing_campo = CampoPlantilla.query.filter_by(
                        plantilla_key=key, 
                        nombre_campo=campo_key, 
                        tenant_id=tenant.id
                    ).first()
                    if not existing_campo:
                        campo = CampoPlantilla(
                            plantilla_key=key,
                            nombre_campo=campo_key[:100],
                            etiqueta=campo_name[:200] if len(campo_name) <= 200 else campo_name[:197] + "...",
                            tipo='text',
                            requerido=True,
                            orden=i,
                            tenant_id=tenant.id
                        )
                        db.session.add(campo)
            
            flash(f"Plantilla creada exitosamente. Se detectaron {len(campos_detectados)} campos.", "success")
        
        db.session.commit()
        
        if campos_detectados and not plantilla_id:
            return redirect(url_for("admin_campos", plantilla_key=key))
        return redirect(url_for("admin"))
    
    return render_template("admin_plantilla.html", plantilla=plantilla, campos_detectados=campos_detectados)


@app.route("/admin/plantilla/eliminar/<int:plantilla_id>", methods=["POST"])
@coordinador_or_admin_required
def eliminar_plantilla(plantilla_id):
    tenant = get_current_tenant()
    plantilla = Plantilla.query.filter_by(id=plantilla_id, tenant_id=tenant.id).first()
    
    if not plantilla:
        flash("No tienes permiso para eliminar esta plantilla.", "error")
        return redirect(url_for("admin"))
    
    db.session.delete(plantilla)
    db.session.commit()
    flash("Plantilla eliminada exitosamente.", "success")
    return redirect(url_for("admin"))


@app.route("/admin/estilo", methods=["GET", "POST"])
@coordinador_or_admin_required
def admin_estilo():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    estilo_id = request.args.get('id', type=int)
    estilo = Estilo.query.filter_by(id=estilo_id, tenant_id=tenant.id).first() if estilo_id else None
    
    if estilo_id and not estilo:
        flash("No tienes permiso para editar este estilo.", "error")
        return redirect(url_for("admin"))
    
    plantillas_db = Plantilla.query.filter_by(tenant_id=tenant.id).all()
    plantillas_keys = list(MODELOS.keys()) + [p.key for p in plantillas_db]
    plantillas_keys = list(set(plantillas_keys))
    
    if request.method == "POST":
        plantilla_key = request.form.get("plantilla_key", "").strip()
        nombre = request.form.get("nombre", "").strip()
        
        archivo = request.files.get('archivo_word')
        contenido = ""
        archivo_path = None
        
        if archivo and archivo.filename and archivo.filename.endswith('.docx'):
            tenant_folder = os.path.join(CARPETA_ESTILOS_SUBIDOS, f"tenant_{tenant.id}")
            os.makedirs(tenant_folder, exist_ok=True)
            
            safe_name = secure_filename(archivo.filename)
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            archivo_name = f"{timestamp}_{safe_name}"
            archivo_path = os.path.join(tenant_folder, archivo_name)
            archivo.save(archivo_path)
            
            contenido = extract_text_from_docx(archivo_path)
        elif estilo:
            contenido = estilo.contenido
        
        if not plantilla_key or not nombre:
            flash("La plantilla asociada y el nombre son obligatorios.", "error")
            return render_template("admin_estilo.html", estilo=estilo, plantillas_keys=plantillas_keys)
        
        if not contenido and not estilo:
            flash("Debes subir un archivo Word con el ejemplo de estilo.", "error")
            return render_template("admin_estilo.html", estilo=estilo, plantillas_keys=plantillas_keys)
        
        if archivo_path and contenido and len(contenido.strip()) < 50:
            flash("El documento Word parece estar vacío o tiene muy poco contenido.", "error")
            if archivo_path and os.path.exists(archivo_path):
                os.remove(archivo_path)
            return render_template("admin_estilo.html", estilo=estilo, plantillas_keys=plantillas_keys)
        
        if estilo:
            estilo.plantilla_key = plantilla_key
            estilo.nombre = nombre
            if contenido:
                estilo.contenido = contenido
            if archivo_path:
                estilo.archivo_original = archivo_path
            flash("Estilo actualizado exitosamente.", "success")
        else:
            estilo = Estilo(
                plantilla_key=plantilla_key, 
                nombre=nombre, 
                contenido=contenido,
                archivo_original=archivo_path,
                tenant_id=tenant.id
            )
            db.session.add(estilo)
            flash("Estilo creado exitosamente.", "success")
        
        db.session.commit()
        return redirect(url_for("admin"))
    
    return render_template("admin_estilo.html", estilo=estilo, plantillas_keys=plantillas_keys)


@app.route("/admin/estilo/eliminar/<int:estilo_id>", methods=["POST"])
@coordinador_or_admin_required
def eliminar_estilo(estilo_id):
    tenant = get_current_tenant()
    estilo = Estilo.query.filter_by(id=estilo_id, tenant_id=tenant.id).first()
    
    if not estilo:
        flash("No tienes permiso para eliminar este estilo.", "error")
        return redirect(url_for("admin"))
    
    db.session.delete(estilo)
    db.session.commit()
    flash("Estilo eliminado exitosamente.", "success")
    return redirect(url_for("admin"))


@app.route("/admin/campos/<plantilla_key>", methods=["GET", "POST"])
@coordinador_or_admin_required
def admin_campos(plantilla_key):
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("index"))
    
    plantilla = Plantilla.query.filter_by(key=plantilla_key, tenant_id=tenant.id).first()
    nombre_plantilla = plantilla.nombre if plantilla else MODELOS.get(plantilla_key, {}).get("nombre", plantilla_key)
    
    if request.method == "POST":
        campo_id = request.form.get("campo_id", type=int)
        nombre_campo = request.form.get("nombre_campo", "").strip()
        etiqueta = request.form.get("etiqueta", "").strip()
        tipo = request.form.get("tipo", "text").strip()
        requerido = request.form.get("requerido") == "on"
        orden = request.form.get("orden", 0, type=int)
        placeholder = request.form.get("placeholder", "").strip()
        opciones = request.form.get("opciones", "").strip()
        
        if not nombre_campo or not etiqueta:
            flash("Nombre del campo y etiqueta son obligatorios.", "error")
        else:
            if campo_id:
                campo = CampoPlantilla.query.get(campo_id)
                if campo and campo.tenant_id == tenant.id:
                    campo.nombre_campo = nombre_campo
                    campo.etiqueta = etiqueta
                    campo.tipo = tipo
                    campo.requerido = requerido
                    campo.orden = orden
                    campo.placeholder = placeholder
                    campo.opciones = opciones
                    flash("Campo actualizado.", "success")
            else:
                campo = CampoPlantilla(
                    plantilla_key=plantilla_key,
                    nombre_campo=nombre_campo,
                    etiqueta=etiqueta,
                    tipo=tipo,
                    requerido=requerido,
                    orden=orden,
                    placeholder=placeholder,
                    opciones=opciones,
                    tenant_id=tenant.id
                )
                db.session.add(campo)
                flash("Campo agregado.", "success")
            db.session.commit()
    
    campos = CampoPlantilla.query.filter_by(plantilla_key=plantilla_key, tenant_id=tenant.id).order_by(CampoPlantilla.orden).all()
    return render_template("admin_campos.html", 
                          plantilla_key=plantilla_key, 
                          nombre_plantilla=nombre_plantilla,
                          campos=campos)


@app.route("/admin/campo/eliminar/<int:campo_id>", methods=["POST"])
@coordinador_or_admin_required
def eliminar_campo(campo_id):
    tenant = get_current_tenant()
    campo = CampoPlantilla.query.get_or_404(campo_id)
    
    if campo.tenant_id != tenant.id:
        flash("No tienes permiso para eliminar este campo.", "error")
        return redirect(url_for("admin"))
    
    plantilla_key = campo.plantilla_key
    db.session.delete(campo)
    db.session.commit()
    flash("Campo eliminado.", "success")
    return redirect(url_for("admin_campos", plantilla_key=plantilla_key))


@app.route("/api/campos/<plantilla_key>")
def get_campos_plantilla(plantilla_key):
    tenant_id = None
    if current_user.is_authenticated:
        tenant = get_current_tenant()
        tenant_id = tenant.id if tenant else None
    
    if tenant_id:
        campos = CampoPlantilla.query.filter_by(plantilla_key=plantilla_key, tenant_id=tenant_id).order_by(CampoPlantilla.orden).all()
    else:
        campos = []
    
    campos_data = [{
        'id': c.id,
        'nombre_campo': c.nombre_campo,
        'etiqueta': c.etiqueta,
        'tipo': c.tipo,
        'requerido': c.requerido,
        'placeholder': c.placeholder or '',
        'opciones': c.opciones.split(',') if c.opciones else []
    } for c in campos]
    
    tablas_data = []
    if tenant_id:
        modelo = Modelo.query.filter_by(key=plantilla_key, tenant_id=tenant_id).first()
        if modelo:
            tablas = ModeloTabla.query.filter_by(modelo_id=modelo.id, tenant_id=tenant_id).order_by(ModeloTabla.orden).all()
            for tabla in tablas:
                tablas_data.append({
                    'id': tabla.id,
                    'nombre': tabla.nombre,
                    'columnas': tabla.columnas,
                    'num_filas': tabla.num_filas,
                    'mostrar_total': tabla.mostrar_total,
                    'columna_total': tabla.columna_total
                })
    
    return jsonify({
        'campos': campos_data,
        'tablas': tablas_data
    })


@app.route("/api/formulario/buscar/<code>")
@login_required
def api_buscar_formulario(code):
    """Busca un formulario por código y retorna sus datos."""
    from models import FormResponse, Modelo
    
    tenant = get_current_tenant()
    if not tenant:
        return jsonify({'error': 'No tienes acceso'}), 403
    
    code = code.strip().upper()
    form_response = FormResponse.query.filter_by(code=code, tenant_id=tenant.id).first()
    
    if not form_response:
        return jsonify({'error': 'Código no encontrado'}), 404
    
    modelo = form_response.template
    if not modelo:
        return jsonify({'error': 'Modelo no encontrado'}), 404
    
    return jsonify({
        'success': True,
        'form_id': form_response.id,
        'code': form_response.code,
        'status': form_response.status,
        'status_label': FormResponse.STATUSES.get(form_response.status, form_response.status),
        'can_use': form_response.can_be_used(),
        'template_key': modelo.key,
        'template_name': modelo.nombre,
        'answers': form_response.answers_json or {},
        'created_at': form_response.created_at.strftime('%d/%m/%Y %H:%M') if form_response.created_at else None
    })


# ==================== GESTIÓN DE CASOS ====================

def case_access_required(f):
    """Decorator to ensure user can access case management."""
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        tenant = get_current_tenant()
        if not tenant:
            flash("No tienes un estudio asociado.", "error")
            return redirect(url_for("dashboard"))
        return f(*args, **kwargs)
    return decorated_function


def case_manage_required(f):
    """Decorator to ensure user can manage cases (admin/coordinador)."""
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not current_user.can_manage_cases():
            flash("No tienes permisos para gestionar casos.", "error")
            return redirect(url_for("casos"))
        tenant = get_current_tenant()
        if not tenant:
            flash("No tienes un estudio asociado.", "error")
            return redirect(url_for("dashboard"))
        return f(*args, **kwargs)
    return decorated_function


@app.route("/casos")
@case_access_required
def casos():
    """Ruta deshabilitada."""
    abort(404)


def _casos_disabled():
    tenant = get_current_tenant()
    
    estado_filter = request.args.get('estado', '')
    prioridad_filter = request.args.get('prioridad', '')
    busqueda = request.args.get('busqueda', '').strip()
    
    query = Case.query.filter_by(tenant_id=tenant.id)
    
    if not current_user.can_manage_cases():
        assigned_case_ids = db.session.query(CaseAssignment.case_id).filter_by(user_id=current_user.id).subquery()
        query = query.filter(
            db.or_(
                Case.created_by_id == current_user.id,
                Case.id.in_(assigned_case_ids)
            )
        )
    
    if estado_filter:
        query = query.filter_by(estado=estado_filter)
    if prioridad_filter:
        query = query.filter_by(prioridad=prioridad_filter)
    if busqueda:
        query = query.filter(
            db.or_(
                Case.titulo.ilike(f'%{busqueda}%'),
                Case.cliente_nombre.ilike(f'%{busqueda}%'),
                Case.numero_expediente.ilike(f'%{busqueda}%')
            )
        )
    
    casos_list = query.order_by(Case.updated_at.desc()).all()
    
    stats = {
        'total': Case.query.filter_by(tenant_id=tenant.id).count(),
        'por_comenzar': Case.query.filter_by(tenant_id=tenant.id, estado='por_comenzar').count(),
        'en_proceso': Case.query.filter_by(tenant_id=tenant.id, estado='en_proceso').count(),
        'en_espera': Case.query.filter_by(tenant_id=tenant.id, estado='en_espera').count(),
        'terminado': Case.query.filter_by(tenant_id=tenant.id, estado='terminado').count(),
    }
    
    return render_template("casos.html", 
                          casos=casos_list,
                          stats=stats,
                          estado_filter=estado_filter,
                          prioridad_filter=prioridad_filter,
                          busqueda=busqueda,
                          estados=Case.ESTADOS,
                          prioridades=Case.PRIORIDADES,
                          now=datetime.now())


@app.route("/casos/nuevo", methods=["GET", "POST"])
@case_manage_required
def caso_nuevo():
    """Ruta deshabilitada."""
    abort(404)


def _caso_nuevo_disabled():
    tenant = get_current_tenant()
    
    if request.method == "POST":
        titulo = request.form.get("titulo", "").strip()
        cliente_nombre = request.form.get("cliente_nombre", "").strip()
        
        if not titulo or not cliente_nombre:
            flash("El título y nombre del cliente son obligatorios.", "error")
            usuarios = User.query.filter_by(tenant_id=tenant.id, activo=True).all()
            tipos_caso = CaseType.query.filter_by(tenant_id=tenant.id, activo=True).order_by(CaseType.orden, CaseType.nombre).all()
            campos_generales = CaseCustomField.query.filter_by(tenant_id=tenant.id, case_type_id=None, activo=True).order_by(CaseCustomField.orden).all()
            return render_template("caso_form.html", caso=None, usuarios=usuarios, estados=Case.ESTADOS, prioridades=Case.PRIORIDADES,
                                  tipos_caso=tipos_caso, campos_generales=campos_generales, tipos_campo=CaseCustomField.TIPOS)
        
        case_type_id_raw = request.form.get("case_type_id", "").strip()
        tipo_caso_texto = request.form.get("tipo_caso", "").strip()
        case_type_id = None
        
        if case_type_id_raw and case_type_id_raw != "otro":
            try:
                case_type_id_int = int(case_type_id_raw)
                case_type = CaseType.query.filter_by(id=case_type_id_int, tenant_id=tenant.id).first()
                if case_type:
                    case_type_id = case_type.id
                    tipo_caso_texto = case_type.nombre
                else:
                    case_type_id = None
            except (ValueError, TypeError):
                case_type_id = None
        
        caso = Case(
            tenant_id=tenant.id,
            titulo=titulo,
            descripcion=request.form.get("descripcion", "").strip(),
            numero_expediente=request.form.get("numero_expediente", "").strip(),
            cliente_nombre=cliente_nombre,
            cliente_email=request.form.get("cliente_email", "").strip(),
            cliente_telefono=request.form.get("cliente_telefono", "").strip(),
            contraparte_nombre=request.form.get("contraparte_nombre", "").strip(),
            demandante=request.form.get("demandante", "").strip(),
            demandado=request.form.get("demandado", "").strip(),
            asesor_externo=request.form.get("asesor_externo", "").strip(),
            abogado_interno=request.form.get("abogado_interno", "").strip(),
            tipo_proceso=request.form.get("tipo_proceso", "").strip(),
            tipo_caso=tipo_caso_texto,
            case_type_id=case_type_id,
            juzgado=request.form.get("juzgado", "").strip(),
            estado=request.form.get("estado", "por_comenzar"),
            prioridad=request.form.get("prioridad", "media"),
            notas=request.form.get("notas", "").strip(),
            created_by_id=current_user.id
        )
        
        fecha_limite = request.form.get("fecha_limite", "").strip()
        if fecha_limite:
            try:
                caso.fecha_limite = datetime.strptime(fecha_limite, "%Y-%m-%d")
            except ValueError:
                pass
        
        db.session.add(caso)
        db.session.flush()
        
        campos_a_guardar = []
        campos_generales = CaseCustomField.query.filter_by(tenant_id=tenant.id, case_type_id=None, activo=True).all()
        campos_a_guardar.extend(campos_generales)
        
        if case_type_id:
            campos_tipo = CaseCustomField.query.filter_by(tenant_id=tenant.id, case_type_id=case_type_id, activo=True).all()
            campos_a_guardar.extend(campos_tipo)
        
        for campo in campos_a_guardar:
            valor = request.form.get(f"custom_{campo.nombre}", "").strip()
            if valor:
                field_value = CaseCustomFieldValue(
                    case_id=caso.id,
                    field_id=campo.id,
                    valor=valor
                )
                db.session.add(field_value)
        
        colaboradores_ids = request.form.getlist("colaboradores")
        responsable_id = request.form.get("responsable_id", type=int)
        
        for colab_id in colaboradores_ids:
            try:
                user_id = int(colab_id)
                es_responsable = (user_id == responsable_id)
                assignment = CaseAssignment(
                    case_id=caso.id,
                    user_id=user_id,
                    rol_en_caso='abogado',
                    es_responsable=es_responsable
                )
                db.session.add(assignment)
            except ValueError:
                pass
        
        if responsable_id and str(responsable_id) not in colaboradores_ids:
            assignment = CaseAssignment(
                case_id=caso.id,
                user_id=responsable_id,
                rol_en_caso='abogado',
                es_responsable=True
            )
            db.session.add(assignment)
        
        archivos = request.files.getlist("archivos")
        if archivos:
            attachments_dir = os.path.join("case_attachments", f"tenant_{tenant.id}", f"case_{caso.id}")
            os.makedirs(attachments_dir, exist_ok=True)
            
            for archivo in archivos:
                if archivo and archivo.filename:
                    filename = secure_filename(archivo.filename)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    unique_filename = f"{timestamp}_{filename}"
                    filepath = os.path.join(attachments_dir, unique_filename)
                    archivo.save(filepath)
                    
                    ext = os.path.splitext(filename)[1].lower()
                    attachment = CaseAttachment(
                        case_id=caso.id,
                        nombre=filename,
                        archivo=filepath,
                        tipo_archivo=ext,
                        uploaded_by_id=current_user.id
                    )
                    db.session.add(attachment)
        
        db.session.commit()
        flash("Caso creado exitosamente.", "success")
        return redirect(url_for("caso_detalle", caso_id=caso.id))
    
    usuarios = User.query.filter_by(tenant_id=tenant.id, activo=True).all()
    tipos_caso = CaseType.query.filter_by(tenant_id=tenant.id, activo=True).order_by(CaseType.orden, CaseType.nombre).all()
    campos_generales = CaseCustomField.query.filter_by(tenant_id=tenant.id, case_type_id=None, activo=True).order_by(CaseCustomField.orden).all()
    
    return render_template("caso_form.html", caso=None, usuarios=usuarios, estados=Case.ESTADOS, prioridades=Case.PRIORIDADES,
                          tipos_caso=tipos_caso, campos_generales=campos_generales, tipos_campo=CaseCustomField.TIPOS)


@app.route("/casos/<int:caso_id>")
@case_access_required
def caso_detalle(caso_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/casos/<int:caso_id>/editar", methods=["GET", "POST"])
@case_manage_required
def caso_editar(caso_id):
    """Ruta deshabilitada."""
    abort(404)


def _caso_editar_disabled(caso_id):
    tenant = get_current_tenant()
    caso = Case.query.filter_by(id=caso_id, tenant_id=tenant.id).first_or_404()
    
    if request.method == "POST":
        caso.titulo = request.form.get("titulo", "").strip()
        caso.descripcion = request.form.get("descripcion", "").strip()
        caso.numero_expediente = request.form.get("numero_expediente", "").strip()
        caso.cliente_nombre = request.form.get("cliente_nombre", "").strip()
        caso.cliente_email = request.form.get("cliente_email", "").strip()
        caso.cliente_telefono = request.form.get("cliente_telefono", "").strip()
        caso.contraparte_nombre = request.form.get("contraparte_nombre", "").strip()
        caso.demandante = request.form.get("demandante", "").strip()
        caso.demandado = request.form.get("demandado", "").strip()
        caso.asesor_externo = request.form.get("asesor_externo", "").strip()
        caso.abogado_interno = request.form.get("abogado_interno", "").strip()
        caso.tipo_proceso = request.form.get("tipo_proceso", "").strip()
        caso.tipo_caso = request.form.get("tipo_caso", "").strip()
        caso.juzgado = request.form.get("juzgado", "").strip()
        caso.estado = request.form.get("estado", "por_comenzar")
        caso.prioridad = request.form.get("prioridad", "media")
        caso.notas = request.form.get("notas", "").strip()
        
        fecha_limite = request.form.get("fecha_limite", "").strip()
        if fecha_limite:
            try:
                caso.fecha_limite = datetime.strptime(fecha_limite, "%Y-%m-%d")
            except ValueError:
                pass
        else:
            caso.fecha_limite = None
        
        if caso.estado == 'terminado' and not caso.fecha_cierre:
            caso.fecha_cierre = datetime.utcnow()
        elif caso.estado != 'terminado':
            caso.fecha_cierre = None
        
        db.session.commit()
        flash("Caso actualizado exitosamente.", "success")
        return redirect(url_for("caso_detalle", caso_id=caso.id))
    
    usuarios = User.query.filter_by(tenant_id=tenant.id, activo=True).all()
    tipos_caso = CaseType.query.filter_by(tenant_id=tenant.id, activo=True).order_by(CaseType.orden, CaseType.nombre).all()
    campos_generales = CaseCustomField.query.filter_by(tenant_id=tenant.id, case_type_id=None, activo=True).order_by(CaseCustomField.orden).all()
    return render_template("caso_form.html", caso=caso, usuarios=usuarios, estados=Case.ESTADOS, prioridades=Case.PRIORIDADES,
                          tipos_caso=tipos_caso, campos_generales=campos_generales, tipos_campo=CaseCustomField.TIPOS)


@app.route("/casos/<int:caso_id>/asignar", methods=["POST"])
@case_manage_required
def caso_asignar(caso_id):
    """Ruta deshabilitada."""
    abort(404)


def _caso_asignar_disabled(caso_id):
    tenant = get_current_tenant()
    caso = Case.query.filter_by(id=caso_id, tenant_id=tenant.id).first_or_404()
    
    user_id = request.form.get("user_id", type=int)
    rol = request.form.get("rol", "abogado")
    es_responsable = request.form.get("es_responsable") == "on"
    
    if not user_id:
        flash("Selecciona un usuario.", "error")
        return redirect(url_for("caso_detalle", caso_id=caso_id))
    
    existing = CaseAssignment.query.filter_by(case_id=caso_id, user_id=user_id).first()
    if existing:
        flash("Este usuario ya está asignado al caso.", "error")
        return redirect(url_for("caso_detalle", caso_id=caso_id))
    
    if es_responsable:
        CaseAssignment.query.filter_by(case_id=caso_id, es_responsable=True).update({'es_responsable': False})
    
    assignment = CaseAssignment(
        case_id=caso_id,
        user_id=user_id,
        rol_en_caso=rol,
        es_responsable=es_responsable
    )
    db.session.add(assignment)
    db.session.commit()
    
    # Send email notification to assigned user
    assigned_user = User.query.get(user_id)
    if assigned_user and assigned_user.email and user_id != current_user.id:
        tenant_name = tenant.nombre if tenant else "el sistema"
        rol_display = {
            'abogado': 'Abogado',
            'asistente': 'Asistente Legal',
            'supervisor': 'Supervisor',
            'practicante': 'Practicante'
        }.get(rol, rol)
        
        html_content = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <h2 style="color: #10b981;">Te han asignado a un caso</h2>
            <p>Hola {assigned_user.username},</p>
            <p><strong>{current_user.username}</strong> te ha asignado al siguiente caso:</p>
            <div style="background: #f3f4f6; padding: 15px; border-radius: 8px; margin: 20px 0;">
                <h3 style="margin: 0 0 10px 0; color: #1f2937;">{caso.titulo}</h3>
                <p><strong>Tu rol:</strong> {rol_display}</p>
                <p><strong>Estado:</strong> {Case.ESTADOS.get(caso.estado, caso.estado)}</p>
                <p><strong>Prioridad:</strong> {Case.PRIORIDADES.get(caso.prioridad, caso.prioridad)}</p>
                {f'<p><strong>Fecha límite:</strong> {caso.fecha_limite.strftime("%d/%m/%Y")}</p>' if caso.fecha_limite else ''}
                {'<p style="color: #059669;"><strong>Eres el responsable principal de este caso.</strong></p>' if es_responsable else ''}
            </div>
            <p>Ya puedes acceder al caso desde tu panel de casos.</p>
            <p style="color: #666; font-size: 12px;">Este correo fue enviado desde {tenant_name}.</p>
        </div>
        """
        try:
            send_notification_email(
                assigned_user.email,
                f"Asignación a caso: {caso.titulo}",
                html_content
            )
            logging.info(f"Case assignment notification sent to {assigned_user.email}")
        except Exception as e:
            logging.error(f"Error sending case assignment email: {e}")
    
    flash("Usuario asignado exitosamente.", "success")
    return redirect(url_for("caso_detalle", caso_id=caso_id))


@app.route("/casos/<int:caso_id>/desasignar/<int:assignment_id>", methods=["POST"])
@case_manage_required
def caso_desasignar(caso_id, assignment_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/casos/<int:caso_id>/estado", methods=["POST"])
@case_access_required
def caso_cambiar_estado(caso_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/casos/<int:caso_id>/evento", methods=["POST"])
@case_access_required
def caso_agregar_evento(caso_id):
    """Ruta deshabilitada."""
    abort(404)


def _caso_agregar_evento_disabled(caso_id):
    """Agregar un evento manualmente a la línea de tiempo del caso."""
    tenant = get_current_tenant()
    caso = Case.query.filter_by(id=caso_id, tenant_id=tenant.id).first_or_404()
    
    tipo_evento = request.form.get("tipo_evento", "otro")
    titulo = request.form.get("titulo", "").strip()
    descripcion = request.form.get("descripcion", "").strip()
    estado_resultado = request.form.get("estado_resultado", "").strip() or None
    fecha_evento_str = request.form.get("fecha_evento", "").strip()
    
    if not titulo:
        flash("El título del evento es requerido.", "error")
        return redirect(url_for("caso_detalle", caso_id=caso_id))
    
    fecha_evento = datetime.utcnow()
    if fecha_evento_str:
        try:
            fecha_evento = datetime.strptime(fecha_evento_str, "%Y-%m-%dT%H:%M")
        except ValueError:
            try:
                fecha_evento = datetime.strptime(fecha_evento_str, "%Y-%m-%d")
            except ValueError:
                pass
    
    CaseEvent.registrar(
        tenant_id=tenant.id,
        case_id=caso_id,
        user_id=current_user.id,
        tipo_evento=tipo_evento,
        titulo=titulo,
        descripcion=descripcion or None,
        estado_resultado=estado_resultado,
        fecha_evento=fecha_evento
    )
    db.session.commit()
    
    flash("Evento agregado a la línea de tiempo.", "success")
    return redirect(url_for("caso_detalle", caso_id=caso_id))


# ==================== GESTIÓN DE TAREAS ====================

@app.route("/tareas")
@case_access_required
def tareas():
    tenant = get_current_tenant()
    
    estado_filter = request.args.get('estado', '')
    tipo_filter = request.args.get('tipo', '')
    mis_tareas = request.args.get('mis_tareas', '')
    
    query = Task.query.filter_by(tenant_id=tenant.id)
    
    if mis_tareas or not current_user.can_manage_cases():
        query = query.filter(
            db.or_(
                Task.assigned_to_id == current_user.id,
                Task.created_by_id == current_user.id
            )
        )
    
    if estado_filter:
        query = query.filter_by(estado=estado_filter)
    if tipo_filter:
        query = query.filter_by(tipo=tipo_filter)
    
    tareas_list = query.order_by(
        db.case(
            (Task.estado == 'pendiente', 1),
            (Task.estado == 'en_curso', 2),
            (Task.estado == 'bloqueado', 3),
            else_=4
        ),
        Task.fecha_vencimiento.asc().nullslast()
    ).all()
    
    tareas_pendientes = Task.query.filter_by(tenant_id=tenant.id, estado='pendiente').count()
    tareas_vencidas = Task.query.filter(
        Task.tenant_id == tenant.id,
        Task.estado.notin_(['completado', 'cancelado']),
        Task.fecha_vencimiento.isnot(None),
        Task.fecha_vencimiento < datetime.utcnow()
    ).count()
    
    return render_template("tareas.html",
                          tareas=tareas_list,
                          tareas_pendientes=tareas_pendientes,
                          tareas_vencidas=tareas_vencidas,
                          estado_filter=estado_filter,
                          tipo_filter=tipo_filter,
                          mis_tareas=mis_tareas,
                          estados=Task.ESTADOS,
                          tipos=Task.TIPOS)


@app.route("/tareas/nueva", methods=["GET", "POST"])
@case_access_required
def tarea_nueva():
    tenant = get_current_tenant()
    
    if request.method == "POST":
        titulo = request.form.get("titulo", "").strip()
        if not titulo:
            flash("El título es obligatorio.", "error")
            return redirect(url_for("tarea_nueva"))
        
        tarea = Task(
            tenant_id=tenant.id,
            titulo=titulo,
            descripcion=request.form.get("descripcion", "").strip(),
            tipo=request.form.get("tipo", "general"),
            prioridad=request.form.get("prioridad", "media"),
            created_by_id=current_user.id
        )
        
        case_id = request.form.get("case_id", type=int)
        if case_id:
            caso = Case.query.filter_by(id=case_id, tenant_id=tenant.id).first()
            if caso:
                tarea.case_id = case_id
        
        assigned_to_id = request.form.get("assigned_to_id", type=int)
        if assigned_to_id:
            user = User.query.filter_by(id=assigned_to_id, tenant_id=tenant.id, activo=True).first()
            if user:
                tarea.assigned_to_id = assigned_to_id
        
        fecha_vencimiento = request.form.get("fecha_vencimiento", "").strip()
        if fecha_vencimiento:
            try:
                tarea.fecha_vencimiento = datetime.strptime(fecha_vencimiento, "%Y-%m-%d")
            except ValueError:
                pass
        
        archivo = request.files.get("archivo")
        if archivo and archivo.filename:
            attachments_dir = os.path.join("task_attachments", f"tenant_{tenant.id}")
            os.makedirs(attachments_dir, exist_ok=True)
            
            filename = secure_filename(archivo.filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_filename = f"{timestamp}_{filename}"
            filepath = os.path.join(attachments_dir, unique_filename)
            archivo.save(filepath)
            
            tarea.archivo = filepath
            tarea.archivo_nombre = filename
        
        db.session.add(tarea)
        db.session.commit()
        
        logging.info(f"Task created: {tarea.id}, assigned_to_id: {tarea.assigned_to_id}")
        
        # Send email notification to assigned user
        if tarea.assigned_to_id:
            assigned_user = User.query.get(tarea.assigned_to_id)
            if assigned_user and assigned_user.email:
                tenant_name = tenant.nombre if tenant else "el sistema"
                caso_info = f"<p><strong>Caso:</strong> {tarea.case.titulo}</p>" if tarea.case else ""
                fecha_info = f"<p><strong>Fecha límite:</strong> {tarea.fecha_vencimiento.strftime('%d/%m/%Y')}</p>" if tarea.fecha_vencimiento else ""
                
                # Check if deadline is soon (1-3 days)
                urgency_warning = ""
                days_until_deadline = None
                if tarea.fecha_vencimiento:
                    from datetime import date
                    today = date.today()
                    deadline_date = tarea.fecha_vencimiento.date() if hasattr(tarea.fecha_vencimiento, 'date') else tarea.fecha_vencimiento
                    days_until_deadline = (deadline_date - today).days
                    
                    if days_until_deadline <= 1:
                        urgency_warning = '<p style="color: #dc2626; font-weight: bold;">⚠️ URGENTE: Esta tarea vence mañana o hoy.</p>'
                    elif days_until_deadline <= 2:
                        urgency_warning = '<p style="color: #f59e0b; font-weight: bold;">⚡ IMPORTANTE: Esta tarea vence en 2 días.</p>'
                    elif days_until_deadline <= 3:
                        urgency_warning = '<p style="color: #3b82f6; font-weight: bold;">📅 RECORDATORIO: Esta tarea vence en 3 días.</p>'
                
                subject_prefix = ""
                if days_until_deadline is not None:
                    if days_until_deadline <= 1:
                        subject_prefix = "⚠️ URGENTE: "
                    elif days_until_deadline <= 2:
                        subject_prefix = "⚡ IMPORTANTE: "
                    elif days_until_deadline <= 3:
                        subject_prefix = "📅 "
                
                html_content = f"""
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                    <h2 style="color: #3b82f6;">Nueva Tarea Asignada</h2>
                    <p>Hola {assigned_user.username},</p>
                    <p><strong>{current_user.username}</strong> te ha asignado una nueva tarea:</p>
                    {urgency_warning}
                    <div style="background: #f3f4f6; padding: 15px; border-radius: 8px; margin: 20px 0;">
                        <h3 style="margin: 0 0 10px 0; color: #1f2937;">{tarea.titulo}</h3>
                        <p><strong>Tipo:</strong> {Task.TIPOS.get(tarea.tipo, tarea.tipo)}</p>
                        <p><strong>Prioridad:</strong> {Task.PRIORIDADES.get(tarea.prioridad, tarea.prioridad)}</p>
                        {fecha_info}
                        {caso_info}
                        {f'<p style="color: #666;">{tarea.descripcion[:300]}{"..." if len(tarea.descripcion) > 300 else ""}</p>' if tarea.descripcion else ''}
                    </div>
                    <p>Por favor, revisa esta tarea en tu bandeja de trabajo.</p>
                    <p style="color: #666; font-size: 12px;">Este correo fue enviado desde {tenant_name}.</p>
                </div>
                """
                try:
                    logging.info(f"Sending task assignment email to {assigned_user.email} for task {tarea.id}")
                    result = send_notification_email(
                        assigned_user.email,
                        f"{subject_prefix}Nueva tarea asignada: {tarea.titulo}",
                        html_content
                    )
                    if result:
                        logging.info(f"Task assignment notification sent successfully to {assigned_user.email}")
                    else:
                        logging.warning(f"Task assignment notification failed for {assigned_user.email}")
                except Exception as e:
                    logging.error(f"Error sending task assignment email: {e}")
        
        flash("Tarea creada exitosamente.", "success")
        return redirect(url_for("tareas"))
    
    casos = Case.query.filter_by(tenant_id=tenant.id).filter(Case.estado.notin_(['terminado', 'archivado'])).all()
    usuarios = User.query.filter_by(tenant_id=tenant.id, activo=True).all()
    preselected_case = request.args.get('caso_id', type=int)
    fecha_preseleccionada = request.args.get('fecha', '')
    
    return render_template("tarea_form.html",
                          tarea=None,
                          casos=casos,
                          usuarios=usuarios,
                          preselected_case=preselected_case,
                          tipos=Task.TIPOS,
                          prioridades=Task.PRIORIDADES,
                          fecha_preseleccionada=fecha_preseleccionada)


@app.route("/tareas/<int:tarea_id>/estado", methods=["POST"])
@case_access_required
def tarea_cambiar_estado(tarea_id):
    tenant = get_current_tenant()
    tarea = Task.query.filter_by(id=tarea_id, tenant_id=tenant.id).first_or_404()
    
    if tarea.assigned_to_id != current_user.id and tarea.created_by_id != current_user.id and not current_user.can_manage_cases():
        flash("No tienes permiso para modificar esta tarea.", "error")
        return redirect(url_for("tareas"))
    
    nuevo_estado = request.form.get("estado")
    if nuevo_estado in Task.ESTADOS:
        tarea.estado = nuevo_estado
        if nuevo_estado == 'completado':
            tarea.fecha_completada = datetime.utcnow()
        else:
            tarea.fecha_completada = None
        db.session.commit()
        flash(f"Tarea actualizada: {Task.ESTADOS[nuevo_estado]}", "success")
    
    next_url = request.form.get("next", url_for("tareas"))
    return redirect(next_url)


@app.route("/tareas/<int:tarea_id>/archivo")
@case_access_required
def descargar_archivo_tarea(tarea_id):
    """Descargar archivo adjunto de una tarea."""
    tenant = get_current_tenant()
    tarea = Task.query.filter_by(id=tarea_id, tenant_id=tenant.id).first_or_404()
    
    if not tarea.archivo or not os.path.exists(tarea.archivo):
        flash("Archivo no encontrado.", "error")
        return redirect(url_for("tareas"))
    
    return send_file(
        tarea.archivo,
        as_attachment=True,
        download_name=tarea.archivo_nombre or os.path.basename(tarea.archivo)
    )


@app.route("/casos/<int:caso_id>/adjunto/<int:attachment_id>")
@case_access_required
def descargar_adjunto_caso(caso_id, attachment_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/tareas/<int:tarea_id>")
@case_access_required
def tarea_detalle(tarea_id):
    """Vista de detalle de una tarea con documentos vinculados."""
    tenant = get_current_tenant()
    tarea = Task.query.filter_by(id=tarea_id, tenant_id=tenant.id).first_or_404()
    
    if not current_user.can_manage_cases():
        if tarea.assigned_to_id != current_user.id and tarea.created_by_id != current_user.id:
            flash("No tienes permiso para ver esta tarea.", "error")
            return redirect(url_for("tareas"))
    
    documentos_vinculados = db.session.query(TaskDocument).join(
        FinishedDocument, TaskDocument.document_id == FinishedDocument.id
    ).filter(
        TaskDocument.task_id == tarea_id,
        FinishedDocument.tenant_id == tenant.id
    ).all()
    documentos_ids = [td.document_id for td in documentos_vinculados]
    
    documentos_disponibles = FinishedDocument.query.filter(
        FinishedDocument.tenant_id == tenant.id,
        FinishedDocument.id.notin_(documentos_ids) if documentos_ids else True
    ).order_by(FinishedDocument.created_at.desc()).limit(50).all()
    
    return render_template("tarea_detalle.html",
                          tarea=tarea,
                          documentos_vinculados=documentos_vinculados,
                          documentos_disponibles=documentos_disponibles,
                          estados=Task.ESTADOS,
                          tipos=Task.TIPOS,
                          prioridades=Task.PRIORIDADES)


@app.route("/tareas/<int:tarea_id>/vincular-documento", methods=["POST"])
@case_access_required
def vincular_documento_tarea(tarea_id):
    """Vincular un documento terminado a una tarea."""
    tenant = get_current_tenant()
    tarea = Task.query.filter_by(id=tarea_id, tenant_id=tenant.id).first_or_404()
    
    if tarea.assigned_to_id != current_user.id and tarea.created_by_id != current_user.id and not current_user.can_manage_cases():
        flash("No tienes permiso para modificar esta tarea.", "error")
        return redirect(url_for("tareas"))
    
    document_id = request.form.get("document_id", type=int)
    if not document_id:
        flash("Debes seleccionar un documento.", "error")
        return redirect(url_for("tarea_detalle", tarea_id=tarea_id))
    
    documento = FinishedDocument.query.filter_by(id=document_id, tenant_id=tenant.id).first()
    if not documento:
        flash("Documento no encontrado.", "error")
        return redirect(url_for("tarea_detalle", tarea_id=tarea_id))
    
    existing = TaskDocument.query.filter_by(task_id=tarea_id, document_id=document_id).first()
    if existing:
        flash("Este documento ya está vinculado a la tarea.", "warning")
        return redirect(url_for("tarea_detalle", tarea_id=tarea_id))
    
    task_doc = TaskDocument(
        task_id=tarea_id,
        document_id=document_id,
        linked_by_id=current_user.id
    )
    db.session.add(task_doc)
    
    documento.task_id = tarea_id
    
    db.session.commit()
    flash("Documento vinculado exitosamente.", "success")
    return redirect(url_for("tarea_detalle", tarea_id=tarea_id))


@app.route("/tareas/<int:tarea_id>/desvincular-documento/<int:doc_id>", methods=["POST"])
@case_access_required
def desvincular_documento_tarea(tarea_id, doc_id):
    """Desvincular un documento de una tarea."""
    tenant = get_current_tenant()
    tarea = Task.query.filter_by(id=tarea_id, tenant_id=tenant.id).first_or_404()
    
    if tarea.assigned_to_id != current_user.id and tarea.created_by_id != current_user.id and not current_user.can_manage_cases():
        flash("No tienes permiso para modificar esta tarea.", "error")
        return redirect(url_for("tareas"))
    
    task_doc = db.session.query(TaskDocument).join(
        FinishedDocument, TaskDocument.document_id == FinishedDocument.id
    ).filter(
        TaskDocument.task_id == tarea_id,
        TaskDocument.document_id == doc_id,
        FinishedDocument.tenant_id == tenant.id
    ).first()
    
    if task_doc:
        db.session.delete(task_doc)
        
        documento = FinishedDocument.query.filter_by(id=doc_id, tenant_id=tenant.id).first()
        if documento and documento.task_id == tarea_id:
            documento.task_id = None
        
        db.session.commit()
        flash("Documento desvinculado.", "success")
    else:
        flash("El documento no estaba vinculado.", "warning")
    
    return redirect(url_for("tarea_detalle", tarea_id=tarea_id))


# ==================== CALENDARIO ====================

@app.route("/calendario")
@case_access_required
def calendario():
    """Vista de calendario con tareas."""
    import calendar
    from datetime import date, timedelta
    
    tenant = get_current_tenant()
    today = date.today()
    
    year = request.args.get('year', today.year, type=int)
    month = request.args.get('month', today.month, type=int)
    view_mode = request.args.get('view', 'month')
    estado_filter = request.args.get('estado', '')
    buscar_filter = request.args.get('buscar', '')
    
    if month < 1:
        month = 12
        year -= 1
    elif month > 12:
        month = 1
        year += 1
    
    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1
    
    cal = calendar.Calendar(firstweekday=0)
    month_days = cal.monthdatescalendar(year, month)
    
    if view_mode == 'week':
        for week in month_days:
            if today in week or (week[0] <= today <= week[-1]):
                month_days = [week]
                break
    
    first_day = month_days[0][0]
    last_day = month_days[-1][-1]
    
    query = Task.query.filter(
        Task.tenant_id == tenant.id,
        Task.fecha_vencimiento.isnot(None),
        Task.fecha_vencimiento >= datetime.combine(first_day, datetime.min.time()),
        Task.fecha_vencimiento <= datetime.combine(last_day, datetime.max.time())
    )
    
    if not current_user.can_manage_cases():
        query = query.filter(
            db.or_(
                Task.assigned_to_id == current_user.id,
                Task.created_by_id == current_user.id
            )
        )
    
    if estado_filter:
        query = query.filter_by(estado=estado_filter)
    if buscar_filter:
        query = query.filter(Task.titulo.ilike(f'%{buscar_filter}%'))
    
    tasks = query.all()
    
    tasks_by_date = {}
    for task in tasks:
        task_date = task.fecha_vencimiento.date()
        if task_date not in tasks_by_date:
            tasks_by_date[task_date] = []
        tasks_by_date[task_date].append(task)
    
    calendar_days = []
    for week in month_days:
        for day in week:
            calendar_days.append({
                'day': day.day,
                'date': day.isoformat(),
                'full_date': day.strftime('%Y-%m-%d'),
                'other_month': day.month != month,
                'is_today': day == today,
                'tasks': tasks_by_date.get(day, [])
            })
    
    upcoming_query = Task.query.filter(
        Task.tenant_id == tenant.id,
        Task.estado.notin_(['completado', 'cancelado']),
        Task.fecha_vencimiento.isnot(None),
        Task.fecha_vencimiento >= datetime.utcnow()
    )
    if not current_user.can_manage_cases():
        upcoming_query = upcoming_query.filter(
            db.or_(Task.assigned_to_id == current_user.id, Task.created_by_id == current_user.id)
        )
    upcoming_tasks = upcoming_query.order_by(Task.fecha_vencimiento.asc()).limit(5).all()
    
    overdue_query = Task.query.filter(
        Task.tenant_id == tenant.id,
        Task.estado.notin_(['completado', 'cancelado']),
        Task.fecha_vencimiento.isnot(None),
        Task.fecha_vencimiento < datetime.utcnow()
    )
    if not current_user.can_manage_cases():
        overdue_query = overdue_query.filter(
            db.or_(Task.assigned_to_id == current_user.id, Task.created_by_id == current_user.id)
        )
    overdue_tasks = overdue_query.order_by(Task.fecha_vencimiento.desc()).limit(5).all()
    
    month_names = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
                   'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']
    day_names = ['Lun', 'Mar', 'Mié', 'Jue', 'Vie', 'Sáb', 'Dom']
    
    # Obtener eventos del calendario
    events_query = CalendarEvent.query.filter(
        CalendarEvent.tenant_id == tenant.id,
        CalendarEvent.fecha_inicio >= datetime.combine(first_day, datetime.min.time()),
        CalendarEvent.fecha_inicio <= datetime.combine(last_day, datetime.max.time())
    )
    
    if not current_user.can_manage_cases():
        events_query = events_query.filter(
            db.or_(
                CalendarEvent.created_by_id == current_user.id,
                CalendarEvent.attendees.any(EventAttendee.user_id == current_user.id)
            )
        )
    
    events = events_query.all()
    
    events_by_date = {}
    for event in events:
        event_date = event.fecha_inicio.date()
        if event_date not in events_by_date:
            events_by_date[event_date] = []
        events_by_date[event_date].append(event)
    
    # Agregar eventos a calendar_days
    for day_data in calendar_days:
        from datetime import date as date_type
        day_date = date_type.fromisoformat(day_data['date'])
        day_data['events'] = events_by_date.get(day_date, [])
    
    # Próximos eventos
    upcoming_events = CalendarEvent.query.filter(
        CalendarEvent.tenant_id == tenant.id,
        CalendarEvent.fecha_inicio >= datetime.utcnow()
    )
    if not current_user.can_manage_cases():
        upcoming_events = upcoming_events.filter(
            db.or_(
                CalendarEvent.created_by_id == current_user.id,
                CalendarEvent.attendees.any(EventAttendee.user_id == current_user.id)
            )
        )
    upcoming_events = upcoming_events.order_by(CalendarEvent.fecha_inicio.asc()).limit(5).all()
    
    return render_template("calendario.html",
                          year=year,
                          month=month,
                          view_mode=view_mode,
                          prev_year=prev_year,
                          prev_month=prev_month,
                          next_year=next_year,
                          next_month=next_month,
                          calendar_days=calendar_days,
                          upcoming_tasks=upcoming_tasks,
                          overdue_tasks=overdue_tasks,
                          upcoming_events=upcoming_events,
                          today=today,
                          month_names=month_names,
                          day_names=day_names,
                          estados=Task.ESTADOS,
                          estado_filter=estado_filter,
                          buscar_filter=buscar_filter,
                          event_tipos=CalendarEvent.TIPOS)


# ==================== EVENTOS DEL CALENDARIO ====================

@app.route("/eventos/nuevo", methods=["GET", "POST"])
@case_access_required
def evento_nuevo():
    """Crear un nuevo evento del calendario."""
    tenant = get_current_tenant()
    
    if request.method == "POST":
        titulo = request.form.get("titulo", "").strip()
        descripcion = request.form.get("descripcion", "").strip()
        tipo = request.form.get("tipo", "reunion")
        ubicacion = request.form.get("ubicacion", "").strip()
        link = request.form.get("link", "").strip() or None
        color = request.form.get("color", "#3b82f6")
        case_id = request.form.get("case_id", type=int)
        todo_el_dia = request.form.get("todo_el_dia") == "on"
        recordatorio_minutos = request.form.get("recordatorio_minutos", 30, type=int)
        
        fecha_inicio_str = request.form.get("fecha_inicio")
        hora_inicio_str = request.form.get("hora_inicio", "09:00")
        fecha_fin_str = request.form.get("fecha_fin")
        hora_fin_str = request.form.get("hora_fin", "10:00")
        
        if not titulo:
            flash("El título es obligatorio.", "error")
            return redirect(url_for("evento_nuevo"))
        
        if not fecha_inicio_str:
            flash("La fecha de inicio es obligatoria.", "error")
            return redirect(url_for("evento_nuevo"))
        
        try:
            if todo_el_dia:
                fecha_inicio = datetime.strptime(fecha_inicio_str, "%Y-%m-%d")
                fecha_fin = datetime.strptime(fecha_fin_str, "%Y-%m-%d") if fecha_fin_str else fecha_inicio
            else:
                fecha_inicio = datetime.strptime(f"{fecha_inicio_str} {hora_inicio_str}", "%Y-%m-%d %H:%M")
                if fecha_fin_str:
                    fecha_fin = datetime.strptime(f"{fecha_fin_str} {hora_fin_str}", "%Y-%m-%d %H:%M")
                else:
                    fecha_fin = fecha_inicio
        except ValueError as e:
            flash(f"Formato de fecha inválido: {e}", "error")
            return redirect(url_for("evento_nuevo"))
        
        evento = CalendarEvent(
            tenant_id=tenant.id,
            titulo=titulo,
            descripcion=descripcion,
            tipo=tipo,
            ubicacion=ubicacion,
            link=link,
            color=color,
            fecha_inicio=fecha_inicio,
            fecha_fin=fecha_fin,
            todo_el_dia=todo_el_dia,
            case_id=case_id if case_id else None,
            created_by_id=current_user.id,
            recordatorio_minutos=recordatorio_minutos
        )
        db.session.add(evento)
        db.session.flush()
        
        # Agregar invitados (solo usuarios del mismo tenant)
        invitados_ids = request.form.getlist("invitados")
        for user_id in invitados_ids:
            if user_id and int(user_id) != current_user.id:
                # Verificar que el usuario pertenece al mismo tenant
                invited_user = User.query.filter_by(id=int(user_id), tenant_id=tenant.id, activo=True).first()
                if invited_user:
                    attendee = EventAttendee(
                        event_id=evento.id,
                        user_id=int(user_id),
                        estado='pendiente'
                    )
                    db.session.add(attendee)
        
        db.session.commit()
        flash("Evento creado correctamente.", "success")
        return redirect(url_for("evento_detalle", evento_id=evento.id))
    
    # GET - mostrar formulario
    cases = Case.query.filter_by(tenant_id=tenant.id).order_by(Case.titulo).all()
    usuarios = User.query.filter_by(tenant_id=tenant.id, activo=True).all()
    
    # Fecha preseleccionada desde el calendario
    fecha_preseleccionada = request.args.get('fecha', '')
    
    return render_template("evento_form.html",
                          evento=None,
                          cases=cases,
                          usuarios=usuarios,
                          tipos=CalendarEvent.TIPOS,
                          fecha_preseleccionada=fecha_preseleccionada)


@app.route("/eventos/<int:evento_id>")
@case_access_required
def evento_detalle(evento_id):
    """Ver detalles de un evento."""
    tenant = get_current_tenant()
    
    evento = CalendarEvent.query.filter_by(id=evento_id, tenant_id=tenant.id).first_or_404()
    
    # Verificar acceso
    if not current_user.can_manage_cases():
        is_creator = evento.created_by_id == current_user.id
        is_attendee = EventAttendee.query.filter_by(event_id=evento_id, user_id=current_user.id).first()
        if not is_creator and not is_attendee:
            flash("No tienes acceso a este evento.", "error")
            return redirect(url_for("calendario"))
    
    # Obtener respuesta del usuario actual
    mi_respuesta = EventAttendee.query.filter_by(event_id=evento_id, user_id=current_user.id).first()
    
    return render_template("evento_detalle.html",
                          evento=evento,
                          mi_respuesta=mi_respuesta)


@app.route("/eventos/<int:evento_id>/editar", methods=["GET", "POST"])
@case_access_required
def evento_editar(evento_id):
    """Editar un evento existente."""
    tenant = get_current_tenant()
    
    evento = CalendarEvent.query.filter_by(id=evento_id, tenant_id=tenant.id).first_or_404()
    
    # Solo el creador o admin puede editar
    if not current_user.can_manage_cases() and evento.created_by_id != current_user.id:
        flash("No tienes permiso para editar este evento.", "error")
        return redirect(url_for("evento_detalle", evento_id=evento_id))
    
    if request.method == "POST":
        evento.titulo = request.form.get("titulo", "").strip()
        evento.descripcion = request.form.get("descripcion", "").strip()
        evento.tipo = request.form.get("tipo", "reunion")
        evento.ubicacion = request.form.get("ubicacion", "").strip()
        evento.link = request.form.get("link", "").strip() or None
        evento.color = request.form.get("color", "#3b82f6")
        evento.todo_el_dia = request.form.get("todo_el_dia") == "on"
        evento.recordatorio_minutos = request.form.get("recordatorio_minutos", 30, type=int)
        
        case_id = request.form.get("case_id", type=int)
        evento.case_id = case_id if case_id else None
        
        fecha_inicio_str = request.form.get("fecha_inicio")
        hora_inicio_str = request.form.get("hora_inicio", "09:00")
        fecha_fin_str = request.form.get("fecha_fin")
        hora_fin_str = request.form.get("hora_fin", "10:00")
        
        try:
            if evento.todo_el_dia:
                evento.fecha_inicio = datetime.strptime(fecha_inicio_str, "%Y-%m-%d")
                evento.fecha_fin = datetime.strptime(fecha_fin_str, "%Y-%m-%d") if fecha_fin_str else evento.fecha_inicio
            else:
                evento.fecha_inicio = datetime.strptime(f"{fecha_inicio_str} {hora_inicio_str}", "%Y-%m-%d %H:%M")
                if fecha_fin_str:
                    evento.fecha_fin = datetime.strptime(f"{fecha_fin_str} {hora_fin_str}", "%Y-%m-%d %H:%M")
                else:
                    evento.fecha_fin = evento.fecha_inicio
        except ValueError as e:
            flash(f"Formato de fecha inválido: {e}", "error")
            return redirect(url_for("evento_editar", evento_id=evento_id))
        
        # Actualizar invitados (solo usuarios del mismo tenant)
        EventAttendee.query.filter_by(event_id=evento.id).delete()
        invitados_ids = request.form.getlist("invitados")
        for user_id in invitados_ids:
            if user_id and int(user_id) != current_user.id:
                # Verificar que el usuario pertenece al mismo tenant
                invited_user = User.query.filter_by(id=int(user_id), tenant_id=tenant.id, activo=True).first()
                if invited_user:
                    attendee = EventAttendee(
                        event_id=evento.id,
                        user_id=int(user_id),
                        estado='pendiente'
                    )
                    db.session.add(attendee)
        
        db.session.commit()
        flash("Evento actualizado correctamente.", "success")
        return redirect(url_for("evento_detalle", evento_id=evento_id))
    
    # GET - mostrar formulario
    cases = Case.query.filter_by(tenant_id=tenant.id).order_by(Case.titulo).all()
    usuarios = User.query.filter_by(tenant_id=tenant.id, activo=True).all()
    invitados_actuales = [a.user_id for a in evento.attendees.all()]
    
    return render_template("evento_form.html",
                          evento=evento,
                          cases=cases,
                          usuarios=usuarios,
                          tipos=CalendarEvent.TIPOS,
                          invitados_actuales=invitados_actuales,
                          fecha_preseleccionada='')


@app.route("/eventos/<int:evento_id>/eliminar", methods=["POST"])
@case_access_required
def evento_eliminar(evento_id):
    """Eliminar un evento."""
    tenant = get_current_tenant()
    
    evento = CalendarEvent.query.filter_by(id=evento_id, tenant_id=tenant.id).first_or_404()
    
    # Solo el creador o admin puede eliminar
    if not current_user.can_manage_cases() and evento.created_by_id != current_user.id:
        flash("No tienes permiso para eliminar este evento.", "error")
        return redirect(url_for("evento_detalle", evento_id=evento_id))
    
    db.session.delete(evento)
    db.session.commit()
    
    flash("Evento eliminado correctamente.", "success")
    return redirect(url_for("calendario"))


@app.route("/eventos/<int:evento_id>/responder", methods=["POST"])
@case_access_required
def evento_responder(evento_id):
    """Responder a una invitación de evento."""
    tenant = get_current_tenant()
    
    # Verificar que el usuario pertenece al tenant
    if current_user.tenant_id != tenant.id:
        flash("No tienes acceso a este evento.", "error")
        return redirect(url_for("calendario"))
    
    evento = CalendarEvent.query.filter_by(id=evento_id, tenant_id=tenant.id).first_or_404()
    respuesta = request.form.get("respuesta", "pendiente")
    
    if respuesta not in ['aceptado', 'rechazado', 'pendiente']:
        flash("Respuesta no válida.", "error")
        return redirect(url_for("evento_detalle", evento_id=evento_id))
    
    # Verificar que el usuario es el creador o fue invitado
    is_creator = evento.created_by_id == current_user.id
    attendee = EventAttendee.query.filter_by(event_id=evento_id, user_id=current_user.id).first()
    
    if not is_creator and not attendee:
        flash("No tienes acceso a este evento.", "error")
        return redirect(url_for("calendario"))
    
    if attendee:
        attendee.estado = respuesta
    elif is_creator:
        # El creador puede agregar su propia respuesta
        attendee = EventAttendee(
            event_id=evento_id,
            user_id=current_user.id,
            estado=respuesta
        )
        db.session.add(attendee)
    
    db.session.commit()
    
    estados_msg = {'aceptado': 'aceptada', 'rechazado': 'rechazada', 'pendiente': 'marcada como pendiente'}
    flash(f"Invitación {estados_msg.get(respuesta, respuesta)}.", "success")
    return redirect(url_for("evento_detalle", evento_id=evento_id))


# ==================== MIS MODELOS (User personal document models) ====================

@app.route("/mis-modelos")
@login_required
def mis_modelos():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    user_id = current_user.id
    
    modelos_usuario = Modelo.query.filter_by(tenant_id=tenant_id, created_by_id=user_id).all()
    estilos_usuario = Estilo.query.filter_by(tenant_id=tenant_id, created_by_id=user_id).all()
    
    return render_template("mis_modelos.html", modelos_usuario=modelos_usuario, estilos_usuario=estilos_usuario, modelos_sistema=MODELOS)


@app.route("/api/detect-campos", methods=["POST"])
@login_required
def api_detect_campos():
    """AJAX endpoint to detect fields from uploaded file with document preview."""
    archivo = request.files.get('archivo')
    
    if not archivo or not archivo.filename:
        return jsonify({'success': False, 'error': 'No se proporcionó archivo'}), 400
    
    ext = os.path.splitext(archivo.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({'success': False, 'error': f'Formato no soportado: {ext}'}), 400
    
    import tempfile
    temp_dir = tempfile.mkdtemp()
    temp_path = os.path.join(temp_dir, secure_filename(archivo.filename))
    
    try:
        archivo.save(temp_path)
        contenido = extract_text_from_file(temp_path)
        
        if not contenido:
            return jsonify({'success': False, 'error': 'No se pudo extraer texto del archivo'}), 400
        
        campos_detectados = detect_placeholders_with_context(contenido)
        
        highlighted_html = generate_highlighted_html(contenido, campos_detectados)
        
        campos_result = []
        for i, campo in enumerate(campos_detectados):
            campos_result.append({
                'nombre': campo['nombre'],
                'etiqueta': campo['etiqueta'],
                'tipo': campo['tipo'],
                'index': i,
                'contexto': campo['contexto'],
                'match_text': campo['match_text'],
                'pattern_type': campo['pattern_type']
            })
        
        return jsonify({
            'success': True, 
            'campos': campos_result,
            'contenido_html': highlighted_html,
            'contenido_raw': contenido[:5000] if len(contenido) > 5000 else contenido
        })
    except Exception as e:
        logging.error(f"Error detecting campos: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)


@app.route("/mi-modelo", methods=["GET", "POST"])
@login_required
def mi_modelo():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("dashboard"))
    
    modelo_id = request.args.get('id', type=int)
    modelo = Modelo.query.filter_by(id=modelo_id, tenant_id=tenant.id, created_by_id=current_user.id).first() if modelo_id else None
    campos_detectados = []
    
    if request.method == "POST":
        key = request.form.get("key", "").strip()
        nombre = request.form.get("nombre", "").strip()
        
        archivo = request.files.get('archivo_word')
        contenido = ""
        archivo_path = None
        
        if archivo and archivo.filename:
            ext = os.path.splitext(archivo.filename)[1].lower()
            if ext not in ALLOWED_EXTENSIONS:
                flash(f"Formato de archivo no soportado ({ext}). Use .docx, .pdf o .txt", "error")
                return render_template("mi_modelo.html", modelo=modelo, campos_detectados=campos_detectados)
            
            user_folder = os.path.join(CARPETA_PLANTILLAS_SUBIDAS, f"user_{current_user.id}")
            os.makedirs(user_folder, exist_ok=True)
            
            safe_name = secure_filename(archivo.filename)
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            archivo_name = f"{timestamp}_{safe_name}"
            archivo_path = os.path.join(user_folder, archivo_name)
            archivo.save(archivo_path)
            
            contenido = extract_text_from_file(archivo_path)
            if not contenido:
                flash("No se pudo extraer texto del archivo. Verifique que el archivo contenga texto.", "error")
                return render_template("mi_modelo.html", modelo=modelo, campos_detectados=campos_detectados)
            campos_detectados = detect_placeholders_from_text(contenido)
        if not contenido and modelo:
            contenido = modelo.contenido
        
        if not key or not nombre:
            flash("La clave y el nombre son obligatorios.", "error")
            return render_template("mi_modelo.html", modelo=modelo, campos_detectados=campos_detectados)
        
        if not contenido and not modelo:
            flash("Debes subir un archivo Word con el modelo.", "error")
            return render_template("mi_modelo.html", modelo=modelo, campos_detectados=campos_detectados)
        
        if modelo:
            modelo.key = key
            modelo.nombre = nombre
            if contenido:
                modelo.contenido = contenido
            if archivo_path:
                modelo.archivo_original = archivo_path
            flash("Modelo actualizado exitosamente.", "success")
        else:
            modelo = Modelo(
                key=f"user_{current_user.id}_{key}",
                nombre=nombre,
                contenido=contenido,
                archivo_original=archivo_path,
                carpeta_estilos=key,
                tenant_id=tenant.id,
                created_by_id=current_user.id
            )
            db.session.add(modelo)
            flash("Modelo creado exitosamente.", "success")
        
        db.session.commit()
        
        campo_nombres = request.form.getlist('campo_nombre[]')
        campo_etiquetas = request.form.getlist('campo_etiqueta[]')
        campo_tipos = request.form.getlist('campo_tipo[]')
        
        if campo_nombres:
            campos_anteriores = {c.nombre_campo: c for c in CampoPlantilla.query.filter_by(tenant_id=tenant.id, plantilla_key=modelo.key).all()}
            CampoPlantilla.query.filter_by(tenant_id=tenant.id, plantilla_key=modelo.key).delete()
            
            for i, nombre_campo in enumerate(campo_nombres):
                if nombre_campo.strip():
                    etiqueta = campo_etiquetas[i] if i < len(campo_etiquetas) else nombre_campo
                    tipo = campo_tipos[i] if i < len(campo_tipos) else 'text'
                    
                    archivo_path_campo = None
                    if tipo == 'file':
                        campo_archivo = request.files.get(f'campo_archivo_{i}')
                        if campo_archivo and campo_archivo.filename:
                            img_ext = os.path.splitext(campo_archivo.filename)[1].lower()
                            if img_ext in ALLOWED_IMAGE_EXTENSIONS:
                                campo_folder = os.path.join(CARPETA_IMAGENES_MODELOS, f"campos_{tenant.id}")
                                os.makedirs(campo_folder, exist_ok=True)
                                
                                safe_img_name = secure_filename(campo_archivo.filename)
                                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                                img_filename = f"{timestamp}_{safe_img_name}"
                                img_full_path = os.path.join(campo_folder, img_filename)
                                campo_archivo.save(img_full_path)
                                archivo_path_campo = f"campos_{tenant.id}/{img_filename}"
                        elif campo_to_key(nombre_campo) in campos_anteriores:
                            archivo_path_campo = campos_anteriores[campo_to_key(nombre_campo)].archivo_path
                    
                    campo = CampoPlantilla(
                        tenant_id=tenant.id,
                        plantilla_key=modelo.key,
                        nombre_campo=campo_to_key(nombre_campo),
                        etiqueta=etiqueta.strip(),
                        tipo=tipo,
                        orden=i,
                        archivo_path=archivo_path_campo
                    )
                    db.session.add(campo)
            
            db.session.commit()
        
        imagen_archivo = request.files.get('imagen_archivo')
        if imagen_archivo and imagen_archivo.filename:
            img_ext = os.path.splitext(imagen_archivo.filename)[1].lower()
            if img_ext in ALLOWED_IMAGE_EXTENSIONS:
                tenant_folder = os.path.join(CARPETA_IMAGENES_MODELOS, str(tenant.id))
                os.makedirs(tenant_folder, exist_ok=True)
                
                safe_img_name = secure_filename(imagen_archivo.filename)
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                img_filename = f"{timestamp}_{safe_img_name}"
                img_path = os.path.join(tenant_folder, img_filename)
                imagen_archivo.save(img_path)
                
                imagen_nombre = request.form.get('imagen_nombre', safe_img_name).strip()
                imagen_posicion = request.form.get('imagen_posicion', 'inline')
                try:
                    imagen_ancho = float(request.form.get('imagen_ancho') or 5.0)
                    if imagen_ancho < 1 or imagen_ancho > 18:
                        imagen_ancho = 5.0
                except (ValueError, TypeError):
                    imagen_ancho = 5.0
                
                nueva_imagen = ImagenModelo(
                    modelo_id=modelo.id,
                    tenant_id=tenant.id,
                    nombre=imagen_nombre,
                    archivo=img_path,
                    posicion=imagen_posicion,
                    ancho_cm=imagen_ancho,
                    orden=modelo.imagenes.count()
                )
                db.session.add(nueva_imagen)
                db.session.commit()
                flash("Imagen agregada al modelo.", "success")
                return redirect(url_for("mi_modelo", id=modelo.id))
            else:
                flash("Formato de imagen no soportado. Use JPG, PNG, GIF o WebP.", "error")
        
        tabla_nombre = request.form.get('tabla_nombre', '').strip()
        tabla_columnas = request.form.get('tabla_columnas', '').strip()
        if tabla_nombre and tabla_columnas:
            columnas_list = [c.strip() for c in tabla_columnas.split(',') if c.strip()]
            if columnas_list:
                try:
                    num_filas = int(request.form.get('tabla_filas', 5))
                    if num_filas < 1:
                        num_filas = 1
                    if num_filas > 50:
                        num_filas = 50
                except ValueError:
                    num_filas = 5
                
                mostrar_total = request.form.get('tabla_mostrar_total') == 'on'
                
                nueva_tabla = ModeloTabla(
                    modelo_id=modelo.id,
                    tenant_id=tenant.id,
                    nombre=tabla_nombre,
                    columnas=columnas_list,
                    num_filas=num_filas,
                    mostrar_total=mostrar_total,
                    columna_total=columnas_list[-1] if mostrar_total and columnas_list else None,
                    orden=modelo.tablas.count() if modelo.tablas else 0
                )
                db.session.add(nueva_tabla)
                db.session.commit()
                flash(f"Cuadro '{tabla_nombre}' agregado al modelo.", "success")
                return redirect(url_for("mi_modelo", id=modelo.id))
        
        return redirect(url_for("mis_modelos"))
    
    campos_guardados = []
    if modelo:
        campos_guardados = CampoPlantilla.query.filter_by(
            tenant_id=tenant.id, 
            plantilla_key=modelo.key
        ).order_by(CampoPlantilla.orden).all()
    
    return render_template("mi_modelo.html", modelo=modelo, campos_detectados=campos_detectados, campos_guardados=campos_guardados)


@app.route("/api/imagen-modelo/<int:imagen_id>", methods=["DELETE"])
@login_required
def eliminar_imagen_modelo(imagen_id):
    tenant = get_current_tenant()
    imagen = ImagenModelo.query.filter_by(id=imagen_id, tenant_id=tenant.id).first()
    
    if not imagen:
        return jsonify({"success": False, "error": "Imagen no encontrada"}), 404
    
    modelo = Modelo.query.get(imagen.modelo_id)
    if not modelo or modelo.created_by_id != current_user.id:
        return jsonify({"success": False, "error": "No tienes permiso"}), 403
    
    if imagen.archivo and os.path.exists(imagen.archivo):
        try:
            os.remove(imagen.archivo)
        except Exception as e:
            logging.error(f"Error deleting image file: {e}")
    
    db.session.delete(imagen)
    db.session.commit()
    return jsonify({"success": True})


@app.route("/api/modelo-tabla/<int:tabla_id>", methods=["DELETE"])
@login_required
def eliminar_tabla_modelo(tabla_id):
    tenant = get_current_tenant()
    tabla = ModeloTabla.query.filter_by(id=tabla_id, tenant_id=tenant.id).first()
    
    if not tabla:
        return jsonify({"success": False, "error": "Cuadro no encontrado"}), 404
    
    modelo = Modelo.query.get(tabla.modelo_id)
    if not modelo or modelo.created_by_id != current_user.id:
        return jsonify({"success": False, "error": "No tienes permiso"}), 403
    
    db.session.delete(tabla)
    db.session.commit()
    return jsonify({"success": True})


@app.route("/mi-modelo/eliminar/<int:modelo_id>", methods=["POST"])
@login_required
def eliminar_mi_modelo(modelo_id):
    tenant = get_current_tenant()
    modelo = Modelo.query.filter_by(id=modelo_id, tenant_id=tenant.id, created_by_id=current_user.id).first()
    
    if not modelo:
        flash("No tienes permiso para eliminar este modelo.", "error")
        return redirect(url_for("mis_modelos"))
    
    db.session.delete(modelo)
    db.session.commit()
    flash("Modelo eliminado exitosamente.", "success")
    return redirect(url_for("mis_modelos"))


# ==================== MIS ESTILOS (User personal styles) ====================

@app.route("/mis-estilos")
@login_required
def mis_estilos():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    user_id = current_user.id
    
    estilos_usuario = Estilo.query.filter_by(tenant_id=tenant_id, created_by_id=user_id).all()
    modelos_usuario = Modelo.query.filter_by(tenant_id=tenant_id, created_by_id=user_id).all()
    
    return render_template("mis_estilos.html", estilos_usuario=estilos_usuario, modelos_usuario=modelos_usuario)


@app.route("/mi-estilo", methods=["GET", "POST"])
@login_required
def mi_estilo():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("dashboard"))
    
    estilo_id = request.args.get('id', type=int)
    estilo = Estilo.query.filter_by(id=estilo_id, tenant_id=tenant.id, created_by_id=current_user.id).first() if estilo_id else None
    modelos_usuario = Modelo.query.filter_by(tenant_id=tenant.id, created_by_id=current_user.id).all()
    
    if request.method == "POST":
        nombre = request.form.get("nombre", "").strip()
        plantilla_key = request.form.get("plantilla_key", "").strip()
        
        archivo = request.files.get('archivo_word')
        contenido = ""
        archivo_path = None
        
        if archivo and archivo.filename:
            ext = os.path.splitext(archivo.filename)[1].lower()
            if ext not in ALLOWED_EXTENSIONS:
                flash(f"Formato de archivo no soportado ({ext}). Use .docx, .pdf o .txt", "error")
                return render_template("mi_estilo.html", estilo=estilo, modelos_usuario=modelos_usuario)
            
            user_folder = os.path.join(CARPETA_ESTILOS_SUBIDOS, f"user_{current_user.id}")
            os.makedirs(user_folder, exist_ok=True)
            
            safe_name = secure_filename(archivo.filename)
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            archivo_name = f"{timestamp}_{safe_name}"
            archivo_path = os.path.join(user_folder, archivo_name)
            archivo.save(archivo_path)
            
            contenido = extract_text_from_file(archivo_path)
            if not contenido:
                flash("No se pudo extraer texto del archivo. Verifique que el archivo contenga texto.", "error")
                return render_template("mi_estilo.html", estilo=estilo, modelos_usuario=modelos_usuario)
        if not contenido and estilo:
            contenido = estilo.contenido
        
        if not nombre:
            flash("El nombre es obligatorio.", "error")
            return render_template("mi_estilo.html", estilo=estilo, modelos_usuario=modelos_usuario)
        
        if not contenido and not estilo:
            flash("Debes subir un archivo Word con el estilo.", "error")
            return render_template("mi_estilo.html", estilo=estilo, modelos_usuario=modelos_usuario)
        
        if estilo:
            estilo.nombre = nombre
            estilo.plantilla_key = plantilla_key
            if contenido:
                estilo.contenido = contenido
            if archivo_path:
                estilo.archivo_original = archivo_path
            flash("Estilo actualizado exitosamente.", "success")
        else:
            estilo = Estilo(
                nombre=nombre,
                plantilla_key=plantilla_key,
                contenido=contenido,
                archivo_original=archivo_path,
                tenant_id=tenant.id,
                created_by_id=current_user.id
            )
            db.session.add(estilo)
            flash("Estilo creado exitosamente.", "success")
        
        db.session.commit()
        return redirect(url_for("mis_estilos"))
    
    return render_template("mi_estilo.html", estilo=estilo, modelos_usuario=modelos_usuario)


@app.route("/mi-estilo/eliminar/<int:estilo_id>", methods=["POST"])
@login_required
def eliminar_mi_estilo(estilo_id):
    tenant = get_current_tenant()
    estilo = Estilo.query.filter_by(id=estilo_id, tenant_id=tenant.id, created_by_id=current_user.id).first()
    
    if not estilo:
        flash("No tienes permiso para eliminar este estilo.", "error")
        return redirect(url_for("mis_estilos"))
    
    db.session.delete(estilo)
    db.session.commit()
    flash("Estilo eliminado exitosamente.", "success")
    return redirect(url_for("mis_estilos"))


# ==================== DOCUMENTOS TERMINADOS ====================

CARPETA_DOCUMENTOS_TERMINADOS = os.path.join(BASE_PERSISTENT, "documentos_terminados")

@app.route("/documentos-terminados")
@login_required
def documentos_terminados():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("dashboard"))
    
    check_and_send_notifications(tenant.id)
    
    documentos = FinishedDocument.query.filter_by(
        tenant_id=tenant.id, 
        user_id=current_user.id
    ).order_by(FinishedDocument.created_at.desc()).all()
    
    casos = Case.query.filter_by(tenant_id=tenant.id).order_by(Case.titulo).all()
    tareas = Task.query.filter_by(tenant_id=tenant.id).filter(
        Task.estado.notin_(['completado', 'cancelado'])
    ).order_by(Task.titulo).all()
    
    return render_template("documentos_terminados.html", documentos=documentos, casos=casos, tareas=tareas)


@app.route("/documentos-terminados/subir", methods=["POST"])
@login_required
def subir_documento_terminado():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes un estudio asociado.", "error")
        return redirect(url_for("documentos_terminados"))
    
    archivo = request.files.get('archivo')
    nombre = request.form.get('nombre', '').strip()
    case_id = request.form.get('case_id', type=int)
    task_id = request.form.get('task_id', type=int)
    descripcion = request.form.get('descripcion', '').strip()
    tipo_documento = request.form.get('tipo_documento', '').strip()
    numero_expediente = request.form.get('numero_expediente', '').strip()
    plazo_entrega_str = request.form.get('plazo_entrega', '').strip()
    
    plazo_entrega = None
    if plazo_entrega_str:
        try:
            plazo_entrega = datetime.strptime(plazo_entrega_str, '%Y-%m-%d')
        except ValueError:
            pass
    
    if not archivo or not archivo.filename:
        flash("Debes seleccionar un archivo.", "error")
        return redirect(url_for("documentos_terminados"))
    
    ext = os.path.splitext(archivo.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        flash(f"Formato de archivo no soportado ({ext}). Use .docx, .pdf o .txt", "error")
        return redirect(url_for("documentos_terminados"))
    
    if not nombre:
        nombre = os.path.splitext(archivo.filename)[0]
    
    tenant_folder = os.path.join(CARPETA_DOCUMENTOS_TERMINADOS, f"tenant_{tenant.id}")
    os.makedirs(tenant_folder, exist_ok=True)
    
    safe_name = secure_filename(archivo.filename)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    archivo_name = f"{timestamp}_{safe_name}"
    archivo_path = os.path.join(tenant_folder, archivo_name)
    archivo.save(archivo_path)
    
    documento = FinishedDocument(
        tenant_id=tenant.id,
        user_id=current_user.id,
        case_id=case_id if case_id else None,
        task_id=task_id if task_id else None,
        nombre=nombre,
        archivo=archivo_path,
        descripcion=descripcion,
        tipo_documento=tipo_documento,
        numero_expediente=numero_expediente if numero_expediente else None,
        plazo_entrega=plazo_entrega
    )
    db.session.add(documento)
    db.session.commit()
    
    if task_id:
        tarea = Task.query.filter_by(id=task_id, tenant_id=tenant.id).first()
        if tarea:
            task_doc = TaskDocument(
                task_id=task_id,
                document_id=documento.id,
                linked_by_id=current_user.id
            )
            db.session.add(task_doc)
            db.session.commit()
    
    flash("Documento subido exitosamente.", "success")
    return redirect(url_for("documentos_terminados"))


@app.route("/documentos-terminados/eliminar/<int:doc_id>", methods=["POST"])
@login_required
def eliminar_documento_terminado(doc_id):
    tenant = get_current_tenant()
    documento = FinishedDocument.query.filter_by(
        id=doc_id, 
        tenant_id=tenant.id, 
        user_id=current_user.id
    ).first()
    
    if not documento:
        flash("No tienes permiso para eliminar este documento.", "error")
        return redirect(url_for("documentos_terminados"))
    
    if documento.archivo and os.path.exists(documento.archivo):
        os.remove(documento.archivo)
    
    db.session.delete(documento)
    db.session.commit()
    flash("Documento eliminado exitosamente.", "success")
    return redirect(url_for("documentos_terminados"))


@app.route("/documentos-terminados/descargar/<int:doc_id>")
@login_required
def descargar_documento_terminado(doc_id):
    tenant = get_current_tenant()
    documento = FinishedDocument.query.filter_by(
        id=doc_id, 
        tenant_id=tenant.id,
        user_id=current_user.id
    ).first()
    
    if not documento or not documento.archivo:
        flash("Documento no encontrado.", "error")
        return redirect(url_for("documentos_terminados"))
    
    if not os.path.exists(documento.archivo):
        flash("El archivo no existe en el servidor.", "error")
        return redirect(url_for("documentos_terminados"))
    
    directory = os.path.dirname(documento.archivo)
    filename = os.path.basename(documento.archivo)
    return send_from_directory(directory, filename, as_attachment=True)


@app.route("/documentos-terminados/editar/<int:doc_id>", methods=["GET", "POST"])
@login_required
def editar_documento_terminado(doc_id):
    tenant = get_current_tenant()
    documento = FinishedDocument.query.filter_by(
        id=doc_id, 
        tenant_id=tenant.id,
        user_id=current_user.id
    ).first()
    
    if not documento:
        flash("No tienes permiso para editar este documento.", "error")
        return redirect(url_for("documentos_terminados"))
    
    casos = Case.query.filter_by(tenant_id=tenant.id).order_by(Case.titulo).all()
    
    contenido_texto = ""
    campos_detectados = []
    if documento.archivo and os.path.exists(documento.archivo):
        contenido_texto = extract_text_from_file(documento.archivo)
        campos_detectados = detect_placeholders_with_context(contenido_texto)
    
    if request.method == "POST":
        action = request.form.get('action', 'save_metadata')
        
        if action == 'save_metadata':
            documento.nombre = request.form.get('nombre', documento.nombre).strip()
            documento.descripcion = request.form.get('descripcion', '').strip()
            documento.tipo_documento = request.form.get('tipo_documento', '').strip()
            documento.numero_expediente = request.form.get('numero_expediente', '').strip() or None
            plazo_str = request.form.get('plazo_entrega', '').strip()
            if plazo_str:
                try:
                    documento.plazo_entrega = datetime.strptime(plazo_str, '%Y-%m-%d')
                except ValueError:
                    pass
            else:
                documento.plazo_entrega = None
            case_id = request.form.get('case_id', type=int)
            documento.case_id = case_id if case_id else None
            db.session.commit()
            flash("Documento actualizado exitosamente.", "success")
            return redirect(url_for("documentos_terminados"))
        
        elif action == 'update_fields':
            field_values = {}
            for key in request.form:
                if key.startswith('field_'):
                    field_name = key.replace('field_', '')
                    field_values[field_name] = request.form[key]
            
            if field_values and documento.archivo and os.path.exists(documento.archivo):
                ext = os.path.splitext(documento.archivo)[1].lower()
                if ext == '.docx':
                    try:
                        doc = Document(documento.archivo)
                        
                        def replace_in_runs(para, search_text, replace_text):
                            """Replace text in runs while preserving formatting."""
                            full_text = para.text
                            if search_text not in full_text:
                                return False
                            
                            runs_text = [(run, run.text) for run in para.runs]
                            combined = ''.join([t for _, t in runs_text])
                            
                            if search_text in combined:
                                start_idx = combined.find(search_text)
                                end_idx = start_idx + len(search_text)
                                
                                char_idx = 0
                                first_run_idx = None
                                first_run_char = None
                                last_run_idx = None
                                last_run_char = None
                                
                                for i, (run, text) in enumerate(runs_text):
                                    run_start = char_idx
                                    run_end = char_idx + len(text)
                                    
                                    if first_run_idx is None and run_end > start_idx:
                                        first_run_idx = i
                                        first_run_char = start_idx - run_start
                                    
                                    if run_end >= end_idx:
                                        last_run_idx = i
                                        last_run_char = end_idx - run_start
                                        break
                                    
                                    char_idx = run_end
                                
                                if first_run_idx is not None and last_run_idx is not None:
                                    if first_run_idx == last_run_idx:
                                        run = para.runs[first_run_idx]
                                        original = run.text
                                        run.text = original[:first_run_char] + replace_text + original[last_run_char:]
                                    else:
                                        first_run = para.runs[first_run_idx]
                                        first_run.text = first_run.text[:first_run_char] + replace_text
                                        
                                        for i in range(first_run_idx + 1, last_run_idx):
                                            para.runs[i].text = ''
                                        
                                        last_run = para.runs[last_run_idx]
                                        last_run.text = last_run.text[last_run_char:]
                                    return True
                            return False
                        
                        for para in doc.paragraphs:
                            for field_name, field_value in field_values.items():
                                if field_value:
                                    for campo in campos_detectados:
                                        if campo['nombre'] == field_name and campo['match_text'] in para.text:
                                            replace_in_runs(para, campo['match_text'], field_value)
                        
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        for field_name, field_value in field_values.items():
                                            if field_value:
                                                for campo in campos_detectados:
                                                    if campo['nombre'] == field_name and campo['match_text'] in para.text:
                                                        replace_in_runs(para, campo['match_text'], field_value)
                        
                        doc.save(documento.archivo)
                        flash("Campos actualizados en el documento.", "success")
                    except Exception as e:
                        logging.error(f"Error updating document fields: {e}")
                        flash("Error al actualizar los campos del documento.", "error")
                elif ext == '.txt':
                    try:
                        with open(documento.archivo, 'r', encoding='utf-8') as f:
                            content = f.read()
                        for field_name, field_value in field_values.items():
                            if field_value:
                                for campo in campos_detectados:
                                    if campo['nombre'] == field_name:
                                        content = content.replace(campo['match_text'], field_value)
                        with open(documento.archivo, 'w', encoding='utf-8') as f:
                            f.write(content)
                        flash("Campos actualizados en el documento.", "success")
                    except Exception as e:
                        logging.error(f"Error updating text document fields: {e}")
                        flash("Error al actualizar los campos del documento.", "error")
                else:
                    flash("La edicion de campos solo esta disponible para archivos .docx y .txt", "warning")
            
            contenido_texto = extract_text_from_file(documento.archivo)
            campos_detectados = detect_placeholders_with_context(contenido_texto)
            contenido_html = ""
            if contenido_texto and campos_detectados:
                contenido_html = generate_highlighted_html(contenido_texto, campos_detectados)
            elif contenido_texto:
                from markupsafe import escape
                contenido_html = str(escape(contenido_texto)).replace('\n', '<br>')
            
            return render_template("editar_documento_terminado.html", 
                                 documento=documento, 
                                 casos=casos,
                                 contenido_texto=contenido_texto,
                                 contenido_html=contenido_html,
                                 campos_detectados=campos_detectados)
    
    contenido_html = ""
    if contenido_texto and campos_detectados:
        contenido_html = generate_highlighted_html(contenido_texto, campos_detectados)
    elif contenido_texto:
        from markupsafe import escape
        contenido_html = str(escape(contenido_texto)).replace('\n', '<br>')
    
    return render_template("editar_documento_terminado.html", 
                         documento=documento, 
                         casos=casos,
                         contenido_texto=contenido_texto,
                         contenido_html=contenido_html,
                         campos_detectados=campos_detectados)


# ==================== ESTADISTICAS INDIVIDUALES ====================

@app.route("/mis-estadisticas")
@login_required
def mis_estadisticas():
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    user_id = current_user.id
    
    from datetime import timedelta
    fecha_inicio_str = request.args.get('fecha_inicio')
    fecha_fin_str = request.args.get('fecha_fin')
    today = datetime.now()
    
    if fecha_inicio_str:
        try:
            fecha_inicio = datetime.strptime(fecha_inicio_str, '%Y-%m-%d')
        except:
            fecha_inicio = today - timedelta(days=30)
    else:
        fecha_inicio = today - timedelta(days=30)
    
    if fecha_fin_str:
        try:
            fecha_fin = datetime.strptime(fecha_fin_str, '%Y-%m-%d')
        except:
            fecha_fin = today
    else:
        fecha_fin = today
    
    assigned_case_ids = db.session.query(CaseAssignment.case_id).filter_by(user_id=user_id).subquery()
    casos_query = Case.query.filter(Case.id.in_(assigned_case_ids))
    
    casos_por_comenzar = casos_query.filter_by(estado='por_comenzar').count()
    casos_en_proceso = casos_query.filter(Case.estado.in_(['en_proceso', 'en_espera'])).count()
    casos_terminados = casos_query.filter_by(estado='terminado').count()
    total_casos = casos_query.count()
    
    docs_query = DocumentRecord.query.filter_by(user_id=user_id, tenant_id=tenant_id)
    total_docs = docs_query.count()
    docs_periodo = docs_query.filter(DocumentRecord.fecha >= fecha_inicio, DocumentRecord.fecha <= fecha_fin).count()
    
    tareas_completadas = Task.query.filter_by(assigned_to_id=user_id, estado='completado').count()
    tareas_pendientes = Task.query.filter_by(assigned_to_id=user_id, estado='pendiente').count()
    tareas_en_curso = Task.query.filter_by(assigned_to_id=user_id, estado='en_curso').count()
    tareas_vencidas = Task.query.filter(
        Task.assigned_to_id == user_id,
        Task.estado.notin_(['completado', 'cancelado']),
        Task.fecha_vencimiento.isnot(None),
        Task.fecha_vencimiento < today
    ).count()
    
    total_tareas_con_fecha = Task.query.filter(
        Task.assigned_to_id == user_id,
        Task.fecha_vencimiento.isnot(None),
        Task.estado == 'completado'
    ).count()
    tareas_a_tiempo = Task.query.filter(
        Task.assigned_to_id == user_id,
        Task.estado == 'completado',
        Task.fecha_completada.isnot(None),
        Task.fecha_vencimiento.isnot(None),
        Task.fecha_completada <= Task.fecha_vencimiento
    ).count()
    cumplimiento = round((tareas_a_tiempo / total_tareas_con_fecha) * 100, 1) if total_tareas_con_fecha > 0 else 100
    
    docs_recientes = docs_query.order_by(DocumentRecord.fecha.desc()).limit(10).all()
    tareas_recientes = Task.query.filter_by(assigned_to_id=user_id).order_by(Task.updated_at.desc()).limit(10).all()
    
    stats = {
        'casos_por_comenzar': casos_por_comenzar,
        'casos_en_proceso': casos_en_proceso,
        'casos_terminados': casos_terminados,
        'total_casos': total_casos,
        'total_docs': total_docs,
        'docs_periodo': docs_periodo,
        'tareas_completadas': tareas_completadas,
        'tareas_pendientes': tareas_pendientes,
        'tareas_en_curso': tareas_en_curso,
        'tareas_vencidas': tareas_vencidas,
        'cumplimiento': cumplimiento
    }
    
    return render_template("mis_estadisticas.html",
                          stats=stats,
                          docs_recientes=docs_recientes,
                          tareas_recientes=tareas_recientes,
                          fecha_inicio=fecha_inicio.strftime('%Y-%m-%d'),
                          fecha_fin=fecha_fin.strftime('%Y-%m-%d'))


# ==================== ESTADISTICAS EQUIPO (Admin/Coordinador) ====================

@app.route("/estadisticas-equipo")
@login_required
def estadisticas_equipo():
    if not current_user.can_manage_cases() and not current_user.is_admin_estudio():
        flash("No tienes acceso a esta seccion.", "error")
        return redirect(url_for("dashboard"))
    
    tenant = get_current_tenant()
    tenant_id = tenant.id if tenant else None
    
    from datetime import timedelta
    today = datetime.now()
    month_ago = today - timedelta(days=30)
    
    usuarios = User.query.filter_by(tenant_id=tenant_id, activo=True).all()
    
    user_stats = []
    for usuario in usuarios:
        assigned_ids = [a.case_id for a in usuario.case_assignments.all()]
        casos_total = len(assigned_ids)
        casos_activos = Case.query.filter(Case.id.in_(assigned_ids), Case.estado.in_(['en_proceso', 'en_espera'])).count() if assigned_ids else 0
        
        docs_total = DocumentRecord.query.filter_by(user_id=usuario.id, tenant_id=tenant_id).count()
        docs_mes = DocumentRecord.query.filter(
            DocumentRecord.user_id == usuario.id,
            DocumentRecord.tenant_id == tenant_id,
            DocumentRecord.fecha >= month_ago
        ).count()
        
        tareas_completadas = Task.query.filter_by(assigned_to_id=usuario.id, estado='completado').count()
        tareas_pendientes = Task.query.filter_by(assigned_to_id=usuario.id, estado='pendiente').count()
        tareas_vencidas = Task.query.filter(
            Task.assigned_to_id == usuario.id,
            Task.estado.notin_(['completado', 'cancelado']),
            Task.fecha_vencimiento.isnot(None),
            Task.fecha_vencimiento < today
        ).count()
        
        total_con_fecha = Task.query.filter(
            Task.assigned_to_id == usuario.id,
            Task.fecha_vencimiento.isnot(None),
            Task.estado == 'completado'
        ).count()
        a_tiempo = Task.query.filter(
            Task.assigned_to_id == usuario.id,
            Task.estado == 'completado',
            Task.fecha_completada.isnot(None),
            Task.fecha_vencimiento.isnot(None),
            Task.fecha_completada <= Task.fecha_vencimiento
        ).count()
        cumplimiento = round((a_tiempo / total_con_fecha) * 100) if total_con_fecha > 0 else 100
        
        user_stats.append({
            'usuario': usuario,
            'casos_total': casos_total,
            'casos_activos': casos_activos,
            'docs_total': docs_total,
            'docs_mes': docs_mes,
            'tareas_completadas': tareas_completadas,
            'tareas_pendientes': tareas_pendientes,
            'tareas_vencidas': tareas_vencidas,
            'cumplimiento': cumplimiento
        })
    
    totales = {
        'casos': sum(s['casos_total'] for s in user_stats),
        'docs': sum(s['docs_total'] for s in user_stats),
        'docs_mes': sum(s['docs_mes'] for s in user_stats),
        'tareas_completadas': sum(s['tareas_completadas'] for s in user_stats),
        'tareas_pendientes': sum(s['tareas_pendientes'] for s in user_stats)
    }
    
    return render_template("estadisticas_equipo.html", user_stats=user_stats, totales=totales)


# ==================== AI ASSISTANT API ====================

@app.route("/api/ai-assistant", methods=["POST"])
@login_required
def ai_assistant_api():
    data = request.get_json()
    message = data.get('message', '')
    
    if not message:
        return jsonify({'response': 'Por favor escribe tu pregunta.'})
    
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": """Eres una guia de ayuda amigable para una plataforma de gestion legal. Tu trabajo es explicar paso a paso como usar la plataforma.

ESTILO DE COMUNICACION:
- Habla de forma muy simple y amigable, como si fueras un amigo explicando
- Usa emojis ocasionalmente para ser mas cercano
- Da instrucciones paso a paso muy claras
- Menciona exactamente donde hacer clic y que botones presionar
- Da recomendaciones utiles

FUNCIONES DE LA PLATAFORMA QUE CONOCES:

1. CASOS:
   - Para crear un caso: Menu lateral izquierdo > "Casos" > boton "Nuevo Caso"
   - Llenar: nombre del caso, cliente, descripcion, fecha
   - Los casos agrupan todos los documentos y tareas de un cliente

2. DOCUMENTOS:
   - Para generar documento: Menu "Generar Documento" o desde un caso
   - Seleccionar modelo/plantilla, llenar los campos, presionar "Generar"
   - Los documentos se guardan automaticamente

3. MIS MODELOS:
   - Crear modelos personalizados en "Mis Modelos" > "Nuevo Modelo"
   - Subir un documento Word como plantilla
   - Agregar campos dinamicos que se llenaran automaticamente

4. ANONIMIZAR DOCUMENTOS:
   - En "Mis Modelos" > boton "Anonimiza tu documento"
   - Subir un documento y el sistema quita datos sensibles automaticamente

5. HISTORIAL:
   - Ver todos los documentos generados en "Historial"
   - Filtrar por fecha, tipo de documento, etc.

6. ADMINISTRACION (solo admins):
   - Agregar usuarios: "Admin" > "Usuarios" > "Nuevo Usuario"
   - Configurar estudio: "Admin" > "Configurar Estudio"

7. TAREAS:
   - Crear tareas desde un caso o menu "Tareas"
   - Asignar a usuarios, poner fechas limite

Siempre responde en espanol de forma clara y amigable."""},
                {"role": "user", "content": message}
            ],
            temperature=0.7,
            max_tokens=600
        )
        ai_response = response.choices[0].message.content
        return jsonify({'response': ai_response})
    except Exception as e:
        logging.error(f"Error AI Assistant: {e}")
        return jsonify({'response': 'Ups! Algo salio mal. Intenta de nuevo por favor.'})


# ==================== DOCUMENT ANONYMIZER ====================

import uuid

CARPETA_ANONIMIZADOS = os.path.join(BASE_PERSISTENT, "documentos_anonimizados")


def get_anonimizados_folder(tenant_id, user_id):
    """Get tenant-scoped and user-scoped folder for anonymized documents."""
    folder = os.path.join(CARPETA_ANONIMIZADOS, f"tenant_{tenant_id}", f"user_{user_id}")
    os.makedirs(folder, exist_ok=True)
    return folder


def anonimizar_con_ia(texto):
    """Use OpenAI to identify and anonymize sensitive information in legal documents."""
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        
        prompt = """Analiza el siguiente documento legal y devuelve un JSON con dos claves:
        
1. "categorias": Un diccionario donde cada clave es una categoria de informacion sensible y el valor es una lista de los elementos encontrados. Las categorias deben ser:
   - "Nombres de personas": nombres completos de personas
   - "DNI/Documentos de identidad": numeros de DNI, RUC, pasaportes
   - "Direcciones": direcciones fisicas, domicilios
   - "Telefonos": numeros de telefono
   - "Montos de dinero": cantidades monetarias (S/, USD, etc)
   - "Fechas especificas": fechas exactas mencionadas
   - "Numeros de expediente": numeros de casos o expedientes
   - "Correos electronicos": direcciones de email
   - "Entidades/Instituciones": nombres de empresas, juzgados, etc.
   - "Otros datos sensibles": cualquier otra informacion identificable

2. "texto_anonimizado": El texto completo del documento con toda la informacion sensible reemplazada por "........................" (24 puntos).

IMPORTANTE:
- Solo incluye categorias que tengan elementos encontrados
- Mantiene la estructura y formato del documento original
- Reemplaza CADA ocurrencia de informacion sensible

DOCUMENTO A ANALIZAR:
"""
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Eres un experto en proteccion de datos personales y anonimizacion de documentos legales. Respondes SOLO con JSON valido, sin texto adicional."},
                {"role": "user", "content": prompt + texto}
            ],
            temperature=0.3,
            max_tokens=4000,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        return result
    except Exception as e:
        logging.error(f"Error anonimizando documento: {e}")
        return None


def guardar_docx_anonimizado(texto, nombre_archivo, tenant_id, user_id):
    """Save anonymized text as Word document in tenant-scoped folder."""
    folder = get_anonimizados_folder(tenant_id, user_id)
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    for para_text in texto.split('\n'):
        if para_text.strip():
            p = doc.add_paragraph(para_text)
            p.paragraph_format.line_spacing = 1.5
        else:
            doc.add_paragraph()
    
    file_path = os.path.join(folder, nombre_archivo)
    doc.save(file_path)
    return file_path


@app.route("/anonimizar-documento", methods=["GET", "POST"])
@login_required
def anonimizar_documento():
    if request.method == "GET":
        return render_template("anonimizar_documento.html")
    
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes acceso a esta funcion.", "error")
        return redirect(url_for('index'))
    
    if 'documento' not in request.files:
        flash("Por favor selecciona un documento.", "error")
        return redirect(url_for('anonimizar_documento'))
    
    archivo = request.files['documento']
    if archivo.filename == '':
        flash("No se selecciono ningun archivo.", "error")
        return redirect(url_for('anonimizar_documento'))
    
    ext = os.path.splitext(archivo.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        flash("Formato no permitido. Usa .docx, .pdf o .txt", "error")
        return redirect(url_for('anonimizar_documento'))
    
    temp_dir = "temp_uploads"
    os.makedirs(temp_dir, exist_ok=True)
    unique_id = uuid.uuid4().hex[:8]
    filename = secure_filename(archivo.filename)
    temp_filename = f"{unique_id}_{filename}"
    temp_path = os.path.join(temp_dir, temp_filename)
    archivo.save(temp_path)
    
    try:
        texto_original = extract_text_from_file(temp_path)
        
        if not texto_original or len(texto_original.strip()) < 50:
            flash("El documento esta vacio o no se pudo extraer el texto.", "error")
            os.remove(temp_path)
            return redirect(url_for('anonimizar_documento'))
        
        resultado = anonimizar_con_ia(texto_original)
        
        if not resultado:
            flash("Error al procesar el documento. Intenta de nuevo.", "error")
            os.remove(temp_path)
            return redirect(url_for('anonimizar_documento'))
        
        categorias = resultado.get('categorias', {})
        texto_anonimizado = resultado.get('texto_anonimizado', texto_original)
        
        categorias_no_vacias = {k: v for k, v in categorias.items() if v}
        
        total_anonimizado = sum(len(items) for items in categorias_no_vacias.values())
        
        random_suffix = uuid.uuid4().hex[:12]
        nombre_base = os.path.splitext(filename)[0][:20]
        nombre_archivo = f"ANON_{random_suffix}.docx"
        
        guardar_docx_anonimizado(texto_anonimizado, nombre_archivo, tenant.id, current_user.id)
        
        os.remove(temp_path)
        
        return render_template("anonimizar_documento.html",
                             resultado=True,
                             categorias=categorias_no_vacias,
                             texto_anonimizado=texto_anonimizado,
                             total_anonimizado=total_anonimizado,
                             nombre_archivo=nombre_archivo)
    
    except Exception as e:
        logging.error(f"Error en anonimizacion: {e}")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        flash("Error al procesar el documento.", "error")
        return redirect(url_for('anonimizar_documento'))


@app.route("/descargar-anonimizado/<nombre>")
@login_required
def descargar_anonimizado(nombre):
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes acceso a este archivo.", "error")
        return redirect(url_for('index'))
    
    folder = get_anonimizados_folder(tenant.id, current_user.id)
    file_path = os.path.join(folder, nombre)
    
    if not os.path.exists(file_path):
        flash("Archivo no encontrado.", "error")
        return redirect(url_for('anonimizar_documento'))
    
    return send_from_directory(folder, nombre, as_attachment=True)


# ==================== BUSQUEDA POR SIMILITUD ====================

@app.route("/buscar-similares")
@login_required
def buscar_similares():
    """Ruta deshabilitada."""
    abort(404)


@app.route("/buscar-similares/consultar", methods=["POST"])
@login_required
def buscar_similares_consultar():
    """Ruta deshabilitada."""
    abort(404)


def _buscar_similares_consultar_disabled():
    tenant = get_current_tenant()
    if not tenant:
        return jsonify({"error": "No tienes acceso"}), 403
    
    consulta = request.form.get("consulta", "").strip()
    if not consulta:
        return jsonify({"error": "Escribe una consulta"}), 400
    
    # Obtener documentos del tenant para buscar similitudes
    documentos = FinishedDocument.query.filter_by(
        tenant_id=tenant.id
    ).order_by(FinishedDocument.created_at.desc()).limit(100).all()
    
    if not documentos:
        return jsonify({
            "respuesta": "No hay documentos guardados todavía. Genera o sube documentos para poder buscar similares.",
            "documentos": []
        })
    
    # Preparar resumen de documentos para la IA
    docs_info = []
    for doc in documentos:
        info = {
            "id": doc.id,
            "nombre": doc.nombre_documento or doc.nombre or "Sin nombre",
            "tipo": doc.tipo_documento or "No especificado",
            "fecha": doc.created_at.strftime("%d/%m/%Y") if doc.created_at else "",
            "numero_expediente": doc.numero_expediente or "",
            "descripcion": (doc.descripcion or "")[:200]
        }
        # Añadir fragmento del contenido si está disponible
        if doc.contenido_texto:
            info["contenido_fragmento"] = doc.contenido_texto[:500]
        docs_info.append(info)
    
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        
        system_prompt = """Eres un asistente legal experto que ayuda a encontrar documentos similares. 
El usuario te hará consultas como "muéstrame casos parecidos a demandas de alimentos" o "busco documentos sobre divorcio".

Tu tarea es:
1. Analizar la consulta del usuario
2. Buscar en la lista de documentos proporcionada aquellos que sean similares o relevantes
3. Responder de forma amigable explicando qué encontraste
4. Devolver los IDs de los documentos relevantes

Responde SIEMPRE en formato JSON con esta estructura:
{
    "respuesta": "Texto explicativo para el usuario sobre lo que encontraste",
    "documentos_ids": [lista de IDs de documentos relevantes, máximo 10]
}

Si no encuentras documentos relevantes, explica por qué y sugiere qué tipo de documentos podría buscar."""

        docs_json = json.dumps(docs_info, ensure_ascii=False)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Documentos disponibles:\n{docs_json}\n\nConsulta del usuario: {consulta}"}
            ],
            temperature=0.3,
            max_tokens=1000,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        
        # Obtener documentos encontrados
        docs_ids = result.get("documentos_ids", [])
        docs_encontrados = []
        for doc_id in docs_ids:
            doc = FinishedDocument.query.filter_by(id=doc_id, tenant_id=tenant.id).first()
            if doc:
                docs_encontrados.append({
                    "id": doc.id,
                    "nombre": doc.nombre_documento or doc.nombre or "Sin nombre",
                    "tipo": doc.tipo_documento or "No especificado",
                    "fecha": doc.created_at.strftime("%d/%m/%Y") if doc.created_at else "",
                    "numero_expediente": doc.numero_expediente or "",
                    "tiene_archivo": bool(doc.archivo_path)
                })
        
        return jsonify({
            "respuesta": result.get("respuesta", ""),
            "documentos": docs_encontrados
        })
        
    except Exception as e:
        logging.error(f"Error en búsqueda de similares: {e}")
        return jsonify({
            "respuesta": "Hubo un error al procesar tu consulta. Por favor intenta de nuevo.",
            "documentos": []
        })


# ==================== REVISOR IA ====================

CARPETA_REVISIONES = os.path.join(BASE_PERSISTENT, "revisiones_temp")

def get_revisiones_folder(tenant_id):
    folder = os.path.join(CARPETA_REVISIONES, f"tenant_{tenant_id}")
    os.makedirs(folder, exist_ok=True)
    return folder


def revisar_documento_con_ia(texto_documento, nombre_documento):
    """Analiza un documento legal con IA para detectar errores y problemas."""
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        
        system_prompt = """Eres un revisor legal experto peruano especializado en documentos jurídicos. Tu tarea es revisar documentos legales y detectar problemas en las siguientes categorías:

1. COHERENCIA JURÍDICA: Verifica que los fundamentos legales sean correctos, que las normas citadas existan y apliquen al caso, y que los argumentos sean lógicamente consistentes.

2. CONTRADICCIONES: Detecta inconsistencias en los datos (nombres, fechas, montos, DNI que no coinciden a lo largo del documento), afirmaciones contradictorias, o hechos que se contradicen entre sí.

3. ESTRUCTURA: Verifica que el documento tenga la estructura correcta para su tipo (demanda, contestación, recurso, etc.), que contenga todas las secciones necesarias y en el orden correcto.

4. CAMPOS INCOMPLETOS: Identifica campos vacíos, placeholders no reemplazados (como {{nombre}}, [COMPLETAR], XXX), datos faltantes obligatorios, o información incompleta.

Responde ÚNICAMENTE en formato JSON válido con la siguiente estructura:
{
    "evaluacion_general": "Breve resumen del estado general del documento (1-2 oraciones)",
    "issues": [
        {
            "severidad": "error|advertencia|sugerencia",
            "tipo": "coherencia|contradiccion|estructura|campo_incompleto",
            "ubicacion": "Sección o parte del documento donde se encuentra el problema",
            "fragmento": "Texto exacto donde se detectó el problema (máximo 100 caracteres)",
            "descripcion": "Descripción clara del problema encontrado",
            "recomendacion": "Cómo corregir el problema"
        }
    ]
}

Severidades:
- "error": Problemas graves que deben corregirse obligatoriamente (contradicciones de datos, campos vacíos críticos)
- "advertencia": Problemas importantes que deberían revisarse (posibles inconsistencias, estructura mejorable)
- "sugerencia": Mejoras opcionales o recomendaciones de estilo

Si el documento está bien, devuelve issues como array vacío."""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Revisa el siguiente documento legal llamado '{nombre_documento}':\n\n{texto_documento[:15000]}"}
            ],
            temperature=0.3,
            max_tokens=4000,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        return result
        
    except Exception as e:
        logging.error(f"Error en revisión IA: {e}")
        return None


@app.route("/revisor-ia")
@login_required
def revisor_ia():
    """Ruta deshabilitada."""
    abort(404)


@app.route("/revisor-ia/analizar", methods=["POST"])
@login_required
def revisor_ia_analizar():
    """Ruta deshabilitada."""
    abort(404)


def _revisor_ia_analizar_disabled():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes acceso a esta función.", "error")
        return redirect(url_for('index'))
    
    texto_documento = None
    nombre_documento = "Documento"
    archivo_path = None
    
    # Opción 1: Subir archivo
    archivo = request.files.get("archivo")
    if archivo and archivo.filename:
        filename = secure_filename(archivo.filename)
        ext = os.path.splitext(filename)[1].lower()
        
        if ext not in ['.docx', '.txt', '.pdf']:
            flash("Formato no soportado. Use archivos .docx, .txt o .pdf", "error")
            return redirect(url_for('revisor_ia'))
        
        folder = get_revisiones_folder(tenant.id)
        unique_name = f"{uuid.uuid4().hex}_{filename}"
        temp_path = os.path.join(folder, unique_name)
        archivo.save(temp_path)
        archivo_path = temp_path
        nombre_documento = filename
        
        # Extraer texto
        if ext == '.docx':
            try:
                doc = Document(temp_path)
                texto_documento = "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                logging.error(f"Error leyendo docx: {e}")
                flash("Error al leer el archivo Word.", "error")
                return redirect(url_for('revisor_ia'))
        elif ext == '.txt':
            with open(temp_path, 'r', encoding='utf-8') as f:
                texto_documento = f.read()
        elif ext == '.pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(temp_path)
                texto_documento = "\n".join([page.extract_text() or "" for page in reader.pages])
            except Exception as e:
                logging.error(f"Error leyendo PDF: {e}")
                flash("Error al leer el archivo PDF.", "error")
                return redirect(url_for('revisor_ia'))
    
    # Opción 2: Seleccionar documento existente
    documento_id = request.form.get("documento_id", type=int)
    if not texto_documento and documento_id:
        doc = FinishedDocument.query.filter_by(
            id=documento_id,
            tenant_id=tenant.id
        ).first()
        if doc:
            nombre_documento = doc.nombre_documento
            # Leer el contenido del archivo
            if doc.archivo_path and os.path.exists(doc.archivo_path):
                try:
                    docx = Document(doc.archivo_path)
                    texto_documento = "\n".join([p.text for p in docx.paragraphs])
                except:
                    pass
            if not texto_documento and doc.contenido_texto:
                texto_documento = doc.contenido_texto
    
    # Opción 3: Texto pegado directamente
    texto_directo = request.form.get("texto_documento", "").strip()
    if not texto_documento and texto_directo:
        texto_documento = texto_directo
        nombre_documento = "Texto pegado"
    
    if not texto_documento:
        flash("No se proporcionó ningún documento para revisar.", "error")
        return redirect(url_for('revisor_ia'))
    
    # Crear sesión de revisión
    review = ReviewSession(
        tenant_id=tenant.id,
        user_id=current_user.id,
        nombre_documento=nombre_documento,
        archivo_path=archivo_path,
        contenido_texto=texto_documento[:50000],
        estado='procesando'
    )
    db.session.add(review)
    db.session.commit()
    
    # Analizar con IA
    resultado = revisar_documento_con_ia(texto_documento, nombre_documento)
    
    if not resultado:
        review.estado = 'error'
        db.session.commit()
        flash("Error al analizar el documento. Intenta de nuevo.", "error")
        return redirect(url_for('revisor_ia'))
    
    # Guardar resultados
    review.evaluacion_general = resultado.get('evaluacion_general', '')
    review.estado = 'completado'
    review.completed_at = datetime.utcnow()
    
    errores = 0
    advertencias = 0
    sugerencias = 0
    
    for i, issue in enumerate(resultado.get('issues', [])):
        severidad = issue.get('severidad', 'sugerencia')
        if severidad == 'error':
            errores += 1
        elif severidad == 'advertencia':
            advertencias += 1
        else:
            sugerencias += 1
        
        review_issue = ReviewIssue(
            session_id=review.id,
            severidad=severidad,
            tipo=issue.get('tipo', 'otro'),
            ubicacion=issue.get('ubicacion', ''),
            fragmento=issue.get('fragmento', ''),
            descripcion=issue.get('descripcion', ''),
            recomendacion=issue.get('recomendacion', ''),
            orden=i
        )
        db.session.add(review_issue)
    
    review.total_errores = errores
    review.total_advertencias = advertencias
    review.total_sugerencias = sugerencias
    db.session.commit()
    
    return redirect(url_for('revisor_ia_resultado', review_id=review.id))


@app.route("/revisor-ia/resultado/<int:review_id>")
@login_required
def revisor_ia_resultado(review_id):
    """Ruta deshabilitada."""
    abort(404)


def _revisor_ia_resultado_disabled(review_id):
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes acceso a esta función.", "error")
        return redirect(url_for('index'))
    
    # Query by tenant and user for security
    query = ReviewSession.query.filter_by(id=review_id, tenant_id=tenant.id)
    
    # Regular users can only see their own reviews
    if not (current_user.is_admin_estudio() or current_user.is_super_admin() or current_user.is_coordinador()):
        query = query.filter_by(user_id=current_user.id)
    
    review = query.first_or_404()
    
    issues = ReviewIssue.query.filter_by(session_id=review.id).order_by(
        db.case(
            (ReviewIssue.severidad == 'error', 1),
            (ReviewIssue.severidad == 'advertencia', 2),
            else_=3
        ),
        ReviewIssue.orden
    ).all()
    
    return render_template("revisor_ia_resultado.html",
                          review=review,
                          issues=issues)


@app.route("/revisor-ia/historial")
@login_required
def revisor_ia_historial():
    """Ruta deshabilitada."""
    abort(404)


def send_task_reminders():
    """Send email reminders for tasks due in 1, 2, or 3 days."""
    from datetime import date, timedelta
    
    today = date.today()
    reminder_days = [1, 2, 3]
    
    with app.app_context():
        for days_before in reminder_days:
            target_date = today + timedelta(days=days_before)
            
            tasks = Task.query.filter(
                Task.estado.notin_(['completado', 'cancelado']),
                Task.fecha_vencimiento.isnot(None),
                db.func.date(Task.fecha_vencimiento) == target_date
            ).all()
            
            for task in tasks:
                if not task.assigned_to or not task.assigned_to.email:
                    continue
                
                existing_reminder = TaskReminder.query.filter_by(
                    task_id=task.id,
                    days_before=days_before
                ).first()
                
                if existing_reminder:
                    continue
                
                reminder_type = 'urgente' if days_before == 1 else ('proximo' if days_before == 2 else 'recordatorio')
                
                subject = f"{'⚠️ URGENTE: ' if days_before == 1 else ''}Recordatorio de tarea - {task.titulo}"
                
                html_content = f"""
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                    <h2 style="color: {'#dc2626' if days_before == 1 else '#f59e0b' if days_before == 2 else '#3b82f6'};">
                        {'⚠️ ' if days_before == 1 else ''}Recordatorio de Tarea
                    </h2>
                    <p>Hola {task.assigned_to.username},</p>
                    <p>Te recordamos que tienes una tarea que vence en <strong>{days_before} día{'s' if days_before > 1 else ''}</strong>:</p>
                    <div style="background: #f3f4f6; padding: 15px; border-radius: 8px; margin: 20px 0;">
                        <h3 style="margin: 0 0 10px 0;">{task.titulo}</h3>
                        <p style="margin: 5px 0; color: #666;">
                            <strong>Fecha de vencimiento:</strong> {task.fecha_vencimiento.strftime('%d/%m/%Y')}
                        </p>
                        {f'<p style="margin: 5px 0; color: #666;"><strong>Caso:</strong> {task.case.titulo}</p>' if task.case else ''}
                        {f'<p style="margin: 5px 0; color: #666;">{task.descripcion[:200]}...</p>' if task.descripcion and len(task.descripcion) > 0 else ''}
                    </div>
                    <p>Por favor, asegúrate de completar esta tarea a tiempo.</p>
                    <p style="color: #888; font-size: 12px;">Este es un recordatorio automático.</p>
                </div>
                """
                
                try:
                    success = send_notification_email(
                        task.assigned_to.email,
                        subject,
                        html_content
                    )
                    
                    reminder = TaskReminder(
                        task_id=task.id,
                        user_id=task.assigned_to.id,
                        reminder_type=reminder_type,
                        days_before=days_before,
                        email_sent=success
                    )
                    db.session.add(reminder)
                    db.session.commit()
                    
                    if success:
                        logging.info(f"Reminder sent for task {task.id} to {task.assigned_to.email}")
                    else:
                        logging.warning(f"Failed to send reminder for task {task.id}")
                except Exception as e:
                    logging.error(f"Error sending reminder for task {task.id}: {e}")


@app.route("/api/send-task-reminders", methods=["POST"])
def api_send_task_reminders():
    """API endpoint to trigger task reminders (for cron jobs)."""
    auth_key = request.headers.get('X-API-Key')
    expected_key = os.environ.get('REMINDER_API_KEY', os.environ.get('SESSION_SECRET'))
    
    if auth_key != expected_key:
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        send_task_reminders()
        return jsonify({"status": "success", "message": "Reminders processed"})
    except Exception as e:
        logging.error(f"Error in reminder API: {e}")
        return jsonify({"error": str(e)}), 500


CARPETA_ARGUMENTACION = os.path.join(BASE_PERSISTENT, "argumentaciones")

def get_argumentacion_folder(tenant_id):
    folder = os.path.join(CARPETA_ARGUMENTACION, f"tenant_{tenant_id}")
    os.makedirs(folder, exist_ok=True)
    return folder


@app.route("/argumentacion")
@login_required
def argumentacion():
    """Ruta deshabilitada."""
    abort(404)


@app.route("/argumentacion/sesion/<int:session_id>")
@login_required
def argumentacion_sesion(session_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/argumentacion/nueva", methods=["POST"])
@login_required
def argumentacion_nueva():
    """Ruta deshabilitada."""
    abort(404)


def _argumentacion_nueva_disabled():
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes acceso a esta función.", "error")
        return redirect(url_for('index'))
    
    texto_documento = None
    archivo_nombre = None
    archivo_tipo = None
    case_id = request.form.get("case_id", type=int)
    
    archivo = request.files.get("archivo")
    if archivo and archivo.filename:
        filename = secure_filename(archivo.filename)
        ext = os.path.splitext(filename)[1].lower()
        
        if ext not in ['.docx', '.doc', '.txt', '.pdf']:
            flash("Formato no soportado. Use archivos .docx, .txt o .pdf", "error")
            return redirect(url_for('argumentacion'))
        
        archivo_nombre = filename
        archivo_tipo = ext
        
        if ext == '.docx' or ext == '.doc':
            try:
                archivo.seek(0)
                doc = Document(archivo)
                texto_documento = "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                logging.error(f"Error leyendo docx: {e}")
                flash("Error al leer el archivo Word.", "error")
                return redirect(url_for('argumentacion'))
        elif ext == '.txt':
            archivo.seek(0)
            texto_documento = archivo.read().decode('utf-8')
        elif ext == '.pdf':
            try:
                from PyPDF2 import PdfReader
                archivo.seek(0)
                reader = PdfReader(archivo)
                texto_documento = "\n".join([page.extract_text() or "" for page in reader.pages])
            except Exception as e:
                logging.error(f"Error leyendo PDF: {e}")
                flash("Error al leer el archivo PDF.", "error")
                return redirect(url_for('argumentacion'))
    
    texto_directo = request.form.get("texto_documento", "").strip()
    if not texto_documento and texto_directo:
        texto_documento = texto_directo
        archivo_nombre = "Texto directo"
        archivo_tipo = "text"
    
    if not texto_documento:
        flash("No se proporcionó ningún documento.", "error")
        return redirect(url_for('argumentacion'))
    
    titulo = archivo_nombre or f"Sesión {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    
    sesion = ArgumentationSession(
        user_id=current_user.id,
        tenant_id=tenant.id,
        case_id=case_id if case_id else None,
        titulo=titulo,
        documento_original=texto_documento,
        archivo_nombre=archivo_nombre,
        archivo_tipo=archivo_tipo
    )
    db.session.add(sesion)
    db.session.commit()
    
    flash("Documento cargado correctamente.", "success")
    return redirect(url_for('argumentacion_sesion', session_id=sesion.id))


@app.route("/argumentacion/mejorar/<int:session_id>", methods=["POST"])
@login_required
def argumentacion_mejorar(session_id):
    """Ruta deshabilitada."""
    abort(404)


def _argumentacion_mejorar_disabled(session_id):
    tenant = get_current_tenant()
    if not tenant:
        return jsonify({"error": "No autorizado"}), 403
    
    sesion = ArgumentationSession.query.filter_by(
        id=session_id,
        user_id=current_user.id,
        tenant_id=tenant.id
    ).first()
    
    if not sesion:
        return jsonify({"error": "Sesión no encontrada"}), 404
    
    instrucciones = request.form.get("instrucciones", "").strip()
    estilo = request.form.get("estilo", "Formal clásico")
    
    if not instrucciones:
        flash("Por favor, indica qué tipo de mejora deseas.", "error")
        return redirect(url_for('argumentacion_sesion', session_id=session_id))
    
    mensaje_usuario = ArgumentationMessage(
        session_id=sesion.id,
        role="user",
        content=instrucciones,
        estilo_aplicado=estilo
    )
    db.session.add(mensaje_usuario)
    
    documento_actual = sesion.ultima_version_mejorada or sesion.documento_original
    
    estilo_instrucciones = ""
    for e in UserArgumentationStyle.ESTILOS_PREDEFINIDOS:
        if e['nombre'] == estilo:
            estilo_instrucciones = e['instrucciones']
            break
    
    if not estilo_instrucciones:
        estilo_custom = UserArgumentationStyle.query.filter_by(
            user_id=current_user.id,
            nombre=estilo,
            activo=True
        ).first()
        if estilo_custom:
            estilo_instrucciones = estilo_custom.instrucciones
    
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=150.0)
        
        system_prompt = f"""Actua como un asistente juridico especializado en redaccion y argumentacion.

Tu tarea es modificar directamente el texto del documento juridico, aplicando EXACTAMENTE las instrucciones del usuario, de forma rapida, directa y sin rodeos.

REGLAS ESTRICTAS:
1. Manten intactos todos los datos facticos (nombres, DNIs, fechas, montos, direcciones, numeros de expediente, numeros de cuenta, porcentajes, acuerdos economicos)
2. Si encuentras incoherencias factuales, solo senala el error; no inventes datos nuevos
3. Puedes anadir parrafos completos si el usuario lo pide
4. Puedes eliminar fragmentos si el usuario lo pide
5. Puedes reorganizar la logica argumentativa si ayuda a la claridad
6. Respeta la estructura general (Hechos - Fundamentos - Petitorio)
7. No inventes hechos ni articulos falsos
8. Aplica el estilo solicitado: {estilo}
9. {estilo_instrucciones}

INSTRUCCIONES DEL USUARIO:
{instrucciones}

Devuelve SIEMPRE el documento modificado completo, listo para copiar, sin comentarios meta, solo el contenido final."""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Documento a mejorar:\n\n{documento_actual[:20000]}"}
            ],
            temperature=0.4,
            max_tokens=8000
        )
        
        resultado = response.choices[0].message.content
        
        mensaje_ia = ArgumentationMessage(
            session_id=sesion.id,
            role="assistant",
            content=resultado,
            estilo_aplicado=estilo
        )
        db.session.add(mensaje_ia)
        
        sesion.ultima_version_mejorada = resultado
        sesion.estilo_usado = estilo
        sesion.updated_at = datetime.utcnow()
        db.session.commit()
        
        flash("Argumentación mejorada correctamente.", "success")
        
    except Exception as e:
        logging.error(f"Error mejorando argumentación: {e}")
        error_msg = str(e).lower()
        if 'timeout' in error_msg or 'timed out' in error_msg:
            flash("El documento es muy extenso y la mejora tarda más de lo esperado. Intenta con instrucciones más específicas o un fragmento más corto.", "warning")
        else:
            flash("Error al procesar la mejora. Intenta nuevamente.", "error")
    
    return redirect(url_for('argumentacion_sesion', session_id=session_id))


@app.route("/argumentacion/start", methods=["POST"])
@login_required
def argumentacion_start_job():
    """Ruta deshabilitada."""
    abort(404)


def _argumentacion_start_job_disabled():
    """Inicia un job asíncrono de argumentación."""
    start_argumentation_worker()
    
    tenant = get_current_tenant()
    if not tenant:
        return jsonify({"success": False, "error": "No autorizado"}), 403
    
    data = request.get_json() or {}
    session_id = data.get('session_id')
    instrucciones = data.get('instrucciones', '').strip()
    estilo = data.get('estilo', 'Formal clásico')
    section = data.get('section', 'full')
    
    if not session_id:
        return jsonify({"success": False, "error": "Sesión no especificada"}), 400
    
    if not instrucciones:
        return jsonify({"success": False, "error": "Instrucciones requeridas"}), 400
    
    sesion = ArgumentationSession.query.filter_by(
        id=session_id,
        user_id=current_user.id,
        tenant_id=tenant.id
    ).first()
    
    if not sesion:
        return jsonify({"success": False, "error": "Sesión no encontrada"}), 404
    
    job_type = detect_intent(instrucciones)
    
    job = ArgumentationJob(
        session_id=sesion.id,
        user_id=current_user.id,
        tenant_id=tenant.id,
        section=section,
        job_type=job_type,
        instructions=instrucciones,
        estilo=estilo,
        status='queued'
    )
    db.session.add(job)
    db.session.commit()
    
    argumentation_job_queue.put(job.id)
    
    return jsonify({
        "success": True,
        "job_id": job.id,
        "job_type": job_type,
        "status": "queued"
    })


@app.route("/argumentacion/jobs/<int:job_id>")
@login_required
def argumentacion_job_status(job_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/argumentacion/descargar/<int:session_id>")
@login_required
def argumentacion_descargar(session_id):
    """Ruta deshabilitada."""
    abort(404)


def _argumentacion_descargar_disabled(session_id):
    tenant = get_current_tenant()
    if not tenant:
        flash("No autorizado", "error")
        return redirect(url_for('argumentacion'))
    
    sesion = ArgumentationSession.query.filter_by(
        id=session_id,
        user_id=current_user.id,
        tenant_id=tenant.id
    ).first_or_404()
    
    texto = sesion.ultima_version_mejorada or sesion.documento_original
    
    doc = Document()
    
    estilo_doc = EstiloDocumento.query.filter_by(tenant_id=tenant.id).first()
    font_name = estilo_doc.fuente if estilo_doc else 'Times New Roman'
    font_size = estilo_doc.tamano_base if estilo_doc else 12
    line_spacing = estilo_doc.interlineado if estilo_doc else 1.5
    
    sections = doc.sections
    for section in sections:
        if estilo_doc:
            section.top_margin = Cm(estilo_doc.margen_superior)
            section.bottom_margin = Cm(estilo_doc.margen_inferior)
            section.left_margin = Cm(estilo_doc.margen_izquierdo)
            section.right_margin = Cm(estilo_doc.margen_derecho)
        else:
            section.top_margin = Cm(3.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)
        
        logo_path = get_tenant_logo_path(tenant)
        if logo_path and os.path.exists(logo_path):
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_para.add_run()
            try:
                run.add_picture(logo_path, width=Cm(4))
            except Exception as e:
                logging.error(f"Error adding logo to argumentation doc: {e}")
            
            info_lines = tenant.get_header_info()
            for linea in info_lines:
                info_para = header.add_paragraph()
                info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                info_run = info_para.add_run(linea)
                info_run.font.name = font_name
                info_run.font.size = Pt(9)
                info_para.paragraph_format.space_after = Pt(0)
                info_para.paragraph_format.space_before = Pt(0)
    
    titulos_principales = ['SUMILLA:', 'PETITORIO:', 'HECHOS:', 'FUNDAMENTOS', 'ANEXOS:', 
                          'POR TANTO:', 'VÍA PROCEDIMENTAL:', 'CONTRACAUTELA:',
                          'FUNDAMENTACION JURÍDICA:', 'FUNDAMENTACIÓN JURÍDICA:']
    titulos_secundarios = ['PRIMERO:', 'SEGUNDO:', 'TERCERO:', 'CUARTO:', 'QUINTO:',
                          'SEXTO:', 'SÉPTIMO:', 'OCTAVO:', 'NOVENO:', 'DÉCIMO:']
    
    for parrafo in texto.split('\n'):
        linea = parrafo.strip()
        if not linea:
            continue
        
        p = doc.add_paragraph()
        run = p.add_run(linea)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        
        es_titulo_principal = any(linea.upper().startswith(t.upper()) for t in titulos_principales)
        es_titulo_secundario = any(linea.upper().startswith(t.upper()) for t in titulos_secundarios)
        
        if es_titulo_principal:
            run.bold = True
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
        elif es_titulo_secundario:
            run.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        
        p.paragraph_format.line_spacing = line_spacing
    
    folder = get_argumentacion_folder(tenant.id)
    nombre_archivo = f"argumentacion_{sesion.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    ruta = os.path.join(folder, nombre_archivo)
    doc.save(ruta)
    
    return send_file(ruta, as_attachment=True, download_name=nombre_archivo)


@app.route("/argumentacion/copiar/<int:session_id>")
@login_required
def argumentacion_copiar_texto(session_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/argumentacion/estilo/nuevo", methods=["POST"])
@login_required
def argumentacion_estilo_nuevo():
    """Ruta deshabilitada."""
    abort(404)


@app.route("/argumentacion/estilo/eliminar/<int:estilo_id>", methods=["POST"])
@login_required
def argumentacion_estilo_eliminar(estilo_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/argumentacion/eliminar/<int:session_id>", methods=["POST"])
@login_required
def argumentacion_eliminar(session_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/argumentacion/historial")
@login_required
def argumentacion_historial():
    """Ruta deshabilitada."""
    abort(404)


# ==================== APC IA AGENT ====================

AGENT_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "obtener_info_caso",
            "description": "Obtiene información básica del caso: partes, materia, juzgado, estado, expediente, etc.",
            "parameters": {
                "type": "object",
                "properties": {
                    "case_id": {"type": "integer", "description": "ID del caso"}
                },
                "required": ["case_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "listar_documentos_del_caso",
            "description": "Lista todos los documentos asociados a un caso (demandas, contratos, resoluciones, escritos, etc.)",
            "parameters": {
                "type": "object",
                "properties": {
                    "case_id": {"type": "integer", "description": "ID del caso"}
                },
                "required": ["case_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "leer_documento",
            "description": "Lee el contenido de texto de un documento específico del caso",
            "parameters": {
                "type": "object",
                "properties": {
                    "doc_id": {"type": "integer", "description": "ID del documento"},
                    "doc_type": {"type": "string", "enum": ["case_document", "finished_document"], "description": "Tipo de documento"}
                },
                "required": ["doc_id", "doc_type"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generar_documento_desde_plantilla",
            "description": "Genera un documento legal completo (demanda, contestación, recurso, escrito) usando las plantillas del estudio",
            "parameters": {
                "type": "object",
                "properties": {
                    "case_id": {"type": "integer", "description": "ID del caso"},
                    "tipo_documento": {"type": "string", "description": "Tipo de documento a generar (demanda, contestación, recurso, escrito, etc.)"},
                    "instrucciones": {"type": "string", "description": "Instrucciones detalladas para la generación del documento"},
                    "plantilla_key": {"type": "string", "description": "Clave de la plantilla a usar (opcional)"}
                },
                "required": ["case_id", "tipo_documento", "instrucciones"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "guardar_borrador_estrategia",
            "description": "Guarda una estrategia legal propuesta como borrador asociada al caso",
            "parameters": {
                "type": "object",
                "properties": {
                    "case_id": {"type": "integer", "description": "ID del caso"},
                    "titulo": {"type": "string", "description": "Título de la estrategia"},
                    "contenido": {"type": "string", "description": "Contenido completo de la estrategia"},
                    "objetivo_principal": {"type": "string", "description": "Objetivo principal de la estrategia"},
                    "argumentos": {"type": "array", "items": {"type": "string"}, "description": "Lista de argumentos principales"},
                    "riesgos": {"type": "array", "items": {"type": "string"}, "description": "Lista de riesgos identificados"},
                    "proximos_pasos": {"type": "array", "items": {"type": "string"}, "description": "Lista de próximos pasos recomendados"}
                },
                "required": ["case_id", "titulo", "contenido"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "crear_tarea",
            "description": "Crea una tarea o recordatorio ligado al caso (vencimientos, audiencias, acciones pendientes)",
            "parameters": {
                "type": "object",
                "properties": {
                    "case_id": {"type": "integer", "description": "ID del caso"},
                    "titulo": {"type": "string", "description": "Título de la tarea"},
                    "descripcion": {"type": "string", "description": "Descripción de la tarea"},
                    "fecha_vencimiento": {"type": "string", "description": "Fecha de vencimiento en formato YYYY-MM-DD"},
                    "tipo": {"type": "string", "enum": ["general", "audiencia", "escrito", "reunion", "vencimiento", "otro"], "description": "Tipo de tarea"},
                    "prioridad": {"type": "string", "enum": ["alta", "media", "baja"], "description": "Prioridad de la tarea"}
                },
                "required": ["case_id", "titulo", "fecha_vencimiento"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "calcular_costos_estimados",
            "description": "Calcula los costos estimados del caso incluyendo honorarios, tasas y otros gastos",
            "parameters": {
                "type": "object",
                "properties": {
                    "case_id": {"type": "integer", "description": "ID del caso"},
                    "tipo_proceso": {"type": "string", "description": "Tipo de proceso judicial"},
                    "num_audiencias": {"type": "integer", "description": "Número estimado de audiencias"},
                    "horas_estimadas": {"type": "number", "description": "Horas de trabajo estimadas"},
                    "tarifa_hora": {"type": "number", "description": "Tarifa por hora en soles"},
                    "duracion_meses": {"type": "integer", "description": "Duración estimada en meses"}
                },
                "required": ["case_id"]
            }
        }
    }
]


def agent_tool_obtener_info_caso(case_id, tenant_id, user_id):
    """Obtiene información del caso."""
    caso = Case.query.filter_by(id=case_id, tenant_id=tenant_id).first()
    if not caso:
        return {"error": "Caso no encontrado"}
    
    assignments = CaseAssignment.query.filter_by(case_id=case_id).all()
    asignados = [{"user_id": a.user_id, "username": a.user.username, "rol": a.rol} for a in assignments if a.user]
    
    return {
        "id": caso.id,
        "titulo": caso.titulo,
        "numero_expediente": caso.numero_expediente,
        "materia": caso.materia,
        "estado": caso.estado,
        "cliente": caso.cliente,
        "contraparte": caso.contraparte,
        "juzgado": caso.juzgado,
        "descripcion": caso.descripcion,
        "fecha_inicio": caso.fecha_inicio.isoformat() if caso.fecha_inicio else None,
        "asignados": asignados
    }


def agent_tool_listar_documentos_del_caso(case_id, tenant_id, user_id):
    """Lista documentos del caso."""
    caso = Case.query.filter_by(id=case_id, tenant_id=tenant_id).first()
    if not caso:
        return {"error": "Caso no encontrado"}
    
    case_docs = CaseDocument.query.filter_by(case_id=case_id).order_by(CaseDocument.created_at.desc()).all()
    finished_docs = FinishedDocument.query.filter_by(case_id=case_id, tenant_id=tenant_id).order_by(FinishedDocument.created_at.desc()).all()
    
    documentos = []
    for doc in case_docs:
        documentos.append({
            "id": doc.id,
            "type": "case_document",
            "nombre": doc.nombre,
            "tipo": doc.tipo,
            "fecha": doc.created_at.isoformat() if doc.created_at else None
        })
    
    for doc in finished_docs:
        documentos.append({
            "id": doc.id,
            "type": "finished_document",
            "nombre": doc.nombre,
            "tipo": doc.tipo_documento,
            "fecha": doc.created_at.isoformat() if doc.created_at else None
        })
    
    return {"documentos": documentos, "total": len(documentos)}


def agent_tool_leer_documento(doc_id, doc_type, tenant_id, user_id):
    """Lee el contenido de un documento."""
    if doc_type == "case_document":
        doc = CaseDocument.query.filter_by(id=doc_id).first()
        if not doc:
            return {"error": "Documento no encontrado"}
        caso = Case.query.filter_by(id=doc.case_id, tenant_id=tenant_id).first()
        if not caso:
            return {"error": "No tienes acceso a este documento"}
        
        contenido = ""
        if doc.archivo_path and os.path.exists(doc.archivo_path):
            contenido = extract_text_from_file(doc.archivo_path)
        
        return {
            "id": doc.id,
            "nombre": doc.nombre,
            "tipo": doc.tipo,
            "contenido": contenido[:15000] if contenido else "No se pudo extraer el contenido"
        }
    
    elif doc_type == "finished_document":
        doc = FinishedDocument.query.filter_by(id=doc_id, tenant_id=tenant_id).first()
        if not doc:
            return {"error": "Documento no encontrado"}
        
        contenido = ""
        if doc.archivo and os.path.exists(doc.archivo):
            contenido = extract_text_from_file(doc.archivo)
        
        return {
            "id": doc.id,
            "nombre": doc.nombre,
            "tipo": doc.tipo_documento,
            "contenido": contenido[:15000] if contenido else "No se pudo extraer el contenido"
        }
    
    return {"error": "Tipo de documento no válido"}


def agent_tool_generar_documento(case_id, tipo_documento, instrucciones, plantilla_key, tenant_id, user_id, session_id):
    """Genera un documento legal usando IA."""
    caso = Case.query.filter_by(id=case_id, tenant_id=tenant_id).first()
    if not caso:
        return {"error": "Caso no encontrado"}
    
    case_docs = CaseDocument.query.filter_by(case_id=case_id).limit(5).all()
    contexto_documentos = ""
    for doc in case_docs:
        if doc.archivo_path and os.path.exists(doc.archivo_path):
            texto = extract_text_from_file(doc.archivo_path)
            if texto:
                contexto_documentos += f"\n\n--- Documento: {doc.nombre} ---\n{texto[:5000]}"
    
    prompt = f"""Eres un abogado experto redactando documentos legales en Perú.

INFORMACIÓN DEL CASO:
- Título: {caso.titulo}
- Expediente: {caso.numero_expediente or 'Sin número'}
- Materia: {caso.materia or 'No especificada'}
- Cliente: {caso.cliente or 'No especificado'}
- Contraparte: {caso.contraparte or 'No especificada'}
- Juzgado: {caso.juzgado or 'No especificado'}
- Descripción: {caso.descripcion or ''}

DOCUMENTOS DEL CASO:
{contexto_documentos if contexto_documentos else 'No hay documentos adjuntos'}

INSTRUCCIONES DEL USUARIO:
{instrucciones}

TIPO DE DOCUMENTO A GENERAR: {tipo_documento}

Genera el documento legal completo con la estructura apropiada:
- Encabezado con datos del expediente
- Sumilla
- Señor Juez / Autoridad
- Datos del solicitante
- Hechos (numerados)
- Fundamentos de Derecho (con citas legales reales de Perú)
- Petitorio
- Anexos
- Firma

El documento debe ser formal, técnico-jurídico y listo para presentar."""

    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=180.0)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=8000
        )
        
        texto_documento = response.choices[0].message.content
        
        tenant = Tenant.query.get(tenant_id)
        fecha_actual = datetime.now()
        nombre_archivo = f"{tipo_documento.replace(' ', '_')}_{case_id}_{fecha_actual.strftime('%Y%m%d_%H%M%S')}.docx"
        
        guardar_docx(texto_documento, nombre_archivo, tenant, None)
        
        finished_doc = FinishedDocument(
            tenant_id=tenant_id,
            user_id=user_id,
            case_id=case_id,
            nombre=f"{tipo_documento} - {caso.titulo}",
            tipo_documento=tipo_documento,
            archivo=os.path.join(get_resultados_folder(tenant), nombre_archivo),
            numero_expediente=caso.numero_expediente
        )
        db.session.add(finished_doc)
        db.session.commit()
        
        return {
            "success": True,
            "documento_id": finished_doc.id,
            "nombre_archivo": nombre_archivo,
            "resumen": texto_documento[:500] + "..." if len(texto_documento) > 500 else texto_documento,
            "download_url": url_for('descargar', nombre_archivo=nombre_archivo)
        }
        
    except Exception as e:
        logging.error(f"Error generando documento: {e}")
        return {"error": f"Error al generar el documento: {str(e)}"}


def agent_tool_guardar_estrategia(case_id, titulo, contenido, objetivo_principal, argumentos, riesgos, proximos_pasos, tenant_id, user_id, session_id):
    """Guarda una estrategia legal."""
    caso = Case.query.filter_by(id=case_id, tenant_id=tenant_id).first()
    if not caso:
        return {"error": "Caso no encontrado"}
    
    estrategia = LegalStrategy(
        tenant_id=tenant_id,
        case_id=case_id,
        session_id=session_id,
        created_by_id=user_id,
        titulo=titulo,
        contenido=contenido,
        objetivo_principal=objetivo_principal,
        argumentos_principales=argumentos,
        riesgos=riesgos,
        proximos_pasos=proximos_pasos,
        status='draft'
    )
    db.session.add(estrategia)
    db.session.commit()
    
    return {
        "success": True,
        "estrategia_id": estrategia.id,
        "mensaje": f"Estrategia '{titulo}' guardada correctamente"
    }


def agent_tool_crear_tarea(case_id, titulo, descripcion, fecha_vencimiento, tipo, prioridad, tenant_id, user_id):
    """Crea una tarea ligada al caso."""
    caso = Case.query.filter_by(id=case_id, tenant_id=tenant_id).first()
    if not caso:
        return {"error": "Caso no encontrado"}
    
    try:
        fecha = datetime.strptime(fecha_vencimiento, "%Y-%m-%d")
    except:
        return {"error": "Formato de fecha inválido. Use YYYY-MM-DD"}
    
    tarea = Task(
        tenant_id=tenant_id,
        case_id=case_id,
        titulo=titulo,
        descripcion=descripcion or "",
        tipo=tipo or "general",
        prioridad=prioridad or "media",
        fecha_vencimiento=fecha,
        created_by_id=user_id,
        assigned_to_id=user_id
    )
    db.session.add(tarea)
    db.session.commit()
    
    return {
        "success": True,
        "tarea_id": tarea.id,
        "mensaje": f"Tarea '{titulo}' creada para el {fecha_vencimiento}"
    }


def agent_tool_calcular_costos(case_id, tipo_proceso, num_audiencias, horas_estimadas, tarifa_hora, duracion_meses, tenant_id, user_id, session_id):
    """Calcula costos estimados del caso."""
    caso = Case.query.filter_by(id=case_id, tenant_id=tenant_id).first()
    if not caso:
        return {"error": "Caso no encontrado"}
    
    tarifa = tarifa_hora or 150.0
    horas = horas_estimadas or 20.0
    audiencias = num_audiencias or 3
    meses = duracion_meses or 12
    
    honorarios = horas * tarifa
    costo_audiencia = 200.0
    costo_audiencias = audiencias * costo_audiencia
    tasas_judiciales = 500.0
    otros_gastos = 300.0
    
    total = honorarios + costo_audiencias + tasas_judiciales + otros_gastos
    
    estimacion = CostEstimate(
        tenant_id=tenant_id,
        case_id=case_id,
        session_id=session_id,
        created_by_id=user_id,
        titulo=f"Estimación de costos - {caso.titulo}",
        assumptions={
            "tipo_proceso": tipo_proceso or caso.materia,
            "horas_estimadas": horas,
            "tarifa_hora": tarifa,
            "num_audiencias": audiencias,
            "duracion_meses": meses
        },
        breakdown={
            "honorarios": honorarios,
            "costo_audiencias": costo_audiencias,
            "tasas_judiciales": tasas_judiciales,
            "otros_gastos": otros_gastos
        },
        honorarios=honorarios,
        tasas_judiciales=tasas_judiciales,
        otros_gastos=otros_gastos + costo_audiencias,
        total_amount=total,
        currency="PEN",
        duracion_estimada_meses=meses,
        num_audiencias=audiencias,
        horas_estimadas=horas,
        tarifa_hora=tarifa
    )
    db.session.add(estimacion)
    db.session.commit()
    
    return {
        "success": True,
        "estimacion_id": estimacion.id,
        "resumen": {
            "honorarios": f"S/. {honorarios:,.2f}",
            "audiencias": f"S/. {costo_audiencias:,.2f}",
            "tasas_judiciales": f"S/. {tasas_judiciales:,.2f}",
            "otros_gastos": f"S/. {otros_gastos:,.2f}",
            "total": f"S/. {total:,.2f}"
        },
        "supuestos": {
            "horas": horas,
            "tarifa_hora": f"S/. {tarifa}",
            "audiencias": audiencias,
            "duracion": f"{meses} meses"
        }
    }


def execute_agent_tool(tool_name, arguments, tenant_id, user_id, session_id):
    """Ejecuta una herramienta del agente."""
    try:
        args = json.loads(arguments) if isinstance(arguments, str) else arguments
        
        if tool_name == "obtener_info_caso":
            return agent_tool_obtener_info_caso(args["case_id"], tenant_id, user_id)
        
        elif tool_name == "listar_documentos_del_caso":
            return agent_tool_listar_documentos_del_caso(args["case_id"], tenant_id, user_id)
        
        elif tool_name == "leer_documento":
            return agent_tool_leer_documento(args["doc_id"], args["doc_type"], tenant_id, user_id)
        
        elif tool_name == "generar_documento_desde_plantilla":
            return agent_tool_generar_documento(
                args["case_id"],
                args["tipo_documento"],
                args["instrucciones"],
                args.get("plantilla_key"),
                tenant_id, user_id, session_id
            )
        
        elif tool_name == "guardar_borrador_estrategia":
            return agent_tool_guardar_estrategia(
                args["case_id"],
                args["titulo"],
                args["contenido"],
                args.get("objetivo_principal"),
                args.get("argumentos", []),
                args.get("riesgos", []),
                args.get("proximos_pasos", []),
                tenant_id, user_id, session_id
            )
        
        elif tool_name == "crear_tarea":
            return agent_tool_crear_tarea(
                args["case_id"],
                args["titulo"],
                args.get("descripcion"),
                args["fecha_vencimiento"],
                args.get("tipo", "general"),
                args.get("prioridad", "media"),
                tenant_id, user_id
            )
        
        elif tool_name == "calcular_costos_estimados":
            return agent_tool_calcular_costos(
                args["case_id"],
                args.get("tipo_proceso"),
                args.get("num_audiencias"),
                args.get("horas_estimadas"),
                args.get("tarifa_hora"),
                args.get("duracion_meses"),
                tenant_id, user_id, session_id
            )
        
        return {"error": f"Herramienta '{tool_name}' no reconocida"}
    
    except Exception as e:
        logging.error(f"Error ejecutando herramienta {tool_name}: {e}")
        return {"error": str(e)}


def run_agent_conversation(session_id, user_message, case_id, tenant_id, user_id):
    """Ejecuta una conversación con el agente."""
    start_time = time.time()
    
    session = AgentSession.query.filter_by(
        id=session_id,
        tenant_id=tenant_id,
        user_id=user_id
    ).first()
    
    if not session:
        return {"error": "Sesión no encontrada"}
    
    user_msg = AgentMessage(
        session_id=session_id,
        tenant_id=tenant_id,
        user_id=user_id,
        role="user",
        content=user_message
    )
    db.session.add(user_msg)
    db.session.commit()
    
    messages_history = AgentMessage.query.filter_by(
        session_id=session_id,
        tenant_id=tenant_id
    ).order_by(AgentMessage.created_at).all()
    
    system_prompt = """Eres APC IA, un agente jurídico avanzado para la plataforma de gestión legal.

Tu rol es asistir a abogados con:
1. Lectura y análisis de documentos legales
2. Redacción de demandas, contestaciones, recursos y escritos
3. Diseño de estrategias legales
4. Evaluación de costos de casos
5. Gestión de tareas y plazos

COMPORTAMIENTO:
- Siempre identifica primero qué necesita el usuario
- Si necesitas información del caso, usa obtener_info_caso
- Si necesitas ver documentos, primero lista con listar_documentos_del_caso y luego lee con leer_documento
- Para generar documentos, usa generar_documento_desde_plantilla
- Para estrategias, genera un análisis completo y guárdalo con guardar_borrador_estrategia
- Para plazos y tareas, usa crear_tarea
- Para costos, usa calcular_costos_estimados

El tono debe ser profesional, técnico-jurídico y orientado a la práctica legal en Perú.
Siempre sé útil y proactivo, ofreciendo sugerencias cuando sea apropiado."""
    
    openai_messages = [{"role": "system", "content": system_prompt}]
    
    for msg in messages_history:
        if msg.role in ["user", "assistant"]:
            openai_messages.append({"role": msg.role, "content": msg.content})
    
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=120.0)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=openai_messages,
            tools=AGENT_TOOLS,
            tool_choice="auto",
            temperature=0.4,
            max_tokens=4000
        )
        
        assistant_message = response.choices[0].message
        tool_calls = assistant_message.tool_calls
        
        tools_used = []
        tool_results = []
        
        if tool_calls:
            for tool_call in tool_calls:
                tool_name = tool_call.function.name
                tool_args = tool_call.function.arguments
                
                result = execute_agent_tool(tool_name, tool_args, tenant_id, user_id, session_id)
                tools_used.append(tool_name)
                tool_results.append({"tool": tool_name, "result": result})
                
                tool_msg = AgentMessage(
                    session_id=session_id,
                    tenant_id=tenant_id,
                    role="tool",
                    content=json.dumps(result, ensure_ascii=False),
                    tool_used=tool_name,
                    tool_result=result
                )
                db.session.add(tool_msg)
            
            openai_messages.append({"role": "assistant", "content": None, "tool_calls": [
                {"id": tc.id, "type": "function", "function": {"name": tc.function.name, "arguments": tc.function.arguments}}
                for tc in tool_calls
            ]})
            
            for i, tool_call in enumerate(tool_calls):
                openai_messages.append({
                    "role": "tool",
                    "tool_call_id": tool_call.id,
                    "content": json.dumps(tool_results[i]["result"], ensure_ascii=False)
                })
            
            final_response = client.chat.completions.create(
                model="gpt-4o",
                messages=openai_messages,
                temperature=0.4,
                max_tokens=4000
            )
            
            final_content = final_response.choices[0].message.content
        else:
            final_content = assistant_message.content
        
        latency = int((time.time() - start_time) * 1000)
        
        assistant_msg = AgentMessage(
            session_id=session_id,
            tenant_id=tenant_id,
            role="assistant",
            content=final_content,
            tool_used=",".join(tools_used) if tools_used else None,
            latency_ms=latency
        )
        db.session.add(assistant_msg)
        
        session.updated_at = datetime.utcnow()
        db.session.commit()
        
        return {
            "success": True,
            "message": final_content,
            "tools_used": tools_used,
            "tool_results": tool_results,
            "latency_ms": latency
        }
        
    except Exception as e:
        logging.error(f"Error en agente: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return {"error": str(e)}


@app.route("/apc-ia")
@login_required
def apc_ia():
    """Ruta deshabilitada."""
    abort(404)


@app.route("/apc-ia/sesion/<int:session_id>")
@login_required
def apc_ia_sesion(session_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/api/apc/sessions", methods=["GET", "POST"])
@login_required
def api_apc_sessions():
    """Ruta deshabilitada."""
    abort(404)


@app.route("/api/apc/agent", methods=["POST"])
@login_required
def api_apc_agent():
    """Ruta deshabilitada."""
    abort(404)


def _api_apc_agent_disabled():
    """API principal del agente - procesa mensajes."""
    tenant = get_current_tenant()
    if not tenant:
        return jsonify({"error": "No autorizado"}), 403
    
    data = request.get_json() or {}
    session_id = data.get("session_id")
    message = data.get("message", "").strip()
    case_id = data.get("case_id")
    
    if not message:
        return jsonify({"error": "Mensaje requerido"}), 400
    
    if not session_id:
        sesion = AgentSession(
            user_id=current_user.id,
            tenant_id=tenant.id,
            case_id=case_id,
            titulo=message[:50] + "..." if len(message) > 50 else message
        )
        db.session.add(sesion)
        db.session.commit()
        session_id = sesion.id
    else:
        sesion = AgentSession.query.filter_by(
            id=session_id,
            user_id=current_user.id,
            tenant_id=tenant.id
        ).first()
        
        if not sesion:
            return jsonify({"error": "Sesión no encontrada"}), 404
        
        if case_id and sesion.case_id != case_id:
            sesion.case_id = case_id
            db.session.commit()
    
    result = run_agent_conversation(session_id, message, case_id, tenant.id, current_user.id)
    
    if "error" in result:
        return jsonify(result), 500
    
    result["session_id"] = session_id
    return jsonify(result)


@app.route("/api/apc/sessions/<int:session_id>/messages")
@login_required
def api_apc_session_messages(session_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/api/apc/sessions/<int:session_id>", methods=["DELETE"])
@login_required
def api_apc_delete_session(session_id):
    """Ruta deshabilitada."""
    abort(404)


@app.route("/f/<public_id>", methods=["GET", "POST"])
def public_form(public_id):
    """Formulario público para que clientes llenen datos de plantilla."""
    from models import Modelo, FormResponse, Tenant, CampoPlantilla
    
    modelo = Modelo.query.filter_by(public_id=public_id, activa=True).first()
    if not modelo or not modelo.is_public_form_enabled:
        return render_template("public_form_not_found.html"), 404
    
    tenant = Tenant.query.get(modelo.tenant_id)
    if not tenant or not tenant.activo:
        return render_template("public_form_not_found.html"), 404
    
    campos = CampoPlantilla.query.filter_by(
        plantilla_key=modelo.key,
        tenant_id=modelo.tenant_id
    ).order_by(CampoPlantilla.orden).all()
    
    if not campos:
        campos = CampoPlantilla.query.filter_by(
            plantilla_key=modelo.key,
            tenant_id=None
        ).order_by(CampoPlantilla.orden).all()
    
    if request.method == "POST":
        accepted = request.form.get('accepted_terms') == 'on'
        if not accepted:
            return render_template("public_form.html", 
                                  modelo=modelo, 
                                  tenant_nombre=tenant.nombre,
                                  campos=campos,
                                  error="Debe aceptar la declaración de veracidad para continuar.")
        
        answers = {}
        for campo in campos:
            answers[campo.nombre_campo] = request.form.get(f"campo_{campo.nombre_campo}", "").strip()
        
        code = FormResponse.generate_code(tenant.id)
        while FormResponse.query.filter_by(code=code).first():
            code = FormResponse.generate_code(tenant.id)
        
        form_response = FormResponse(
            tenant_id=tenant.id,
            template_id=modelo.id,
            code=code,
            answers_json=answers,
            status='NEW',
            accepted_terms=True,
            accepted_terms_at=datetime.utcnow()
        )
        db.session.add(form_response)
        db.session.commit()
        
        log_audit(tenant.id, 'FORM_SUBMITTED', f'Formulario enviado: {code}', 
                 {'template_id': modelo.id, 'code': code}, user_id=None)
        
        qr_base64 = generate_qr_code_base64(code)
        
        return render_template("public_form_success.html",
                              code=code,
                              qr_base64=qr_base64,
                              tenant_nombre=tenant.nombre)
    
    return render_template("public_form.html",
                          modelo=modelo,
                          tenant_nombre=tenant.nombre,
                          campos=campos,
                          error=None)


@app.route("/cargar_formulario", methods=["GET", "POST"])
@login_required
def cargar_formulario_por_codigo():
    """Cargar un formulario por código para generar documento."""
    from models import FormResponse, Modelo
    
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes acceso.", "error")
        return redirect(url_for('dashboard'))
    
    if request.method == "POST":
        code = request.form.get('code', '').strip().upper()
        if not code:
            flash("Ingrese un código de formulario.", "error")
            return render_template("cargar_formulario.html")
        
        form_response = FormResponse.query.filter_by(code=code, tenant_id=tenant.id).first()
        if not form_response:
            flash("Código no encontrado. Verifique e intente nuevamente.", "error")
            return render_template("cargar_formulario.html")
        
        log_audit(tenant.id, 'FORM_LOADED_BY_STAFF', f'Formulario cargado: {code}', 
                 {'code': code, 'user_id': current_user.id})
        
        return redirect(url_for('vista_previa_formulario', form_id=form_response.id))
    
    formularios = FormResponse.query.filter_by(tenant_id=tenant.id).order_by(FormResponse.created_at.desc()).limit(20).all()
    return render_template("cargar_formulario.html", formularios=formularios)


@app.route("/formulario/<int:form_id>/preview", methods=["GET", "POST"])
@login_required
def vista_previa_formulario(form_id):
    """Vista previa del formulario antes de generar documento."""
    from models import FormResponse, Modelo
    
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes acceso.", "error")
        return redirect(url_for('dashboard'))
    
    form_response = FormResponse.query.filter_by(id=form_id, tenant_id=tenant.id).first()
    if not form_response:
        flash("Formulario no encontrado.", "error")
        return redirect(url_for('cargar_formulario_por_codigo'))
    
    modelo = form_response.template
    can_generate = form_response.can_be_used()
    is_admin = current_user.is_admin
    
    if request.method == "POST":
        action = request.form.get('action', '')
        
        if action == 'generate' and can_generate:
            return redirect(url_for('generar_desde_formulario', form_id=form_id))
        
        elif action == 'reopen' and is_admin and form_response.status == 'USED':
            form_response.reopen(current_user.id)
            db.session.commit()
            log_audit(tenant.id, 'FORM_REOPENED', f'Formulario reabierto: {form_response.code}',
                     {'code': form_response.code, 'user_id': current_user.id})
            flash("Formulario reabierto exitosamente.", "success")
            return redirect(url_for('vista_previa_formulario', form_id=form_id))
    
    return render_template("vista_previa_formulario.html",
                          form_response=form_response,
                          modelo=modelo,
                          can_generate=can_generate,
                          is_admin=is_admin)


@app.route("/formulario/<int:form_id>/generar", methods=["GET", "POST"])
@login_required
def generar_desde_formulario(form_id):
    """Genera documento final desde un formulario completado."""
    from models import FormResponse, Modelo
    
    tenant = get_current_tenant()
    if not tenant:
        flash("No tienes acceso.", "error")
        return redirect(url_for('dashboard'))
    
    form_response = FormResponse.query.filter_by(id=form_id, tenant_id=tenant.id).first()
    if not form_response:
        flash("Formulario no encontrado.", "error")
        return redirect(url_for('cargar_formulario_por_codigo'))
    
    if not form_response.can_be_used():
        flash("Este formulario ya fue utilizado. Contacte al administrador para reabrirlo.", "warning")
        return redirect(url_for('vista_previa_formulario', form_id=form_id))
    
    modelo = form_response.template
    answers = form_response.answers_json or {}
    
    if request.method == "POST":
        try:
            fecha_actual = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            nombre_base = modelo.nombre.replace(" ", "_")[:30]
            nombre_archivo = f"{nombre_base}_{form_response.code}_{fecha_actual}.docx"
            
            if modelo.archivo_original and os.path.exists(modelo.archivo_original):
                from docx import Document
                doc = Document(modelo.archivo_original)
                
                for para in doc.paragraphs:
                    for key, value in answers.items():
                        placeholder = "{{" + key + "}}"
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for key, value in answers.items():
                                    placeholder = "{{" + key + "}}"
                                    if placeholder in para.text:
                                        para.text = para.text.replace(placeholder, str(value))
                
                output_path = os.path.join(CARPETA_RESULTADOS, nombre_archivo)
                doc.save(output_path)
            else:
                texto_contenido = modelo.contenido or ""
                for key, value in answers.items():
                    placeholder = "{{" + key + "}}"
                    texto_contenido = texto_contenido.replace(placeholder, str(value))
                
                guardar_docx(texto_contenido, nombre_archivo, tenant, None)
            
            demandante = answers.get('DEMANDANTE', answers.get('NOMBRE', answers.get('CLIENTE', 'N/A')))
            
            record = DocumentRecord(
                user_id=current_user.id,
                tenant_id=tenant.id,
                fecha=datetime.utcnow(),
                tipo_documento=modelo.nombre,
                tipo_documento_key=modelo.key,
                demandante=demandante[:200] if demandante else 'Formulario',
                archivo=nombre_archivo,
                texto_generado=str(answers),
                datos_caso=answers
            )
            db.session.add(record)
            db.session.flush()
            
            form_response.mark_as_used(current_user.id, record.id)
            db.session.commit()
            
            log_audit(tenant.id, 'DOCUMENT_CREATED', f'Documento creado desde formulario: {form_response.code}',
                     {'code': form_response.code, 'document_id': record.id, 'archivo': nombre_archivo})
            
            flash(f"Documento generado exitosamente: {nombre_archivo}", "success")
            return redirect(url_for('terminados'))
            
        except Exception as e:
            logging.error(f"Error generando documento desde formulario: {e}")
            db.session.rollback()
            flash(f"Error al generar el documento: {str(e)}", "error")
            return redirect(url_for('vista_previa_formulario', form_id=form_id))
    
    return render_template("confirmar_generacion.html",
                          form_response=form_response,
                          modelo=modelo,
                          answers=answers)


@app.route("/modelo/<int:modelo_id>/toggle_public_form", methods=["POST"])
@login_required
def toggle_public_form(modelo_id):
    """Habilitar/deshabilitar formulario público para un modelo."""
    tenant = get_current_tenant()
    if not tenant:
        return jsonify({"error": "No autorizado"}), 403
    
    modelo = Modelo.query.filter_by(id=modelo_id, tenant_id=tenant.id).first()
    if not modelo:
        return jsonify({"error": "Modelo no encontrado"}), 404
    
    modelo.is_public_form_enabled = not modelo.is_public_form_enabled
    
    if modelo.is_public_form_enabled and not modelo.public_id:
        modelo.generate_public_id()
        if modelo.archivo_original and os.path.exists(modelo.archivo_original):
            modelo.placeholders_json = extract_placeholders_from_docx(modelo.archivo_original)
        elif modelo.contenido:
            modelo.placeholders_json = extract_placeholders_from_text(modelo.contenido)
    
    db.session.commit()
    
    return jsonify({
        "success": True,
        "is_enabled": modelo.is_public_form_enabled,
        "public_id": modelo.public_id
    })


@app.route("/modelo/<int:modelo_id>/get_public_link")
@login_required
def get_modelo_public_link(modelo_id):
    """Obtener el link público de un modelo."""
    tenant = get_current_tenant()
    if not tenant:
        return jsonify({"error": "No autorizado"}), 403
    
    modelo = Modelo.query.filter_by(id=modelo_id, tenant_id=tenant.id).first()
    if not modelo:
        return jsonify({"error": "Modelo no encontrado"}), 404
    
    if not modelo.public_id:
        modelo.generate_public_id()
        db.session.commit()
    
    public_url = url_for('public_form', public_id=modelo.public_id, _external=True)
    
    return jsonify({
        "success": True,
        "public_id": modelo.public_id,
        "public_url": public_url,
        "is_enabled": modelo.is_public_form_enabled
    })


with app.app_context():
    db.create_all()
    os.makedirs(CARPETA_MODELOS, exist_ok=True)
    os.makedirs(CARPETA_ESTILOS, exist_ok=True)
    os.makedirs(CARPETA_RESULTADOS, exist_ok=True)
    os.makedirs(CARPETA_PLANTILLAS_SUBIDAS, exist_ok=True)
    os.makedirs(CARPETA_ESTILOS_SUBIDOS, exist_ok=True)
    os.makedirs(CARPETA_IMAGENES_MODELOS, exist_ok=True)
    os.makedirs(CARPETA_DOCUMENTOS_TERMINADOS, exist_ok=True)
    os.makedirs(CARPETA_ANONIMIZADOS, exist_ok=True)
    os.makedirs(CARPETA_REVISIONES, exist_ok=True)
    os.makedirs(CARPETA_ARGUMENTACION, exist_ok=True)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port, debug=False)
