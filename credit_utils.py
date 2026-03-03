import os
import math
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

TRIAL_PAGES_DEFAULT = int(os.environ.get("TRIAL_PAGES_DEFAULT", "40"))


def get_or_create_credits(user_id):
    from models import db, UserCredits
    credits = UserCredits.query.filter_by(user_id=user_id).first()
    if not credits:
        credits = UserCredits(user_id=user_id, pages_balance=0, pages_used_total=0)
        db.session.add(credits)
        db.session.commit()
        logger.info(f"CREDITS_CREATED | user={user_id}")
    return credits


def ensure_trial(user_id):
    from models import db, UserCredits, PageUsageLog
    credits = get_or_create_credits(user_id)
    if credits.trial_granted_at is not None:
        return credits
    credits.pages_balance += TRIAL_PAGES_DEFAULT
    credits.trial_granted_at = datetime.utcnow()
    log_entry = PageUsageLog(
        user_id=user_id,
        job_id="trial",
        stage="trial",
        pages=TRIAL_PAGES_DEFAULT,
        action="trial_granted",
        detail=f"Trial: +{TRIAL_PAGES_DEFAULT} páginas gratis"
    )
    db.session.add(log_entry)
    db.session.commit()
    logger.info(f"TRIAL_GRANTED | user={user_id} | pages={TRIAL_PAGES_DEFAULT}")
    return credits


def count_pages(file_path, ext):
    ext = ext.lower().strip('.')
    if ext == 'pdf':
        return _count_pages_pdf(file_path)
    elif ext == 'docx':
        return _count_pages_docx(file_path)
    elif ext == 'txt':
        return _count_pages_txt(file_path)
    elif ext == 'doc':
        return _count_pages_docx(file_path)
    return 1


def _count_pages_pdf(file_path):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = len(reader.pages)
        logger.info(f"PAGE_COUNT_PDF | path={file_path} | pages={pages}")
        return max(1, pages)
    except Exception as e:
        logger.warning(f"PAGE_COUNT_PDF_FAIL | error={e}")
        try:
            import fitz
            doc = fitz.open(file_path)
            pages = len(doc)
            doc.close()
            return max(1, pages)
        except Exception as e2:
            logger.error(f"PAGE_COUNT_PDF_FALLBACK_FAIL | error={e2}")
            return 1


def _count_pages_docx(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        word_count = 0
        for para in doc.paragraphs:
            word_count += len(para.text.split())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    word_count += len(cell.text.split())
        pages = max(1, math.ceil(word_count / 500))
        logger.info(f"PAGE_COUNT_DOCX | path={file_path} | words={word_count} | pages_equiv={pages}")
        return pages
    except Exception as e:
        logger.error(f"PAGE_COUNT_DOCX_FAIL | error={e}")
        return 1


def _count_pages_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        word_count = len(text.split())
        pages = max(1, math.ceil(word_count / 500))
        logger.info(f"PAGE_COUNT_TXT | path={file_path} | words={word_count} | pages_equiv={pages}")
        return pages
    except Exception as e:
        logger.error(f"PAGE_COUNT_TXT_FAIL | error={e}")
        return 1


def is_unlimited_user(user_id):
    """Retorna True si el usuario es super_admin o tiene unlimited_access."""
    from models import User
    user = User.query.get(user_id)
    if not user:
        return False
    return user.role == 'super_admin' or bool(getattr(user, 'unlimited_access', False))


def check_and_reserve_pages(user_id, job_id, pages_needed):
    from models import db, PageReservation, PageUsageLog
    if is_unlimited_user(user_id):
        logger.info(f"RESERVE_UNLIMITED | user={user_id} | job={job_id} | pages={pages_needed}")
        return True, pages_needed
    existing = PageReservation.query.filter_by(job_id=job_id).first()
    if existing:
        if existing.status == 'reserved' and existing.user_id == user_id:
            logger.info(f"RESERVE_IDEMPOTENT | user={user_id} | job={job_id} | pages={existing.pages_reserved}")
            return True, existing.pages_reserved
        elif existing.status == 'charged':
            logger.warning(f"RESERVE_ALREADY_CHARGED | user={user_id} | job={job_id}")
            return True, existing.pages_reserved
    credits = get_or_create_credits(user_id)
    if credits.pages_balance < pages_needed:
        log_entry = PageUsageLog(
            user_id=user_id,
            job_id=job_id,
            stage="process",
            pages=pages_needed,
            action="blocked",
            detail=f"Saldo insuficiente: necesita {pages_needed}, tiene {credits.pages_balance}"
        )
        db.session.add(log_entry)
        db.session.commit()
        logger.warning(f"RESERVE_BLOCKED | user={user_id} | job={job_id} | needed={pages_needed} | balance={credits.pages_balance}")
        return False, credits.pages_balance
    credits.pages_balance -= pages_needed
    reservation = PageReservation(
        user_id=user_id,
        job_id=job_id,
        pages_reserved=pages_needed,
        status='reserved'
    )
    db.session.add(reservation)
    log_entry = PageUsageLog(
        user_id=user_id,
        job_id=job_id,
        stage="process",
        pages=pages_needed,
        action="reserved",
        detail=f"Reservadas {pages_needed} páginas. Saldo restante: {credits.pages_balance}"
    )
    db.session.add(log_entry)
    db.session.commit()
    logger.info(f"RESERVE_OK | user={user_id} | job={job_id} | pages={pages_needed} | remaining={credits.pages_balance}")
    return True, pages_needed


def charge_pages(user_id, job_id, stage="apply"):
    from models import db, AnonymizerJob, PageReservation, PageUsageLog
    reservation = PageReservation.query.filter_by(job_id=job_id, user_id=user_id).first()
    if not reservation:
        logger.error(f"CHARGE_NO_RESERVATION | user={user_id} | job={job_id}")
        return False
    if reservation.status == 'charged':
        logger.info(f"CHARGE_IDEMPOTENT | user={user_id} | job={job_id} | already charged")
        return True
    pages_to_charge = reservation.pages_reserved
    credits = get_or_create_credits(user_id)
    credits.pages_used_total += pages_to_charge
    reservation.status = 'charged'
    job = AnonymizerJob.query.filter_by(job_id=job_id).first()
    if job:
        job.pages_charged = pages_to_charge
        job.status = 'success'
    log_entry = PageUsageLog(
        user_id=user_id,
        job_id=job_id,
        stage=stage,
        pages=pages_to_charge,
        action="charged",
        detail=f"Cobradas {pages_to_charge} páginas. Total usado: {credits.pages_used_total}"
    )
    db.session.add(log_entry)
    db.session.commit()
    logger.info(f"CHARGE_OK | user={user_id} | job={job_id} | pages={pages_to_charge} | total_used={credits.pages_used_total}")
    return True


def release_reservation(user_id, job_id):
    from models import db, PageReservation, PageUsageLog
    reservation = PageReservation.query.filter_by(job_id=job_id, user_id=user_id).first()
    if not reservation:
        return
    if reservation.status != 'reserved':
        return
    credits = get_or_create_credits(user_id)
    credits.pages_balance += reservation.pages_reserved
    reservation.status = 'released'
    log_entry = PageUsageLog(
        user_id=user_id,
        job_id=job_id,
        stage="release",
        pages=reservation.pages_reserved,
        action="released",
        detail=f"Liberadas {reservation.pages_reserved} páginas. Saldo: {credits.pages_balance}"
    )
    db.session.add(log_entry)
    db.session.commit()
    logger.info(f"RELEASE_OK | user={user_id} | job={job_id} | pages={reservation.pages_reserved}")
