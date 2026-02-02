"""
Procesador DOCX Run-Aware
=========================
Implementa reemplazo de texto en DOCX preservando formato,
incluso cuando el texto está partido en múltiples runs.
"""

import re
import os
import logging
from typing import Dict, List, Tuple, Optional, Any
from collections import defaultdict
from copy import deepcopy

from detector_capas import Entity, detect_all_pii, post_scan_final


class EntityMapping:
    """Mantiene mapeo consistente de valores a tokens."""
    
    def __init__(self):
        self.mappings: Dict[str, Dict[str, str]] = defaultdict(dict)
        self.counters: Dict[str, int] = defaultdict(int)
        self.reverse_mappings: Dict[str, str] = {}
    
    def get_token(self, entity_type: str, value: str) -> str:
        """Obtiene o crea un token para un valor."""
        normalized = value.strip().upper()
        
        if normalized in self.mappings[entity_type]:
            return self.mappings[entity_type][normalized]
        
        self.counters[entity_type] += 1
        token = f"{{{{{entity_type}_{self.counters[entity_type]}}}}}"
        
        self.mappings[entity_type][normalized] = token
        self.reverse_mappings[token] = self._mask_value(value, entity_type)
        
        return token
    
    def _mask_value(self, value: str, entity_type: str) -> str:
        """Crea versión enmascarada del valor para el reporte."""
        if len(value) <= 4:
            return '*' * len(value)
        
        if entity_type in ['DNI', 'RUC']:
            return value[:2] + '*' * (len(value) - 4) + value[-2:]
        elif entity_type == 'EMAIL':
            parts = value.split('@')
            if len(parts) == 2:
                return f"{parts[0][:2]}***@{parts[1]}"
        elif entity_type == 'TELEFONO':
            return value[:3] + '***' + value[-2:]
        elif entity_type == 'PERSONA':
            words = value.split()
            if len(words) >= 2:
                return words[0][:2] + '*** ' + words[-1][:2] + '***'
            return value[:2] + '***'
        
        if len(value) > 10:
            return value[:3] + '...' + value[-3:]
        return value[:2] + '***'
    
    def get_summary(self) -> Dict[str, int]:
        """Resumen de entidades por tipo."""
        return {t: len(m) for t, m in self.mappings.items() if m}


def replace_in_runs_aware(paragraph, replacements: List[Tuple[str, str]]) -> int:
    """
    Reemplaza texto en un párrafo preservando formato.
    Maneja texto partido en múltiples runs.
    
    Args:
        paragraph: Objeto paragraph de python-docx
        replacements: Lista de tuplas (valor_original, token)
    
    Returns:
        Número de reemplazos realizados
    """
    count = 0
    MAX_ITERATIONS = 100
    
    try:
        full_text = paragraph.text
        if not full_text:
            return 0
    except Exception:
        return 0
    
    for original, token in replacements:
        if not original or original not in full_text:
            continue
        
        iterations = 0
        # Construir mapa de posiciones de runs
        run_map = []  # [(start, end, run_idx, run)]
        pos = 0
        for idx, run in enumerate(paragraph.runs):
            try:
                run_text = run.text or ''
            except Exception:
                run_text = ''
            run_map.append((pos, pos + len(run_text), idx, run))
            pos += len(run_text)
        
        # Buscar todas las ocurrencias
        start = 0
        while iterations < MAX_ITERATIONS:
            iterations += 1
            idx = full_text.find(original, start)
            if idx == -1:
                break
            
            end_idx = idx + len(original)
            
            # Encontrar runs afectados
            affected_runs = []
            for run_start, run_end, run_idx, run in run_map:
                if run_start < end_idx and run_end > idx:
                    affected_runs.append((run_start, run_end, run_idx, run))
            
            if affected_runs:
                # Calcular la porción a reemplazar en cada run
                first_run = affected_runs[0]
                last_run = affected_runs[-1]
                
                first_run_start = first_run[0]
                first_run_obj = first_run[3]
                local_start = idx - first_run_start
                
                if len(affected_runs) == 1:
                    local_end = local_start + len(original)
                    old_text = first_run_obj.text or ''
                    first_run_obj.text = old_text[:local_start] + token + old_text[local_end:]
                else:
                    old_text = first_run_obj.text or ''
                    first_run_obj.text = old_text[:local_start] + token
                    
                    for _, _, _, run in affected_runs[1:-1]:
                        run.text = ''
                    
                    last_run_obj = last_run[3]
                    last_run_start = last_run[0]
                    local_end_in_last = end_idx - last_run_start
                    last_text = last_run_obj.text or ''
                    last_run_obj.text = last_text[local_end_in_last:]
                
                count += 1
            
            # Actualizar texto completo y mapa para siguiente iteración
            try:
                full_text = paragraph.text or ''
            except Exception:
                break
            run_map = []
            pos = 0
            for r_idx, run in enumerate(paragraph.runs):
                try:
                    rt = run.text or ''
                except Exception:
                    rt = ''
                run_map.append((pos, pos + len(rt), r_idx, run))
                pos += len(rt)
            
            start = 0
    
    return count


def apply_replacements_to_docx(doc, replacements: List[Tuple[str, str]]) -> int:
    """
    Aplica reemplazos a un documento DOCX completo.
    Cubre: párrafos, tablas (recursivas), headers/footers y sus tablas.
    
    Args:
        doc: Documento python-docx
        replacements: Lista de tuplas (valor_original, token)
    
    Returns:
        Número total de reemplazos realizados
    """
    if not replacements:
        return 0
    
    total_replacements = 0
    
    sorted_replacements = sorted(replacements, key=lambda x: len(x[0]), reverse=True)
    
    for para in doc.paragraphs:
        total_replacements += replace_in_runs_aware(para, sorted_replacements)
    
    def process_table(table):
        count = 0
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += replace_in_runs_aware(para, sorted_replacements)
                for nested_table in cell.tables:
                    count += process_table(nested_table)
        return count
    
    for table in doc.tables:
        total_replacements += process_table(table)
    
    try:
        for section in doc.sections:
            header_footer_elements = [
                section.header, section.footer,
                section.first_page_header, section.first_page_footer,
                section.even_page_header, section.even_page_footer
            ]
            
            for element in header_footer_elements:
                if element is None:
                    continue
                try:
                    for para in element.paragraphs:
                        total_replacements += replace_in_runs_aware(para, sorted_replacements)
                    for table in element.tables:
                        total_replacements += process_table(table)
                except Exception:
                    continue
    except Exception:
        pass
    
    return total_replacements


def process_docx_run_aware(doc, entities: List[Entity], mapping: EntityMapping) -> Dict[str, Any]:
    """
    Procesa documento DOCX con reemplazo run-aware.
    
    Args:
        doc: Documento python-docx
        entities: Lista de entidades detectadas
        mapping: Objeto EntityMapping para tokens
    
    Returns:
        Dict con estadísticas del proceso
    """
    stats = {
        'paragraphs_processed': 0,
        'tables_processed': 0,
        'headers_processed': 0,
        'footers_processed': 0,
        'replacements': 0,
        'entities_replaced': defaultdict(int)
    }
    
    # Crear lista de reemplazos (ordenar por longitud descendente para evitar conflictos)
    replacements = []
    for entity in sorted(entities, key=lambda e: len(e.value), reverse=True):
        token = mapping.get_token(entity.type, entity.value)
        replacements.append((entity.value, token))
    
    # Procesar párrafos del cuerpo
    for para in doc.paragraphs:
        count = replace_in_runs_aware(para, replacements)
        stats['replacements'] += count
        stats['paragraphs_processed'] += 1
    
    # Procesar tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count = replace_in_runs_aware(para, replacements)
                    stats['replacements'] += count
        stats['tables_processed'] += 1
    
    # Procesar headers y footers
    try:
        for section in doc.sections:
            # Header
            if section.header:
                for para in section.header.paragraphs:
                    count = replace_in_runs_aware(para, replacements)
                    stats['replacements'] += count
                stats['headers_processed'] += 1
            
            # Footer
            if section.footer:
                for para in section.footer.paragraphs:
                    count = replace_in_runs_aware(para, replacements)
                    stats['replacements'] += count
                stats['footers_processed'] += 1
    except Exception as e:
        logging.warning(f"Error processing headers/footers: {e}")
    
    # Contar entidades por tipo
    for entity in entities:
        stats['entities_replaced'][entity.type] += 1
    
    return stats


def extract_full_text_docx(doc) -> str:
    """Extrae todo el texto del documento DOCX."""
    text_parts = []
    
    # Párrafos principales
    for para in doc.paragraphs:
        text_parts.append(para.text)
    
    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_parts.append(para.text)
    
    # Headers y footers
    try:
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    text_parts.append(para.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    text_parts.append(para.text)
    except:
        pass
    
    return '\n'.join(text_parts)


def anonymize_docx_complete(file_path: str, output_path: str, strict_mode: bool = True) -> Dict[str, Any]:
    """
    Anonimiza un documento DOCX completo con las 4 capas + post-scan.
    
    Args:
        file_path: Ruta al archivo DOCX original
        output_path: Ruta para guardar el archivo anonimizado
        strict_mode: Si True, ejecuta post-scan y marca needs_review
    
    Returns:
        Dict con resultado completo del proceso
    """
    from docx import Document
    
    result = {
        'ok': True,
        'needs_review': False,
        'entities': [],
        'detection_metadata': {},
        'replacement_stats': {},
        'post_scan_results': [],
        'mapping': {},
        'error': None
    }
    
    try:
        # Cargar documento
        doc = Document(file_path)
        
        # Extraer texto completo
        full_text = extract_full_text_docx(doc)
        
        # Detectar PII con las 4 capas
        entities, metadata = detect_all_pii(full_text)
        result['detection_metadata'] = metadata
        result['entities'] = [
            {'type': e.type, 'start': e.start, 'end': e.end, 'source': e.source}
            for e in entities
        ]
        
        # Crear mapping y procesar reemplazos
        mapping = EntityMapping()
        stats = process_docx_run_aware(doc, entities, mapping)
        result['replacement_stats'] = dict(stats)
        result['replacement_stats']['entities_replaced'] = dict(stats['entities_replaced'])
        result['mapping'] = mapping.reverse_mappings
        
        # Guardar documento
        doc.save(output_path)
        
        # POST-SCAN obligatorio
        if strict_mode:
            # Recargar y extraer texto del documento anonimizado
            doc_check = Document(output_path)
            final_text = extract_full_text_docx(doc_check)
            
            needs_review, detected = post_scan_final(final_text)
            result['needs_review'] = needs_review
            result['post_scan_results'] = detected
        
    except Exception as e:
        logging.error(f"Error in anonymize_docx_complete: {e}")
        result['ok'] = False
        result['error'] = str(e)
    
    return result


HARD_REDACT_PATTERNS = [
    (re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', re.IGNORECASE), '{{EMAIL_REDACT}}'),
    (re.compile(r'\b\d{8}\b'), '{{DNI_REDACT}}'),
    (re.compile(r'\b(?:10|20)\d{9}\b'), '{{RUC_REDACT}}'),
    (re.compile(r'\+51[\s\-]?9[\d\s\-]{8,12}|9\d{8}', re.IGNORECASE), '{{TEL_REDACT}}'),
    (re.compile(r'(?:C\.?A\.?L\.?|CAL|CMP|CIP)[\s:N°º]*\d{4,6}', re.IGNORECASE), '{{COLEGIATURA_REDACT}}'),
]


def hard_redact_patterns(doc) -> int:
    """
    Aplica redacción forzada de patrones sensibles como último recurso.
    Esto es un safety net cuando la detección normal falla.
    Usa el mismo enfoque run-aware que replace_in_runs_aware.
    """
    total_fixes = 0
    
    def redact_paragraph_run_aware(paragraph):
        fixes = 0
        MAX_ITERATIONS = 50
        
        try:
            runs = paragraph.runs
            if not runs:
                return 0
        except Exception:
            return 0
        
        full_text = ''.join(run.text or '' for run in runs)
        if not full_text:
            return 0
        
        for pattern, replacement in HARD_REDACT_PATTERNS:
            iterations = 0
            while iterations < MAX_ITERATIONS:
                iterations += 1
                full_text = ''.join(run.text or '' for run in runs)
                match = pattern.search(full_text)
                if not match:
                    break
                
                original = match.group(0)
                start_pos = match.start()
                end_pos = match.end()
                
                run_map = []
                pos = 0
                for idx, run in enumerate(runs):
                    run_text = run.text or ''
                    run_map.append((pos, pos + len(run_text), idx, run))
                    pos += len(run_text)
                
                start_runs = [(s, e, i, r) for s, e, i, r in run_map if s <= start_pos < e]
                end_runs = [(s, e, i, r) for s, e, i, r in run_map if s < end_pos <= e]
                
                if not start_runs or not end_runs:
                    break
                
                start_run_info = start_runs[0]
                end_run_info = end_runs[0]
                
                if start_run_info[2] == end_run_info[2]:
                    run = start_run_info[3]
                    run_start = start_run_info[0]
                    local_start = start_pos - run_start
                    local_end = end_pos - run_start
                    run.text = run.text[:local_start] + replacement + run.text[local_end:]
                    fixes += 1
                else:
                    first_run = start_run_info[3]
                    first_run_start = start_run_info[0]
                    local_start = start_pos - first_run_start
                    first_run.text = first_run.text[:local_start] + replacement
                    
                    for s, e, i, r in run_map:
                        if start_run_info[2] < i < end_run_info[2]:
                            r.text = ''
                    
                    last_run = end_run_info[3]
                    last_run_start = end_run_info[0]
                    local_end = end_pos - last_run_start
                    last_run.text = last_run.text[local_end:]
                    fixes += 1
        
        return fixes
    
    for para in doc.paragraphs:
        total_fixes += redact_paragraph_run_aware(para)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total_fixes += redact_paragraph_run_aware(para)
    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    total_fixes += redact_paragraph_run_aware(para)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    total_fixes += redact_paragraph_run_aware(para)
    
    return total_fixes
