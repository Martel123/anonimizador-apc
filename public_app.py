"""
Anonimizador Legal - Blueprint con autenticación y créditos por páginas
=======================================================================
Flujo: LOGIN → SUBIR → CONTAR PÁGINAS → RESERVAR CRÉDITOS → REVIEW → APLICAR → COBRAR → DESCARGAR
"""

import os
import re
import uuid
import json
import logging
import tempfile
import traceback
import zipfile
import hashlib
import secrets
from io import BytesIO
from datetime import datetime, timedelta
from collections import defaultdict
from functools import lru_cache
from flask import Blueprint, render_template, request, send_file, jsonify, redirect, url_for, flash
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    from final_auditor import audit_document, log_audit_result
    FINAL_AUDITOR_AVAILABLE = True
except ImportError:
    FINAL_AUDITOR_AVAILABLE = False
    logger.warning("Final auditor not available")

from credit_utils import (
    get_or_create_credits, ensure_trial, count_pages,
    check_and_reserve_pages, charge_pages, release_reservation
)

anonymizer_bp = Blueprint("anonymizer", __name__)

ALLOWED_EXTENSIONS = {'doc', 'docx', 'pdf', 'txt'}

# ============================================================================
# UTILIDADES
# ============================================================================

@lru_cache(maxsize=1)
def check_openai_available():
    """Verifica que OPENAI_API_KEY esté configurada. OBLIGATORIO por negocio."""
    key = os.environ.get("OPENAI_API_KEY", "").strip()
    return bool(key and len(key) > 10)


def safe_remove(path):
    """Elimina archivo de forma segura."""
    if path and os.path.exists(path):
        try:
            os.remove(path)
            logger.info(f"CLEANUP | removed={path}")
        except Exception as e:
            logger.warning(f"CLEANUP_FAIL | path={path} | error={e}")


def allowed_file(filename):
    if not filename or '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext in ALLOWED_EXTENSIONS


def get_extension(filename):
    if '.' in filename:
        return filename.rsplit('.', 1)[1].lower()
    return ''


def get_output_extension(input_ext):
    """
    Normaliza extensión de salida:
    - DOCX → DOCX
    - TXT → TXT
    - PDF → TXT (texto plano)
    - DOC → TXT (texto plano)
    """
    if input_ext == 'docx':
        return 'docx'
    elif input_ext == 'txt':
        return 'txt'
    else:
        return 'txt'


def validate_file_format(file_path, ext):
    """Valida formato real del archivo."""
    try:
        with open(file_path, 'rb') as f:
            header = f.read(16)
        
        if ext == 'pdf':
            if not header.startswith(b'%PDF'):
                return False, "El archivo no parece ser un PDF válido"
        
        if ext == 'docx':
            if not zipfile.is_zipfile(file_path):
                return False, "El archivo no parece ser un DOCX válido"
        
        return True, None
    except Exception as e:
        return False, f"Error validando archivo: {str(e)}"


def render_error(message, status_code=400):
    """Renderiza página de error limpia sin stacktrace."""
    return render_template("anonymizer_standalone.html",
                           error=message,
                           openai_available=check_openai_available()), status_code


# ============================================================================
# NORMALIZACIÓN DE ENTIDADES CON CANDIDATES
# ============================================================================

def normalize_entity(ent_dict):
    """
    Normaliza una entidad y genera candidates.
    Returns dict con: value, candidates, type, confidence, source
    """
    value_base = ent_dict.get('value') or ent_dict.get('text') or ''
    if not value_base:
        return None
    
    original = value_base
    
    normalized = value_base.strip()
    normalized = re.sub(r'[\r\n\t]+', ' ', normalized)
    normalized = re.sub(r'\s+', ' ', normalized)
    
    no_newlines = re.sub(r'\n+', ' ', original)
    no_newlines = re.sub(r'\s+', ' ', no_newlines).strip()
    
    ent_type = ent_dict.get('type') or ent_dict.get('entity_type') or ''
    # Soft types: sólo valor exacto, sin expansión de candidates.
    # Debe coincidir con SOFT_MATCH_TYPES y ALWAYS_REVIEW_TYPES.
    SOFT_TYPES = {
        'PERSONA', 'ENTIDAD', 'DIRECCION', 'PLACA',
        'RESOLUCION', 'PARTIDA', 'JUZGADO', 'SALA', 'TRIBUNAL',
    }
    include_no_spaces = ent_type.upper() not in SOFT_TYPES
    no_spaces = normalized.replace(' ', '') if (len(normalized) >= 4 and include_no_spaces) else None

    candidates = []
    seen_lower = set()

    for c in [original, normalized, no_newlines, no_spaces]:
        if c and len(c) >= 4:
            c_lower = c.lower()
            if c_lower not in seen_lower:
                seen_lower.add(c_lower)
                candidates.append(c)
    
    if not candidates:
        return None
    
    return {
        'value': normalized,
        'text': normalized,
        'original': original,
        'candidates': candidates,
        'type': ent_dict.get('type') or ent_dict.get('entity_type') or 'UNKNOWN',
        'confidence': float(ent_dict.get('confidence', 1.0)),
        'source': ent_dict.get('source', 'detector'),
        'start': ent_dict.get('start', 0),
        'end': ent_dict.get('end', 0)
    }


def normalize_entities(items):
    """Normaliza lista de entidades, agrega candidates."""
    if not items:
        return []
    result = []
    for e in items:
        if isinstance(e, dict):
            normalized = normalize_entity(e)
        else:
            d = {
                'type': getattr(e, 'type', 'UNKNOWN'),
                'value': getattr(e, 'value', ''),
                'start': getattr(e, 'start', 0),
                'end': getattr(e, 'end', 0),
                'confidence': getattr(e, 'confidence', 1.0),
                'source': getattr(e, 'source', 'detector')
            }
            normalized = normalize_entity(d)
        
        if normalized:
            result.append(normalized)
    return result


def deduplicate_entities(entities):
    """Deduplica entidades por (type, value.lower())."""
    seen = set()
    result = []
    for e in entities:
        key = (e['type'].upper(), e['value'].lower())
        if key not in seen:
            seen.add(key)
            result.append(e)
    return result


# ============================================================================
# EXTRACCIÓN DE TEXTO
# ============================================================================

def extract_full_text_docx(doc):
    """Extrae todo el texto del documento DOCX incluyendo tablas, headers, footers."""
    text_parts = []
    
    for para in doc.paragraphs:
        text_parts.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_parts.append(para.text)
    
    try:
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    text_parts.append(para.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    text_parts.append(para.text)
    except:
        pass
    
    return '\n'.join(text_parts)


def extract_text(file_path, ext):
    """Extrae texto de un archivo."""
    if ext == 'docx':
        from docx import Document
        doc = Document(file_path)
        return extract_full_text_docx(doc)
    
    elif ext == 'pdf':
        from processor_pdf import extract_text_pdf
        result = extract_text_pdf(file_path)
        if result.get('success'):
            return result.get('text', '')
        raise ValueError(result.get('error', 'Error extrayendo texto del PDF'))
    
    elif ext == 'txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    elif ext == 'doc':
        raise ValueError("Formato DOC no soportado directamente. Por favor convierta a DOCX.")
    
    return ''


# ============================================================================
# APLICACIÓN DE ANONIMIZACIÓN
# ============================================================================

def expand_entities_with_candidates(entity_dicts):
    """
    Expande entidades con todos sus candidates para reemplazo MUY ALTO.
    Cada candidate genera una Entity separada.
    """
    from detector_capas import Entity
    
    entities = []
    seen = set()
    
    SOFT_EXPAND = {'PERSONA', 'ENTIDAD', 'DIRECCION', 'PLACA'}

    for d in entity_dicts:
        ent_type = d.get('type', 'UNKNOWN')
        confidence = d.get('confidence', 1.0)
        source = d.get('source', 'detector')

        candidates = d.get('candidates', [])
        value = d.get('value', '')

        if ent_type.upper() in SOFT_EXPAND:
            all_values = {value} if value and len(value) >= 2 else set()
        else:
            all_values = set(candidates) if candidates else set()
            if value and len(value) >= 4:
                all_values.add(value)
        
        for candidate in all_values:
            if not candidate or len(candidate) < 4:
                continue
            
            key = (ent_type.lower(), candidate.lower())
            if key in seen:
                continue
            seen.add(key)
            
            entities.append(Entity(
                type=ent_type,
                value=candidate,
                start=0,
                end=len(candidate),
                source=source,
                confidence=confidence
            ))
    
    return entities


# Tipos que requieren matching estricto (bordes de palabra, longitud mínima)
# SOFT_MATCH_TYPES: tipos ambiguos que requieren matching estricto en apply.
# Deben coincidir exactamente con ALWAYS_REVIEW_TYPES para consistencia.
SOFT_MATCH_TYPES = {
    'PERSONA', 'ENTIDAD', 'DIRECCION', 'PLACA',
    'RESOLUCION', 'PARTIDA', 'JUZGADO', 'SALA', 'TRIBUNAL',
}
# Prioridad de aplicación: los tipos estructurados primero para proteger emails, RUC, etc.
# Orden de aplicación: estructurados primero para proteger sus valores.
# Prioridad 0 = se aplica primero (AUTO); 10 = se aplica último (REVIEW/SOFT).
TYPE_APPLY_PRIORITY = {
    # AUTO - structured PII (confianza alta, sin ambigüedad)
    'EMAIL': 0, 'DNI': 0, 'RUC': 0, 'TELEFONO': 0,
    'CUENTA': 0, 'CCI': 0, 'COLEGIATURA': 0, 'CASILLA': 0,
    'EXPEDIENTE': 1, 'ACTA': 1,
    # REVIEW - ambiguous / soft types (requieren confirmación explícita)
    'PLACA': 10, 'PARTIDA': 10, 'RESOLUCION': 10,
    'DIRECCION': 11, 'JUZGADO': 11, 'SALA': 11, 'TRIBUNAL': 11,
    'ENTIDAD': 12, 'PERSONA': 13,
}


def apply_entities_to_docx(input_path, output_path, entity_dicts):
    """
    Aplica anonimización a DOCX con soporte para tokens predefinidos (manual entities).
    Usa reemplazo run-aware para manejar texto partido.
    - Tipos suaves (PERSONA/ENTIDAD/DIRECCION/PLACA): word-boundary estricto,
      longitud mínima, solo valor exacto (sin candidatos expandidos).
    - EMAIL/DNI/RUC/…: se aplican primero para proteger sus valores.
    """
    from docx import Document

    doc = Document(input_path)

    replacements = []
    reverse_mapping = {}
    type_counters = {}
    value_to_token = {}

    for d in entity_dicts:
        ent_type = d.get('type', 'UNKNOWN')
        value = d.get('value', '')
        token = d.get('token', '')
        replace_all = d.get('replace_all', True)
        candidates = d.get('candidates', [])
        soft = ent_type.upper() in SOFT_MATCH_TYPES

        if not value:
            continue

        # ── Filtro de longitud/estructura para tipos suaves ──────────────────
        # Soft types: sólo valor exacto; no se expanden candidatos.
        # Filtros estructurales adicionales para evitar replacements ambiguos.
        if soft:
            v_stripped = value.strip()
            if len(v_stripped) < 4:
                continue
            # PERSONA, ENTIDAD, JUZGADO, SALA, TRIBUNAL de una sola palabra
            # son demasiado ambiguos: se exige al menos 2 tokens.
            if ent_type.upper() in ('PERSONA', 'ENTIDAD', 'JUZGADO', 'SALA', 'TRIBUNAL')                     and ' ' not in v_stripped:
                continue
            # DIRECCION muy corta: descartar (probablemente captura parcial)
            if ent_type.upper() == 'DIRECCION' and len(v_stripped) < 8:
                continue
            all_values = {v_stripped}
        else:
            all_values = set()
            if value:
                all_values.add(value)
            if candidates:
                all_values.update(candidates)

        for v in all_values:
            if not v or len(v) < 2:
                continue

            normalized = v.strip().lower()
            key = f"{ent_type}|{normalized}"

            if key in value_to_token:
                t = value_to_token[key]
            elif token:
                t = token
                value_to_token[key] = t
            else:
                type_counters[ent_type] = type_counters.get(ent_type, 0) + 1
                t = f"{{{{{ent_type}_{type_counters[ent_type]}}}}}"
                value_to_token[key] = t

            replacements.append((v, t, replace_all, soft, ent_type.upper()))

            if t not in reverse_mapping:
                masked = v[:3] + '...' + v[-2:] if len(v) > 8 else v[:2] + '***'
                reverse_mapping[t] = masked

    # Estructurados primero (priority 0), luego por longitud desc
    # Tupla: (value, token, replace_all, soft, ent_type)
    replacements.sort(key=lambda x: (int(x[3]) * 10, -len(x[0])))

    replaced_count = apply_replacements_to_docx(doc, replacements)

    doc.save(output_path)

    return replaced_count, reverse_mapping


def _is_word_char(c):
    """Retorna True si el carácter es alfanumérico o guión bajo (parte de una palabra)."""
    return c.isalnum() or c == '_'


def apply_replacements_to_docx(doc, replacements):
    """
    Aplica lista de reemplazos (value, token, replace_all, soft) al documento DOCX.
    - soft=True: verifica bordes de palabra antes de reemplazar (evita sub-palabras).
    - Maneja párrafos, tablas, headers y footers con reemplazo run-aware.
    """
    total_count = 0
    done_once = set()

    def replace_in_paragraph(paragraph, replacements):
        import re as _re_inner
        nonlocal done_once
        count = 0
        for original, token, replace_all, soft, _ent_type in replacements:
            if not replace_all and original in done_once:
                continue
            full_text = paragraph.text
            if original not in full_text:
                continue

            run_map = []
            pos = 0
            for idx, run in enumerate(paragraph.runs):
                run_text = run.text
                run_map.append((pos, pos + len(run_text), idx, run))
                pos += len(run_text)

            start = 0
            iterations = 0
            max_iterations = 100

            while iterations < max_iterations:
                iterations += 1
                full_text = paragraph.text
                # EMAIL: búsqueda case-insensitive porque el mismo correo puede
                # aparecer en distintas capitalizaciones en el documento.
                if _ent_type == 'EMAIL':
                    _ci_m = _re_inner.search(_re_inner.escape(original), full_text[start:], _re_inner.IGNORECASE)
                    if _ci_m is None:
                        break
                    idx = start + _ci_m.start()
                    original = full_text[idx:idx + len(original)]  # use exact case from doc
                else:
                    idx = full_text.find(original, start)
                    if idx == -1:
                        break

                end_idx = idx + len(original)

                # ── Word-boundary check para tipos suaves ────────────────────
                if soft:
                    char_before = full_text[idx - 1] if idx > 0 else ' '
                    char_after = full_text[end_idx] if end_idx < len(full_text) else ' '
                    if _is_word_char(char_before) or _is_word_char(char_after):
                        start = idx + 1
                        continue

                affected_runs = []
                for run_start, run_end, run_idx, run in run_map:
                    if run_start < end_idx and run_end > idx:
                        affected_runs.append((run_start, run_end, run_idx, run))

                if not affected_runs:
                    start = idx + 1
                    continue

                first_run = affected_runs[0]
                first_run_obj = first_run[3]
                local_start = idx - first_run[0]

                if len(affected_runs) == 1:
                    local_end = local_start + len(original)
                    old_text = first_run_obj.text
                    first_run_obj.text = old_text[:local_start] + token + old_text[local_end:]
                else:
                    old_text = first_run_obj.text
                    first_run_obj.text = old_text[:local_start] + token

                    for _, _, _, run in affected_runs[1:-1]:
                        run.text = ''

                    last_run = affected_runs[-1]
                    last_run_obj = last_run[3]
                    local_end_in_last = end_idx - last_run[0]
                    last_run_obj.text = last_run_obj.text[local_end_in_last:]

                count += 1

                if not replace_all:
                    done_once.add(original)
                    break

                run_map = []
                pos = 0
                for r_idx, run in enumerate(paragraph.runs):
                    run_map.append((pos, pos + len(run.text), r_idx, run))
                    pos += len(run.text)
                start = 0

        return count
    
    for para in doc.paragraphs:
        total_count += replace_in_paragraph(para, replacements)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total_count += replace_in_paragraph(para, replacements)
    
    try:
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    total_count += replace_in_paragraph(para, replacements)
            if section.footer:
                for para in section.footer.paragraphs:
                    total_count += replace_in_paragraph(para, replacements)
    except Exception as e:
        logger.warning(f"Error processing headers/footers: {e}")
    
    return total_count


def apply_entities_to_text(input_path, output_path, entity_dicts, ext='txt'):
    """
    Aplica anonimización a texto plano con soporte para tokens predefinidos.
    - Tipos suaves (PERSONA/ENTIDAD/DIRECCION/PLACA): word-boundary estricto y
      filtro de longitud mínima.
    - EMAIL/DNI/RUC/…: se aplican primero.
    """
    import re as _re

    text = extract_text(input_path, ext)

    type_counters = {}
    value_to_token = {}
    reverse_mapping = {}
    all_replacements = []

    for d in entity_dicts:
        ent_type = d.get('type', 'UNKNOWN')
        value = d.get('value', '')
        token = d.get('token', '')
        replace_all = d.get('replace_all', True)
        candidates = d.get('candidates', [])
        soft = ent_type.upper() in SOFT_MATCH_TYPES

        if not value:
            continue

        # ── Filtro longitud/estructura tipos suaves ──────────────────────────
        if soft:
            v_stripped = value.strip()
            if len(v_stripped) < 4:
                continue
            if ent_type.upper() in ('PERSONA', 'ENTIDAD', 'JUZGADO', 'SALA', 'TRIBUNAL') \
                    and ' ' not in v_stripped:
                continue
            if ent_type.upper() == 'DIRECCION' and len(v_stripped) < 8:
                continue
            all_values = {v_stripped}
        else:
            all_values = set()
            if value:
                all_values.add(value)
            if candidates:
                all_values.update(candidates)

        for v in all_values:
            if not v or len(v) < 2:
                continue

            normalized = v.strip().lower()
            key = f"{ent_type}|{normalized}"

            if key in value_to_token:
                t = value_to_token[key]
            elif token:
                t = token
                value_to_token[key] = t
            else:
                type_counters[ent_type] = type_counters.get(ent_type, 0) + 1
                t = f"{{{{{ent_type}_{type_counters[ent_type]}}}}}"
                value_to_token[key] = t

            all_replacements.append((v, t, replace_all, soft, ent_type.upper()))

            if t not in reverse_mapping:
                masked = v[:3] + '...' + v[-2:] if len(v) > 8 else v[:2] + '***'
                reverse_mapping[t] = masked

    # Estructurados primero, luego por longitud desc
    all_replacements.sort(key=lambda x: (int(x[3]) * 10, -len(x[0])))

    replaced_count = 0
    for value, token, replace_all, soft, ent_type in all_replacements:
        if soft:
            # Word-boundary estricto: excluye guión y caracteres latinos acentuados
            # (?<![...]) = no precedido por ninguno de esos chars
            _WB_INNER = r'A-Za-z0-9_À-ɏ\-'
            pattern = r'(?<![' + _WB_INNER + r'])' + _re.escape(value) + r'(?![' + _WB_INNER + r'])'
            try:
                if replace_all:
                    new_text, n = _re.subn(pattern, token, text)
                else:
                    new_text, n = _re.subn(pattern, token, text, count=1)
                replaced_count += n
                text = new_text
            except _re.error:
                pass  # Si el valor tiene chars especiales, lo saltamos
        else:
            # EMAIL usa reemplazo case-insensitive porque el mismo correo puede
            # aparecer en distintas capitalizaciones en el documento (ej. PDF
            # que exporta dominios institucionales en mayúsculas).
            if ent_type == 'EMAIL':
                try:
                    ci_pattern = _re.compile(_re.escape(value), _re.IGNORECASE)
                    if replace_all:
                        new_text, n = ci_pattern.subn(lambda m: token, text)
                    else:
                        new_text, n = ci_pattern.subn(lambda m: token, text, count=1)
                    if n > 0:
                        replaced_count += n
                        text = new_text
                except _re.error:
                    pass
            else:
                if value in text:
                    if replace_all:
                        count = text.count(value)
                        text = text.replace(value, token)
                        replaced_count += count
                    else:
                        text = text.replace(value, token, 1)
                        replaced_count += 1

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

    return replaced_count, reverse_mapping


# ============================================================================
# ENDPOINTS
# ============================================================================

@anonymizer_bp.route("/health")
def health():
    return "ok"


@anonymizer_bp.route("/")
@login_required
def index():
    openai_available = check_openai_available()
    from detector_openai import USE_AI_SEMANTIC_FILTER
    credits = ensure_trial(current_user.id)
    return render_template("anonymizer_standalone.html",
                           openai_available=openai_available,
                           ai_semantic_active=USE_AI_SEMANTIC_FILTER,
                           credits=credits)


@anonymizer_bp.route("/anonymizer")
@login_required
def anonymizer_home():
    openai_available = check_openai_available()
    from detector_openai import USE_AI_SEMANTIC_FILTER
    credits = ensure_trial(current_user.id)
    return render_template("anonymizer_standalone.html",
                           openai_available=openai_available,
                           ai_semantic_active=USE_AI_SEMANTIC_FILTER,
                           credits=credits)


@anonymizer_bp.route("/anonymizer/onboarding", methods=["GET"])
@login_required
def anonymizer_onboarding():
    """Onboarding post-registro para usuarios del anonimizador."""
    from credit_utils import get_or_create_credits
    credits = get_or_create_credits(current_user.id)
    trial_claimed = credits.trial_granted_at is not None
    return render_template("anonymizer_onboarding.html", trial_claimed=trial_claimed)


@anonymizer_bp.route("/anonymizer/onboarding/free", methods=["POST"])
@login_required
def anonymizer_onboarding_free():
    """Aplica el trial gratuito y redirige al anonimizador."""
    ensure_trial(current_user.id)
    return redirect(url_for('anonymizer.anonymizer_home'))


@anonymizer_bp.route("/account")
@login_required
def account():
    """Página de cuenta del usuario: saldo, compras y consumos."""
    from models import AnonymizerPurchase, PageUsageLog
    from credit_utils import get_or_create_credits

    credits = get_or_create_credits(current_user.id)
    purchases = AnonymizerPurchase.query.filter_by(user_id=current_user.id)\
        .order_by(AnonymizerPurchase.created_at.desc()).limit(10).all()
    usage_logs = PageUsageLog.query.filter_by(user_id=current_user.id)\
        .order_by(PageUsageLog.created_at.desc()).limit(10).all()

    balance = credits.pages_balance
    if balance < 20:
        suggested_package_pages = 300
    elif balance <= 200:
        suggested_package_pages = 800
    else:
        suggested_package_pages = 2000

    return render_template("account.html",
        credits=credits,
        purchases=purchases,
        usage_logs=usage_logs,
        suggested_package_pages=suggested_package_pages,
        user_email=getattr(current_user, 'email', '') or '',
    )


@anonymizer_bp.route("/account/security")
@login_required
def account_security():
    """Historial de los últimos 10 logins del usuario."""
    from models import LoginAttempt
    logs = LoginAttempt.query.filter_by(email=current_user.email)\
        .order_by(LoginAttempt.created_at.desc()).limit(10).all()
    return render_template("account_security.html", login_logs=logs)


@anonymizer_bp.route("/anonymizer/process", methods=["GET", "POST"])
@login_required
def anonymizer_process():
    """
    Procesa archivo: contar páginas, reservar créditos, detectar PII, mostrar review.
    """
    from models import db, AnonymizerJob

    if request.method == "GET":
        return redirect(url_for('anonymizer.anonymizer_home'))

    try:
        credits = ensure_trial(current_user.id)
    except Exception as e:
        logger.error(f"ENSURE_TRIAL_ERROR | user={current_user.id} | error={e}")
        credits = None

    if 'file' not in request.files:
        return render_error("No se seleccionó ningún archivo")

    file = request.files['file']
    if not file or not file.filename:
        return render_error("El archivo está vacío")

    filename = secure_filename(file.filename)
    ext = get_extension(filename)

    if not allowed_file(filename):
        return render_error(f"Formato no soportado: .{ext}. Use DOCX, PDF o TXT")

    strict_mode = request.form.get('strict_mode', 'true').lower() == 'true'
    export_csv = request.form.get('export_csv', 'false').lower() == 'true'

    job_id = str(uuid.uuid4())
    temp_input = os.path.join(tempfile.gettempdir(), f"in_{job_id}_{filename}")

    try:
        file.save(temp_input)
        file_size = os.path.getsize(temp_input)
        logger.info(f"UPLOAD | job={job_id} | user={current_user.id} | file={filename} | ext={ext} | size={file_size}")

        valid, error_msg = validate_file_format(temp_input, ext)
        if not valid:
            safe_remove(temp_input)
            return render_error(error_msg)

        pages_needed = count_pages(temp_input, ext)

        job = AnonymizerJob(
            job_id=job_id,
            user_id=current_user.id,
            filename_original=filename,
            ext=ext,
            pages_counted=pages_needed,
            status='created',
            input_path=temp_input
        )
        db.session.add(job)
        db.session.commit()

        can_reserve, current_balance = check_and_reserve_pages(current_user.id, job_id, pages_needed)

        if not can_reserve:
            job.status = 'blocked'
            db.session.commit()
            safe_remove(temp_input)
            pages_missing = pages_needed - current_balance
            page_method = "páginas reales" if ext == "pdf" else "equivalente por contenido (500 palabras = 1 página)"
            return render_template("insufficient_pages.html",
                pages_needed=pages_needed,
                pages_balance=current_balance,
                pages_missing=pages_missing,
                filename=filename,
                ext=ext,
                page_method=page_method
            )

        job.status = 'reserved'
        db.session.commit()

        full_text = extract_text(temp_input, ext)

        if not full_text or len(full_text.strip()) < 10:
            release_reservation(current_user.id, job_id)
            job.status = 'failed'
            db.session.commit()
            safe_remove(temp_input)
            return render_error("No se pudo leer el contenido del documento")

        from detector_capas import detect_all_pii
        entities, detect_meta = detect_all_pii(full_text)
        all_entities = normalize_entities(entities)
        all_entities = deduplicate_entities(all_entities)

        # ── Capa IA semántica (validador de candidatos ambiguos) ─────────────
        # Solo activa si USE_AI_SEMANTIC_FILTER=1.
        # NO envía el documento completo: recibe candidatos ya detectados,
        # clasifica los ambiguos con UN llamado a la API y descarta falsos positivos.
        # PII estructurada (DNI, RUC, EMAIL, etc.) no pasa por este filtro.
        try:
            from detector_openai import (
                USE_AI_SEMANTIC_FILTER, is_openai_available,
                validate_ambiguous_candidates,
            )
            if USE_AI_SEMANTIC_FILTER and is_openai_available():
                pre_filter_count = len(all_entities)
                all_entities = validate_ambiguous_candidates(all_entities, full_text)
                all_entities = deduplicate_entities(all_entities)
                logger.info(
                    f"AI_SEMANTIC_FILTER | job={job_id} "
                    f"| before={pre_filter_count} | after={len(all_entities)}"
                )
        except Exception as e:
            logger.warning(f"AI_SEMANTIC_FILTER_FAIL | job={job_id} | error={str(e)}")

        confirmed = []
        needs_review = []

        # ══════════════════════════════════════════════════════════════════════
        # POLÍTICA FINAL DE CLASIFICACIÓN POR TIPO
        # ──────────────────────────────────────────────────────────────────────
        # AUTO (se aplican automáticamente si confidence >= 0.80):
        #   DNI, RUC, EMAIL, TELEFONO, CASILLA, COLEGIATURA,
        #   CUENTA/CCI (con contexto), EXPEDIENTE (con trigger), ACTA (con trigger)
        #
        # ALWAYS REVIEW (jamás se aplican sin confirmación explícita del usuario,
        #   independientemente de la confidence):
        #   PERSONA, ENTIDAD, DIRECCION, PLACA,
        #   RESOLUCION, PARTIDA, JUZGADO, SALA, TRIBUNAL
        # ══════════════════════════════════════════════════════════════════════
        ALWAYS_REVIEW_TYPES = {
            'PERSONA', 'ENTIDAD', 'DIRECCION', 'PLACA',
            'RESOLUCION', 'PARTIDA', 'JUZGADO', 'SALA', 'TRIBUNAL',
        }

        for i, ent in enumerate(all_entities):
            ent['index'] = i
            conf = ent.get('confidence', 1.0)
            ent_type = ent.get('type', '').upper()
            if ent_type in ALWAYS_REVIEW_TYPES:
                ent['status'] = 'needs_review'
                needs_review.append(ent)
            elif conf >= 0.80:
                ent['status'] = 'confirmed'
                confirmed.append(ent)
            elif conf >= 0.50:
                ent['status'] = 'needs_review'
                needs_review.append(ent)

        entities_all = confirmed + needs_review

        confirmed_by_type = {}
        needs_review_by_type = {}

        for e in confirmed:
            t = e['type']
            if t not in confirmed_by_type:
                confirmed_by_type[t] = []
            confirmed_by_type[t].append(e)

        for e in needs_review:
            t = e['type']
            if t not in needs_review_by_type:
                needs_review_by_type[t] = []
            needs_review_by_type[t].append(e)

        job.status = 'reviewed'
        db.session.commit()

        logger.info(f"DETECT_OK | job={job_id} | confirmed={len(confirmed)} | needs_review={len(needs_review)}")

        credits_refreshed = get_or_create_credits(current_user.id)

        return render_template("anonymizer_review.html",
            ext=ext,
            original_filename=filename,
            job_id=job_id,
            full_text=full_text,
            preview_text=full_text,
            confirmed=confirmed,
            needs_review=needs_review,
            confirmed_by_type=confirmed_by_type,
            needs_review_by_type=needs_review_by_type,
            entities_all=entities_all,
            confirmed_count=len(confirmed),
            needs_review_count=len(needs_review),
            strict_mode=strict_mode,
            export_csv=export_csv,
            pages_counted=pages_needed,
            pages_balance=credits_refreshed.pages_balance
        )

    except Exception as e:
        try:
            release_reservation(current_user.id, job_id)
            job_rec = AnonymizerJob.query.filter_by(job_id=job_id).first()
            if job_rec and job_rec.status not in ('success', 'charged'):
                job_rec.status = 'failed'
                db.session.commit()
        except Exception:
            pass
        safe_remove(temp_input)
        logger.error(f"PROCESS_ERROR | job={job_id} | error={e}")
        logger.error(traceback.format_exc())
        return render_error("No se pudo procesar el documento. Verifique el formato e intente nuevamente.")


def truncate_value(value, max_len=10):
    """Truncate value for display: first 3 chars ... last 3 chars"""
    if len(value) <= max_len:
        return value
    return value[:3] + "..." + value[-3:]

def get_result_paths(job_id):
    """Get file paths for storing results."""
    base = os.path.join(tempfile.gettempdir(), f"result_{job_id}")
    return {
        'doc': f"{base}.doc",
        'meta': f"{base}.meta.json"
    }

@anonymizer_bp.route("/anonymizer/apply", methods=["POST"])
@login_required
def anonymizer_apply():
    """
    Aplica anonimización, cobra créditos y muestra resultados.
    """
    import html as html_lib
    from models import db, AnonymizerJob

    ext = request.form.get('ext', 'docx')
    original_filename = request.form.get('original_filename', 'documento')
    selected_entities_json = request.form.get('selected_entities_json', '[]')
    export_csv = request.form.get('export_csv', 'false').lower() == 'true'
    job_id = request.form.get('job_id', '')

    if not job_id:
        return render_error("Sesión inválida. Suba el documento nuevamente.")

    job = AnonymizerJob.query.filter_by(job_id=job_id).first()
    if not job or job.user_id != current_user.id:
        logger.warning(f"APPLY_OWNERSHIP_FAIL | job={job_id} | user={current_user.id}")
        return render_error("No tiene permiso para procesar este documento.", 403)

    if job.pages_charged > 0 or job.status == 'success':
        logger.warning(f"APPLY_ALREADY_CHARGED | job={job_id} | user={current_user.id}")
        return render_error("Este documento ya fue procesado. Suba uno nuevo.")

    from models import PageReservation
    from credit_utils import is_unlimited_user
    _unlimited = is_unlimited_user(current_user.id)
    reservation = PageReservation.query.filter_by(job_id=job_id, user_id=current_user.id).first()
    if not _unlimited:
        if not reservation or reservation.status != 'reserved':
            logger.warning(f"APPLY_NO_RESERVATION | job={job_id} | user={current_user.id} | reservation={reservation.status if reservation else 'none'}")
            return render_error("Reserva de créditos no encontrada. Suba el documento nuevamente.")
        if reservation.pages_reserved != job.pages_counted:
            logger.warning(f"APPLY_PAGES_MISMATCH | job={job_id} | reserved={reservation.pages_reserved} | counted={job.pages_counted}")
            release_reservation(current_user.id, job_id)
            job.status = 'failed'
            db.session.commit()
            return render_error("Error de consistencia en créditos. Suba el documento nuevamente.")

    # Obtener ruta del input desde DB (nunca del cliente)
    temp_input = job.input_path or ''
    _tmpdir = os.path.realpath(tempfile.gettempdir())
    _input_real = os.path.realpath(temp_input) if temp_input else ''
    if not temp_input or not _input_real.startswith(_tmpdir) or not os.path.exists(temp_input):
        logger.warning(f"APPLY_INPUT_MISSING | job={job_id} | path={temp_input!r}")
        release_reservation(current_user.id, job_id)
        job.status = 'failed'
        db.session.commit()
        return render_error("Sesión expirada. Suba el documento nuevamente.")

    output_ext = get_output_extension(ext)
    temp_output = os.path.join(tempfile.gettempdir(), f"out_{job_id}.{output_ext}")
    
    logger.info(f"APPLY_START | job={job_id} | ext={ext} | output_ext={output_ext}")
    
    try:
        selected_entities_json = html_lib.unescape(selected_entities_json)
        selected_entities = json.loads(selected_entities_json)
        
        entity_count = len(selected_entities) if selected_entities else 0
        logger.info(f"APPLY_START | job={job_id} | entities={entity_count}")
        
        if not selected_entities:
            safe_remove(temp_input)
            release_reservation(current_user.id, job_id)
            job.status = 'failed'
            db.session.commit()
            return render_error("No se seleccionaron entidades para anonimizar.")
        
        if ext == 'docx':
            replaced_count, mapping = apply_entities_to_docx(temp_input, temp_output, selected_entities)
        else:
            replaced_count, mapping = apply_entities_to_text(temp_input, temp_output, selected_entities, ext)
        
        if not os.path.exists(temp_output):
            logger.error(f"APPLY_FAIL | job={job_id} | reason=output_not_created")
            release_reservation(current_user.id, job_id)
            job.status = 'failed'
            db.session.commit()
            return render_error("No se pudo generar el archivo final. Intente nuevamente.")

        output_size = os.path.getsize(temp_output)
        if output_size == 0:
            logger.error(f"APPLY_FAIL | job={job_id} | reason=output_empty")
            safe_remove(temp_output)
            release_reservation(current_user.id, job_id)
            job.status = 'failed'
            db.session.commit()
            return render_error("No se pudo generar el archivo final. Intente nuevamente.")
        
        logger.info(f"APPLY_DONE | job={job_id} | replaced={replaced_count} | output_size={output_size}")
        
        # ETAPA 8: AUDITOR FINAL OBLIGATORIO - Garantizar 0 fugas
        post_scan_text = ""
        if output_ext == 'docx':
            from docx import Document
            doc_check = Document(temp_output)
            post_scan_text = extract_full_text_docx(doc_check)
        else:
            with open(temp_output, 'r', encoding='utf-8', errors='ignore') as f:
                post_scan_text = f.read()
        
        residual_warning = None
        
        if FINAL_AUDITOR_AVAILABLE:
            # Usar auditor con auto-corrección
            existing_counters = {}
            for token in mapping.keys():
                match = re.match(r'\{\{([A-Z]+)_(\d+)\}\}', token)
                if match:
                    etype, num = match.groups()
                    existing_counters[etype] = max(existing_counters.get(etype, 0), int(num))
            
            audit_result = audit_document(post_scan_text, auto_fix=True, existing_counters=existing_counters)
            log_audit_result(audit_result)
            
            if audit_result.leaks_found:
                logger.warning(f"AUDIT | job={job_id} | leaks_found={len(audit_result.leaks_found)} | auto_fixed={audit_result.leaks_auto_fixed}")
            
            # PERSISTIR AUTO-FIXES: Escribir texto corregido al archivo de salida
            if output_ext == 'docx':
                from docx import Document
                from processor_docx import apply_replacements_to_docx
                
                MAX_ITERATIONS = 2
                current_iteration = 0
                
                while current_iteration < MAX_ITERATIONS:
                    current_iteration += 1
                    
                    if not audit_result.replacements:
                        break
                    
                    logger.info(f"AUDIT_AUTOFIX | job={job_id} | iteration={current_iteration} | applying {len(audit_result.replacements)} replacements")
                    
                    doc_fix = Document(temp_output)
                    fixes_applied = apply_replacements_to_docx(doc_fix, audit_result.replacements)
                    doc_fix.save(temp_output)
                    logger.info(f"AUDIT_AUTOFIX_DOCX | job={job_id} | iteration={current_iteration} | fixes_applied={fixes_applied}")
                    
                    doc_recheck = Document(temp_output)
                    recheck_text = extract_full_text_docx(doc_recheck)
                    audit_result = audit_document(recheck_text, auto_fix=True, existing_counters=existing_counters)
                    
                    if audit_result.is_safe:
                        logger.info(f"AUDIT_RECHECK_PASSED | job={job_id} | iteration={current_iteration} | document is safe")
                        break
                    else:
                        logger.warning(f"AUDIT_RECHECK | job={job_id} | iteration={current_iteration} | remaining={audit_result.remaining_leaks}")
            
            elif audit_result.leaks_auto_fixed > 0 and audit_result.fixed_text:
                with open(temp_output, 'w', encoding='utf-8') as f:
                    f.write(audit_result.fixed_text)
                logger.info(f"AUDIT_AUTOFIX_TXT | job={job_id} | saved corrected text")
            
            if not audit_result.is_safe:
                logger.error(f"AUDIT_UNSAFE | job={job_id} | remaining_leaks={audit_result.remaining_leaks}")
                safe_remove(temp_output)
                release_reservation(current_user.id, job_id)
                job.status = 'failed'
                db.session.commit()

                leak_types = list(set(l['type'] for l in audit_result.leaks_found))
                types_list = [f"{t} (posibles fugas)" for t in leak_types]

                return render_template("anonymizer_blocked.html",
                    residual_types=types_list,
                    total_residual=audit_result.remaining_leaks,
                    original_filename=original_filename
                )
            
            # Si hubo auto-correcciones exitosas, advertir al usuario
            if audit_result.leaks_auto_fixed > 0:
                residual_warning = f"NOTA: El auditor detectó y corrigió automáticamente {audit_result.leaks_auto_fixed} posibles fugas de PII adicionales."
        
        else:
            # Fallback: usar post_scan_final si el auditor no está disponible
            from detector_capas import post_scan_final
            text_without_tokens = re.sub(r'\{\{[A-Z]+_\d+\}\}', '', post_scan_text)
            _, residual_pii_clean = post_scan_final(text_without_tokens)
            
            if residual_pii_clean:
                real_residual = [r for r in residual_pii_clean if r['count'] > 0]
                if real_residual:
                    logger.warning(f"POST_SCAN | job={job_id} | residual_pii={real_residual}")
                    types_found = ', '.join([f"{r['type']} ({r['count']})" for r in real_residual])
                    residual_warning = f"ATENCIÓN: Se detectó posible PII residual en el documento: {types_found}. Revise el documento antes de compartirlo."
        
        with open(temp_output, 'rb') as f:
            output_bytes = f.read()
        
        base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
        download_name = f"{base_name}_anonimizado.{output_ext}"
        
        type_counts = {}
        replacements_by_type = {}
        
        for token, original in mapping.items():
            clean_token = token.replace('{{', '').replace('}}', '')
            parts = clean_token.split('_')
            if len(parts) >= 2:
                entity_type = '_'.join(parts[:-1])
            else:
                entity_type = clean_token
            
            type_counts[entity_type] = type_counts.get(entity_type, 0) + 1
            
            if entity_type not in replacements_by_type:
                replacements_by_type[entity_type] = []
            
            replacements_by_type[entity_type].append({
                'token': clean_token,
                'original': original,
                'original_truncated': truncate_value(original)
            })
        
        warnings = []
        has_direccion = 'DIRECCION' in type_counts
        has_persona = 'PERSONA' in type_counts
        manual_count = sum(1 for e in selected_entities if e.get('source') == 'manual')
        review_count = sum(1 for e in selected_entities if e.get('status') == 'needs_review')
        
        # Agregar warning de PII residual si existe (más importante)
        if residual_warning:
            warnings.insert(0, residual_warning)
        
        if has_direccion:
            warnings.append("Las direcciones fueron detectadas mediante heurísticas. Revise el documento para confirmar la correcta anonimización.")
        if has_persona:
            warnings.append("Los nombres fueron detectados mediante NER y patrones. Pueden existir nombres adicionales no detectados.")
        if review_count > 0:
            warnings.append(f"Se detectaron {review_count} entidades que requieren revisión manual.")
        if manual_count > 0:
            warnings.append(f"Se anonimizaron {manual_count} entidades adicionales tras la revisión manual.")
        
        report_data = {
            'total_replaced': replaced_count,
            'type_counts': type_counts,
            'mapping': mapping,
            'warnings': warnings,
            'original_filename': original_filename
        }
        report_json = json.dumps(report_data, ensure_ascii=False, indent=2)
        
        result_paths = get_result_paths(job_id)
        
        with open(result_paths['doc'], 'wb') as f:
            f.write(output_bytes)
        
        charge_pages(current_user.id, job_id, stage="apply")

        meta_data = {
            'download_name': download_name,
            'output_ext': output_ext,
            'report_json': report_json,
            'created_at': datetime.now().isoformat(),
            'user_id': current_user.id
        }
        with open(result_paths['meta'], 'w', encoding='utf-8') as f:
            json.dump(meta_data, f, ensure_ascii=False)

        safe_remove(temp_input)
        safe_remove(temp_output)

        logger.info(f"RESULTS_PAGE | job={job_id} | user={current_user.id} | replaced={replaced_count}")

        return render_template("anonymizer_results.html",
            job_id=job_id,
            total_replaced=replaced_count,
            type_counts=type_counts,
            replacements_by_type=replacements_by_type,
            warnings=warnings,
            download_url=f"/anonymizer/download/{job_id}",
            report_url=f"/anonymizer/report/{job_id}",
            output_ext=output_ext
        )

    except json.JSONDecodeError as e:
        logger.error(f"APPLY_FAIL | job={job_id} | reason=json_error | error={e}")
        logger.error(traceback.format_exc())
        release_reservation(current_user.id, job_id)
        job.status = 'failed'
        db.session.commit()
        safe_remove(temp_input)
        return render_error("Error procesando la selección de entidades. Intente nuevamente.")

    except Exception as e:
        logger.error(f"APPLY_FAIL | job={job_id} | error={e}")
        logger.error(traceback.format_exc())
        try:
            release_reservation(current_user.id, job_id)
            job.status = 'failed'
            db.session.commit()
        except Exception:
            pass
        safe_remove(temp_input)
        safe_remove(temp_output)
        return render_error("No se pudo generar el archivo final. Verifique el documento e intente nuevamente.")


@anonymizer_bp.route("/anonymizer/download/<job_id>")
@login_required
def anonymizer_download(job_id):
    """
    Download the anonymized document with FINAL GUARANTEE.
    Validates ownership before serving.
    """
    from models import AnonymizerJob
    job = AnonymizerJob.query.filter_by(job_id=job_id).first()
    if not job or job.user_id != current_user.id:
        logger.warning(f"DOWNLOAD_OWNERSHIP_FAIL | job={job_id} | user={current_user.id}")
        return render_error("No tiene permiso para descargar este documento.", 403)

    result_paths = get_result_paths(job_id)
    
    if not os.path.exists(result_paths['doc']) or not os.path.exists(result_paths['meta']):
        return render_error("El enlace ha expirado. Por favor procese el documento nuevamente.")
    
    try:
        with open(result_paths['meta'], 'r', encoding='utf-8') as f:
            meta = json.load(f)
        
        # =========================================================================
        # GARANTÍA FINAL: Auditoría JUSTO ANTES de servir el archivo
        # =========================================================================
        if meta['output_ext'] == 'docx' and FINAL_AUDITOR_AVAILABLE:
            from docx import Document
            from processor_docx import apply_replacements_to_docx
            
            MAX_ITERATIONS = 2
            doc_path = result_paths['doc']
            
            for iteration in range(1, MAX_ITERATIONS + 1):
                doc = Document(doc_path)
                full_text = extract_full_text_docx(doc)
                
                audit_result = audit_document(full_text, auto_fix=True)
                
                if audit_result.is_safe:
                    logger.info(f"DOWNLOAD_AUDIT_SAFE | job={job_id} | iteration={iteration}")
                    break
                
                if audit_result.replacements:
                    fixes = apply_replacements_to_docx(doc, audit_result.replacements)
                    doc.save(doc_path)
                    logger.warning(f"DOWNLOAD_AUDIT_FIX | job={job_id} | iteration={iteration} | fixes={fixes}")
                else:
                    break
            
            # Re-verificar seguridad final después de todas las iteraciones
            doc_final = Document(doc_path)
            final_text = extract_full_text_docx(doc_final)
            final_audit = audit_document(final_text, auto_fix=False)
            
            if not final_audit.is_safe:
                strict_mode = os.environ.get('STRICT_ZERO_LEAKS', '1') == '1'
                force_download = request.args.get('force', '0') == '1'
                
                if strict_mode and not force_download:
                    from processor_docx import hard_redact_patterns
                    doc_hard = Document(doc_path)
                    hard_fixes = hard_redact_patterns(doc_hard)
                    if hard_fixes > 0:
                        doc_hard.save(doc_path)
                        logger.warning(f"HARD_REDACT | job={job_id} | fixes={hard_fixes}")
                        doc_recheck = Document(doc_path)
                        recheck_text = extract_full_text_docx(doc_recheck)
                        final_audit = audit_document(recheck_text, auto_fix=False)
                
                if not final_audit.is_safe:
                    if force_download:
                        logger.warning(f"DOWNLOAD_FORCED | job={job_id} | remaining_leaks={final_audit.remaining_leaks} | user_accepted_risk=true")
                    else:
                        logger.warning(f"DOWNLOAD_WARNING | job={job_id} | remaining_leaks={final_audit.remaining_leaks}")
                        leak_types = list(set(l['type'] for l in final_audit.leaks_found))
                        types_list = [f"{t} (posibles fugas)" for t in leak_types]
                        
                        return render_template("anonymizer_blocked.html",
                            residual_types=types_list,
                            total_residual=final_audit.remaining_leaks,
                            original_filename=meta.get('download_name', 'documento'),
                            job_id=job_id
                        )
        
        elif meta['output_ext'] == 'txt' and FINAL_AUDITOR_AVAILABLE:
            with open(result_paths['doc'], 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            audit_result = audit_document(text_content, auto_fix=True)
            
            if audit_result.leaks_auto_fixed > 0 and audit_result.fixed_text:
                with open(result_paths['doc'], 'w', encoding='utf-8') as f:
                    f.write(audit_result.fixed_text)
                logger.info(f"DOWNLOAD_AUDIT_FIX_TXT | job={job_id} | fixes={audit_result.leaks_auto_fixed}")
            
            if not audit_result.is_safe:
                force_download = request.args.get('force', '0') == '1'
                
                if force_download:
                    logger.warning(f"DOWNLOAD_FORCED_TXT | job={job_id} | remaining={audit_result.remaining_leaks} | user_accepted_risk=true")
                else:
                    logger.warning(f"DOWNLOAD_WARNING_TXT | job={job_id} | remaining={audit_result.remaining_leaks}")
                    return render_template("anonymizer_blocked.html",
                        residual_types=[f"{l['type']} (posibles fugas)" for l in audit_result.leaks_found[:5]],
                        total_residual=audit_result.remaining_leaks,
                        original_filename=meta.get('download_name', 'documento'),
                        job_id=job_id
                    )
        
        # =========================================================================
        # Leer archivo FINAL (después de auditoría) y servir
        # =========================================================================
        with open(result_paths['doc'], 'rb') as f:
            output_bytes = f.read()
        
        output_buffer = BytesIO(output_bytes)
        
        if meta['output_ext'] == 'docx':
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            mimetype = 'text/plain; charset=utf-8'
        
        logger.info(f"DOWNLOAD | job={job_id} | filename={meta['download_name']} | SAFE")
        
        return send_file(
            output_buffer,
            as_attachment=True,
            download_name=meta['download_name'],
            mimetype=mimetype
        )
    except Exception as e:
        logger.error(f"DOWNLOAD_ERROR | job={job_id} | error={e}")
        return render_error("Error al descargar el documento. Por favor procese el documento nuevamente.")


@anonymizer_bp.route("/anonymizer/report/<job_id>")
@login_required
def anonymizer_report(job_id):
    """Download the anonymization report."""
    from models import AnonymizerJob
    job = AnonymizerJob.query.filter_by(job_id=job_id).first()
    if not job or job.user_id != current_user.id:
        return render_error("No tiene permiso para descargar este reporte.", 403)

    result_paths = get_result_paths(job_id)
    
    if not os.path.exists(result_paths['meta']):
        return render_error("El enlace ha expirado. Por favor procese el documento nuevamente.")
    
    try:
        with open(result_paths['meta'], 'r', encoding='utf-8') as f:
            meta = json.load(f)
        
        report_buffer = BytesIO(meta['report_json'].encode('utf-8'))
        
        base_name = meta['download_name'].rsplit('.', 1)[0] if '.' in meta['download_name'] else meta['download_name']
        report_name = f"{base_name}_reporte.json"
        
        logger.info(f"REPORT | job={job_id} | filename={report_name}")
        
        return send_file(
            report_buffer,
            as_attachment=True,
            download_name=report_name,
            mimetype='application/json'
        )
    except Exception as e:
        logger.error(f"REPORT_ERROR | job={job_id} | error={e}")
        return render_error("Error al descargar el reporte. Por favor procese el documento nuevamente.")


@anonymizer_bp.route("/anonymizer/plans")
@login_required
def anonymizer_plans():
    """Página de compra de paquetes de páginas."""
    from models import AnonymizerPackage, UserCredits
    from credit_utils import get_or_create_credits
    pkgs = AnonymizerPackage.query.filter_by(is_active=True)\
        .order_by(AnonymizerPackage.display_order.asc(), AnonymizerPackage.id.asc()).all()
    credits = get_or_create_credits(current_user.id)
    culqi_public_key = os.environ.get('CULQI_PUBLIC_KEY', '')
    return render_template(
        "anonymizer_plans.html",
        packages=pkgs,
        pages_balance=credits.pages_balance,
        culqi_public_key=culqi_public_key,
        user_email=getattr(current_user, 'email', '') or '',
    )


@anonymizer_bp.route("/redeem-code", methods=["POST"])
@login_required
def redeem_credit_code():
    """Canjear un código manual de crédito."""
    from models import db, CreditCode, CreditRedemption, UserCredits, PageUsageLog
    code_str = request.form.get("code", "").strip().upper()
    if not code_str:
        flash("Por favor ingresa un código.", "error")
        return redirect(url_for('anonymizer.account'))
    code = CreditCode.query.filter_by(code=code_str).first()
    if not code:
        flash("Código inválido.", "error")
        return redirect(url_for('anonymizer.account'))
    valid, msg = code.is_valid()
    if not valid:
        flash(msg, "error")
        return redirect(url_for('anonymizer.account'))
    already = CreditRedemption.query.filter_by(
        user_id=current_user.id, credit_code_id=code.id
    ).first()
    if already:
        flash("Ya canjeaste este código anteriormente.", "error")
        return redirect(url_for('anonymizer.account'))
    redemption = CreditRedemption(
        user_id=current_user.id,
        credit_code_id=code.id,
        ip=request.remote_addr,
        user_agent=(request.headers.get('User-Agent', '') or '')[:500],
    )
    code.uses_count += 1
    credits = get_or_create_credits(current_user.id)
    credits.pages_balance += code.credit_amount
    log_entry = PageUsageLog(
        user_id=current_user.id,
        job_id=f"code_{code.code}",
        stage="code_redeem",
        pages=code.credit_amount,
        action="credited",
        detail=f"Código {code.code}: +{code.credit_amount} páginas",
    )
    db.session.add(redemption)
    db.session.add(log_entry)
    db.session.commit()
    logger.info(f"CODE_REDEEMED | user={current_user.id} | code={code.code} | pages={code.credit_amount}")
    flash(f"¡Código canjeado! +{code.credit_amount} páginas agregadas a tu cuenta.", "success")
    return redirect(url_for('anonymizer.account'))


_reward_rate_limits = defaultdict(list)

def _check_reward_rate_limit(ip, limit=5, window_secs=60):
    now = datetime.utcnow()
    cutoff = now - timedelta(seconds=window_secs)
    _reward_rate_limits[ip] = [t for t in _reward_rate_limits[ip] if t > cutoff]
    if len(_reward_rate_limits[ip]) >= limit:
        return False
    _reward_rate_limits[ip].append(now)
    return True


def _hash_token(token):
    return hashlib.sha256(token.encode()).hexdigest()


@anonymizer_bp.route("/api/rewards/issue", methods=["POST"])
def rewards_issue():
    """Emitir token de recompensa para WordPress/Tutor LMS."""
    from models import db, User, RewardToken, UserCredits
    reward_api_key = os.environ.get("REWARD_API_KEY", "")
    if not reward_api_key:
        return jsonify({"error": "rewards_not_configured"}), 503
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer ") or auth[7:] != reward_api_key:
        logger.warning(f"REWARD_ISSUE_UNAUTH | ip={request.remote_addr}")
        return jsonify({"error": "unauthorized"}), 401
    data = request.get_json(silent=True) or {}
    external_key = (data.get("external_user_key") or "").strip().lower()
    lesson_id = (data.get("lesson_id") or "").strip()
    credit_amount = data.get("credit_amount")
    if not external_key or not lesson_id or not isinstance(credit_amount, int) or credit_amount <= 0:
        return jsonify({"error": "invalid_payload", "detail": "external_user_key, lesson_id y credit_amount (int>0) son requeridos"}), 400
    user = User.query.filter_by(email=external_key, activo=True).first()
    if not user:
        return jsonify({"error": "user_not_found"}), 404
    existing = RewardToken.query.filter_by(user_id=user.id, lesson_id=lesson_id).first()
    public_app_url = os.environ.get("PUBLIC_APP_URL", request.host_url.rstrip("/"))
    if existing:
        raw_token = None
        if existing.status == 'issued' and datetime.utcnow() < existing.expires_at:
            raw_note = "(token existente, no se puede revelar el texto plano)"
            redeem_url = f"{public_app_url}/redeem?lesson={lesson_id}&user={user.id}"
        else:
            raw_note = "(token ya usado o expirado)"
            redeem_url = f"{public_app_url}/redeem"
        return jsonify({
            "status": existing.status,
            "note": raw_note,
            "redeem_url": redeem_url,
            "expires_at": existing.expires_at.isoformat() + "Z",
        }), 200
    raw_token = secrets.token_urlsafe(32)
    token_hash = _hash_token(raw_token)
    expires_at = datetime.utcnow() + timedelta(minutes=30)
    reward = RewardToken(
        user_id=user.id,
        external_user_key=external_key,
        lesson_id=lesson_id,
        credit_amount=credit_amount,
        token_hash=token_hash,
        status='issued',
        expires_at=expires_at,
        issued_ip=request.remote_addr,
    )
    db.session.add(reward)
    db.session.commit()
    redeem_url = f"{public_app_url}/redeem?token={raw_token}"
    logger.info(f"REWARD_ISSUED | user={user.id} | lesson={lesson_id} | pages={credit_amount}")
    return jsonify({
        "redeem_url": redeem_url,
        "token": raw_token,
        "expires_at": expires_at.isoformat() + "Z",
    }), 200


# ---------------------------------------------------------------------------
# POST /api/rewards/issue-code
# Server-to-server: emite un CreditCode tipo APC-XXXX-XXXX para un usuario.
# Deduplicación por (email, lesson_id) via RewardToken(lesson_id="code:<id>").
# ---------------------------------------------------------------------------
_CODE_ALPHABET = "ABCDEFGHJKLMNPQRSTUVWXYZ23456789"  # sin 0/O/I/1 para evitar confusión
_CODE_EXPIRY_DAYS = int(os.environ.get("ISSUE_CODE_EXPIRY_DAYS", "30"))


def _generate_credit_code():
    """Genera un código tipo APC-XXXX-XXXX con el alfabeto sin ambiguos."""
    part = lambda: "".join(secrets.choice(_CODE_ALPHABET) for _ in range(4))
    return f"APC-{part()}-{part()}"


@anonymizer_bp.route("/api/rewards/issue-code", methods=["POST"])
def rewards_issue_code():
    """
    Emite un CreditCode canjeble para un usuario identificado por email.
    Requiere: Authorization: Bearer <REWARD_API_KEY>
    Body JSON: { "external_user_key": "<email>", "lesson_id": "<str>", "credit_amount": <int> }
    Deduplicación: 1 código por (email, lesson_id). Retorna el mismo estado si ya fue emitido.
    """
    from models import db, User, CreditCode, RewardToken

    reward_api_key = os.environ.get("REWARD_API_KEY", "")
    if not reward_api_key:
        return jsonify({"error": "rewards_not_configured"}), 503

    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer ") or auth[7:] != reward_api_key:
        logger.warning("ISSUE_CODE_UNAUTH | ip=%s", request.remote_addr)
        return jsonify({"error": "unauthorized"}), 401

    data = request.get_json(silent=True) or {}
    external_key = (data.get("external_user_key") or "").strip().lower()
    lesson_id    = (data.get("lesson_id") or "").strip()
    credit_amount = data.get("credit_amount")

    if not external_key or not lesson_id:
        return jsonify({"error": "invalid_payload", "detail": "external_user_key y lesson_id son requeridos"}), 400
    if not isinstance(credit_amount, int) or credit_amount <= 0:
        return jsonify({"error": "invalid_payload", "detail": "credit_amount debe ser un entero mayor a 0"}), 400

    user = User.query.filter_by(email=external_key, activo=True).first()
    if not user:
        return jsonify({"error": "user_not_found", "detail": f"No existe usuario activo con email {external_key}"}), 404

    namespaced_lesson = f"code:{lesson_id}"

    existing_tracker = RewardToken.query.filter_by(
        user_id=user.id, lesson_id=namespaced_lesson
    ).first()

    if existing_tracker:
        existing_code = CreditCode.query.filter_by(
            id=int(existing_tracker.external_user_key)
        ).first() if existing_tracker.external_user_key.isdigit() else None
        expires_str = existing_code.expires_at.isoformat() + "Z" if (existing_code and existing_code.expires_at) else None
        return jsonify({
            "status": "already_issued",
            "code": existing_code.code if existing_code else None,
            "credit_amount": existing_code.credit_amount if existing_code else credit_amount,
            "expires_at": expires_str,
        }), 200

    expires_at = datetime.utcnow() + timedelta(days=_CODE_EXPIRY_DAYS)

    for attempt in range(10):
        candidate = _generate_credit_code()
        if not CreditCode.query.filter_by(code=candidate).first():
            break
    else:
        logger.error("ISSUE_CODE | could not generate unique code after 10 attempts")
        return jsonify({"error": "server_error", "detail": "No se pudo generar un código único"}), 500

    new_code = CreditCode(
        code=candidate,
        credit_amount=credit_amount,
        max_uses=1,
        uses_count=0,
        is_active=True,
        expires_at=expires_at,
        created_by_id=None,
    )
    db.session.add(new_code)
    db.session.flush()

    tracker = RewardToken(
        user_id=user.id,
        external_user_key=str(new_code.id),
        lesson_id=namespaced_lesson,
        credit_amount=credit_amount,
        token_hash=hashlib.sha256(candidate.encode()).hexdigest(),
        status="code_issued",
        issued_at=datetime.utcnow(),
        expires_at=expires_at,
        issued_ip=request.remote_addr,
    )
    db.session.add(tracker)
    db.session.commit()

    logger.info(
        "ISSUE_CODE | user=%s lesson=%s code=%s amount=%d expires=%s",
        user.id, lesson_id, candidate, credit_amount, expires_at.date()
    )
    return jsonify({
        "status": "issued",
        "code": candidate,
        "credit_amount": credit_amount,
        "expires_at": expires_at.isoformat() + "Z",
    }), 200


@anonymizer_bp.route("/api/rewards/redeem", methods=["POST"])
def rewards_redeem():
    """Canjear un token de recompensa."""
    from models import db, RewardToken, UserCredits, PageUsageLog
    ip = request.remote_addr or "unknown"
    if not _check_reward_rate_limit(ip):
        logger.warning(f"REWARD_REDEEM_RATELIMIT | ip={ip}")
        return jsonify({"error": "rate_limited", "detail": "Máximo 5 intentos por minuto"}), 429
    data = request.get_json(silent=True) or {}
    raw_token = (data.get("token") or "").strip()
    if not raw_token:
        return jsonify({"error": "token_required"}), 400
    token_hash = _hash_token(raw_token)
    reward = RewardToken.query.filter_by(token_hash=token_hash).first()
    if not reward:
        return jsonify({"error": "token_invalid"}), 404
    if reward.status != 'issued':
        return jsonify({"error": "token_already_used", "status": reward.status}), 409
    if datetime.utcnow() > reward.expires_at:
        reward.status = 'expired'
        db.session.commit()
        return jsonify({"error": "token_expired"}), 410
    credits = get_or_create_credits(reward.user_id)
    credits.pages_balance += reward.credit_amount
    reward.status = 'used'
    reward.used_at = datetime.utcnow()
    reward.used_ip = ip
    log_entry = PageUsageLog(
        user_id=reward.user_id,
        job_id=f"reward_{reward.id}",
        stage="reward_redeem",
        pages=reward.credit_amount,
        action="credited",
        detail=f"Recompensa lección {reward.lesson_id}: +{reward.credit_amount} páginas",
    )
    db.session.add(log_entry)
    db.session.commit()
    logger.info(f"REWARD_REDEEMED | user={reward.user_id} | lesson={reward.lesson_id} | pages={reward.credit_amount}")
    return jsonify({
        "credited_amount": reward.credit_amount,
        "new_balance": credits.pages_balance,
    }), 200


@anonymizer_bp.route("/redeem")
def redeem_page():
    """Página HTML simple para canjear tokens de recompensa de WordPress."""
    token = request.args.get("token", "")
    return render_template("redeem_reward.html", token=token)


# ─── Custom Labels API ────────────────────────────────────────────────────────

@anonymizer_bp.route("/api/user-labels", methods=["GET"])
@login_required
def api_user_labels_list():
    from models import UserCustomLabel
    labels = UserCustomLabel.query.filter_by(user_id=current_user.id)\
        .order_by(UserCustomLabel.created_at.asc()).all()
    return jsonify([{"id": l.id, "label_name": l.label_name} for l in labels])


@anonymizer_bp.route("/api/user-labels", methods=["POST"])
@login_required
def api_user_labels_create():
    from models import db, UserCustomLabel
    data = request.get_json(silent=True) or {}
    raw = (data.get("label_name") or "").strip().upper()
    # Solo letras, números y guión bajo; máximo 30 chars
    import re as _re
    if not raw or not _re.match(r'^[A-Z0-9_]{1,30}$', raw):
        return jsonify({"error": "Nombre inválido. Usa letras, números o _ (máx. 30 caracteres)."}), 400
    if UserCustomLabel.query.filter_by(user_id=current_user.id, label_name=raw).first():
        return jsonify({"error": f"Ya tienes una etiqueta llamada '{raw}'."}), 409
    count = UserCustomLabel.query.filter_by(user_id=current_user.id).count()
    if count >= 50:
        return jsonify({"error": "Límite de 50 etiquetas personalizadas alcanzado."}), 400
    label = UserCustomLabel(user_id=current_user.id, label_name=raw)
    db.session.add(label)
    db.session.commit()
    logger.info("CUSTOM_LABEL_CREATED | user=%s | label=%s", current_user.id, raw)
    return jsonify({"id": label.id, "label_name": label.label_name}), 201


@anonymizer_bp.route("/api/user-labels/<int:label_id>", methods=["DELETE"])
@login_required
def api_user_labels_delete(label_id):
    from models import db, UserCustomLabel
    label = UserCustomLabel.query.filter_by(id=label_id, user_id=current_user.id).first()
    if not label:
        return jsonify({"error": "Etiqueta no encontrada o sin permiso."}), 404
    db.session.delete(label)
    db.session.commit()
    logger.info("CUSTOM_LABEL_DELETED | user=%s | label=%s", current_user.id, label.label_name)
    return jsonify({"ok": True})


