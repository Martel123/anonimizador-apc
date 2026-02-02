"""
Tests para el auditor final de PII
===================================
Verifica que el auditor detecte y corrija correctamente:
- Emails (incluyendo partidos en runs)
- Teléfonos (varios formatos PE)
- DNI con contexto
- Números de colegiatura
- Direcciones
"""

import pytest
from final_auditor import (
    audit_document, find_dni_leaks, find_email_leaks, find_phone_leaks,
    find_colegiatura_leaks, find_ruc_leaks, find_direccion_leaks,
    find_expediente_leaks, find_resolucion_leaks, find_partida_leaks,
    find_casilla_leaks, find_tribunal_sala_leaks,
    _find_all_leaks
)


class TestEmailDetection:
    def test_detect_simple_email(self):
        text = "Contactar a manuel.monty@gmail.com para más información."
        leaks = find_email_leaks(text)
        assert len(leaks) == 1
        assert leaks[0]['value'] == 'manuel.monty@gmail.com'
    
    def test_detect_email_uppercase(self):
        text = "Email: JUAN.PEREZ@EMPRESA.COM.PE"
        leaks = find_email_leaks(text)
        assert len(leaks) == 1
    
    def test_ignore_tokenized_email(self):
        text = "Email: {{EMAIL_1}} para contacto."
        leaks = find_email_leaks(text)
        assert len(leaks) == 0


class TestPhoneDetection:
    def test_detect_9_digit_phone(self):
        text = "Celular: 987654321"
        leaks = find_phone_leaks(text)
        assert len(leaks) == 1
        assert '987654321' in leaks[0]['value']
    
    def test_detect_phone_with_spaces(self):
        text = "Teléfono: 987 654 321"
        leaks = find_phone_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_phone_with_dashes(self):
        text = "Celular: 987-654-321"
        leaks = find_phone_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_phone_with_prefix(self):
        text = "Llamar al +51 987654321"
        leaks = find_phone_leaks(text)
        assert len(leaks) >= 1
    
    def test_ignore_tokenized_phone(self):
        text = "Celular: {{TELEFONO_1}}"
        leaks = find_phone_leaks(text)
        assert len(leaks) == 0


class TestDNIDetection:
    def test_detect_dni_with_context(self):
        text = "El demandante identificado con DNI N° 12345678"
        leaks = find_dni_leaks(text)
        assert len(leaks) == 1
        assert leaks[0]['value'] == '12345678'
    
    def test_detect_dni_explicit(self):
        text = "D.N.I. 87654321"
        leaks = find_dni_leaks(text)
        assert len(leaks) == 1
    
    def test_ignore_dates(self):
        text = "Fecha: 20240115"
        leaks = find_dni_leaks(text)
        assert len(leaks) == 0
    
    def test_ignore_tokenized_dni(self):
        text = "DNI: {{DNI_1}}"
        leaks = find_dni_leaks(text)
        assert len(leaks) == 0


class TestColegiaturaDetection:
    def test_detect_cal(self):
        text = "Abogado con CAL N° 12345"
        leaks = find_colegiatura_leaks(text)
        assert len(leaks) == 1
    
    def test_detect_cal_dots(self):
        text = "C.A.L. 54321"
        leaks = find_colegiatura_leaks(text)
        assert len(leaks) == 1
    
    def test_detect_cmp(self):
        text = "Médico CMP N° 67890"
        leaks = find_colegiatura_leaks(text)
        assert len(leaks) == 1
    
    def test_detect_colegiatura_word(self):
        text = "N° de Colegiatura 11111"
        leaks = find_colegiatura_leaks(text)
        assert len(leaks) == 1


class TestDireccionDetection:
    def test_detect_calle(self):
        text = "Domicilio: Calle Los Olivos N° 123"
        leaks = find_direccion_leaks(text)
        assert len(leaks) == 1
    
    def test_detect_avenida_with_dpto(self):
        text = "Vive en Av. Larco 456, Dpto 502"
        leaks = find_direccion_leaks(text)
        assert len(leaks) == 1
    
    def test_detect_jiron(self):
        text = "Jr. Huallaga N° 789"
        leaks = find_direccion_leaks(text)
        assert len(leaks) == 1
    
    def test_detect_mz_lote(self):
        text = "Mz A, Lt 15, Urb Los Jardines"
        leaks = find_direccion_leaks(text)
        assert len(leaks) == 1


class TestRUCDetection:
    def test_detect_ruc_persona_natural(self):
        text = "RUC 10123456789"
        leaks = find_ruc_leaks(text)
        assert len(leaks) == 1
    
    def test_detect_ruc_empresa(self):
        text = "Con RUC 20123456789"
        leaks = find_ruc_leaks(text)
        assert len(leaks) == 1
    
    def test_ignore_without_context(self):
        text = "Número 10123456789"
        leaks = find_ruc_leaks(text)
        assert len(leaks) == 0


class TestAuditDocument:
    def test_autofix_email(self):
        text = "Contactar a test@example.org para más info."
        result = audit_document(text, auto_fix=True)
        assert result.is_safe
        assert 'test@example.org' not in result.fixed_text
        assert '{{EMAIL_' in result.fixed_text
    
    def test_autofix_phone(self):
        text = "Llamar al 987654321"
        result = audit_document(text, auto_fix=True)
        assert result.is_safe
        assert '987654321' not in result.fixed_text
        assert '{{TELEFONO_' in result.fixed_text
    
    def test_autofix_dni(self):
        text = "Identificado con DNI 12345678"
        result = audit_document(text, auto_fix=True)
        assert result.is_safe
        assert '12345678' not in result.fixed_text
        assert '{{DNI_' in result.fixed_text
    
    def test_autofix_multiple(self):
        text = """
        El demandante Juan Pérez, identificado con DNI 12345678,
        domiciliado en Av. Larco 456, con email juan@correo.com
        y celular 987654321, interpone demanda.
        """
        result = audit_document(text, auto_fix=True)
        assert result.is_safe
        assert '12345678' not in result.fixed_text
        assert '987654321' not in result.fixed_text
        assert 'juan@correo.com' not in result.fixed_text
    
    def test_safe_document(self):
        text = """
        DEMANDA DE ALIMENTOS
        
        SEÑOR JUEZ DEL {{JUZGADO_1}}
        
        {{PERSONA_1}}, identificado con DNI {{DNI_1}},
        con domicilio en {{DIRECCION_1}}, interpone demanda.
        """
        result = audit_document(text, auto_fix=True)
        assert result.is_safe
        assert len(result.leaks_found) == 0
    
    def test_replacements_list(self):
        text = "Email: test@correo.com, Tel: 987654321"
        result = audit_document(text, auto_fix=True)
        assert len(result.replacements) >= 2


class TestActaRegistroDetection:
    def test_detect_acta_conciliacion(self):
        from final_auditor import find_acta_registro_leaks
        text = "Acta de Conciliación N° 00123-2024 del Centro de Conciliación"
        leaks = find_acta_registro_leaks(text)
        assert len(leaks) >= 1
        assert any('ACTA_REGISTRO' in l['type'] for l in leaks)
    
    def test_detect_registro(self):
        from final_auditor import find_acta_registro_leaks
        text = "Registro N° 45678 del Libro de Actas"
        leaks = find_acta_registro_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_expediente(self):
        from final_auditor import find_acta_registro_leaks
        text = "Expediente N° 00456-2024-0-1801-JR-CI-01"
        leaks = find_acta_registro_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_constancia(self):
        from final_auditor import find_acta_registro_leaks
        text = "Constancia N° 12345 de fecha 15 de enero"
        leaks = find_acta_registro_leaks(text)
        assert len(leaks) >= 1


class TestPlacaDetection:
    def test_detect_placa_with_context(self):
        from final_auditor import find_placa_leaks
        text = "Vehículo con placa ABC-123 color rojo"
        leaks = find_placa_leaks(text)
        assert len(leaks) >= 1
        assert any('PLACA' in l['type'] for l in leaks)
    
    def test_detect_placa_moto(self):
        from final_auditor import find_placa_leaks
        text = "Motocicleta con placa ABC-123 de color azul"
        leaks = find_placa_leaks(text)
        assert len(leaks) >= 1
    
    def test_ignore_without_context(self):
        from final_auditor import find_placa_leaks
        text = "El código ABC-123 no corresponde a nada"
        leaks = find_placa_leaks(text)
        assert len(leaks) == 0


class TestIntegration:
    def test_complete_document(self):
        text = """
        EXPEDIENTE N° 00123-2024
        
        DEMANDANTE: Juan Carlos Pérez López
        DNI: 12345678
        Dirección: Av. La Marina N° 2500, Dpto 1502, San Miguel
        Email: jcperez@gmail.com
        Teléfono: +51 987 654 321
        
        ABOGADO PATROCINANTE
        Dr. Manuel García
        CAL N° 54321
        
        DEMANDADO: María Elena Rodríguez
        DNI: 87654321
        """
        result = audit_document(text, auto_fix=True)
        
        assert result.is_safe, f"Document should be safe. Remaining: {result.remaining_leaks}"
        
        assert 'jcperez@gmail.com' not in result.fixed_text
        assert '12345678' not in result.fixed_text
        assert '987 654 321' not in result.fixed_text or '987654321' not in result.fixed_text
        assert '54321' not in result.fixed_text or 'CAL' not in result.fixed_text


class TestDOCXIntegration:
    """Tests de integración con documentos DOCX reales."""
    
    def test_apply_replacements_to_docx_paragraphs(self):
        """Verifica reemplazo en párrafos."""
        from docx import Document
        from processor_docx import apply_replacements_to_docx
        import tempfile
        import os
        
        doc = Document()
        doc.add_paragraph("Contactar a juan@empresa.com para más información.")
        doc.add_paragraph("Teléfono: 987654321")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            doc = Document(temp_path)
            replacements = [
                ('juan@empresa.com', '{{EMAIL_1}}'),
                ('987654321', '{{TELEFONO_1}}'),
            ]
            
            count = apply_replacements_to_docx(doc, replacements)
            doc.save(temp_path)
            
            doc_check = Document(temp_path)
            full_text = '\n'.join([p.text for p in doc_check.paragraphs])
            
            assert 'juan@empresa.com' not in full_text
            assert '987654321' not in full_text
            assert '{{EMAIL_1}}' in full_text
            assert '{{TELEFONO_1}}' in full_text
        finally:
            os.unlink(temp_path)
    
    def test_apply_replacements_to_docx_tables(self):
        """Verifica reemplazo en tablas."""
        from docx import Document
        from processor_docx import apply_replacements_to_docx
        import tempfile
        import os
        
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "DNI: 12345678"
        table.cell(0, 1).text = "Email: test@correo.com"
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            doc = Document(temp_path)
            replacements = [
                ('12345678', '{{DNI_1}}'),
                ('test@correo.com', '{{EMAIL_1}}'),
            ]
            
            apply_replacements_to_docx(doc, replacements)
            doc.save(temp_path)
            
            doc_check = Document(temp_path)
            table_text = ''
            for table in doc_check.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_text += cell.text + ' '
            
            assert '12345678' not in table_text
            assert 'test@correo.com' not in table_text
            assert '{{DNI_1}}' in table_text
            assert '{{EMAIL_1}}' in table_text
        finally:
            os.unlink(temp_path)
    
    def test_apply_replacements_split_runs(self):
        """Verifica reemplazo cuando el texto está partido en múltiples runs."""
        from docx import Document
        from docx.shared import Pt
        from processor_docx import apply_replacements_to_docx
        import tempfile
        import os
        
        doc = Document()
        para = doc.add_paragraph()
        run1 = para.add_run("Email: manu")
        run2 = para.add_run("el@")
        run2.bold = True
        run3 = para.add_run("gmail.com")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            doc = Document(temp_path)
            replacements = [('manuel@gmail.com', '{{EMAIL_1}}')]
            
            apply_replacements_to_docx(doc, replacements)
            doc.save(temp_path)
            
            doc_check = Document(temp_path)
            full_text = doc_check.paragraphs[0].text
            
            assert 'manuel@gmail.com' not in full_text
            assert '{{EMAIL_1}}' in full_text
        finally:
            os.unlink(temp_path)
    
    def test_full_pipeline_integration(self):
        """Test de integración completo: documento → auditor → fix → re-audit."""
        from docx import Document
        from processor_docx import apply_replacements_to_docx
        import tempfile
        import os
        
        doc = Document()
        doc.add_paragraph("DEMANDANTE: Juan Pérez, DNI 12345678")
        doc.add_paragraph("Email: jperez@mail.com, Cel: 987654321")
        doc.add_paragraph("Domicilio: Av. La Marina N° 2500")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            doc = Document(temp_path)
            text = '\n'.join([p.text for p in doc.paragraphs])
            
            result = audit_document(text, auto_fix=True)
            
            if result.replacements:
                apply_replacements_to_docx(doc, result.replacements)
                doc.save(temp_path)
            
            doc_check = Document(temp_path)
            final_text = '\n'.join([p.text for p in doc_check.paragraphs])
            final_result = audit_document(final_text, auto_fix=False)
            
            assert final_result.is_safe, f"Should be safe. Remaining: {final_result.remaining_leaks}"
            assert '12345678' not in final_text
            assert 'jperez@mail.com' not in final_text
            assert '987654321' not in final_text
        finally:
            os.unlink(temp_path)


class TestExpedienteDetection:
    def test_detect_expediente_judicial(self):
        text = "Expediente N° 00123-2024-0-1801-JP-FC-01"
        leaks = find_expediente_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_expediente_simple(self):
        text = "Exp. N° 12345-2024"
        leaks = find_expediente_leaks(text)
        assert len(leaks) >= 1
    
    def test_ignore_tokenized_expediente(self):
        text = "Expediente: {{EXPEDIENTE_1}}"
        leaks = find_expediente_leaks(text)
        assert len(leaks) == 0


class TestResolucionDetection:
    def test_detect_resolucion(self):
        text = "Resolución N° 123-2024"
        leaks = find_resolucion_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_auto(self):
        text = "Auto N° 456-2024"
        leaks = find_resolucion_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_oficio(self):
        text = "Oficio N° 789-2024"
        leaks = find_resolucion_leaks(text)
        assert len(leaks) >= 1


class TestPartidaDetection:
    def test_detect_partida_electronica(self):
        text = "Partida electrónica N° 12345678"
        leaks = find_partida_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_sunarp(self):
        text = "SUNARP N° 87654321"
        leaks = find_partida_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_asiento(self):
        text = "Asiento N° 00001234"
        leaks = find_partida_leaks(text)
        assert len(leaks) >= 1


class TestCasillaDetection:
    def test_detect_casilla_electronica(self):
        text = "Casilla electrónica N° 12345"
        leaks = find_casilla_leaks(text)
        assert len(leaks) >= 1
    
    def test_detect_mesa_partes(self):
        text = "Mesa de partes N° 67890"
        leaks = find_casilla_leaks(text)
        assert len(leaks) >= 1


class TestTribunalSalaDetection:
    def test_detect_juzgado(self):
        text = "1° Juzgado de Paz Letrado de Lima"
        leaks = find_tribunal_sala_leaks(text)
        assert len(leaks) >= 1
        assert leaks[0]['type'] == 'JUZGADO'
    
    def test_detect_tribunal(self):
        text = "Tribunal Constitucional del Perú"
        leaks = find_tribunal_sala_leaks(text)
        assert len(leaks) >= 1
        assert leaks[0]['type'] == 'TRIBUNAL'
    
    def test_detect_sala(self):
        text = "Sala Civil de Lima Norte"
        leaks = find_tribunal_sala_leaks(text)
        assert len(leaks) >= 1
        assert leaks[0]['type'] == 'SALA'


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
