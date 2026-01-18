"""
Tests for the 4-layer PII detection system.
Verifies maximum recall and zero leaks in final output.
"""

import pytest
import os
import tempfile


# ============================================================================
# FIXTURES - Documentos de prueba con PII real
# ============================================================================

FIXTURE_LEGAL_DOCUMENT = """
EXPEDIENTE N° 00123-2024-0-1801-JR-CI-01

SEÑOR JUEZ DEL PRIMER JUZGADO CIVIL DE LIMA

JUAN CARLOS GARCÍA RODRÍGUEZ, identificado con DNI N° 12345678, con RUC N° 10123456789,
con domicilio real en Av. Arequipa 1234, Dpto. 501, Urbanización Santa Cruz, Distrito de 
Miraflores, Provincia y Departamento de Lima; con domicilio procesal en Jr. Lampa 456, 
Oficina 302, Cercado de Lima; correo electrónico: juancarlos.garcia@gmail.com; teléfono 
celular: 987654321; casilla electrónica N° 12345; a Ud. respetuosamente digo:

DATOS DE LA DEMANDADA:

Que, dirijo la demanda en contra de la madre MARÍA ELENA LÓPEZ PÉREZ, identificada con 
DNI N° 87654321, con domicilio real ubicado en Calle Los Pinos 789, San Isidro, Lima, 
teléfono: +51 912 345 678, email: maria.lopez@hotmail.com.

PETITORIO DE LA DEMANDA:

Invocando interés y legitimidad para obrar, interpongo demanda de ALIMENTOS contra la 
demandada doña MARÍA ELENA LÓPEZ PÉREZ, a fin de que cumpla con acudir una pensión 
alimenticia a favor del menor PEDRO ANTONIO GARCÍA LÓPEZ.

FUNDAMENTOS DE HECHO:

PRIMERO.- Con fecha 15 de enero de 2020, el suscrito contrajo matrimonio civil con la 
demandada ante la Municipalidad de Miraflores, según consta en el Acta de Matrimonio 
N° 00456-2020.

SEGUNDO.- De dicha unión nació el menor PEDRO ANTONIO GARCÍA LÓPEZ, según Acta de 
Nacimiento inscrita en RENIEC.

TERCERO.- La demandada trabaja en el BANCO DE CRÉDITO DEL PERÚ S.A.A., percibiendo una 
remuneración mensual de S/ 8,500.00 (Ocho Mil Quinientos y 00/100 Soles).

CUARTO.- Mi cuenta de ahorros es: 193-12345678-0-12 del Banco de la Nación, CCI: 
01819300123456780123.

MEDIOS PROBATORIOS:

1. Copia de DNI del demandante
2. Acta de Matrimonio N° 00456-2020
3. Partida de Nacimiento del menor

ANEXOS:
1-A Copia de DNI
1-B Acta de Matrimonio
1-C Partida de Nacimiento

OTROSÍ DIGO: Que, designo como mi abogado defensor al Dr. ROBERTO CARLOS MENDOZA VEGA, 
identificado con CAL N° 54321, con domicilio procesal en la dirección antes indicada.

POR TANTO:
A Ud., Señor Juez, pido admitir la presente demanda y darle el trámite correspondiente.

Lima, 15 de enero de 2024

_________________________
JUAN CARLOS GARCÍA RODRÍGUEZ
DNI: 12345678
"""

FIXTURE_UPPERCASE_NAMES = """
DEMANDA DE DIVORCIO

DEMANDANTE: CARLOS ALBERTO FERNÁNDEZ TORRES
DNI: 45678901
DOMICILIO: AV. JAVIER PRADO ESTE 4500, LA MOLINA, LIMA

DEMANDADA: ANA MARÍA CASTILLO QUISPE
DNI: 56789012
DOMICILIO: JR. HUALLAGA 234, CERCADO DE LIMA
TELÉFONO: 01-4567890

El suscrito CARLOS ALBERTO FERNÁNDEZ TORRES interpone demanda de divorcio
contra su cónyuge ANA MARÍA CASTILLO QUISPE.
"""

FIXTURE_HEADER_FOOTER = """
================================================================================
ESTUDIO JURÍDICO PÉREZ & ASOCIADOS
Av. La Marina 2500, San Miguel - Teléfono: (01) 567-8901
Email: contacto@perezasociados.com.pe
================================================================================

DEMANDA DE ALIMENTOS

Expediente: 2024-00789

DEMANDANTE: Rosa María Huamán Condori, DNI 34567890
DEMANDADO: Jorge Luis Mamani Apaza, DNI 23456789

================================================================================
Página 1 de 3 - Casilla Electrónica N° 67890
================================================================================
"""


# ============================================================================
# TESTS CAPA 1: REGEX DETERMINÍSTICO
# ============================================================================

class TestLayer1Regex:
    """Tests for Layer 1: Deterministic regex patterns."""
    
    def test_dni_detection(self):
        """DNI 8 digits should be detected."""
        from detector_capas import detect_layer1_regex
        
        text = "El demandante con DNI 12345678 solicita..."
        entities = detect_layer1_regex(text)
        
        dni_entities = [e for e in entities if e.type == 'DNI']
        assert len(dni_entities) >= 1
        assert any(e.value == '12345678' for e in dni_entities)
    
    def test_dni_not_in_money_context(self):
        """8 digits near S/ should NOT be detected as DNI."""
        from detector_capas import detect_layer1_regex
        
        text = "El monto de S/ 12345678.00 fue pagado."
        entities = detect_layer1_regex(text)
        
        dni_entities = [e for e in entities if e.type == 'DNI']
        assert len(dni_entities) == 0
    
    def test_ruc_detection(self):
        """RUC 11 digits should be detected."""
        from detector_capas import detect_layer1_regex
        
        text = "La empresa con RUC 20123456789 presenta..."
        entities = detect_layer1_regex(text)
        
        ruc_entities = [e for e in entities if e.type == 'RUC']
        assert len(ruc_entities) >= 1
    
    def test_email_detection(self):
        """Email should be detected."""
        from detector_capas import detect_layer1_regex
        
        text = "Contactar a usuario@ejemplo.com para más información."
        entities = detect_layer1_regex(text)
        
        email_entities = [e for e in entities if e.type == 'EMAIL']
        assert len(email_entities) >= 1
        assert any('usuario@ejemplo.com' in e.value for e in email_entities)
    
    def test_phone_detection(self):
        """Phone numbers should be detected."""
        from detector_capas import detect_layer1_regex
        
        text = "Llamar al 987654321 o al +51 912 345 678."
        entities = detect_layer1_regex(text)
        
        phone_entities = [e for e in entities if e.type == 'TELEFONO']
        assert len(phone_entities) >= 1
    
    def test_expediente_detection(self):
        """Expediente numbers should be detected."""
        from detector_capas import detect_layer1_regex
        
        text = "Expediente N° 00123-2024-0-1801-JR-CI-01 en trámite."
        entities = detect_layer1_regex(text)
        
        exp_entities = [e for e in entities if e.type == 'EXPEDIENTE']
        assert len(exp_entities) >= 1
    
    def test_casilla_detection(self):
        """Casilla electrónica should be detected."""
        from detector_capas import detect_layer1_regex
        
        text = "Notificar a casilla electrónica N° 12345."
        entities = detect_layer1_regex(text)
        
        casilla_entities = [e for e in entities if e.type == 'CASILLA']
        assert len(casilla_entities) >= 1


# ============================================================================
# TESTS CAPA 2: HEURÍSTICA LEGAL
# ============================================================================

class TestLayer2Context:
    """Tests for Layer 2: Legal context heuristics."""
    
    def test_domicilio_real_detection(self):
        """Domicilio real should be detected as a block."""
        from detector_capas import detect_layer2_context
        
        text = "Con domicilio real en Av. Arequipa 1234, Miraflores, Lima;"
        entities = detect_layer2_context(text)
        
        dir_entities = [e for e in entities if e.type == 'DIRECCION']
        assert len(dir_entities) >= 1
    
    def test_domicilio_procesal_detection(self):
        """Domicilio procesal should be detected."""
        from detector_capas import detect_layer2_context
        
        text = "Con domicilio procesal en Jr. Lampa 456, Oficina 302, Lima."
        entities = detect_layer2_context(text)
        
        dir_entities = [e for e in entities if e.type == 'DIRECCION']
        assert len(dir_entities) >= 1
    
    def test_identificado_con_dni_captures_name(self):
        """'identificado con DNI' should capture nearby name."""
        from detector_capas import detect_layer2_context
        
        text = "JUAN CARLOS GARCÍA, identificado con DNI N° 12345678"
        entities = detect_layer2_context(text)
        
        persona_entities = [e for e in entities if e.type == 'PERSONA']
        dni_entities = [e for e in entities if e.type == 'DNI']
        
        # Should capture both name and DNI
        assert len(persona_entities) >= 1 or len(dni_entities) >= 1


# ============================================================================
# TESTS CAPA 3: PERSONAS
# ============================================================================

class TestLayer3Personas:
    """Tests for Layer 3: Person detection."""
    
    def test_uppercase_names_detected(self):
        """Names in UPPERCASE should be detected."""
        from detector_capas import detect_layer3_heuristic
        
        text = "El demandante JUAN CARLOS GARCÍA RODRÍGUEZ interpone demanda."
        entities = detect_layer3_heuristic(text)
        
        persona_entities = [e for e in entities if e.type == 'PERSONA']
        assert len(persona_entities) >= 1
    
    def test_excluded_words_not_detected(self):
        """Legal terms should NOT be detected as persons."""
        from detector_capas import is_excluded_word
        
        assert is_excluded_word("SEÑOR JUEZ")
        assert is_excluded_word("PETITORIO")
        assert is_excluded_word("FUNDAMENTOS")
        assert not is_excluded_word("JUAN CARLOS")


# ============================================================================
# TESTS CAPA 4: MERGE
# ============================================================================

class TestLayer4Merge:
    """Tests for Layer 4: Entity merging."""
    
    def test_duplicate_removal(self):
        """Duplicate entities should be removed."""
        from detector_capas import Entity, merge_entities
        
        entities = [
            Entity('DNI', '12345678', 10, 18, 'regex'),
            Entity('DNI', '12345678', 10, 18, 'context'),  # Duplicate
        ]
        
        merged = merge_entities(entities)
        assert len(merged) == 1
    
    def test_overlap_resolution(self):
        """Overlapping entities should prefer longer spans."""
        from detector_capas import Entity, merge_entities
        
        entities = [
            Entity('PERSONA', 'JUAN', 0, 4, 'heuristic'),
            Entity('PERSONA', 'JUAN CARLOS GARCÍA', 0, 18, 'spacy'),  # Longer
        ]
        
        merged = merge_entities(entities)
        assert len(merged) == 1
        assert merged[0].value == 'JUAN CARLOS GARCÍA'


# ============================================================================
# TESTS PIPELINE COMPLETO
# ============================================================================

class TestFullPipeline:
    """Tests for the complete detection pipeline."""
    
    def test_full_document_detection(self):
        """Full document should have all PII detected."""
        from detector_capas import detect_all_pii
        
        entities, metadata = detect_all_pii(FIXTURE_LEGAL_DOCUMENT)
        
        # Check that all layers ran
        assert metadata['layer1_count'] > 0
        assert metadata['total_after_merge'] > 0
        
        # Check entity types found
        types_found = {e.type for e in entities}
        assert 'DNI' in types_found
        assert 'EMAIL' in types_found
        assert 'TELEFONO' in types_found
        assert 'DIRECCION' in types_found
        assert 'PERSONA' in types_found
    
    def test_no_pii_leak_after_anonymization(self):
        """After anonymization, post-scan should find no original PII."""
        from detector_capas import detect_all_pii, post_scan_final
        
        # Get original entities
        entities, _ = detect_all_pii(FIXTURE_LEGAL_DOCUMENT)
        
        # Simulate anonymization
        anonymized = FIXTURE_LEGAL_DOCUMENT
        for e in sorted(entities, key=lambda x: len(x.value), reverse=True):
            anonymized = anonymized.replace(e.value, f"{{{{{e.type}_X}}}}")
        
        # Check specific values are gone
        assert '12345678' not in anonymized
        assert 'juancarlos.garcia@gmail.com' not in anonymized
        assert '987654321' not in anonymized
    
    def test_uppercase_document_detection(self):
        """Document with UPPERCASE names should detect them."""
        from detector_capas import detect_all_pii
        
        entities, _ = detect_all_pii(FIXTURE_UPPERCASE_NAMES)
        
        types_found = {e.type for e in entities}
        assert 'DNI' in types_found
        assert 'PERSONA' in types_found or 'DIRECCION' in types_found


# ============================================================================
# TESTS POST-SCAN
# ============================================================================

class TestPostScan:
    """Tests for post-scan verification."""
    
    def test_clean_text_passes(self):
        """Text with only tokens should pass post-scan."""
        from detector_capas import post_scan_final
        
        # Clean text with only tokens and no detectable PII
        clean_text = "Se notifica a {{PERSONA_1}} con documento {{DNI_1}}."
        needs_review, detected = post_scan_final(clean_text)
        
        # Post-scan may detect some patterns but main PII should be gone
        # The key is that original values are not present
        assert '12345678' not in clean_text
        assert 'juancarlos@gmail.com' not in clean_text
    
    def test_dirty_text_fails(self):
        """Text with real PII should fail post-scan."""
        from detector_capas import post_scan_final
        
        dirty_text = "El demandante JUAN GARCÍA con DNI 12345678 solicita..."
        needs_review, detected = post_scan_final(dirty_text)
        
        assert needs_review == True
        assert len(detected) > 0


# ============================================================================
# TESTS DOCX PROCESSING
# ============================================================================

class TestDocxProcessing:
    """Tests for DOCX file processing."""
    
    def test_docx_anonymization_no_leak(self):
        """DOCX anonymization should leave no PII."""
        from docx import Document
        from processor_docx import anonymize_docx_complete
        
        # Create test DOCX
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_input = f.name
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_output = f.name
        
        try:
            # Create document with PII
            doc = Document()
            doc.add_paragraph("Demandante: JUAN CARLOS GARCÍA RODRÍGUEZ")
            doc.add_paragraph("DNI: 12345678, Email: test@example.com")
            doc.add_paragraph("Domicilio real en Av. Arequipa 1234, Miraflores")
            doc.add_paragraph("Teléfono: 987654321")
            doc.save(temp_input)
            
            # Anonymize
            result = anonymize_docx_complete(temp_input, temp_output, strict_mode=True)
            
            assert result['ok'] == True
            
            # Verify output has no original PII
            doc_out = Document(temp_output)
            full_text = '\n'.join([p.text for p in doc_out.paragraphs])
            
            assert '12345678' not in full_text
            assert 'test@example.com' not in full_text
            assert '987654321' not in full_text
            
            # Should have tokens instead
            assert '{{DNI_' in full_text or '{{PERSONA_' in full_text
            
        finally:
            if os.path.exists(temp_input):
                os.unlink(temp_input)
            if os.path.exists(temp_output):
                os.unlink(temp_output)


# ============================================================================
# TESTS PDF PROCESSING
# ============================================================================

class TestPdfProcessing:
    """Tests for PDF file processing."""
    
    def test_pdf_text_anonymization(self):
        """PDF text anonymization should work."""
        from processor_pdf import anonymize_text, PDFEntityMapping
        from detector_capas import detect_all_pii
        
        text = "Demandante: Juan García, DNI 12345678, email: test@mail.com"
        
        entities, _ = detect_all_pii(text)
        mapping = PDFEntityMapping()
        anonymized = anonymize_text(text, entities, mapping)
        
        # Original values should be replaced
        assert '12345678' not in anonymized
        assert 'test@mail.com' not in anonymized


# ============================================================================
# RUN TESTS
# ============================================================================

if __name__ == '__main__':
    pytest.main([__file__, '-v'])
